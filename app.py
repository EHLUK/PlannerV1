"""
P6 XER Project Manager Planning Tool
=====================================
A Streamlit app for interrogating Primavera P6 XER schedules
without needing to open P6. Designed for Project Managers.
"""

import io
import re
import math
import warnings
from collections import defaultdict, deque
from datetime import datetime, timedelta

import networkx as nx
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

warnings.filterwarnings("ignore")

# -----------------------------------------------------------------------------
# PAGE CONFIG
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title="PlanTrace",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded",
)

# -----------------------------------------------------------------------------
# -----------------------------------------------------------------------------
# -----------------------------------------------------------------------------
# CUSTOM CSS  --  PlanTrace v5  Clean Production
# ONE colour story: #0B1929 navy  |  #FFFFFF white  |  #E8951D amber  |  #EBF0F5 page bg
# Red only for risk. Green only for healthy. Everything else: navy or white.
# -----------------------------------------------------------------------------
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Sans:wght@300;400;500;600;700&family=IBM+Plex+Mono:wght@400;500;600&display=swap');

/* -- RESET -- */
#MainMenu,footer,header{visibility:hidden}
[data-testid="stToolbar"]{display:none!important}
[data-testid="stDecoration"]{display:none!important}
[data-testid="stStatusWidget"]{display:none!important}
.stDeployButton{display:none!important}
section[data-testid="stSidebarNav"]{display:none!important}
.viewerBadge_container__1QSob{display:none!important}
[data-testid="InputInstructions"]{display:none!important}

/* -- ROOT -- */
html,body,.stApp{
    background:#EBF0F5!important;
    font-family:'IBM Plex Sans','Segoe UI',system-ui,sans-serif!important;
    font-size:13px;color:#1C2B3A;
    -webkit-font-smoothing:antialiased;
    text-rendering:optimizeLegibility}
[data-testid="stAppViewContainer"]{background:#EBF0F5!important}
[data-testid="block-container"]{
    padding-top:0!important;
    padding-bottom:48px!important;
    padding-left:32px!important;
    padding-right:32px!important;
    max-width:100%!important}

/* -- SIDEBAR -- */
[data-testid="stSidebar"]{
    background:#0B1929!important;
    border-right:1px solid #0d2035;
    min-width:220px!important;
    max-width:232px!important}
/* Override ALL text inside sidebar */
[data-testid="stSidebar"],[data-testid="stSidebar"] *{
    font-family:'IBM Plex Sans','Segoe UI',sans-serif!important}
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] div,
[data-testid="stSidebar"] label{color:#4A6070!important}
/* File uploader */
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"]{
    background:#0d2035!important;
    border:1px dashed #1a3550!important;
    border-radius:4px!important}
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] *{
    font-size:11px!important;color:#2D4255!important}
[data-testid="stSidebar"] [data-testid="stFileUploadDropzoneInstructions"] div{
    display:none!important}
/* Slider */
[data-testid="stSidebar"] .stSlider label{
    color:#2D4255!important;font-size:10px!important;
    font-weight:700!important;text-transform:uppercase;letter-spacing:1px}
[data-testid="stSidebar"] [data-testid="stSliderThumb"]{background:#E8951D!important}
/* Selectbox (mode toggle) */
[data-testid="stSidebar"] .stSelectbox label{display:none!important}
[data-testid="stSidebar"] .stSelectbox>div>div{
    background:#0d2035!important;border:1px solid #1a3550!important;
    border-radius:4px!important;font-size:12px!important;color:#8AAABF!important}
/* Nav buttons */
[data-testid="stSidebar"] .stButton>button{
    background:transparent!important;
    color:#4A6070!important;
    border:none!important;
    border-left:2px solid transparent!important;
    border-radius:0!important;
    text-align:left!important;
    font-size:13px!important;
    font-weight:500!important;
    padding:8px 16px 8px 14px!important;
    width:100%!important;
    box-shadow:none!important;
    letter-spacing:0.1px;
    transition:background 0.1s,color 0.1s}
[data-testid="stSidebar"] .stButton>button:hover{
    background:#0d2035!important;
    color:#8AAABF!important;
    border-left-color:#1a3550!important;
    transform:none!important;box-shadow:none!important}
[data-testid="stSidebar"] .stButton>button:focus{
    box-shadow:none!important;outline:none!important}

/* -- CONTROL BAR -- */
.ctrl-bar{
    background:#0B1929;
    padding:16px 32px 14px;
    margin:-28px -32px 24px -32px;
    border-bottom:2px solid #E8951D;
    display:flex;align-items:center;
    justify-content:space-between;
    gap:20px;flex-wrap:wrap}
.ctrl-bar-title{
    font-size:20px;font-weight:700;color:#FFFFFF;
    letter-spacing:-0.3px;line-height:1}
.ctrl-bar-desc{
    font-size:11.5px;color:#4A6070;margin-top:4px}
.ctrl-bar-meta{display:flex;gap:24px;align-items:center;flex-wrap:wrap}
.ctrl-meta-item{text-align:right}
.ctrl-meta-label{
    font-size:9px;color:#253A50;text-transform:uppercase;
    letter-spacing:1px;font-weight:700;font-family:'IBM Plex Mono',monospace}
.ctrl-meta-value{
    font-size:13px;font-weight:600;color:#4A6070;margin-top:1px;
    font-family:'IBM Plex Mono',monospace}
.ctrl-meta-value.loaded{color:#E8951D}

/* -- TABS -- */
.stTabs [data-baseweb="tab-list"]{
    background:#FFFFFF;gap:0;padding:0 4px;
    border-bottom:1px solid #D4DCE4;border-radius:0;
    box-shadow:none}
.stTabs [data-baseweb="tab"]{
    font-family:'IBM Plex Sans',sans-serif!important;
    font-size:13px;font-weight:600;color:#5A7080;
    padding:11px 20px;
    border-bottom:2px solid transparent;
    margin-bottom:-1px;
    background:transparent;border-radius:0}
.stTabs [aria-selected="true"]{
    color:#0B1929!important;
    border-bottom:2px solid #E8951D!important;
    background:transparent!important}
.stTabs [data-baseweb="tab"]:hover{
    color:#0B1929!important;background:#F3F6F9!important}
.stTabs [data-baseweb="tab-panel"]{padding:22px 0 0 0}

/* -- KPI CARDS -- */
.kpi{background:#FFFFFF;border:1px solid #D4DCE4;border-radius:6px;
     padding:16px 18px;box-shadow:0 1px 2px rgba(11,25,41,0.05)}
.kpi::before{content:'';display:block;height:3px;
             margin:-16px -18px 14px -18px;border-radius:6px 6px 0 0}
.kpi-border-top-navy::before  {background:#0B1929}
.kpi-border-top-red::before   {background:#C0392B}
.kpi-border-top-amber::before {background:#E8951D}
.kpi-border-top-green::before {background:#1E7A4E}
.kpi-border-top-blue::before  {background:#0B1929}
.kpi-label{font-size:10px;font-weight:700;color:#8A9DB0;
           letter-spacing:1.2px;text-transform:uppercase;
           margin-bottom:8px;font-family:'IBM Plex Mono',monospace}
.kpi-num{font-size:30px;font-weight:700;color:#0B1929;
         line-height:1;margin-bottom:4px;
         font-family:'IBM Plex Mono',monospace;letter-spacing:-1px}
.kpi-sub{font-size:11px;color:#7A8FA0}
.kpi-red .kpi-num   {color:#C0392B}
.kpi-amber .kpi-num {color:#B07010}
.kpi-green .kpi-num {color:#1E7A4E}
.kpi-blue .kpi-num  {color:#0B1929}

/* -- CHIPS -- */
.chip{display:inline-flex;align-items:center;gap:3px;
      padding:2px 8px;border-radius:3px;
      font-size:11px;font-weight:600;letter-spacing:0.2px;
      white-space:nowrap;font-family:'IBM Plex Sans',sans-serif}
.chip-red   {background:#FAEAEA;color:#9E2828;border:1px solid #EABABA}
.chip-amber {background:#FDF3E0;color:#8A5A00;border:1px solid #EED090}
.chip-green {background:#E4F5EC;color:#145C38;border:1px solid #A0D8B8}
.chip-blue  {background:#E8F0F8;color:#1A3A5C;border:1px solid #AABDD8}
.chip-grey  {background:#EEF1F5;color:#3A5068;border:1px solid #C8D4E0}
.chip-navy  {background:#0d2035;color:#4A6070;border:1px solid #1a3550}

/* -- DATAFRAMES -- */
.stDataFrame{border-radius:6px;border:1px solid #D4DCE4;
             overflow:hidden;box-shadow:0 1px 2px rgba(11,25,41,0.04)}
.stDataFrame thead tr th{
    background:#0B1929!important;color:#8AAABF!important;
    font-family:'IBM Plex Mono',monospace!important;
    font-size:10px!important;font-weight:600!important;
    padding:10px 12px!important;letter-spacing:0.8px!important;
    text-transform:uppercase!important;white-space:nowrap}
.stDataFrame tbody tr td{
    font-size:12.5px!important;padding:8px 12px!important;
    color:#1C2B3A!important}
.stDataFrame tbody tr:nth-child(even){background:#F5F8FB}
.stDataFrame tbody tr:hover{background:#EBF0F7}
div[data-testid="metric-container"]{
    background:#FFFFFF;border-radius:6px;padding:16px;
    border:1px solid #D4DCE4;box-shadow:0 1px 2px rgba(11,25,41,0.05)}

/* -- MAIN BUTTONS -- */
[data-testid="stAppViewContainer"] .stButton>button{
    background:#0B1929;color:#FFFFFF;
    border:none;border-radius:4px;
    font-family:'IBM Plex Sans',sans-serif!important;
    font-weight:600;font-size:13px;
    padding:8px 18px;letter-spacing:0.1px;
    transition:all 0.12s ease}
[data-testid="stAppViewContainer"] .stButton>button:hover{
    background:#1a3550;
    box-shadow:0 2px 8px rgba(11,25,41,0.25);
    transform:translateY(-1px)}
.stDownloadButton>button{
    background:#FFFFFF!important;color:#0B1929!important;
    border:1.5px solid #0B1929!important;
    border-radius:4px!important;font-weight:600!important;
    font-family:'IBM Plex Sans',sans-serif!important;
    font-size:13px!important;padding:8px 18px!important;
    transition:all 0.12s ease!important}
.stDownloadButton>button:hover{
    background:#0B1929!important;color:#FFFFFF!important;
    transform:translateY(-1px)!important}

/* -- INPUTS -- */
.stSelectbox>div>div,.stMultiSelect>div>div{
    border-radius:4px;border:1px solid #C4CED8;
    font-size:13px;background:#FFFFFF;color:#1C2B3A}
.stTextInput>div>div>input{
    border-radius:4px;border:1px solid #C4CED8;
    font-size:13px;padding:8px 12px;
    background:#FFFFFF;color:#1C2B3A}
.stTextInput>div>div>input:focus{
    border-color:#E8951D;
    box-shadow:0 0 0 3px rgba(232,149,29,0.12);outline:none}

/* -- EXPANDER -- */
.streamlit-expanderHeader{
    background:#FFFFFF;border:1px solid #D4DCE4;
    border-radius:4px;font-weight:600;font-size:13px;color:#1C2B3A;
    padding:10px 14px;font-family:'IBM Plex Sans',sans-serif!important}

/* -- ALERTS -- */
div[data-testid="stAlert"]{border-radius:4px;font-size:13px}

/* -- TYPOGRAPHY -- */
h1{display:none!important}
h2{font-size:17px!important;font-weight:700!important;color:#0B1929!important;
   margin:20px 0 10px 0!important;letter-spacing:-0.2px;
   font-family:'IBM Plex Sans',sans-serif!important}
h3{font-size:14px!important;font-weight:600!important;color:#1C2B3A!important;
   margin:14px 0 8px 0!important;font-family:'IBM Plex Sans',sans-serif!important}
p,li{font-size:13px;color:#374151;line-height:1.65;
     font-family:'IBM Plex Sans',sans-serif!important}
.stCaption,.stCaption p{font-size:11.5px!important;color:#8A9DB0!important}
hr{border:none;border-top:1px solid #D4DCE4;margin:18px 0}

/* -- UTILITY -- */
.section-label{font-size:10px;font-weight:700;color:#7A8FA0;
               letter-spacing:1.5px;text-transform:uppercase;
               margin-bottom:10px;font-family:'IBM Plex Mono',monospace}
.card{background:#FFFFFF;border:1px solid #D4DCE4;border-radius:6px;
      padding:20px;box-shadow:0 1px 2px rgba(11,25,41,0.05)}
.card-flat{background:#FFFFFF;border:1px solid #D4DCE4;border-radius:6px;padding:16px}
.attn{background:#FFFFFF;border:1px solid #D4DCE4;
      border-left:3px solid #C0392B;
      border-radius:0 4px 4px 0;padding:10px 14px;margin-bottom:6px}
.attn-amber{border-left-color:#E8951D}
.attn-blue{border-left-color:#0B1929}
.empty-state{background:#FFFFFF;border:2px dashed #C4CED8;
             border-radius:6px;padding:48px 24px;text-align:center}
.dq-row{display:flex;align-items:center;justify-content:space-between;
        padding:7px 0;border-bottom:1px solid #EEF1F5;font-size:12px}
.dq-row:last-child{border-bottom:none}
.dot-green{display:inline-block;width:6px;height:6px;border-radius:50%;
           background:#1E7A4E;margin-right:5px;vertical-align:middle}
.dot-grey{display:inline-block;width:6px;height:6px;border-radius:50%;
          background:#6B7280;margin-right:5px;vertical-align:middle}
.dot-red{display:inline-block;width:6px;height:6px;border-radius:50%;
         background:#C0392B;margin-right:5px;vertical-align:middle}
</style>
""", unsafe_allow_html=True)










# -----------------------------------------------------------------------------
# XER PARSING  (xerparser + manual fallback)
# -----------------------------------------------------------------------------

def parse_xer_fallback(raw_text: str) -> dict:
    """
    Manual fallback parser that reads XER table format:
    %T TABLE_NAME  /  %F col1 col2 ...  /  %R val1 val2 ...
    Returns dict of {table_name: list_of_dicts}
    """
    tables = {}
    current_table = None
    current_fields = []

    for line in raw_text.splitlines():
        line = line.rstrip("\r")
        if line.startswith("%T\t"):
            current_table = line[3:].strip()
            current_fields = []
            tables[current_table] = []
        elif line.startswith("%F\t") and current_table:
            current_fields = line[3:].split("\t")
        elif line.startswith("%R\t") and current_table and current_fields:
            values = line[3:].split("\t")
            # Pad values if shorter than fields
            while len(values) < len(current_fields):
                values.append("")
            row = {current_fields[i]: values[i] for i in range(len(current_fields))}
            tables[current_table].append(row)

    return tables


def hours_to_days(hours, hours_per_day=8.0):
    """Convert hours to working days."""
    if hours is None:
        return None
    try:
        return round(float(hours) / hours_per_day, 1)
    except (TypeError, ValueError):
        return None


def safe_float(val, default=None):
    try:
        return float(val)
    except (TypeError, ValueError):
        return default


def safe_date(val):
    if val is None or str(val).strip() in ("", "None"):
        return None
    if isinstance(val, datetime):
        return val
    for fmt in ("%Y-%m-%d %H:%M", "%Y-%m-%d", "%d/%m/%Y %H:%M", "%d/%m/%Y"):
        try:
            return datetime.strptime(str(val).strip(), fmt)
        except ValueError:
            pass
    return None


def parse_xer(file_bytes: bytes):
    """
    Parse an XER file. Uses xerparser library first; falls back to manual parsing.
    Returns a dict with keys: tasks_df, relationships_df, wbs_df, resources_df,
    task_resources_df, project_info, calendars_df, parse_method
    """
    # Try to decode the file
    for codec in ("cp1252", "utf-8", "latin-1"):
        try:
            raw_text = file_bytes.decode(codec)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise ValueError("Cannot decode XER file. Please check the file encoding.")

    result = {
        "tasks_df": pd.DataFrame(),
        "relationships_df": pd.DataFrame(),
        "wbs_df": pd.DataFrame(),
        "resources_df": pd.DataFrame(),
        "task_resources_df": pd.DataFrame(),
        "project_info": {},
        "calendars_df": pd.DataFrame(),
        "parse_method": "unknown",
    }

    # -- Try xerparser library -------------------------------------------------
    try:
        from xerparser.src.xer import Xer
        xer = Xer(raw_text)

        # Project info
        proj = None
        if xer.projects:
            proj_id = next(iter(xer.projects))
            proj = xer.projects[proj_id]
            result["project_info"] = {
                "name": getattr(proj, "name", ""),
                "data_date": getattr(proj, "last_recalc_date", None),
                "project_id": proj_id,
                "plan_start": getattr(proj, "plan_start_date", None),
                "scd_end": getattr(proj, "scd_end_date", None),
            }

        # Tasks DataFrame
        rows = []
        for uid, task in xer.tasks.items():
            tf = task.total_float_hr_cnt
            ff = task.free_float_hr_cnt
            # Effective start/finish (actual if done, early if not)
            eff_start = task.act_start_date or task.early_start_date or task.target_start_date
            eff_finish = task.act_end_date or task.early_end_date or task.target_end_date

            # WBS path
            wbs_node = xer.wbs_nodes.get(task.wbs_id)
            wbs_path = ""
            if wbs_node:
                parts = []
                n = wbs_node
                while n:
                    parts.append(getattr(n, "name", ""))
                    n = getattr(n, "parent", None)
                wbs_path = " > ".join(reversed(parts))

            # Calendar name
            cal = xer.calendars.get(task.clndr_id)
            cal_name = getattr(cal, "name", "") if cal else ""

            rows.append({
                "task_id": uid,
                "task_code": task.task_code,
                "task_name": task.name,
                "wbs_id": task.wbs_id,
                "wbs_path": wbs_path,
                "status": task.status.value if task.status else "",
                "task_type": task.type.value if task.type else "",
                "calendar": cal_name,
                "early_start": task.early_start_date,
                "early_finish": task.early_end_date,
                "late_start": task.late_start_date,
                "late_finish": task.late_end_date,
                "act_start": task.act_start_date,
                "act_finish": task.act_end_date,
                "target_start": task.target_start_date,
                "target_finish": task.target_end_date,
                "eff_start": eff_start,
                "eff_finish": eff_finish,
                "orig_dur_days": hours_to_days(task.target_drtn_hr_cnt),
                "rem_dur_days": hours_to_days(task.remain_drtn_hr_cnt),
                "total_float_days": hours_to_days(tf),
                "free_float_days": hours_to_days(ff),
                "total_float_hrs": tf,
                "is_longest_path": task.is_longest_path,
                "cstr_type": task.cstr_type,
                "cstr_date": task.cstr_date,
                "cstr_type2": task.cstr_type2,
                "cstr_date2": task.cstr_date2,
                "phys_pct": round(task.phys_complete_pct * 100, 1),
                "float_path": task.float_path,
            })

        result["tasks_df"] = pd.DataFrame(rows)

        # Relationships DataFrame
        rel_rows = []
        for uid, rel in xer.relationships.items():
            rel_rows.append({
                "pred_id": uid,
                "pred_task_id": rel.predecessor.uid if rel.predecessor else "",
                "pred_task_code": rel.predecessor.task_code if rel.predecessor else "",
                "pred_task_name": rel.predecessor.name if rel.predecessor else "",
                "succ_task_id": rel.successor.uid if rel.successor else "",
                "succ_task_code": rel.successor.task_code if rel.successor else "",
                "succ_task_name": rel.successor.name if rel.successor else "",
                "rel_type": rel.link,
                "lag_days": rel.lag,
                "lag_hrs": rel.lag_hr_cnt,
            })
        result["relationships_df"] = pd.DataFrame(rel_rows)

        # WBS DataFrame
        wbs_rows = []
        for uid, wbs in xer.wbs_nodes.items():
            wbs_rows.append({
                "wbs_id": uid,
                "wbs_code": getattr(wbs, "short_name", ""),
                "wbs_name": getattr(wbs, "name", ""),
                "parent_wbs_id": getattr(wbs, "parent_wbs_id", ""),
                "proj_id": getattr(wbs, "proj_id", ""),
            })
        result["wbs_df"] = pd.DataFrame(wbs_rows)

        # Resources & task resources
        if xer.resources:
            res_rows = []
            for uid, r in xer.resources.items():
                res_rows.append({
                    "rsrc_id": uid,
                    "rsrc_name": getattr(r, "name", ""),
                    "rsrc_short": getattr(r, "rsrc_short_name", ""),
                    "rsrc_type": getattr(r, "rsrc_type", ""),
                })
            result["resources_df"] = pd.DataFrame(res_rows)

        # Task resources (loading)
        taskrsrc_rows = []
        for uid, task in xer.tasks.items():
            for tr in getattr(task, "resources", []):
                taskrsrc_rows.append({
                    "task_id": uid,
                    "task_code": task.task_code,
                    "rsrc_id": getattr(tr, "rsrc_id", ""),
                    "target_qty": safe_float(getattr(tr, "target_qty", 0), 0),
                    "remain_qty": safe_float(getattr(tr, "remain_qty", 0), 0),
                    "act_reg_qty": safe_float(getattr(tr, "act_reg_qty", 0), 0),
                    "target_start": safe_date(getattr(tr, "target_start_date", None)),
                    "target_finish": safe_date(getattr(tr, "target_end_date", None)),
                })
        result["task_resources_df"] = pd.DataFrame(taskrsrc_rows)

        result["parse_method"] = "xerparser"
        return result

    except Exception as e:
        st.warning(f"xerparser failed ({e}), using fallback parser...")

    # -- Manual fallback -------------------------------------------------------
    try:
        tables = parse_xer_fallback(raw_text)
        return _build_from_raw_tables(tables)
    except Exception as e2:
        raise ValueError(f"Both parsers failed. Last error: {e2}")


def _build_from_raw_tables(tables: dict) -> dict:
    """Build result dict from raw parsed tables (fallback)."""
    result = {
        "tasks_df": pd.DataFrame(),
        "relationships_df": pd.DataFrame(),
        "wbs_df": pd.DataFrame(),
        "resources_df": pd.DataFrame(),
        "task_resources_df": pd.DataFrame(),
        "project_info": {},
        "calendars_df": pd.DataFrame(),
        "parse_method": "manual_fallback",
    }

    # Project info
    if "PROJECT" in tables and tables["PROJECT"]:
        proj = tables["PROJECT"][0]
        result["project_info"] = {
            "name": proj.get("proj_short_name", proj.get("proj_id", "")),
            "data_date": safe_date(proj.get("last_recalc_date")),
            "plan_start": safe_date(proj.get("plan_start_date")),
            "scd_end": safe_date(proj.get("scd_end_date")),
        }

    # Tasks
    if "TASK" in tables:
        df = pd.DataFrame(tables["TASK"])
        # Normalise date columns
        for col in ["early_start_date", "early_end_date", "late_start_date",
                    "late_end_date", "act_start_date", "act_end_date",
                    "target_start_date", "target_end_date", "cstr_date", "cstr_date2"]:
            if col in df.columns:
                df[col] = df[col].apply(safe_date)
        # Float
        for col in ["total_float_hr_cnt", "free_float_hr_cnt",
                    "target_drtn_hr_cnt", "remain_drtn_hr_cnt"]:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors="coerce")

        # Build normalised columns
        df["eff_start"] = df.get("act_start_date", df.get("early_start_date"))
        df["eff_finish"] = df.get("act_end_date", df.get("early_end_date"))
        df["total_float_days"] = df.get("total_float_hr_cnt", pd.Series(dtype=float)).apply(hours_to_days)
        df["free_float_days"] = df.get("free_float_hr_cnt", pd.Series(dtype=float)).apply(hours_to_days)
        df["orig_dur_days"] = df.get("target_drtn_hr_cnt", pd.Series(dtype=float)).apply(hours_to_days)
        df["rem_dur_days"] = df.get("remain_drtn_hr_cnt", pd.Series(dtype=float)).apply(hours_to_days)

        # Rename for consistency
        rename = {
            "task_id": "task_id", "task_code": "task_code",
            "task_name": "task_name", "wbs_id": "wbs_id",
            "status_code": "status", "task_type": "task_type",
            "early_start_date": "early_start", "early_end_date": "early_finish",
            "late_start_date": "late_start", "late_end_date": "late_finish",
            "act_start_date": "act_start", "act_end_date": "act_finish",
            "target_start_date": "target_start", "target_end_date": "target_finish",
            "cstr_type": "cstr_type", "cstr_date": "cstr_date",
            "cstr_type2": "cstr_type2", "cstr_date2": "cstr_date2",
            "driving_path_flag": "is_longest_path_flag",
            "phys_complete_pct": "phys_pct",
        }
        df = df.rename(columns={k: v for k, v in rename.items() if k in df.columns})
        if "is_longest_path_flag" in df.columns:
            df["is_longest_path"] = df["is_longest_path_flag"] == "Y"
        df["wbs_path"] = df.get("wbs_id", "")
        result["tasks_df"] = df

    # Relationships
    if "TASKPRED" in tables:
        df = pd.DataFrame(tables["TASKPRED"])
        if "lag_hr_cnt" in df.columns:
            df["lag_days"] = pd.to_numeric(df["lag_hr_cnt"], errors="coerce").apply(hours_to_days)
        rename_r = {"pred_type": "rel_type", "task_id": "succ_task_id",
                    "pred_task_id": "pred_task_id"}
        df = df.rename(columns={k: v for k, v in rename_r.items() if k in df.columns})
        result["relationships_df"] = df

    # WBS
    if "PROJWBS" in tables:
        df = pd.DataFrame(tables["PROJWBS"])
        rename_w = {"wbs_id": "wbs_id", "wbs_short_name": "wbs_code",
                    "wbs_name": "wbs_name", "parent_wbs_id": "parent_wbs_id"}
        df = df.rename(columns={k: v for k, v in rename_w.items() if k in df.columns})
        result["wbs_df"] = df

    # Resources
    if "RSRC" in tables:
        df = pd.DataFrame(tables["RSRC"])
        rename_rs = {"rsrc_id": "rsrc_id", "rsrc_name": "rsrc_name",
                     "rsrc_short_name": "rsrc_short"}
        df = df.rename(columns={k: v for k, v in rename_rs.items() if k in df.columns})
        result["resources_df"] = df

    # Task resources
    if "TASKRSRC" in tables:
        df = pd.DataFrame(tables["TASKRSRC"])
        for col in ["target_qty", "remain_qty", "act_reg_qty"]:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0)
        for col in ["target_start_date", "target_end_date"]:
            if col in df.columns:
                df[col] = df[col].apply(safe_date)
                df = df.rename(columns={col: col.replace("_date", "")})
        result["task_resources_df"] = df

    return result


# -----------------------------------------------------------------------------
# GRAPH BUILDING
# -----------------------------------------------------------------------------

def build_graph(tasks_df: pd.DataFrame, rels_df: pd.DataFrame) -> nx.DiGraph:
    """Build a networkx directed graph from tasks and relationships."""
    G = nx.DiGraph()
    for _, row in tasks_df.iterrows():
        G.add_node(row["task_id"], **row.to_dict())
    for _, row in rels_df.iterrows():
        if row.get("pred_task_id") and row.get("succ_task_id"):
            G.add_edge(
                row["pred_task_id"],
                row["succ_task_id"],
                rel_type=row.get("rel_type", "FS"),
                lag_days=row.get("lag_days", 0),
            )
    return G


# -----------------------------------------------------------------------------
# CRITICAL PATH HELPERS
# -----------------------------------------------------------------------------

def get_critical_threshold(tasks_df: pd.DataFrame, near_crit_days: float = 10.0):
    """Classify activities as critical / near-critical / float."""
    df = tasks_df.copy()
    df["is_critical"] = df["total_float_days"].apply(
        lambda f: f is not None and f <= 0
    )
    df["is_near_critical"] = df["total_float_days"].apply(
        lambda f: f is not None and 0 < f <= near_crit_days
    )
    return df


def float_status_badge(f):
    if f is None:
        return "-"
    elif f <= 0:
        return "🔴 Critical"
    elif f <= 10:
        return "🟡 Near-Critical"
    else:
        return "🟢 Float"


# -----------------------------------------------------------------------------
# LOGIC TRACE HELPERS
# -----------------------------------------------------------------------------

def trace_predecessors(G: nx.DiGraph, task_id: str, max_depth=100) -> list:
    """BFS backwards through predecessors. Returns list of (task_id, depth)."""
    visited = {}
    queue = deque([(task_id, 0)])
    result = []
    while queue:
        node, depth = queue.popleft()
        if node in visited or depth > max_depth:
            continue
        visited[node] = depth
        if node != task_id:
            result.append((node, depth))
        for pred in G.predecessors(node):
            if pred not in visited:
                queue.append((pred, depth + 1))
    return result


def trace_successors(G: nx.DiGraph, task_id: str, max_depth=100) -> list:
    """BFS forwards through successors."""
    visited = {}
    queue = deque([(task_id, 0)])
    result = []
    while queue:
        node, depth = queue.popleft()
        if node in visited or depth > max_depth:
            continue
        visited[node] = depth
        if node != task_id:
            result.append((node, depth))
        for succ in G.successors(node):
            if succ not in visited:
                queue.append((succ, depth + 1))
    return result


def driving_path_to_activity(
    G: nx.DiGraph,
    tasks_df: pd.DataFrame,
    rels_df: pd.DataFrame,
    target_id: str,
) -> list:
    """
    Identify the most likely driving predecessor chain into a target activity.

    Driving predecessor selection priority (in order):
      1. Lowest total float  (most constrained activity wins)
      2. On P6 longest-path / driving flag where available
      3. Latest early-finish date  (latest predecessor is usually the driver)
      4. Highest lag on the connecting relationship (more constraining)

    Returns ordered list of task_ids, from chain start -> target.
    """
    task_lookup = tasks_df.set_index("task_id").to_dict("index") if not tasks_df.empty else {}

    def _score(pred_id, succ_id):
        t  = task_lookup.get(pred_id, {})
        tf = safe_float(t.get("total_float_days"), 9999)
        finish = t.get("eff_finish")
        if finish is not None:
            try:
                finish_score = -(finish.timestamp() / 86400)
            except Exception:
                finish_score = 0
        else:
            finish_score = 0
        driving_bonus = 0 if t.get("is_longest_path", False) else 1
        lag_score = 0
        if not rels_df.empty:
            rel = rels_df[
                (rels_df.get("pred_task_id", pd.Series(dtype=str)) == pred_id) &
                (rels_df.get("succ_task_id", pd.Series(dtype=str)) == succ_id)
            ]
            if not rel.empty and "lag_days" in rel.columns:
                lag_val = safe_float(rel["lag_days"].iloc[0], 0)
                lag_score = -lag_val
        return (tf, driving_bonus, finish_score, lag_score)

    path    = [target_id]
    visited = {target_id}
    current = target_id

    for _ in range(500):
        preds = list(G.predecessors(current))
        if not preds:
            break
        unvisited = [p for p in preds if p not in visited]
        if not unvisited:
            break
        best = min(unvisited, key=lambda p: _score(p, current))
        path.insert(0, best)
        visited.add(best)
        current = best

    return path


def _all_pred_paths(
    G: nx.DiGraph,
    tasks_df: pd.DataFrame,
    target_id: str,
    max_paths: int = 8,
) -> list:
    """
    Find up to max_paths predecessor chains into target_id.
    Each chain is a list of task_ids ordered start -> target.
    Only returns chains that begin at an activity with no predecessors.
    """
    task_lookup = tasks_df.set_index("task_id").to_dict("index") if not tasks_df.empty else {}

    def _float(tid):
        return safe_float(task_lookup.get(tid, {}).get("total_float_days"), 9999)

    found_paths = []

    def dfs(node, current_path, visited_set):
        if len(found_paths) >= max_paths:
            return
        preds = [p for p in G.predecessors(node) if p not in visited_set]
        if not preds:
            found_paths.append(list(reversed(current_path)))
            return
        for pred in sorted(preds, key=_float)[:4]:
            dfs(pred, current_path + [pred], visited_set | {pred})

    dfs(target_id, [target_id], {target_id})
    return found_paths


# -----------------------------------------------------------------------------
# EXPORT HELPERS
# -----------------------------------------------------------------------------

def style_header_row(ws, row_idx, fill_color="1e3a5f", font_color="FFFFFF"):
    fill = PatternFill("solid", start_color=fill_color, fgColor=fill_color)
    font = Font(bold=True, color=font_color)
    for cell in ws[row_idx]:
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center", vertical="center")


def df_to_sheet(ws, df, sheet_title=None):
    """Write a DataFrame to an openpyxl worksheet with formatting."""
    if sheet_title:
        ws.title = sheet_title[:31]
    ws.append(list(df.columns))
    style_header_row(ws, 1)
    for r in df.itertuples(index=False):
        ws.append(list(r))
    # Auto-width
    for col_cells in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col_cells), default=10)
        ws.column_dimensions[get_column_letter(col_cells[0].column)].width = min(max_len + 4, 50)


def export_df_to_excel(sheets: dict) -> bytes:
    """sheets = {sheet_name: dataframe}. Returns Excel bytes."""
    wb = Workbook()
    first = True
    for name, df in sheets.items():
        if first:
            ws = wb.active
            first = False
        else:
            ws = wb.create_sheet()
        df_to_sheet(ws, df, name)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def format_date(d):
    if d is None:
        return "-"
    try:
        return d.strftime("%d %b %Y")
    except Exception:
        return str(d)


# -----------------------------------------------------------------------------
# PAGE: PROJECT SUMMARY
# -----------------------------------------------------------------------------

def page_project_summary(data: dict, near_crit_days: float):
    st.title("📊 Project Summary")

    proj = data["project_info"]
    tasks = data["tasks_df"]
    rels = data["relationships_df"]

    if tasks.empty:
        st.warning("No activities found in this file.")
        return

    tasks = get_critical_threshold(tasks, near_crit_days)

    # Header metrics
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Project Name", proj.get("name", "Unknown"))
    c2.metric("Data Date", format_date(proj.get("data_date")))
    c3.metric("Parse Method", data.get("parse_method", "-"))
    c4.metric("Activities", len(tasks))

    st.divider()
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("🔴 Critical", int(tasks["is_critical"].sum()))
    c2.metric("🟡 Near-Critical", int(tasks["is_near_critical"].sum()))
    neg_float = tasks["total_float_days"].apply(lambda f: f is not None and f < 0).sum()
    c3.metric("⚠️ Negative Float", int(neg_float))
    c4.metric("🔗 Relationships", len(rels))

    # Open-ended
    if not rels.empty and "pred_task_id" in rels.columns:
        tasks_with_pred = set(rels["succ_task_id"].dropna())
        tasks_with_succ = set(rels["pred_task_id"].dropna())
        task_ids = set(tasks["task_id"])
        no_pred = len(task_ids - tasks_with_pred)
        no_succ = len(task_ids - tasks_with_succ)
        c5.metric("Open-Ended Activities", no_pred + no_succ)
    else:
        c5.metric("Open-Ended Activities", "-")

    # Date range
    valid_starts = tasks["eff_start"].dropna()
    valid_finishes = tasks["eff_finish"].dropna()
    if not valid_starts.empty and not valid_finishes.empty:
        earliest = min(valid_starts)
        latest = max(valid_finishes)
        st.info(f"**Schedule Span:** {format_date(earliest)} -> {format_date(latest)}")

    # Constraint count
    constrained = tasks["cstr_type"].apply(lambda x: bool(x) and str(x).strip() not in ("", "None")).sum() if "cstr_type" in tasks.columns else 0
    st.info(f"**Constrained Activities:** {int(constrained)}")

    st.divider()

    # Charts
    tab1, tab2, tab3 = st.tabs(["Float Distribution", "Activities by WBS", "Status Breakdown"])

    with tab1:
        float_vals = tasks["total_float_days"].dropna()
        if not float_vals.empty:
            fig = px.histogram(
                float_vals, nbins=40, title="Total Float Distribution (Days)",
                labels={"value": "Float (days)", "count": "Activities"},
                color_discrete_sequence=["#2563eb"],
            )
            fig.add_vline(x=0, line_dash="dash", line_color="red", annotation_text="Critical")
            fig.add_vline(x=near_crit_days, line_dash="dot", line_color="orange",
                          annotation_text=f"Near-Critical ({near_crit_days}d)")
            st.plotly_chart(fig)

    with tab2:
        if "wbs_path" in tasks.columns:
            # Show top-level WBS only
            tasks["wbs_top"] = tasks["wbs_path"].apply(
                lambda x: str(x).split(" > ")[0] if pd.notna(x) and x else "Unknown"
            )
            wbs_counts = tasks.groupby("wbs_top").size().reset_index(name="count")
            wbs_counts = wbs_counts.sort_values("count", ascending=False).head(20)
            fig = px.bar(wbs_counts, x="count", y="wbs_top", orientation="h",
                         title="Activities by Top-Level WBS",
                         color_discrete_sequence=["#1e3a5f"])
            fig.update_layout(yaxis_title="", xaxis_title="Activity Count")
            st.plotly_chart(fig)

    with tab3:
        if "status" in tasks.columns:
            status_counts = tasks["status"].value_counts().reset_index()
            status_counts.columns = ["Status", "Count"]
            fig = px.pie(status_counts, values="Count", names="Status",
                         title="Activity Status Breakdown",
                         color_discrete_sequence=px.colors.qualitative.Set2)
            st.plotly_chart(fig)

    # Summary table
    st.subheader("Activity Summary Table")
    display_cols = ["task_code", "task_name", "wbs_path", "eff_start", "eff_finish",
                    "total_float_days", "status", "is_critical"]
    avail = [c for c in display_cols if c in tasks.columns]
    st.dataframe(tasks[avail].head(100))


# -----------------------------------------------------------------------------
# PAGE: ACTIVITY SEARCH
# -----------------------------------------------------------------------------

def _col(df: pd.DataFrame, col: str, default="-"):
    """Safely get a column value; return default if column missing or null."""
    if col not in df.index:
        return default
    val = df.get(col, default)
    if val is None or (isinstance(val, float) and math.isnan(val)):
        return default
    return val


def _float_color(f) -> str:
    """Return a hex colour string for a float value."""
    if f is None:
        return "#6b7280"
    try:
        f = float(f)
    except (TypeError, ValueError):
        return "#6b7280"
    if f < 0:
        return "#991b1b"   # dark red  - negative float
    if f == 0:
        return "#dc2626"   # red       - critical
    if f <= 10:
        return "#d97706"   # amber     - near-critical
    return "#15803d"       # green     - has float


def _status_label(status: str) -> str:
    """Convert P6 status code to a readable label."""
    mapping = {
        "TK_NotStart": "Not Started",
        "TK_Active":   "In Progress",
        "TK_Complete": "Complete",
        "Not Started": "Not Started",
        "In Progress": "In Progress",
        "Complete":    "Complete",
    }
    return mapping.get(str(status).strip(), str(status).strip() or "-")


def _status_colour(status: str) -> str:
    s = _status_label(status)
    if s == "Complete":    return "#15803d"
    if s == "In Progress": return "#2563eb"
    return "#6b7280"


def page_activity_search(data: dict, near_crit_days: float):
    """
    Activity Search page.
    Lets the user filter activities by ID, name, WBS, date range and
    critical status, then shows a full detail panel for the selected activity.
    """
    st.title("🔍 Activity Search")
    st.caption("Search and filter activities, then click any row to view its full detail.")

    tasks = data["tasks_df"]
    rels  = data["relationships_df"]

    if tasks.empty:
        st.warning("No activities found. Please upload an XER file first.")
        return

    tasks = get_critical_threshold(tasks, near_crit_days)

    # -------------------------------------------------------------------------
    # SEARCH / FILTER PANEL
    # -------------------------------------------------------------------------
    with st.expander("🔎  Search & Filter", expanded=True):
        r1c1, r1c2, r1c3 = st.columns(3)
        search_code = r1c1.text_input("Activity ID", placeholder="e.g. A1000")
        search_name = r1c2.text_input("Activity Name", placeholder="partial match")
        search_wbs  = r1c3.text_input("WBS", placeholder="partial match")

        r2c1, r2c2, r2c3 = st.columns(3)

        # Critical / float filter
        crit_filter = r2c1.selectbox(
            "Float Status",
            ["All", "Critical (float <= 0)", "Near-Critical", "Positive Float", "Negative Float"],
        )

        # Status filter - only show if column exists
        if "status" in tasks.columns:
            status_opts = ["All"] + sorted(
                [s for s in tasks["status"].dropna().unique() if str(s).strip()]
            )
        else:
            status_opts = ["All"]
        status_filter = r2c2.selectbox("Activity Status", status_opts)

        # Activity type filter
        if "task_type" in tasks.columns:
            type_opts = ["All"] + sorted(
                [t for t in tasks["task_type"].dropna().unique() if str(t).strip()]
            )
        else:
            type_opts = ["All"]
        type_filter = r2c3.selectbox("Activity Type", type_opts)

        # Date range - only show if date columns exist and have data
        date_from = date_to = None
        valid_starts  = tasks["eff_start"].dropna()  if "eff_start"  in tasks.columns else pd.Series(dtype=object)
        valid_finishes = tasks["eff_finish"].dropna() if "eff_finish" in tasks.columns else pd.Series(dtype=object)

        if not valid_starts.empty and not valid_finishes.empty:
            try:
                min_d = min(valid_starts).date()
                max_d = max(valid_finishes).date()
                dc1, dc2 = st.columns(2)
                date_from = dc1.date_input("Finish on or After", value=min_d,
                                           min_value=min_d, max_value=max_d)
                date_to   = dc2.date_input("Finish on or Before", value=max_d,
                                           min_value=min_d, max_value=max_d)
            except Exception:
                pass  # silently skip date filter if dates are unusable

    # -------------------------------------------------------------------------
    # APPLY FILTERS
    # -------------------------------------------------------------------------
    filtered = tasks.copy()

    if search_code.strip():
        if "task_code" in filtered.columns:
            filtered = filtered[
                filtered["task_code"].astype(str).str.contains(
                    search_code.strip(), case=False, na=False
                )
            ]
    if search_name.strip():
        if "task_name" in filtered.columns:
            filtered = filtered[
                filtered["task_name"].astype(str).str.contains(
                    search_name.strip(), case=False, na=False
                )
            ]
    if search_wbs.strip():
        if "wbs_path" in filtered.columns:
            filtered = filtered[
                filtered["wbs_path"].astype(str).str.contains(
                    search_wbs.strip(), case=False, na=False
                )
            ]

    if crit_filter == "Critical (float <= 0)" and "is_critical" in filtered.columns:
        filtered = filtered[filtered["is_critical"] == True]
    elif crit_filter == "Near-Critical" and "is_near_critical" in filtered.columns:
        filtered = filtered[filtered["is_near_critical"] == True]
    elif crit_filter == "Positive Float" and "total_float_days" in filtered.columns:
        filtered = filtered[filtered["total_float_days"].apply(
            lambda f: f is not None and safe_float(f, 1) > 0
        )]
    elif crit_filter == "Negative Float" and "total_float_days" in filtered.columns:
        filtered = filtered[filtered["total_float_days"].apply(
            lambda f: f is not None and safe_float(f, 0) < 0
        )]

    if status_filter != "All" and "status" in filtered.columns:
        filtered = filtered[filtered["status"] == status_filter]

    if type_filter != "All" and "task_type" in filtered.columns:
        filtered = filtered[filtered["task_type"] == type_filter]

    if date_from is not None and "eff_finish" in filtered.columns:
        filtered = filtered[
            filtered["eff_finish"].apply(
                lambda d: d is not None and hasattr(d, "date") and d.date() >= date_from
            )
        ]
    if date_to is not None and "eff_finish" in filtered.columns:
        filtered = filtered[
            filtered["eff_finish"].apply(
                lambda d: d is not None and hasattr(d, "date") and d.date() <= date_to
            )
        ]

    # -------------------------------------------------------------------------
    # RESULTS TABLE
    # -------------------------------------------------------------------------
    n_found = len(filtered)
    n_total = len(tasks)

    if n_found == 0:
        st.warning("No activities match your filters. Try broadening your search.")
        return

    # Build a clean display version of the table
    TABLE_COLS = {
        "task_code":        "Activity ID",
        "task_name":        "Activity Name",
        "wbs_path":         "WBS",
        "eff_start":        "Start",
        "eff_finish":       "Finish",
        "orig_dur_days":    "Orig Dur (d)",
        "rem_dur_days":     "Rem Dur (d)",
        "total_float_days": "Float (d)",
        "status":           "Status",
        "task_type":        "Type",
        "is_critical":      "Critical",
    }

    present_cols = {k: v for k, v in TABLE_COLS.items() if k in filtered.columns}
    display_df = filtered[list(present_cols.keys())].copy()
    display_df = display_df.rename(columns=present_cols)

    # Format date columns for readability
    for col in ["Start", "Finish"]:
        if col in display_df.columns:
            display_df[col] = display_df[col].apply(format_date)

    # Format critical flag
    if "Critical" in display_df.columns:
        display_df["Critical"] = display_df["Critical"].apply(
            lambda x: "Yes" if x else ""
        )

    # Friendly status labels
    if "Status" in display_df.columns:
        display_df["Status"] = display_df["Status"].apply(_status_label)

    st.markdown(
        f"<p style='color:#6b7280;font-size:13px;'>"
        f"Showing <strong>{n_found}</strong> of <strong>{n_total}</strong> activities"
        f"</p>",
        unsafe_allow_html=True,
    )

    st.dataframe(
        display_df,
        height=min(400, 45 + n_found * 35),
        hide_index=True,
    )

    # -------------------------------------------------------------------------
    # ACTIVITY SELECTOR  -  pick from the filtered results
    # -------------------------------------------------------------------------
    st.divider()

    st.markdown(
        '<div style="font-size:10px;font-weight:700;color:#9CA3AF;' +
        'text-transform:uppercase;letter-spacing:1px;margin-bottom:8px;">' +
        'Activity Detail</div>',
        unsafe_allow_html=True,
    )

    if not filtered.empty:
        render_selected_activity_panel(
            tasks_df=filtered,
            rels_df=rels,
            near_crit_days=near_crit_days,
            context_key="activity_search",
        )
    else:
        st.info("No activities match your filters.")


def _rel_label(code: str) -> str:
    return REL_TYPE_LABELS.get(str(code).strip(), str(code).strip() or "-")


def _crit_flag(tf) -> str:
    """Return a text critical flag suitable for display and export."""
    if tf is None:
        return "-"
    try:
        f = float(tf)
    except (TypeError, ValueError):
        return "-"
    if f < 0:
        return "Negative Float"
    if f == 0:
        return "Critical"
    if f <= 10:
        return "Near-Critical"
    return "Float"


def _build_full_trace_df(
    G: nx.DiGraph,
    rels_df: pd.DataFrame,
    task_lookup: dict,
    selected_id: str,
    trace_list: list,     # [(task_id, depth), ...]
    direction: str,       # "pred" | "succ" | "both"
) -> pd.DataFrame:
    """
    Build the trace results DataFrame.

    For each activity in trace_list, look up the relationship that
    connects it to the selected activity (or to the activity one step
    closer in the chain) so we can display the correct link type and lag.

    direction:
        "pred"  - activities are predecessors; depth counts backwards  (1 = direct pred)
        "succ"  - activities are successors;   depth counts forwards   (1 = direct succ)
        "both"  - mixed: negative depth = pred side, positive = succ side
    """
    rows = []

    for tid, depth in trace_list:
        t = task_lookup.get(tid, {})
        tf = t.get("total_float_days")

        # ---- find the relationship between this activity and the one at
        #      depth-1 in the chain (i.e. the step that led here) ----------
        rel_type_str = "-"
        lag_val = 0

        if not rels_df.empty:
            if direction in ("pred", "both") and depth > 0:
                # this activity is a predecessor of something; find its
                # outgoing relationship toward the selected direction
                mask = (
                    (rels_df.get("pred_task_id", pd.Series(dtype=str)) == tid) |
                    (rels_df.get("succ_task_id", pd.Series(dtype=str)) == tid)
                )
                rel_candidates = rels_df[mask]
            elif direction == "succ":
                mask = (
                    (rels_df.get("pred_task_id", pd.Series(dtype=str)) == tid) |
                    (rels_df.get("succ_task_id", pd.Series(dtype=str)) == tid)
                )
                rel_candidates = rels_df[mask]
            else:
                rel_candidates = pd.DataFrame()

            if not rel_candidates.empty and "rel_type" in rel_candidates.columns:
                rel_type_str = _rel_label(rel_candidates["rel_type"].iloc[0])
                lag_val = rel_candidates["lag_days"].iloc[0] if "lag_days" in rel_candidates.columns else 0
                try:
                    lag_val = float(lag_val) if lag_val is not None else 0
                except (TypeError, ValueError):
                    lag_val = 0

        rows.append({
            "Depth": depth,
            "Direction": "Predecessor" if depth < 0 or direction == "pred"
                         else ("Successor" if direction == "succ" else "Both"),
            "Activity ID":    t.get("task_code", tid),
            "Activity Name":  t.get("task_name", ""),
            "Link Type":      rel_type_str,
            "Lag (days)":     lag_val,
            "Start":          format_date(t.get("eff_start")),
            "Finish":         format_date(t.get("eff_finish")),
            "Total Float (d)": tf if tf is not None else "-",
            "Status":         _status_label(str(t.get("status", ""))),
            "Critical Flag":  _crit_flag(tf),
        })

    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.sort_values("Depth").reset_index(drop=True)
    return df


def _summary_bar(label: str, value: int, colour: str) -> str:
    return (
        f'<div style="display:inline-block;background:{colour};color:white;'
        f'border-radius:8px;padding:8px 18px;margin:4px 6px 4px 0;font-size:13px;">'
        f'<strong>{value}</strong> {label}</div>'
    )


def page_logic_trace(data: dict, near_crit_days: float):
    """
    Logic Trace page.
    Search and select any activity, then trace its predecessor and successor
    chains through the schedule network. Results show depth, link type, lag,
    dates, float, WBS and critical status with filters and Excel export.
    """
    st.title("🔗 Logic Trace")
    st.caption(
        "Select an activity then trace its logic through the schedule network. "
        "Predecessors flow INTO the activity. Successors flow OUT of it."
    )

    tasks = data.get("tasks_df", pd.DataFrame())
    rels  = data.get("relationships_df", pd.DataFrame())

    # Quick-access SAP panel in expander
    with st.expander("Selected Activity Summary", expanded=False, key="sap_exp_logic"):
        render_selected_activity_panel(
            tasks_df=tasks,
            rels_df=rels,
            near_crit_days=near_crit_days,
            context_key="logic_trace_sap",
        )

    tasks = data["tasks_df"]
    rels  = data["relationships_df"]

    # -------------------------------------------------------------------------
    # GUARD RAILS
    # -------------------------------------------------------------------------
    if tasks.empty:
        st.warning("No activities found in this programme.")
        return

    tasks = get_critical_threshold(tasks, near_crit_days)

    if rels.empty:
        st.warning(
            "**No relationship data found in this XER file.** "
            "Logic tracing requires both activities and relationships. "
            "Check the XER was exported with relationships included."
        )
        st.dataframe(
            tasks[["task_code", "task_name", "total_float_days", "status"]].head(50),
            hide_index=True,
        )
        return

    # -------------------------------------------------------------------------
    # BUILD NETWORK GRAPH
    # -------------------------------------------------------------------------
    G           = build_graph(tasks, rels)
    task_lookup = tasks.set_index("task_id").to_dict("index")

    # -------------------------------------------------------------------------
    # ACTIVITY SEARCH & SELECTOR
    # -------------------------------------------------------------------------
    st.markdown(
        '<div style="font-size:12px;font-weight:700;color:#94A3B8;'
        'letter-spacing:1px;text-transform:uppercase;margin-bottom:6px;">'
        'Find Activity</div>',
        unsafe_allow_html=True,
    )

    search_col, select_col = st.columns([1, 2])

    with search_col:
        search_text = st.text_input(
            "Search by ID or name",
            placeholder="e.g. A1000 or Excavation",
            key="lt_search",
            label_visibility="collapsed",
        )

    # Filter task list by search text
    if search_text.strip():
        mask = (
            tasks["task_code"].astype(str).str.contains(search_text.strip(), case=False, na=False) |
            tasks["task_name"].astype(str).str.contains(search_text.strip(), case=False, na=False)
        )
        filtered_tasks = tasks[mask]
        if filtered_tasks.empty:
            st.warning(f"No activities match '{search_text}'. Showing all activities.")
            filtered_tasks = tasks
    else:
        filtered_tasks = tasks

    with select_col:
        def _act_label(r):
            tf   = r.get("total_float_days")
            flag = " [CRITICAL]" if (tf is not None and safe_float(tf, 1) <= 0) else ""
            return f"{r.get('task_code','?')}  --  {r.get('task_name','?')}{flag}"

        act_labels = filtered_tasks.apply(_act_label, axis=1).tolist()

        if not act_labels:
            st.warning("No activities to select.")
            return

        selected_label = st.selectbox(
            "Select activity",
            options=act_labels,
            key="logic_trace_selector",
            label_visibility="collapsed",
        )

    sel_idx      = act_labels.index(selected_label)
    selected_row = filtered_tasks.iloc[sel_idx]
    selected_id  = selected_row["task_id"]
    sel_code     = str(selected_row.get("task_code", "-"))
    sel_name     = str(selected_row.get("task_name", "-"))

    # Clear results when activity changes
    if st.session_state.get("_trace_last_id") != selected_id:
        for k in ("trace_df", "trace_label", "trace_direction"):
            st.session_state.pop(k, None)
        st.session_state["_trace_last_id"] = selected_id

    # -------------------------------------------------------------------------
    # SELECTED ACTIVITY SUMMARY CARD
    # -------------------------------------------------------------------------
    sel_tf   = safe_float(selected_row.get("total_float_days"), None) if "total_float_days" in selected_row.index else None
    sel_ff   = safe_float(selected_row.get("free_float_days"),  None) if "free_float_days"  in selected_row.index else None
    sel_fcol = _float_color(sel_tf)
    sel_crit = bool(selected_row.get("is_critical", False)) if "is_critical" in selected_row.index else False
    sel_stat = _status_label(str(selected_row.get("status", "")))
    sel_scol = _status_colour(str(selected_row.get("status", "")))
    sel_wbs  = str(selected_row.get("wbs_path", "-")) if "wbs_path" in selected_row.index else "-"
    sel_start  = format_date(selected_row.get("eff_start")  if "eff_start"  in selected_row.index else None)
    sel_finish = format_date(selected_row.get("eff_finish") if "eff_finish" in selected_row.index else None)

    crit_pill = (
        '<span style="background:#dc2626;color:white;padding:2px 10px;'
        'border-radius:12px;font-size:11px;font-weight:700;margin-left:8px;">CRITICAL</span>'
        if sel_crit else ""
    )

    # Card layout: left = identity, right = schedule metrics
    card_left = f"""
        <div style="flex:1;min-width:0;">
            <div style="font-size:11px;color:#93c5fd;font-weight:600;
                        letter-spacing:1px;text-transform:uppercase;
                        margin-bottom:6px;">Selected Activity</div>
            <div style="font-size:20px;font-weight:800;white-space:nowrap;
                        overflow:hidden;text-overflow:ellipsis;">
                {sel_code}{crit_pill}
            </div>
            <div style="font-size:14px;color:#bfdbfe;margin-top:4px;
                        white-space:nowrap;overflow:hidden;text-overflow:ellipsis;"
                 title="{sel_name}">{sel_name}</div>
            <div style="font-size:12px;color:#64748B;margin-top:4px;
                        white-space:nowrap;overflow:hidden;text-overflow:ellipsis;"
                 title="{sel_wbs}">{sel_wbs}</div>
            <div style="margin-top:10px;display:flex;gap:6px;flex-wrap:wrap;">
                <span style="background:{sel_scol};color:white;padding:3px 10px;
                             border-radius:12px;font-size:11px;">{sel_stat}</span>
                <span style="background:{sel_fcol};color:white;padding:3px 10px;
                             border-radius:12px;font-size:11px;">
                    Float: {sel_tf if sel_tf is not None else "-"} days
                </span>
            </div>
        </div>"""

    card_right = f"""
        <div style="display:grid;grid-template-columns:1fr 1fr;gap:8px;
                    min-width:200px;align-self:center;">
            <div style="background:#0d2640;border-radius:6px;padding:8px 12px;">
                <div style="font-size:9px;color:#64748B;text-transform:uppercase;
                            letter-spacing:0.8px;">Start</div>
                <div style="font-size:13px;font-weight:600;color:#e2e8f0;
                            margin-top:2px;">{sel_start}</div>
            </div>
            <div style="background:#0d2640;border-radius:6px;padding:8px 12px;">
                <div style="font-size:9px;color:#64748B;text-transform:uppercase;
                            letter-spacing:0.8px;">Finish</div>
                <div style="font-size:13px;font-weight:600;color:#e2e8f0;
                            margin-top:2px;">{sel_finish}</div>
            </div>
            <div style="background:#0d2640;border-radius:6px;padding:8px 12px;">
                <div style="font-size:9px;color:#64748B;text-transform:uppercase;
                            letter-spacing:0.8px;">Total Float</div>
                <div style="font-size:15px;font-weight:700;color:#E8951D;
                            margin-top:2px;">{(str(sel_tf) + "d") if sel_tf is not None else "-"}</div>
            </div>
            <div style="background:#0d2640;border-radius:6px;padding:8px 12px;">
                <div style="font-size:9px;color:#64748B;text-transform:uppercase;
                            letter-spacing:0.8px;">Free Float</div>
                <div style="font-size:13px;font-weight:600;color:#e2e8f0;
                            margin-top:2px;">{(str(sel_ff) + "d") if sel_ff is not None else "-"}</div>
            </div>
        </div>"""

    st.markdown(
        f"""
        <div style="background:#1e3a5f;color:white;border-radius:12px;
                    padding:20px 24px;margin:10px 0 20px 0;
                    display:flex;gap:24px;align-items:flex-start;flex-wrap:wrap;">
            {card_left}
            {card_right}
        </div>
        """,
        unsafe_allow_html=True,
    )

    # -------------------------------------------------------------------------
    # OPEN-END WARNINGS
    # -------------------------------------------------------------------------
    direct_preds = list(G.predecessors(selected_id))
    direct_succs = list(G.successors(selected_id))
    has_preds    = len(direct_preds) > 0
    has_succs    = len(direct_succs) > 0

    if not has_preds and not has_succs:
        st.error(
            f"**{sel_code} has no logic connections.** "
            "This activity is completely isolated in the programme network - "
            "it has neither predecessors nor successors."
        )
        return

    if not has_preds:
        st.warning(
            f"**Open Start:** {sel_code} has no predecessors. "
            "It is not driven by any logic and may have artificially high float."
        )
    if not has_succs:
        st.warning(
            f"**Open Finish:** {sel_code} has no successors. "
            "Nothing in the programme is dependent on it completing."
        )

    # -------------------------------------------------------------------------
    # QUICK STATS ROW
    # -------------------------------------------------------------------------
    all_pred_ids = [tid for tid, _ in trace_predecessors(G, selected_id)]
    all_succ_ids = [tid for tid, _ in trace_successors(G,  selected_id)]

    def count_crit_in(id_list):
        return sum(
            1 for tid in id_list
            if safe_float(task_lookup.get(tid, {}).get("total_float_days"), 1) <= 0
        )

    st.markdown(
        _summary_bar(f"direct predecessors",   len(direct_preds), "#1d4ed8") +
        _summary_bar(f"direct successors",     len(direct_succs), "#1d4ed8") +
        _summary_bar(f"total predecessors",    len(all_pred_ids), "#4338ca") +
        _summary_bar(f"total successors",      len(all_succ_ids), "#4338ca") +
        (_summary_bar("critical in pred network", count_crit_in(all_pred_ids), "#dc2626") if all_pred_ids else "") +
        (_summary_bar("critical in succ network", count_crit_in(all_succ_ids), "#dc2626") if all_succ_ids else ""),
        unsafe_allow_html=True,
    )
    st.markdown("<br>", unsafe_allow_html=True)

    # -------------------------------------------------------------------------
    # TRACE BUTTONS
    # -------------------------------------------------------------------------
    st.markdown(
        '<div style="font-size:12px;font-weight:700;color:#94A3B8;'
        'letter-spacing:1px;text-transform:uppercase;margin-bottom:8px;">'
        'Choose what to trace</div>',
        unsafe_allow_html=True,
    )

    b1, b2, b3, b4 = st.columns(4)
    btn_dir_pred = b1.button(
        "◀  Direct Predecessors",
        key="btn_dp", disabled=not has_preds,
        help="Show only the activities that directly link into this activity (depth 1).",
    )
    btn_dir_succ = b2.button(
        "▶  Direct Successors",
        key="btn_ds", disabled=not has_succs,
        help="Show only the activities this activity directly links into (depth 1).",
    )
    btn_all_pred = b3.button(
        "◀◀  All Predecessors",
        key="btn_ap", disabled=not has_preds,
        help="Trace the full predecessor network back through all levels.",
    )
    btn_all_succ = b4.button(
        "▶▶  All Successors",
        key="btn_as", disabled=not has_succs,
        help="Trace the full successor network forward through all levels.",
    )

    # -------------------------------------------------------------------------
    # HANDLE BUTTONS
    # -------------------------------------------------------------------------
    new_trace = None
    new_label = None
    new_dir   = None

    if btn_dir_pred:
        new_trace = [(p, 1) for p in direct_preds]
        new_label = f"Direct Predecessors of {sel_code}"
        new_dir   = "pred"
    elif btn_dir_succ:
        new_trace = [(s, 1) for s in direct_succs]
        new_label = f"Direct Successors of {sel_code}"
        new_dir   = "succ"
    elif btn_all_pred:
        new_trace = trace_predecessors(G, selected_id)
        new_label = f"All Predecessors of {sel_code}"
        new_dir   = "pred"
    elif btn_all_succ:
        new_trace = trace_successors(G, selected_id)
        new_label = f"All Successors of {sel_code}"
        new_dir   = "succ"

    if new_trace is not None:
        raw_df = _build_full_trace_df(G, rels, task_lookup, selected_id, new_trace, new_dir)

        # Enrich with WBS column
        if not raw_df.empty and "Activity ID" in raw_df.columns:
            code_to_wbs = tasks.set_index("task_code")["wbs_path"].to_dict() if "wbs_path" in tasks.columns else {}
            raw_df.insert(
                raw_df.columns.get_loc("Activity Name") + 1,
                "WBS",
                raw_df["Activity ID"].map(code_to_wbs).fillna("-"),
            )

        st.session_state["trace_df"]        = raw_df
        st.session_state["trace_label"]     = new_label
        st.session_state["trace_direction"] = new_dir

    # -------------------------------------------------------------------------
    # RESULTS
    # -------------------------------------------------------------------------
    if "trace_df" not in st.session_state or st.session_state["trace_df"].empty:
        st.markdown(
            '<div style="background:#f0f9ff;border:2px dashed #93c5fd;border-radius:10px;'
            'padding:32px;text-align:center;color:#1e40af;margin-top:20px;">'
            '<div style="font-size:28px;margin-bottom:10px;">🔗</div>'
            '<strong style="font-size:15px;">Press a button above to trace the logic</strong><br>'
            '<span style="font-size:13px;color:#64748B;margin-top:6px;display:block;">'
            'Results will appear here.</span>'
            '</div>',
            unsafe_allow_html=True,
        )
        return

    trace_df  = st.session_state["trace_df"]
    trace_lbl = st.session_state.get("trace_label", "Trace Results")
    trace_dir = st.session_state.get("trace_direction", "pred")

    st.divider()

    # Counts
    n_res  = len(trace_df)
    n_crit = int((trace_df["Critical Flag"] == "Critical").sum())
    n_neg  = int((trace_df["Critical Flag"] == "Negative Float").sum())
    n_near = int((trace_df["Critical Flag"] == "Near-Critical").sum())

    st.markdown(
        f"<h4 style='color:#0B1F33;margin-bottom:6px;'>{trace_lbl}</h4>",
        unsafe_allow_html=True,
    )
    st.markdown(
        _summary_bar(f"activities",      n_res,  "#374151") +
        (_summary_bar("critical",        n_crit, "#dc2626") if n_crit else "") +
        (_summary_bar("negative float",  n_neg,  "#7f1d1d") if n_neg  else "") +
        (_summary_bar("near-critical",   n_near, "#d97706") if n_near else ""),
        unsafe_allow_html=True,
    )
    st.markdown("<br>", unsafe_allow_html=True)

    if trace_dir == "pred":
        st.caption("Depth = steps back from the selected activity. Depth 1 = directly connected.")
    elif trace_dir == "succ":
        st.caption("Depth = steps forward from the selected activity. Depth 1 = directly connected.")

    # -------------------------------------------------------------------------
    # RESULT FILTERS
    # -------------------------------------------------------------------------
    with st.expander("Filter results", expanded=False):
        fc1, fc2, fc3 = st.columns(3)

        filt_flag = fc1.selectbox(
            "Float status",
            ["All", "Critical & Negative Float", "Near-Critical", "Has Float"],
            key="lt_filt_flag",
        )

        filt_wbs = fc2.text_input(
            "WBS contains",
            placeholder="e.g. Civil",
            key="lt_filt_wbs",
        )

        filt_float_max = fc3.number_input(
            "Max float (days)",
            value=9999,
            step=1,
            min_value=-9999,
            key="lt_filt_float",
            help="Show only activities with total float at or below this value.",
        )

        # Date range filter
        fd1, fd2 = st.columns(2)
        filt_finish_from = fd1.text_input(
            "Finish on or after (dd/mm/yyyy)",
            placeholder="01/01/2024",
            key="lt_filt_df",
        )
        filt_finish_to = fd2.text_input(
            "Finish on or before (dd/mm/yyyy)",
            placeholder="31/12/2025",
            key="lt_filt_dt",
        )

    # Apply filters to a copy of trace_df
    display_df = trace_df.copy()

    if filt_flag == "Critical & Negative Float":
        display_df = display_df[display_df["Critical Flag"].isin(["Critical", "Negative Float"])]
    elif filt_flag == "Near-Critical":
        display_df = display_df[display_df["Critical Flag"] == "Near-Critical"]
    elif filt_flag == "Has Float":
        display_df = display_df[display_df["Critical Flag"] == "Float"]

    if "WBS" in display_df.columns and filt_wbs.strip():
        display_df = display_df[
            display_df["WBS"].astype(str).str.contains(filt_wbs.strip(), case=False, na=False)
        ]

    if filt_float_max < 9999 and "Total Float (d)" in display_df.columns:
        display_df = display_df[
            display_df["Total Float (d)"].apply(
                lambda v: safe_float(v, 9999) <= filt_float_max
            )
        ]

    def _parse_date_filter(s):
        for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"):
            try:
                return datetime.strptime(s.strip(), fmt)
            except Exception:
                pass
        return None

    if filt_finish_from.strip() and "Finish" in display_df.columns:
        d_from = _parse_date_filter(filt_finish_from)
        if d_from:
            # Finish column is a formatted string - compare via task lookup
            finish_dates = {}
            for _, row in tasks.iterrows():
                finish_dates[str(row.get("task_code",""))] = row.get("eff_finish")
            display_df = display_df[
                display_df["Activity ID"].apply(
                    lambda code: (lambda fd: fd is not None and fd >= d_from)(finish_dates.get(code))
                )
            ]

    if filt_finish_to.strip() and "Finish" in display_df.columns:
        d_to = _parse_date_filter(filt_finish_to)
        if d_to:
            finish_dates = {str(r.get("task_code","")): r.get("eff_finish") for _, r in tasks.iterrows()}
            display_df = display_df[
                display_df["Activity ID"].apply(
                    lambda code: (lambda fd: fd is not None and fd <= d_to)(finish_dates.get(code))
                )
            ]

    n_display = len(display_df)
    if n_display < n_res:
        st.caption(f"Showing {n_display} of {n_res} activities after filters.")

    if display_df.empty:
        st.info("No activities match the current filters.")
    else:
        # Styled table
        def _colour_flag(val):
            return {
                "Critical":       "background-color:#fee2e2;color:#991b1b;font-weight:600;",
                "Negative Float": "background-color:#fecaca;color:#7f1d1d;font-weight:700;",
                "Near-Critical":  "background-color:#fef3c7;color:#92400e;font-weight:600;",
                "Float":          "background-color:#dcfce7;color:#166534;",
            }.get(val, "")

        styled = display_df.style.map(_colour_flag, subset=["Critical Flag"])
        st.dataframe(styled, hide_index=True,
                     height=min(600, 45 + n_display * 35))

    # -------------------------------------------------------------------------
    # TABS: FILTERED VIEWS
    # -------------------------------------------------------------------------
    if n_res > 5:
        st.markdown("<br>", unsafe_allow_html=True)
        tab_all, tab_crit, tab_near, tab_open = st.tabs(
            ["All Results", "Critical & Neg Float", "Near-Critical", "Open Ends in Chain"]
        )

        with tab_all:
            st.dataframe(trace_df, hide_index=True)

        with tab_crit:
            crit_df = trace_df[trace_df["Critical Flag"].isin(["Critical", "Negative Float"])]
            if crit_df.empty:
                st.success("No critical or negative float activities in this trace.")
            else:
                st.dataframe(crit_df, hide_index=True)

        with tab_near:
            near_df = trace_df[trace_df["Critical Flag"] == "Near-Critical"]
            if near_df.empty:
                st.success("No near-critical activities in this trace.")
            else:
                st.dataframe(near_df, hide_index=True)

        with tab_open:
            open_rows = []
            for act_code in trace_df["Activity ID"].unique():
                match = tasks[tasks["task_code"] == act_code]
                if match.empty:
                    continue
                mid  = match.iloc[0]["task_id"]
                no_p = len(list(G.predecessors(mid))) == 0
                no_s = len(list(G.successors(mid))) == 0
                if no_p or no_s:
                    t = task_lookup.get(mid, {})
                    open_rows.append({
                        "Activity ID":   t.get("task_code", mid),
                        "Activity Name": t.get("task_name", ""),
                        "Issue": ("No predecessors" if no_p else "") +
                                 (" | No successors" if no_s else ""),
                        "Float (d)":     t.get("total_float_days", "-"),
                    })
            if open_rows:
                st.dataframe(pd.DataFrame(open_rows), hide_index=True)
            else:
                st.success("No open-ended activities found in this trace.")

    # -------------------------------------------------------------------------
    # GANTT CHART
    # -------------------------------------------------------------------------
    st.divider()
    st.markdown(
        '<div style="font-size:12px;font-weight:700;color:#94A3B8;'
        'letter-spacing:1px;text-transform:uppercase;margin-bottom:8px;">'
        'Trace Timeline</div>',
        unsafe_allow_html=True,
    )

    gantt_data = trace_df.copy().merge(
        tasks[["task_code","eff_start","eff_finish"]].rename(columns={"task_code":"Activity ID"}),
        on="Activity ID", how="left",
    ).dropna(subset=["eff_start","eff_finish"])

    if not gantt_data.empty:
        gantt_data["Label"] = (
            gantt_data["Activity ID"].astype(str) + "  " +
            gantt_data["Activity Name"].astype(str).str[:38]
        )
        gantt_data["Bar Colour"] = gantt_data["Critical Flag"].map({
            "Critical":       "Critical",
            "Negative Float": "Critical",
            "Near-Critical":  "Near-Critical",
            "Float":          "Has Float",
        }).fillna("Has Float")

        fig = px.timeline(
            gantt_data.head(80),
            x_start="eff_start", x_end="eff_finish", y="Label",
            color="Bar Colour",
            color_discrete_map={
                "Critical":     "#dc2626",
                "Near-Critical":"#d97706",
                "Has Float":    "#2563eb",
            },
            title=f"Timeline: {trace_lbl}",
            labels={"Label": ""},
        )
        fig.update_yaxes(autorange="reversed")
        fig.add_vline(
            x=datetime.now().strftime("%Y-%m-%dT%H:%M:%S"), line_dash="dot", line_color="#94A3B8",
            annotation_text="Today", annotation_position="top left",
        )
        fig.update_layout(
            legend_title_text="Float Status",
            height=max(300, min(700, 60 + len(gantt_data.head(80)) * 28)),
            margin=dict(l=10, r=10, t=40, b=10),
            plot_bgcolor="#ffffff",
            paper_bgcolor="#ffffff",
        )
        st.plotly_chart(fig)
        if len(gantt_data) > 80:
            st.caption(f"Showing first 80 of {len(gantt_data)} activities on the timeline.")
    else:
        st.info("No date data available to build the timeline.")

    # -------------------------------------------------------------------------
    # EXCEL EXPORT
    # -------------------------------------------------------------------------
    st.divider()

    summary_data = {
        "Item": [
            "Selected Activity ID", "Activity Name", "WBS",
            "Trace Type", "Total in chain",
            "Critical in chain", "Near-Critical in chain", "Negative Float in chain",
        ],
        "Value": [
            sel_code, sel_name, sel_wbs,
            trace_lbl, n_res, n_crit, n_near, n_neg,
        ],
    }

    export_sheets = {"Summary": pd.DataFrame(summary_data), "Full Trace": trace_df}

    crit_exp = trace_df[trace_df["Critical Flag"].isin(["Critical","Negative Float"])]
    near_exp = trace_df[trace_df["Critical Flag"] == "Near-Critical"]
    if not crit_exp.empty:
        export_sheets["Critical Activities"] = crit_exp
    if not near_exp.empty:
        export_sheets["Near-Critical"]       = near_exp
    if n_display < n_res and not display_df.empty:
        export_sheets["Filtered View"] = display_df

    xls_bytes = export_df_to_excel(export_sheets)

    st.download_button(
        label="📥  Export Logic Trace to Excel",
        data=xls_bytes,
        file_name=f"logic_trace_{sel_code}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        help="Exports Summary, Full Trace, Critical Activities, Near-Critical and any filtered view.",
        key="dl_001",
    )


# -----------------------------------------------------------------------------
# PAGE: CRITICAL PATH ANALYSIS
# -----------------------------------------------------------------------------

def page_critical_path(data: dict, near_crit_days: float):
    st.title("🚨 Critical Path Analysis")

    tasks = data.get("tasks_df", pd.DataFrame())
    rels  = data.get("relationships_df", pd.DataFrame())

    # Quick-access SAP panel in expander
    with st.expander("Activity Detail Panel", expanded=False, key="sap_exp_crit"):
        render_selected_activity_panel(
            tasks_df=tasks,
            rels_df=rels,
            near_crit_days=near_crit_days,
            context_key="critical_path_sap",
        )

    tasks = data["tasks_df"]
    rels = data["relationships_df"]

    if tasks.empty:
        st.warning("No activities loaded.")
        return

    tasks = get_critical_threshold(tasks, near_crit_days)

    tab1, tab2, tab3, tab4 = st.tabs(
        ["Critical Activities", "Near-Critical", "Negative Float", "By WBS / Package"]
    )

    with tab1:
        critical = tasks[tasks["is_critical"]].sort_values("total_float_days")
        st.metric("Critical Activities", len(critical))
        disp = ["task_code", "task_name", "wbs_path", "eff_start", "eff_finish",
                "total_float_days", "status"]
        avail = [c for c in disp if c in critical.columns]
        st.dataframe(critical[avail])

        if not critical.empty and "eff_start" in critical.columns:
            st.subheader("Critical Path Gantt")
            gantt_df = critical.dropna(subset=["eff_start", "eff_finish"]).copy()
            gantt_df["Start"] = gantt_df["eff_start"]
            gantt_df["Finish"] = gantt_df["eff_finish"]
            gantt_df["Task"] = gantt_df["task_code"] + " - " + gantt_df["task_name"]
            if len(gantt_df) > 0:
                fig = px.timeline(
                    gantt_df.head(50),
                    x_start="Start", x_end="Finish", y="Task",
                    title="Critical Path Activities (top 50)",
                    color_discrete_sequence=["#dc2626"],
                )
                fig.update_yaxes(autorange="reversed")
                st.plotly_chart(fig)

        xls = export_df_to_excel({"Critical Path": critical[avail]})
        st.download_button("📥 Export Critical Path", xls,
                           "critical_path.xlsx",
                           key="dl_002",
                           )

    with tab2:
        near_crit = tasks[tasks["is_near_critical"]].sort_values("total_float_days")
        st.metric(f"Near-Critical (0 < float <= {near_crit_days}d)", len(near_crit))
        avail = [c for c in ["task_code","task_name","wbs_path","eff_start","eff_finish",
                              "total_float_days","status"] if c in near_crit.columns]
        st.dataframe(near_crit[avail])

    with tab3:
        neg = tasks[tasks["total_float_days"].apply(lambda f: f is not None and f < 0)].sort_values("total_float_days")
        st.metric("Negative Float Activities", len(neg))
        if not neg.empty:
            st.warning("⚠️ Activities with negative float indicate the schedule cannot be met -- investigate immediately.")
            avail = [c for c in ["task_code","task_name","total_float_days","eff_start",
                                  "eff_finish","status"] if c in neg.columns]
            st.dataframe(neg[avail])

    with tab4:
        if "wbs_path" not in tasks.columns:
            st.info("WBS data not available.")
            return
        tasks["wbs_top"] = tasks["wbs_path"].apply(
            lambda x: str(x).split(" > ")[0] if pd.notna(x) else "Unknown"
        )
        wbs_crit = tasks.groupby("wbs_top").agg(
            total=("task_id", "count"),
            critical=("is_critical", "sum"),
            near_critical=("is_near_critical", "sum"),
        ).reset_index()
        wbs_crit["crit_%"] = (wbs_crit["critical"] / wbs_crit["total"] * 100).round(1)
        fig = px.bar(
            wbs_crit, x="wbs_top", y=["critical", "near_critical"],
            title="Critical & Near-Critical by WBS",
            labels={"value": "Activities", "wbs_top": "WBS"},
            color_discrete_map={"critical": "#dc2626", "near_critical": "#f59e0b"},
            barmode="group",
        )
        st.plotly_chart(fig)
        st.dataframe(wbs_crit)


# -----------------------------------------------------------------------------
# PAGE: CRITICAL PATH TO SELECTED ACTIVITY
# -----------------------------------------------------------------------------

def _network_diagram_html(
    path_ids: list,
    all_pred_ids: list,
    task_lookup: dict,
    rels_df: pd.DataFrame,
) -> str:
    """
    Build a lightweight SVG-based network diagram showing the driving path
    as a horizontal chain, with branch predecessors shown above/below.
    Returns an HTML string ready for st.components.v1.html().
    """
    if not path_ids:
        return ""

    # Colour helpers
    def node_colour(tid):
        t  = task_lookup.get(tid, {})
        tf = safe_float(t.get("total_float_days"), 9999)
        if tf < 0:  return "#7f1d1d", "#fecaca"   # bg, border
        if tf == 0: return "#dc2626", "#fee2e2"
        if tf <= 10:return "#d97706", "#fef3c7"
        return       "#1d4ed8", "#dbeafe"

    BOX_W, BOX_H = 160, 54
    H_GAP, V_GAP = 40, 70
    n = len(path_ids)
    canvas_w = n * (BOX_W + H_GAP) + H_GAP
    canvas_h = 200

    svg_parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" '
        f'width="{canvas_w}" height="{canvas_h}" '
        f'style="font-family:Arial,sans-serif;font-size:11px;">'
    ]

    # Draw boxes for driving path (middle row y=70)
    mid_y = 80
    box_centres = {}
    for i, tid in enumerate(path_ids):
        t    = task_lookup.get(tid, {})
        x    = H_GAP + i * (BOX_W + H_GAP)
        y    = mid_y
        cx   = x + BOX_W // 2
        cy   = y + BOX_H // 2
        box_centres[tid] = (cx, cy)
        fc, bc = node_colour(tid)
        code = str(t.get("task_code", tid))[:14]
        name = str(t.get("task_name", ""))[:20]
        tf   = t.get("total_float_days")
        tf_s = f"Float: {tf}d" if tf is not None else ""
        svg_parts.append(
            f'<rect x="{x}" y="{y}" width="{BOX_W}" height="{BOX_H}" '
            f'rx="6" fill="{bc}" stroke="{fc}" stroke-width="2"/>'
            f'<text x="{cx}" y="{y+16}" text-anchor="middle" '
            f'font-weight="bold" fill="{fc}">{code}</text>'
            f'<text x="{cx}" y="{y+30}" text-anchor="middle" fill="#374151">{name}</text>'
            f'<text x="{cx}" y="{y+44}" text-anchor="middle" fill="{fc}" '
            f'font-size="10">{tf_s}</text>'
        )

        # Arrow from previous box
        if i > 0:
            prev_tid = path_ids[i - 1]
            px2, py2 = box_centres[prev_tid]
            # Get rel type
            rel_label = "FS"
            lag_label = ""
            if not rels_df.empty:
                rel = rels_df[
                    (rels_df.get("pred_task_id", pd.Series(dtype=str)) == prev_tid) &
                    (rels_df.get("succ_task_id", pd.Series(dtype=str)) == tid)
                ]
                if not rel.empty:
                    rel_label = str(rel["rel_type"].iloc[0])[-2:] if "rel_type" in rel.columns else "FS"
                    lag = safe_float(rel["lag_days"].iloc[0] if "lag_days" in rel.columns else 0, 0)
                    lag_label = f" +{int(lag)}d" if lag > 0 else (f" {int(lag)}d" if lag < 0 else "")
            ax1 = px2 + BOX_W // 2
            ax2 = x
            ay  = cy
            svg_parts.append(
                f'<line x1="{ax1}" y1="{ay}" x2="{ax2}" y2="{ay}" '
                f'stroke="#6b7280" stroke-width="2" marker-end="url(#arr)"/>'
                f'<text x="{(ax1+ax2)//2}" y="{ay-5}" text-anchor="middle" '
                f'fill="#6b7280" font-size="10">{rel_label}{lag_label}</text>'
            )

    # Arrow marker def
    svg_parts.insert(1,
        '<defs><marker id="arr" markerWidth="8" markerHeight="8" '
        'refX="6" refY="3" orient="auto">'
        '<path d="M0,0 L0,6 L8,3 z" fill="#6b7280"/>'
        '</marker></defs>'
    )

    svg_parts.append('</svg>')
    html = (
        '<div style="overflow-x:auto;background:#f8fafc;border:1px solid #e2e8f0;'
        'border-radius:8px;padding:12px;">'
        + "".join(svg_parts) +
        '</div>'
    )
    return html


def page_critical_path_to_activity(data: dict, near_crit_days: float):
    """
    Critical Path to Selected Activity page.

    Traces backwards through predecessor logic to identify the most likely
    driving chain into any selected activity or milestone.
    Uses float, finish dates, relationship type and lag to determine the driver.
    """
    st.title("🎯 Critical Path to Selected Activity")

    tasks = data.get("tasks_df", pd.DataFrame())
    rels  = data.get("relationships_df", pd.DataFrame())

    # Quick-access SAP panel in expander
    with st.expander("Selected Activity Summary", expanded=False, key="sap_exp_cpta"):
        render_selected_activity_panel(
            tasks_df=tasks,
            rels_df=rels,
            near_crit_days=near_crit_days,
            context_key="cpta_sap",
        )

    # -------------------------------------------------------------------------
    # EXPLANATION BANNER
    # -------------------------------------------------------------------------
    st.markdown(
        '<div style="background:#eff6ff;border-left:4px solid #3b82f6;'
        'border-radius:6px;padding:14px 18px;margin-bottom:18px;">'
        '<strong>How this works</strong><br>'
        'This shows the likely chain of activities driving the selected activity. '
        'It traces backwards through predecessor logic and identifies the most critical '
        'path based on total float, latest finish dates, relationship types and lag. '
        '<br><br>'
        '<em>This is based on available XER logic and float values and should be '
        'reviewed with the planner before being used for decision-making.</em>'
        '</div>',
        unsafe_allow_html=True,
    )

    tasks = data["tasks_df"]
    rels  = data["relationships_df"]

    # Guard rails
    if tasks.empty:
        st.warning("No activities found. Please upload an XER file first.")
        return
    if rels.empty:
        st.warning(
            "No relationship data found in this XER file. "
            "This page requires both activities and relationships."
        )
        return

    tasks = get_critical_threshold(tasks, near_crit_days)
    G           = build_graph(tasks, rels)
    task_lookup = tasks.set_index("task_id").to_dict("index")

    # -------------------------------------------------------------------------
    # ACTIVITY SELECTOR
    # -------------------------------------------------------------------------
    def _label(r):
        code = str(r.get("task_code", "?"))
        name = str(r.get("task_name", "?"))
        tf   = r.get("total_float_days")
        flag = " [CRITICAL]" if (tf is not None and safe_float(tf, 1) <= 0) else ""
        return f"{code}  --  {name}{flag}"

    act_labels = tasks.apply(_label, axis=1).tolist()
    selected_label = st.selectbox(
        "Select target activity or milestone",
        options=act_labels,
        key="cpta_selector",
        help="Choose the activity you want to understand the driving path for.",
    )
    sel_idx    = act_labels.index(selected_label)
    target_row = tasks.iloc[sel_idx]
    target_id  = target_row["task_id"]
    tgt_code   = str(target_row.get("task_code", "-"))
    tgt_name   = str(target_row.get("task_name", "-"))
    tgt_tf     = safe_float(target_row.get("total_float_days"), None) if "total_float_days" in target_row.index else None
    tgt_fcol   = _float_color(tgt_tf)
    tgt_crit   = bool(target_row.get("is_critical", False)) if "is_critical" in target_row.index else False
    tgt_stat   = _status_label(str(target_row.get("status", "")))
    tgt_scol   = _status_colour(str(target_row.get("status", "")))

    # Clear cached results when target changes
    if st.session_state.get("_cpta_last_id") != target_id:
        for k in ("cpta_path", "cpta_all_preds"):
            st.session_state.pop(k, None)
        st.session_state["_cpta_last_id"] = target_id

    # -------------------------------------------------------------------------
    # TARGET ACTIVITY BANNER
    # -------------------------------------------------------------------------
    crit_pill = (
        '<span style="background:#dc2626;color:white;padding:2px 10px;'
        'border-radius:12px;font-size:11px;font-weight:700;margin-left:8px;">CRITICAL</span>'
        if tgt_crit else ""
    )
    st.markdown(
        f"""
        <div style="background:#1e3a5f;color:white;border-radius:10px;
                    padding:16px 22px;margin:8px 0 18px 0;">
            <div style="font-size:12px;color:#93c5fd;font-weight:600;
                        letter-spacing:1px;text-transform:uppercase;">Target Activity</div>
            <div style="font-size:20px;font-weight:700;margin-top:4px;">
                {tgt_code}{crit_pill}
            </div>
            <div style="font-size:14px;color:#bfdbfe;margin-top:2px;">{tgt_name}</div>
            <div style="margin-top:10px;">
                <span style="background:{tgt_scol};color:white;padding:3px 10px;
                             border-radius:12px;font-size:12px;">{tgt_stat}</span>
                <span style="background:{tgt_fcol};color:white;padding:3px 10px;
                             border-radius:12px;font-size:12px;margin-left:6px;">
                    Float: {tgt_tf if tgt_tf is not None else "-"} days
                </span>
                <span style="color:#93c5fd;font-size:12px;margin-left:12px;">
                    Finish: {format_date(target_row.get("eff_finish") if "eff_finish" in target_row.index else None)}
                </span>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Check if target has any predecessors at all
    direct_preds = list(G.predecessors(target_id))
    if not direct_preds:
        st.warning(
            f"**{tgt_code} has no predecessors.** This activity has an open start "
            "and is not driven by any logic in the programme. "
            "Nothing can be identified as the driving path."
        )
        return

    # -------------------------------------------------------------------------
    # RUN BUTTON
    # -------------------------------------------------------------------------
    run_col, _ = st.columns([1, 3])
    run_btn = run_col.button(
        "🔍  Find Driving Path",
        key="cpta_run",
        type="primary",
    )

    if run_btn:
        with st.spinner("Tracing predecessor network..."):
            driving_path   = driving_path_to_activity(G, tasks, rels, target_id)
            all_pred_pairs = trace_predecessors(G, target_id)
            all_pred_ids   = [p for p, _ in all_pred_pairs]
        st.session_state["cpta_path"]      = driving_path
        st.session_state["cpta_all_preds"] = all_pred_ids

    # -------------------------------------------------------------------------
    # RESULTS
    # -------------------------------------------------------------------------
    if "cpta_path" not in st.session_state:
        st.markdown(
            '<div style="background:#f0f9ff;border:1px dashed #93c5fd;border-radius:8px;'
            'padding:24px;text-align:center;color:#1e40af;margin-top:16px;">'
            '<strong>Press "Find Driving Path" above to run the analysis.</strong>'
            '</div>',
            unsafe_allow_html=True,
        )
        return

    driving_path = st.session_state["cpta_path"]
    all_pred_ids = st.session_state["cpta_all_preds"]

    # ---- KEY METRICS --------------------------------------------------------
    chain_tasks = tasks[tasks["task_id"].isin(driving_path)]
    n_chain     = len(driving_path)
    min_float   = chain_tasks["total_float_days"].min() if "total_float_days" in chain_tasks.columns else None
    n_crit_chain = int((chain_tasks["total_float_days"].apply(
        lambda f: safe_float(f, 1) <= 0
    )).sum()) if "total_float_days" in chain_tasks.columns else 0
    n_neg_chain  = int((chain_tasks["total_float_days"].apply(
        lambda f: safe_float(f, 0) < 0
    )).sum()) if "total_float_days" in chain_tasks.columns else 0

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Activities in Driving Chain",  n_chain)
    m2.metric("Lowest Float in Chain",         f"{min_float:.1f} days" if min_float is not None else "-")
    m3.metric("Critical in Chain",             n_crit_chain)
    m4.metric("Total Predecessor Network",     len(all_pred_ids))

    if n_neg_chain > 0:
        st.error(
            f"⚠️ **{n_neg_chain} activit{'y' if n_neg_chain == 1 else 'ies'} with negative float** "
            "in the driving chain. The current schedule cannot meet its target dates for this path."
        )

    st.divider()

    # ---- TABS ---------------------------------------------------------------
    tab_path, tab_network, tab_all_preds, tab_constraints = st.tabs([
        "Driving Path", "Network Diagram", "All Predecessors", "Constraints & Issues"
    ])

    # =========================================================================
    # TAB 1: DRIVING PATH TABLE
    # =========================================================================
    with tab_path:
        st.markdown(
            "The table below shows the most likely chain of activities driving "
            f"**{tgt_code}**, ordered from the earliest activity to the target. "
            "Activities are selected based on lowest float, latest finish date "
            "and relationship constraints."
        )

        path_rows = []
        for i, tid in enumerate(driving_path):
            t         = task_lookup.get(tid, {})
            tf        = t.get("total_float_days")
            is_target = (tid == target_id)

            # Relationship to next activity in chain
            rel_label = "-"
            lag_val   = 0
            if i < len(driving_path) - 1:
                next_tid = driving_path[i + 1]
                if not rels.empty:
                    rel = rels[
                        (rels.get("pred_task_id", pd.Series(dtype=str)) == tid) &
                        (rels.get("succ_task_id", pd.Series(dtype=str)) == next_tid)
                    ]
                    if not rel.empty:
                        rel_label = _rel_label(rel["rel_type"].iloc[0] if "rel_type" in rel.columns else "FS")
                        lag_val   = safe_float(rel["lag_days"].iloc[0] if "lag_days" in rel.columns else 0, 0)

            cstr = str(t.get("cstr_type", "")) if "cstr_type" in t else ""
            has_cstr = cstr.strip() not in ("", "None", "nan")

            path_rows.append({
                "Step":            i + 1,
                "Activity ID":     t.get("task_code", tid),
                "Activity Name":   t.get("task_name", ""),
                "Start":           format_date(t.get("eff_start")),
                "Finish":          format_date(t.get("eff_finish")),
                "Orig Dur (d)":    t.get("orig_dur_days", "-"),
                "Total Float (d)": tf if tf is not None else "-",
                "Link to Next":    rel_label if not is_target else "-",
                "Lag (d)":         lag_val if not is_target else "-",
                "Critical Flag":   _crit_flag(tf),
                "Constraint":      cstr if has_cstr else "",
                "Status":          _status_label(str(t.get("status", ""))),
                "Target":          "TARGET" if is_target else "",
            })

        path_df = pd.DataFrame(path_rows)

        # Colour code
        def _style_path_row(row):
            flag = row.get("Critical Flag", "")
            is_tgt = row.get("Target", "") == "TARGET"
            if is_tgt:
                return ["background-color:#1e3a5f;color:white;font-weight:700;"] * len(row)
            colour_map = {
                "Negative Float": "background-color:#fecaca;",
                "Critical":       "background-color:#fee2e2;",
                "Near-Critical":  "background-color:#fef3c7;",
            }
            style = colour_map.get(flag, "")
            return [style] * len(row)

        styled_path = path_df.style.apply(_style_path_row, axis=1)
        st.dataframe(styled_path, hide_index=True)

        # Gantt for driving path
        st.markdown("**Driving Path Timeline**")
        gantt_src = chain_tasks.dropna(subset=["eff_start","eff_finish"]).copy() if "eff_start" in chain_tasks.columns else pd.DataFrame()
        if not gantt_src.empty:
            gantt_src = gantt_src.merge(
                tasks[["task_id","task_code","task_name"]],
                on="task_id", how="left", suffixes=("","_t")
            )
            gantt_src["Label"]   = gantt_src["task_code"].astype(str) + "  " + gantt_src["task_name"].astype(str).str[:35]
            gantt_src["Colour"]  = gantt_src["task_id"].apply(
                lambda t: "Target" if t == target_id else (
                    "Critical" if safe_float(task_lookup.get(t,{}).get("total_float_days"), 1) <= 0
                    else "Near-Critical" if safe_float(task_lookup.get(t,{}).get("total_float_days"), 11) <= 10
                    else "Has Float"
                )
            )
            fig = px.timeline(
                gantt_src,
                x_start="eff_start", x_end="eff_finish", y="Label",
                color="Colour",
                color_discrete_map={
                    "Target":       "#1e3a5f",
                    "Critical":     "#dc2626",
                    "Near-Critical":"#d97706",
                    "Has Float":    "#2563eb",
                },
                title=f"Driving Path to {tgt_code}",
            )
            fig.update_yaxes(autorange="reversed")
            fig.add_vline(
                x=datetime.now().strftime("%Y-%m-%dT%H:%M:%S"), line_dash="dot", line_color="#6b7280",
                annotation_text="Today", annotation_position="top left",
            )
            fig.update_layout(
                height=max(280, 50 + len(gantt_src) * 30),
                margin=dict(l=10, r=10, t=40, b=10),
                legend_title_text="Float Status",
            )
            st.plotly_chart(fig)
        else:
            st.info("No date data available for the Gantt chart.")

    # =========================================================================
    # TAB 2: NETWORK DIAGRAM
    # =========================================================================
    with tab_network:
        st.markdown(
            "A simple left-to-right network diagram of the driving path. "
            "Each box shows the activity ID, name and float. "
            "Colours: **red** = critical/negative float, **amber** = near-critical, **blue** = has float, "
            "**navy** = target activity."
        )

        if len(driving_path) > 0:
            diagram_html = _network_diagram_html(
                driving_path, all_pred_ids, task_lookup, rels
            )
            if diagram_html:
                import streamlit.components.v1 as components
                n_boxes = len(driving_path)
                diagram_w = n_boxes * 200 + 80
                components.html(diagram_html, height=220, scrolling=True)
            else:
                st.info("Could not generate network diagram.")

            st.caption(
                "Note: The diagram shows the identified driving path only. "
                "Use the All Predecessors tab to see the full predecessor network."
            )
        else:
            st.info("No path data to display.")

    # =========================================================================
    # TAB 3: ALL PREDECESSORS
    # =========================================================================
    with tab_all_preds:
        all_pred_tasks = tasks[tasks["task_id"].isin(all_pred_ids)].copy()

        if all_pred_tasks.empty:
            st.info("No predecessor activities found.")
        else:
            n_ap     = len(all_pred_tasks)
            n_ap_crit = int((all_pred_tasks["total_float_days"].apply(
                lambda f: safe_float(f, 1) <= 0
            )).sum()) if "total_float_days" in all_pred_tasks.columns else 0

            st.markdown(
                _summary_bar(f"total predecessors", n_ap, "#374151") +
                (_summary_bar("critical", n_ap_crit, "#dc2626") if n_ap_crit else ""),
                unsafe_allow_html=True,
            )
            st.markdown("<br>", unsafe_allow_html=True)
            st.caption(
                "All activities in the predecessor network of the target activity, "
                "sorted by float (most critical first)."
            )

            all_pred_tasks = all_pred_tasks.sort_values("total_float_days")

            AP_COLS = {
                "task_code":        "Activity ID",
                "task_name":        "Activity Name",
                "wbs_path":         "WBS",
                "eff_start":        "Start",
                "eff_finish":       "Finish",
                "total_float_days": "Float (d)",
                "status":           "Status",
                "is_critical":      "Critical",
            }
            ap_show = {k: v for k, v in AP_COLS.items() if k in all_pred_tasks.columns}
            ap_df   = all_pred_tasks[list(ap_show.keys())].rename(columns=ap_show).copy()
            for col in ["Start","Finish"]:
                if col in ap_df.columns:
                    ap_df[col] = ap_df[col].apply(format_date)
            if "Critical" in ap_df.columns:
                ap_df["Critical"] = ap_df["Critical"].apply(lambda x: "Yes" if x else "")
            if "Status" in ap_df.columns:
                ap_df["Status"] = ap_df["Status"].apply(_status_label)

            st.dataframe(ap_df, hide_index=True, height=400)

    # =========================================================================
    # TAB 4: CONSTRAINTS & ISSUES
    # =========================================================================
    with tab_constraints:
        st.markdown(
            "Activities in the driving path or predecessor network that have "
            "constraints, negative float, or other schedule quality issues."
        )

        issues_found = False

        # --- Negative float in driving path ---
        neg_in_path = chain_tasks[
            chain_tasks["total_float_days"].apply(lambda f: safe_float(f, 0) < 0)
        ] if "total_float_days" in chain_tasks.columns else pd.DataFrame()

        if not neg_in_path.empty:
            issues_found = True
            st.markdown(
                '<div style="background:#fef2f2;border-left:4px solid #dc2626;'
                'border-radius:6px;padding:10px 14px;margin-bottom:12px;">'
                f'<strong>⚠️ Negative Float in Driving Chain ({len(neg_in_path)} activities)</strong><br>'
                'These activities are beyond their target dates. '
                'The driving chain cannot currently meet its schedule.'
                '</div>',
                unsafe_allow_html=True,
            )
            neg_cols = {k: v for k, v in {
                "task_code":"Activity ID","task_name":"Activity Name",
                "total_float_days":"Float (d)","eff_finish":"Finish","status":"Status"
            }.items() if k in neg_in_path.columns}
            neg_disp = neg_in_path[list(neg_cols.keys())].rename(columns=neg_cols).copy()
            if "Finish" in neg_disp.columns:
                neg_disp["Finish"] = neg_disp["Finish"].apply(format_date)
            st.dataframe(neg_disp, hide_index=True)

        # --- Constraints in driving path ---
        if "cstr_type" in chain_tasks.columns:
            constrained = chain_tasks[
                chain_tasks["cstr_type"].apply(
                    lambda x: bool(x) and str(x).strip() not in ("","None","nan")
                )
            ]
            if not constrained.empty:
                issues_found = True
                st.markdown(
                    '<div style="background:#fffbeb;border-left:4px solid #f59e0b;'
                    'border-radius:6px;padding:10px 14px;margin-bottom:12px;">'
                    f'<strong>Constraints in Driving Chain ({len(constrained)} activities)</strong><br>'
                    'Constraints override schedule logic and can cause artificial float or '
                    'negative float. Each one should be reviewed with the planner.'
                    '</div>',
                    unsafe_allow_html=True,
                )
                cstr_cols = {k: v for k, v in {
                    "task_code":"Activity ID","task_name":"Activity Name",
                    "cstr_type":"Constraint Type","cstr_date":"Constraint Date",
                    "total_float_days":"Float (d)"
                }.items() if k in constrained.columns}
                cstr_disp = constrained[list(cstr_cols.keys())].rename(columns=cstr_cols).copy()
                if "Constraint Date" in cstr_disp.columns:
                    cstr_disp["Constraint Date"] = cstr_disp["Constraint Date"].apply(format_date)
                st.dataframe(cstr_disp, hide_index=True)

        # --- Constraints in full predecessor network ---
        all_pred_tasks_full = tasks[tasks["task_id"].isin(all_pred_ids)].copy()
        if "cstr_type" in all_pred_tasks_full.columns:
            all_constrained = all_pred_tasks_full[
                all_pred_tasks_full["cstr_type"].apply(
                    lambda x: bool(x) and str(x).strip() not in ("","None","nan")
                )
            ]
            if not all_constrained.empty:
                issues_found = True
                with st.expander(f"Constraints in Full Predecessor Network ({len(all_constrained)})"):
                    cstr_cols2 = {k: v for k, v in {
                        "task_code":"Activity ID","task_name":"Activity Name",
                        "cstr_type":"Constraint Type","cstr_date":"Constraint Date",
                        "total_float_days":"Float (d)"
                    }.items() if k in all_constrained.columns}
                    cstr_disp2 = all_constrained[list(cstr_cols2.keys())].rename(columns=cstr_cols2).copy()
                    if "Constraint Date" in cstr_disp2.columns:
                        cstr_disp2["Constraint Date"] = cstr_disp2["Constraint Date"].apply(format_date)
                    st.dataframe(cstr_disp2, hide_index=True)

        # --- High lag in driving path relationships ---
        if not rels.empty and "lag_days" in rels.columns:
            path_set  = set(driving_path)
            path_rels = rels[
                rels.get("pred_task_id", pd.Series(dtype=str)).isin(path_set) &
                rels.get("succ_task_id", pd.Series(dtype=str)).isin(path_set)
            ]
            high_lag = path_rels[
                path_rels["lag_days"].apply(lambda l: abs(safe_float(l, 0)) > 5)
            ] if not path_rels.empty else pd.DataFrame()

            if not high_lag.empty:
                issues_found = True
                st.markdown(
                    '<div style="background:#eff6ff;border-left:4px solid #3b82f6;'
                    'border-radius:6px;padding:10px 14px;margin-bottom:12px;">'
                    f'<strong>Significant Lag in Driving Path ({len(high_lag)} relationships)</strong><br>'
                    'Lag of more than 5 days can hide logic issues and affect float calculations.'
                    '</div>',
                    unsafe_allow_html=True,
                )
                lag_cols = {k: v for k, v in {
                    "pred_task_code":"From","pred_task_name":"From Name",
                    "succ_task_code":"To","succ_task_name":"To Name",
                    "rel_type":"Link","lag_days":"Lag (d)"
                }.items() if k in high_lag.columns}
                if lag_cols:
                    st.dataframe(
                        high_lag[list(lag_cols.keys())].rename(columns=lag_cols),
                        hide_index=True,
                    )

        if not issues_found:
            st.success(
                "No constraints, negative float or significant lag found in the driving path. "
                "The chain appears logically sound."
            )

    # =========================================================================
    # EXCEL EXPORT
    # =========================================================================
    st.divider()

    # Rebuild path_df for export (already built above in tab_path)
    export_path_df = pd.DataFrame(path_rows) if path_rows else pd.DataFrame()

    summary_rows = {
        "Item":  [
            "Target Activity ID", "Target Activity Name", "Target Finish",
            "Target Float (days)", "Activities in Driving Chain",
            "Lowest Float in Chain", "Critical in Chain", "Negative Float in Chain",
            "Total Predecessor Network",
        ],
        "Value": [
            tgt_code, tgt_name,
            format_date(target_row.get("eff_finish") if "eff_finish" in target_row.index else None),
            tgt_tf,
            n_chain, min_float, n_crit_chain, n_neg_chain, len(all_pred_ids),
        ],
    }

    export_sheets = {
        "Summary":        pd.DataFrame(summary_rows),
        "Driving Path":   export_path_df,
        "All Predecessors": ap_df if not all_pred_tasks.empty else pd.DataFrame(columns=["No data"]),
    }

    if not neg_in_path.empty:
        export_sheets["Negative Float"] = neg_disp
    if "cstr_type" in chain_tasks.columns and not constrained.empty:
        export_sheets["Constraints"] = cstr_disp

    xls_bytes = export_df_to_excel(export_sheets)

    dl_col, _ = st.columns([1, 3])
    dl_col.download_button(
        label="📥  Export Driving Path Report to Excel",
        data=xls_bytes,
        file_name=f"driving_path_{tgt_code}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        help="Exports Summary, Driving Path, All Predecessors, Negative Float and Constraints sheets.",
        key="dl_003",
    )

# -----------------------------------------------------------------------------
# PAGE: LABOUR HISTOGRAM
# -----------------------------------------------------------------------------

def page_labour_histogram(data: dict):
    st.title("👷 Labour Histogram")

    task_res = data["task_resources_df"]
    tasks = data["tasks_df"]
    resources = data["resources_df"]

    if task_res.empty:
        st.markdown("""
        <div class="warn-box">
        ⚠️ <strong>No resource loading found in this XER file.</strong><br>
        This usually means the programme was not resourced in P6, or resource data was not exported.
        <br><br>
        You can upload a separate resource CSV or Excel file below.
        </div>
        """, unsafe_allow_html=True)

        st.subheader("Upload Resource Loading File")
        res_file = st.file_uploader("Upload CSV or Excel (columns: task_code, rsrc_name, target_qty, target_start, target_finish)", type=["csv","xlsx"])
        if res_file:
            try:
                if res_file.name.endswith(".csv"):
                    task_res = pd.read_csv(res_file)
                else:
                    task_res = pd.read_excel(res_file)
                for col in ["target_start","target_finish"]:
                    if col in task_res.columns:
                        task_res[col] = pd.to_datetime(task_res[col], errors="coerce")
                st.success(f"Loaded {len(task_res)} resource rows.")
            except Exception as e:
                st.error(f"Could not read resource file: {e}")
                return
        else:
            return

    # Merge with task info and resource names
    if not tasks.empty and "task_id" in task_res.columns:
        _task_cols = ["task_id","task_code","task_name","wbs_path"]
        if "is_critical" in tasks.columns:
            _task_cols.append("is_critical")
        _tasks_slim = tasks[_task_cols].drop_duplicates(subset="task_id").reset_index(drop=True)
        task_res = task_res.merge(_tasks_slim, on="task_id", how="left", suffixes=("","_task"))
    if not resources.empty and "rsrc_id" in task_res.columns:
        task_res = task_res.merge(resources[["rsrc_id","rsrc_name"]], on="rsrc_id", how="left", suffixes=("","_res"))
        if "rsrc_name_res" in task_res.columns:
            task_res["rsrc_name"] = task_res["rsrc_name_res"].fillna(task_res.get("rsrc_name",""))

    # Expand resource loading to weekly intervals
    def expand_to_weeks(df):
        rows = []
        for _, r in df.iterrows():
            s = pd.to_datetime(r.get("target_start") or r.get("target_start_date"))
            e = pd.to_datetime(r.get("target_finish") or r.get("target_end_date"))
            if pd.isna(s) or pd.isna(e) or s > e:
                continue
            qty = safe_float(r.get("target_qty", 0), 0)
            if qty == 0:
                continue
            weeks = max(1, math.ceil((e - s).days / 7))
            qty_per_week = qty / weeks
            current = s
            for _ in range(weeks):
                rows.append({
                    "week": current.to_period("W").start_time,
                    "month": current.to_period("M").start_time,
                    "qty": qty_per_week,
                    "rsrc_name": r.get("rsrc_name","Unknown"),
                    "task_code": r.get("task_code",""),
                    "task_name": r.get("task_name",""),
                    "wbs_path": r.get("wbs_path",""),
                })
                current += timedelta(weeks=1)
        return pd.DataFrame(rows)

    weekly = expand_to_weeks(task_res)

    if weekly.empty:
        st.warning("Could not generate histogram -- resource dates or quantities may be missing.")
        return

    # Filters
    st.sidebar.divider()
    st.sidebar.subheader("Labour Filters")
    all_resources = sorted(weekly["rsrc_name"].unique().tolist())
    sel_res = st.sidebar.multiselect("Resource / Trade", all_resources, default=all_resources)
    if sel_res:
        weekly = weekly[weekly["rsrc_name"].isin(sel_res)]

    # Metrics
    c1, c2, c3 = st.columns(3)
    c1.metric("Total Planned Hours", f"{weekly['qty'].sum():,.0f}")
    weekly_totals = weekly.groupby("week")["qty"].sum()
    c2.metric("Peak Week (hrs)", f"{weekly_totals.max():,.0f}" if not weekly_totals.empty else "-")
    c3.metric("Average Week (hrs)", f"{weekly_totals.mean():,.0f}" if not weekly_totals.empty else "-")

    tab1, tab2, tab3, tab4 = st.tabs(["By Week", "By Month", "By Resource", "By WBS"])

    with tab1:
        weekly_sum = weekly.groupby("week")["qty"].sum().reset_index()
        fig = px.bar(weekly_sum, x="week", y="qty",
                     title="Labour Loading by Week (Hours)",
                     labels={"week":"Week","qty":"Hours"},
                     color_discrete_sequence=["#2563eb"])
        st.plotly_chart(fig)

    with tab2:
        monthly_sum = weekly.groupby("month")["qty"].sum().reset_index()
        fig = px.bar(monthly_sum, x="month", y="qty",
                     title="Labour Loading by Month (Hours)",
                     labels={"month":"Month","qty":"Hours"},
                     color_discrete_sequence=["#1e3a5f"])
        st.plotly_chart(fig)

    with tab3:
        res_sum = weekly.groupby("rsrc_name")["qty"].sum().reset_index().sort_values("qty", ascending=False)
        fig = px.bar(res_sum, x="rsrc_name", y="qty",
                     title="Total Hours by Resource / Trade",
                     labels={"rsrc_name":"Resource","qty":"Hours"},
                     color_discrete_sequence=["#7c3aed"])
        st.plotly_chart(fig)

        # By week and resource stacked
        if len(sel_res) <= 10:
            by_res_week = weekly.groupby(["week","rsrc_name"])["qty"].sum().reset_index()
            fig2 = px.bar(by_res_week, x="week", y="qty", color="rsrc_name",
                          title="Weekly Labour by Resource",
                          labels={"week":"Week","qty":"Hours","rsrc_name":"Resource"})
            st.plotly_chart(fig2)

    with tab4:
        if "wbs_path" in weekly.columns:
            weekly["wbs_top"] = weekly["wbs_path"].apply(
                lambda x: str(x).split(" > ")[0] if pd.notna(x) and x else "Unknown"
            )
            wbs_sum = weekly.groupby("wbs_top")["qty"].sum().reset_index().sort_values("qty", ascending=False)
            fig = px.bar(wbs_sum, x="qty", y="wbs_top", orientation="h",
                         title="Total Hours by WBS",
                         color_discrete_sequence=["#059669"])
            st.plotly_chart(fig)

    # Export
    xls = export_df_to_excel({
        "Weekly Labour": weekly.groupby(["week","rsrc_name"])["qty"].sum().reset_index(),
        "Monthly Labour": weekly.groupby(["month","rsrc_name"])["qty"].sum().reset_index(),
        "By Resource": res_sum,
    })
    st.download_button("📥 Export Labour Data", xls, "labour_histogram.xlsx",
        key="dl_004",
                       )


# -----------------------------------------------------------------------------
# PAGE: SCHEDULE HEALTH CHECK
# -----------------------------------------------------------------------------

def page_health_check(data: dict, near_crit_days: float):
    st.title("🩺 Schedule Health Check")
    st.markdown("> Automated quality checks to identify common schedule issues.")

    tasks = data["tasks_df"]
    rels = data["relationships_df"]

    if tasks.empty:
        st.warning("No activities loaded.")
        return

    tasks = get_critical_threshold(tasks, near_crit_days)

    # Build predecessor/successor sets
    tasks_with_pred = set()
    tasks_with_succ = set()
    if not rels.empty:
        tasks_with_pred = set(rels["succ_task_id"].dropna()) if "succ_task_id" in rels.columns else set()
        tasks_with_succ = set(rels["pred_task_id"].dropna()) if "pred_task_id" in rels.columns else set()

    # Define checks
    checks = []

    # 1. No predecessors (excl. milestones at start)
    no_pred = tasks[~tasks["task_id"].isin(tasks_with_pred)]
    checks.append({
        "Check": "No Predecessors",
        "Count": len(no_pred),
        "Severity": "⚠️ Warning",
        "Why It Matters": "Activities with no predecessors are open-ended. They cannot be driven by logic and may cause float calculation issues.",
        "df": no_pred,
    })

    # 2. No successors
    no_succ = tasks[~tasks["task_id"].isin(tasks_with_succ)]
    checks.append({
        "Check": "No Successors",
        "Count": len(no_succ),
        "Severity": "⚠️ Warning",
        "Why It Matters": "Activities with no successors are open-ended and may have artificially high float.",
        "df": no_succ,
    })

    # 3. Negative float
    neg_float = tasks[tasks["total_float_days"].apply(lambda f: f is not None and f < 0)]
    checks.append({
        "Check": "Negative Float",
        "Count": len(neg_float),
        "Severity": "🔴 Critical",
        "Why It Matters": "Negative float means the current schedule cannot meet its target dates. Immediate attention required.",
        "df": neg_float,
    })

    # 4. High float (> 60 days)
    high_float = tasks[tasks["total_float_days"].apply(lambda f: f is not None and f > 60)]
    checks.append({
        "Check": "Very High Float (>60 days)",
        "Count": len(high_float),
        "Severity": "ℹ️ Info",
        "Why It Matters": "Activities with very high float may have missing logic or may not be properly constrained.",
        "df": high_float,
    })

    # 5. Excessive duration (> 60 working days)
    excess_dur = tasks[tasks["orig_dur_days"].apply(lambda d: d is not None and d > 60)]
    checks.append({
        "Check": "Excessive Duration (>60 days)",
        "Count": len(excess_dur),
        "Severity": "⚠️ Warning",
        "Why It Matters": "Very long activities are difficult to control and should usually be broken down into smaller work packages.",
        "df": excess_dur,
    })

    # 6. Constraints
    constrained = tasks[tasks["cstr_type"].apply(
        lambda x: bool(x) and str(x).strip() not in ("", "None")
    )] if "cstr_type" in tasks.columns else pd.DataFrame()
    checks.append({
        "Check": "Constrained Activities",
        "Count": len(constrained),
        "Severity": "⚠️ Warning",
        "Why It Matters": "Constraints override schedule logic and can create artificial float or negative float. Each constraint should be justified.",
        "df": constrained,
    })

    # 7. Excessive lag (> 10 days)
    if not rels.empty and "lag_days" in rels.columns:
        high_lag = rels[rels["lag_days"].apply(lambda l: l is not None and abs(safe_float(l,0)) > 10)]
        checks.append({
            "Check": "Excessive Lag (|lag| > 10 days)",
            "Count": len(high_lag),
            "Severity": "⚠️ Warning",
            "Why It Matters": "Excessive lag can hide critical path issues. Lag should be replaced with properly sequenced activities.",
            "df": high_lag,
        })

    # 8. Missing dates
    missing_dates = tasks[tasks["eff_start"].isna() | tasks["eff_finish"].isna()]
    checks.append({
        "Check": "Missing Start or Finish Dates",
        "Count": len(missing_dates),
        "Severity": "🔴 Critical",
        "Why It Matters": "Activities with no dates cannot be scheduled or reported on.",
        "df": missing_dates,
    })

    # 9. Actual dates in future
    now = datetime.now()
    future_actuals = tasks[
        tasks["act_start"].apply(lambda d: d is not None and d > now) |
        tasks["act_finish"].apply(lambda d: d is not None and d > now)
    ] if "act_start" in tasks.columns else pd.DataFrame()
    checks.append({
        "Check": "Future Actual Dates",
        "Count": len(future_actuals),
        "Severity": "🔴 Critical",
        "Why It Matters": "Actual start/finish dates should not be in the future. This indicates data entry errors.",
        "df": future_actuals,
    })

    # 10. Critical not started
    crit_not_started = tasks[
        tasks["is_critical"] &
        tasks["status"].apply(lambda s: str(s) in ("TK_NotStart", "Not Started") if pd.notna(s) else False)
    ] if "status" in tasks.columns else pd.DataFrame()
    checks.append({
        "Check": "Critical Activities Not Started",
        "Count": len(crit_not_started),
        "Severity": "🔴 Critical",
        "Why It Matters": "Critical activities that haven't started need immediate attention to avoid slippage.",
        "df": crit_not_started,
    })

    # 11. Near-critical due in 8 weeks
    eight_weeks = now + timedelta(weeks=8)
    near_due = tasks[
        tasks["is_near_critical"] &
        tasks["eff_finish"].apply(lambda d: d is not None and d <= eight_weeks)
    ] if "eff_finish" in tasks.columns else pd.DataFrame()
    checks.append({
        "Check": "Near-Critical Due in 8 Weeks",
        "Count": len(near_due),
        "Severity": "⚠️ Warning",
        "Why It Matters": "Near-critical activities finishing soon may become critical if not progressed.",
        "df": near_due,
    })

    # Scorecard
    st.subheader("Health Check Scorecard")
    score_data = [
        {"Check": c["Check"], "Count": c["Count"], "Severity": c["Severity"]}
        for c in checks
    ]
    score_df = pd.DataFrame(score_data)
    st.dataframe(score_df)

    # Detail per check
    st.divider()
    for chk in checks:
        with st.expander(f"{chk['Severity']} -- {chk['Check']} ({chk['Count']})"):
            st.markdown(f"**Why it matters:** {chk['Why It Matters']}")
            df = chk["df"]
            if not df.empty:
                disp = [c for c in ["task_code","task_name","wbs_path","eff_start",
                                     "eff_finish","total_float_days","status",
                                     "cstr_type","lag_days"] if c in df.columns]
                st.dataframe(df[disp].head(100))
                # Export individual check
                xls = export_df_to_excel({chk["Check"][:31]: df[disp]})
                st.download_button(
                    f"📥 Export: {chk['Check']}", xls,
                    f"health_{chk['Check'][:20].replace(' ','_')}.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="dl_005",
                )
            else:
                st.success("✅ No issues found for this check.")

    # Full export
    all_export = {chk["Check"][:31]: chk["df"][[c for c in ["task_code","task_name","total_float_days","status"] if c in chk["df"].columns]] if not chk["df"].empty else pd.DataFrame(columns=["No issues"]) for chk in checks}
    xls_all = export_df_to_excel(all_export)
    st.download_button("📥 Export Full Health Check Report", xls_all, "schedule_health_check.xlsx",
        key="dl_006",
                       )


# -----------------------------------------------------------------------------
# PAGE: PLANNING NOTES
# -----------------------------------------------------------------------------

HIGHLIGHT_WORDS = [
    "risk", "delay", "delayed", "blocked", "constraint", "access",
    "design", "procurement", "client", "instruction", "CE", "EWN",
    "change", "issue", "hold", "pending", "late", "overrun",
]

def highlight_text(text: str) -> str:
    """Wrap highlight words in HTML span."""
    for word in HIGHLIGHT_WORDS:
        pattern = re.compile(r"\b(" + re.escape(word) + r")\b", re.IGNORECASE)
        text = pattern.sub(r'<span style="background:#fef08a;font-weight:bold;">\1</span>', text)
    return text


def page_planning_notes(data: dict):
    st.title("📝 Planning Notes")
    st.markdown("> Upload planning notes and link them to activities in the programme.")

    tasks = data["tasks_df"]
    notes_file = st.file_uploader("Upload Planning Notes (CSV, Excel, TXT, or DOCX)",
                                   type=["csv","xlsx","txt","docx"])

    if notes_file is None:
        st.info("Upload a notes file to get started. The file should contain free-text notes referencing activity IDs.")
        return

    # Read notes
    notes_text = ""
    notes_rows = []

    try:
        if notes_file.name.endswith(".csv"):
            df = pd.read_csv(notes_file)
            notes_text = " ".join(df.astype(str).values.flatten())
            notes_rows = df.to_dict("records")
        elif notes_file.name.endswith(".xlsx"):
            df = pd.read_excel(notes_file)
            notes_text = " ".join(df.astype(str).values.flatten())
            notes_rows = df.to_dict("records")
        elif notes_file.name.endswith(".txt"):
            notes_text = notes_file.read().decode("utf-8", errors="replace")
            notes_rows = [{"line": i+1, "text": line} for i, line in enumerate(notes_text.splitlines()) if line.strip()]
        elif notes_file.name.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(notes_file.read()))
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            notes_text = "\n".join(lines)
            notes_rows = [{"paragraph": i+1, "text": line} for i, line in enumerate(lines)]
        else:
            st.error("Unsupported file format.")
            return
        st.success(f"Loaded notes file: {notes_file.name}")
    except Exception as e:
        st.error(f"Could not read notes file: {e}")
        return

    # Find activity IDs mentioned in notes
    if not tasks.empty and "task_code" in tasks.columns:
        task_codes = tasks["task_code"].dropna().tolist()
        found_codes = [code for code in task_codes if code in notes_text]

        st.subheader(f"Activity IDs Found in Notes: {len(found_codes)}")
        if found_codes:
            matched_tasks = tasks[tasks["task_code"].isin(found_codes)][
                ["task_code","task_name","eff_start","eff_finish","total_float_days","status"]
            ]
            st.dataframe(matched_tasks)
        else:
            st.info("No activity IDs from the programme were found in the notes.")

        # Not found
        not_found = [code for code in task_codes if code not in notes_text]
        st.caption(f"{len(not_found)} activities not mentioned in notes.")

    # Keyword search
    st.divider()
    st.subheader("Keyword Search")
    keyword = st.text_input("Search notes for keyword")

    display_rows = notes_rows
    if keyword:
        display_rows = [r for r in notes_rows if keyword.lower() in str(r).lower()]
        st.caption(f"{len(display_rows)} matching entries")

    # Display with highlights
    for row in display_rows[:100]:
        text = str(row.get("text","") or list(row.values())[-1])
        highlighted = highlight_text(text)
        st.markdown(f"<div style='background:#f8fafc;border-left:3px solid #2563eb;padding:8px;margin:4px 0;font-size:13px;'>{highlighted}</div>", unsafe_allow_html=True)

    # Full highlighted dump
    st.divider()
    st.subheader("Full Notes (with keyword highlighting)")
    highlighted_full = highlight_text(notes_text.replace("\n","<br>"))
    st.markdown(f"<div style='background:white;border:1px solid #e2e8f0;padding:16px;border-radius:8px;max-height:400px;overflow-y:auto;font-size:12px;'>{highlighted_full}</div>", unsafe_allow_html=True)


# -----------------------------------------------------------------------------
# PAGE: PROGRAMME COMPARISON / MOVEMENT INTELLIGENCE
# -----------------------------------------------------------------------------

def _mi_card(title: str, value, subtitle: str = "", colour: str = "#0B1F33",
             bg: str = "#ffffff", border: str = "#E2E8F0") -> str:
    """Render a compact metric card as HTML."""
    return (
        f'<div style="background:{bg};border:1px solid {border};border-radius:10px;'
        f'padding:16px 18px;box-shadow:0 1px 4px rgba(11,31,51,0.07);">'
        f'<div style="font-size:10px;font-weight:700;color:#94A3B8;letter-spacing:1px;'
        f'text-transform:uppercase;margin-bottom:6px;">{title}</div>'
        f'<div style="font-size:26px;font-weight:800;color:{colour};line-height:1;">{value}</div>'
        f'{"" if not subtitle else f"<div style=font-size:11px;color:#64748B;margin-top:4px;>{subtitle}</div>"}'
        f'</div>'
    )


def _commentary_box(heading: str, body: str, colour: str = "#0B1F33",
                    bg: str = "#eff6ff", border: str = "#3b82f6") -> str:
    return (
        f'<div style="background:{bg};border-left:4px solid {border};border-radius:6px;'
        f'padding:14px 18px;margin-bottom:10px;">'
        f'<div style="font-weight:700;color:{colour};margin-bottom:4px;">{heading}</div>'
        f'<div style="font-size:14px;color:#334155;line-height:1.6;">{body}</div>'
        f'</div>'
    )


def _compute_movement(prev_tasks: pd.DataFrame, curr_tasks: pd.DataFrame,
                      near_crit: float = 10.0) -> dict:
    """
    Merge previous and current task sets and compute all movement metrics.
    Returns a dict of DataFrames and scalar stats.
    """
    prev = get_critical_threshold(prev_tasks.copy(), near_crit)
    curr = get_critical_threshold(curr_tasks.copy(), near_crit)

    # Common activities (matched on task_code)
    merged = prev.merge(curr, on="task_code", how="outer", suffixes=("_p", "_c"))

    added   = curr[~curr["task_code"].isin(prev["task_code"])].copy()
    removed = prev[~prev["task_code"].isin(curr["task_code"])].copy()
    common  = merged.dropna(subset=["task_code"]).copy()

    # ---- Date movement -------------------------------------------------------
    def _days_diff(a, b):
        """b minus a in calendar days. Positive = slipped."""
        try:
            if pd.isna(a) or pd.isna(b):
                return None
            return int((pd.Timestamp(b) - pd.Timestamp(a)).days)
        except Exception:
            return None

    common["finish_move"] = common.apply(
        lambda r: _days_diff(r.get("eff_finish_p"), r.get("eff_finish_c")), axis=1
    )
    common["start_move"] = common.apply(
        lambda r: _days_diff(r.get("eff_start_p"), r.get("eff_start_c")), axis=1
    )

    # ---- Float movement ------------------------------------------------------
    common["float_move"] = common.apply(
        lambda r: (
            safe_float(r.get("total_float_days_c"), None) is not None and
            safe_float(r.get("total_float_days_p"), None) is not None
        ) and safe_float(r.get("total_float_days_c"), 0) - safe_float(r.get("total_float_days_p"), 0)
        if (
            safe_float(r.get("total_float_days_c"), None) is not None and
            safe_float(r.get("total_float_days_p"), None) is not None
        ) else None,
        axis=1,
    )

    # ---- Critical transitions ------------------------------------------------
    def _is_crit(row, suffix):
        tf = safe_float(row.get(f"total_float_days_{suffix}"), 9999)
        return tf <= 0

    common["was_crit"]  = common.apply(lambda r: _is_crit(r, "p"), axis=1)
    common["now_crit"]  = common.apply(lambda r: _is_crit(r, "c"), axis=1)
    became_crit  = common[~common["was_crit"] & common["now_crit"]]
    no_longer_crit = common[common["was_crit"] & ~common["now_crit"]]

    # ---- Slipped / improved --------------------------------------------------
    delayed  = common[common["finish_move"].apply(lambda x: x is not None and x > 0)]
    improved = common[common["finish_move"].apply(lambda x: x is not None and x < 0)]
    float_lost   = common[common["float_move"].apply(lambda x: x is not None and x < 0)]
    float_gained = common[common["float_move"].apply(lambda x: x is not None and x > 0)]

    # ---- Project finish movement --------------------------------------------
    prev_end = prev["eff_finish"].dropna().max() if "eff_finish" in prev.columns else None
    curr_end = curr["eff_finish"].dropna().max() if "eff_finish" in curr.columns else None
    proj_finish_move = _days_diff(prev_end, curr_end)

    # ---- WBS-level summary --------------------------------------------------
    wbs_col_p = "wbs_path_p" if "wbs_path_p" in common.columns else None
    wbs_col_c = "wbs_path_c" if "wbs_path_c" in common.columns else None
    wbs_col   = wbs_col_c or wbs_col_p

    if wbs_col and not common.empty:
        common["wbs_top"] = common[wbs_col].apply(
            lambda x: str(x).split(" > ")[0] if pd.notna(x) and str(x).strip() not in ("","nan") else "Unknown"
        )
        wbs_summary = common.groupby("wbs_top").agg(
            total_activities=("task_code", "count"),
            delayed_count=("finish_move", lambda x: (x > 0).sum()),
            improved_count=("finish_move", lambda x: (x < 0).sum()),
            avg_finish_move=("finish_move", lambda x: round(x.dropna().mean(), 1) if x.dropna().any() else 0),
            max_finish_slip=("finish_move", lambda x: x.dropna().max() if not x.dropna().empty else 0),
            crit_gained=("now_crit", lambda x: (x & ~common.loc[x.index, "was_crit"]).sum()),
            crit_lost=("now_crit", lambda x: (~x & common.loc[x.index, "was_crit"]).sum()),
        ).reset_index()
    else:
        wbs_summary = pd.DataFrame()

    return {
        "merged":        common,
        "added":         added,
        "removed":       removed,
        "delayed":       delayed.sort_values("finish_move", ascending=False),
        "improved":      improved.sort_values("finish_move"),
        "float_lost":    float_lost.sort_values("float_move"),
        "float_gained":  float_gained.sort_values("float_move", ascending=False),
        "became_crit":   became_crit,
        "no_longer_crit":no_longer_crit,
        "wbs_summary":   wbs_summary,
        "proj_finish_move": proj_finish_move,
        "prev_end":      prev_end,
        "curr_end":      curr_end,
        "n_added":       len(added),
        "n_removed":     len(removed),
        "n_delayed":     len(delayed),
        "n_improved":    len(improved),
        "n_became_crit": len(became_crit),
        "n_no_longer_crit": len(no_longer_crit),
        "n_float_lost":  len(float_lost),
        "n_float_gained":len(float_gained),
        "n_common":      len(common),
    }


def _build_commentary(m: dict, prev_name: str, curr_name: str) -> list:
    """
    Generate plain-English commentary bullets from movement data.
    Returns list of (heading, body, colour, bg, border) tuples.
    """
    items = []
    pfm = m["proj_finish_move"]

    # -- What changed? ---------------------------------------------------------
    changes = []
    if pfm is not None and pfm > 0:
        changes.append(f"The project finish date has slipped by <strong>{pfm} days</strong>.")
    elif pfm is not None and pfm < 0:
        changes.append(f"The project finish date has improved by <strong>{abs(pfm)} days</strong>.")
    elif pfm == 0:
        changes.append("The project finish date is unchanged between revisions.")

    if m["n_delayed"] > 0:
        changes.append(
            f"<strong>{m['n_delayed']}</strong> activities have a later finish date in the current programme."
        )
    if m["n_improved"] > 0:
        changes.append(f"<strong>{m['n_improved']}</strong> activities have an earlier finish date.")
    if m["n_added"] > 0:
        changes.append(f"<strong>{m['n_added']}</strong> new activities have been added.")
    if m["n_removed"] > 0:
        changes.append(f"<strong>{m['n_removed']}</strong> activities have been removed.")

    if changes:
        items.append((
            "What changed?",
            " ".join(changes),
            "#1e3a5f", "#eff6ff", "#3b82f6"
        ))

    # -- Where is the risk? ----------------------------------------------------
    risks = []
    if m["n_became_crit"] > 0:
        risks.append(
            f"<strong>{m['n_became_crit']} activities became critical</strong> in the current revision. "
            "These activities now have zero or negative float and must be monitored closely."
        )
    neg_float_curr = m["merged"]["total_float_days_c"].apply(
        lambda f: safe_float(f, 0) < 0
    ).sum() if "total_float_days_c" in m["merged"].columns else 0
    if neg_float_curr > 0:
        risks.append(
            f"<strong>{int(neg_float_curr)} activities have negative float</strong> in the current programme. "
            "This means the schedule cannot currently achieve its target dates on these paths."
        )
    if m["n_float_lost"] > 0:
        worst_loss = m["float_lost"]["float_move"].min()
        risks.append(
            f"<strong>{m['n_float_lost']} activities lost float</strong>. "
            f"The worst case is a loss of {abs(worst_loss):.1f} days."
        )
    if not m["delayed"].empty:
        worst_slip = m["delayed"]["finish_move"].max()
        worst_act  = m["delayed"].iloc[0].get("task_name_c", m["delayed"].iloc[0].get("task_code",""))
        risks.append(
            f"The biggest finish date slip is <strong>{int(worst_slip)} days</strong> "
            f"({str(worst_act)[:60]})."
        )

    if risks:
        items.append((
            "Where is the risk?",
            " ".join(risks),
            "#7f1d1d", "#fef2f2", "#dc2626"
        ))
    else:
        items.append((
            "Where is the risk?",
            "No significant new risks identified between these two revisions.",
            "#166534", "#f0fdf4", "#16a34a"
        ))

    # -- What needs PM attention? ----------------------------------------------
    attention = []
    if m["n_became_crit"] > 0:
        codes = m["became_crit"]["task_code"].head(3).tolist()
        attention.append(
            f"Review the {m['n_became_crit']} newly critical activities, "
            f"including: {', '.join(str(c) for c in codes)}."
        )
    if neg_float_curr > 0:
        attention.append(
            "Investigate all activities with negative float. "
            "These require a recovery plan or a revised target date."
        )
    if not m["delayed"].empty:
        attention.append(
            "Review the biggest finish date slips in the Biggest Slips section below "
            "and confirm whether recovery actions are in place."
        )
    if m["n_removed"] > 0:
        attention.append(
            f"{m['n_removed']} activities were removed. "
            "Confirm these were intentional scope reductions and not accidental deletions."
        )

    if attention:
        items.append((
            "What needs PM attention?",
            " ".join(attention),
            "#92400e", "#fffbeb", "#F5A623"
        ))

    # -- What improved? --------------------------------------------------------
    good = []
    if m["n_no_longer_crit"] > 0:
        good.append(
            f"<strong>{m['n_no_longer_crit']} activities are no longer critical</strong>, "
            "indicating improved schedule recovery on those paths."
        )
    if m["n_improved"] > 0:
        best = abs(m["improved"]["finish_move"].min())
        good.append(
            f"<strong>{m['n_improved']} activities finished earlier</strong> than the previous revision. "
            f"The best improvement is {int(best)} days."
        )
    if m["n_float_gained"] > 0:
        good.append(f"<strong>{m['n_float_gained']} activities gained float</strong>.")
    if pfm is not None and pfm < 0:
        good.append(
            f"The overall project finish has <strong>improved by {abs(pfm)} days</strong>."
        )

    if good:
        items.append((
            "What improved?",
            " ".join(good),
            "#166534", "#f0fdf4", "#16a34a"
        ))

    return items


def _display_cols(df: pd.DataFrame, col_map: dict) -> pd.DataFrame:
    """Return a display-ready copy with renamed columns and formatted dates."""
    avail = {k: v for k, v in col_map.items() if k in df.columns}
    out = df[list(avail.keys())].copy().rename(columns=avail)
    for col in out.columns:
        if any(kw in col.lower() for kw in ("start","finish","date")):
            try:
                out[col] = out[col].apply(format_date)
            except Exception:
                pass
    return out


def page_programme_comparison():
    """
    Programme Movement Intelligence page.
    Upload two XER files (previous and current) and get a plain-English
    analysis of what changed, where the risk is, and what improved.
    """
    st.title("📅 Movement Intelligence")
    st.caption(
        "Upload the previous and current version of your XER programme to see "
        "a plain-English analysis of what has changed, where the risk is, "
        "and what needs your attention."
    )

    # ---- File uploads --------------------------------------------------------
    up_col1, up_col2 = st.columns(2)
    with up_col1:
        st.markdown(
            '<div style="font-size:12px;font-weight:700;color:#94A3B8;'
            'letter-spacing:1px;text-transform:uppercase;margin-bottom:6px;">'
            'Previous Programme</div>',
            unsafe_allow_html=True,
        )
        prev_file = st.file_uploader("Previous XER", type=["xer"], key="mi_prev_xer",
                                     label_visibility="collapsed")
    with up_col2:
        st.markdown(
            '<div style="font-size:12px;font-weight:700;color:#94A3B8;'
            'letter-spacing:1px;text-transform:uppercase;margin-bottom:6px;">'
            'Current Programme</div>',
            unsafe_allow_html=True,
        )
        curr_file = st.file_uploader("Current XER", type=["xer"], key="mi_curr_xer",
                                     label_visibility="collapsed")

    if not prev_file or not curr_file:
        st.markdown(
            '<div style="background:#f8fafc;border:2px dashed #CBD5E1;border-radius:10px;'
            'padding:32px;text-align:center;color:#64748B;margin-top:20px;">'
            '<div style="font-size:28px;margin-bottom:10px;">📅</div>'
            '<strong style="font-size:15px;">Upload both programmes above to begin</strong><br>'
            '<span style="font-size:13px;margin-top:6px;display:block;">'
            'Upload the previous revision on the left and the current revision on the right.</span>'
            '</div>',
            unsafe_allow_html=True,
        )
        return

    # ---- Parse & cache -------------------------------------------------------
    cache_key = f"mi_{prev_file.name}_{prev_file.size}_{curr_file.name}_{curr_file.size}"
    if st.session_state.get("_mi_cache_key") != cache_key:
        with st.spinner("Parsing both programmes..."):
            try:
                prev_data = parse_xer(prev_file.read())
                curr_data = parse_xer(curr_file.read())
            except Exception as e:
                st.error(f"Could not parse XER files: {e}")
                return
        st.session_state["_mi_prev"] = prev_data
        st.session_state["_mi_curr"] = curr_data
        st.session_state["_mi_cache_key"] = cache_key

    prev_data = st.session_state["_mi_prev"]
    curr_data = st.session_state["_mi_curr"]

    prev_tasks = prev_data["tasks_df"]
    curr_tasks = curr_data["tasks_df"]

    if prev_tasks.empty or curr_tasks.empty:
        st.error("Could not extract activities from one or both files. Check the XER exports.")
        return

    near_crit_days = 10.0

    # ---- Run movement analysis -----------------------------------------------
    m = _compute_movement(prev_tasks, curr_tasks, near_crit_days)

    prev_name = prev_data.get("project_info", {}).get("name", prev_file.name)
    curr_name = curr_data.get("project_info", {}).get("name", curr_file.name)

    # ---- File header strip ---------------------------------------------------
    pfm = m["proj_finish_move"]
    pfm_str = (
        f"+{pfm}d (slipped)" if pfm and pfm > 0 else
        f"{pfm}d (improved)" if pfm and pfm < 0 else
        "No change" if pfm == 0 else "N/A"
    )
    pfm_colour = "#dc2626" if pfm and pfm > 0 else "#16a34a" if pfm and pfm < 0 else "#6b7280"

    st.markdown(
        f"""
        <div style="background:#0B1F33;border-radius:12px;padding:20px 24px;
                    margin-bottom:24px;display:flex;gap:32px;flex-wrap:wrap;
                    align-items:center;">
            <div>
                <div style="font-size:10px;color:#64748B;text-transform:uppercase;
                            letter-spacing:1px;">Previous</div>
                <div style="font-size:14px;font-weight:600;color:#CBD5E1;margin-top:2px;">
                    {prev_name}</div>
                <div style="font-size:11px;color:#475569;">
                    {format_date(m["prev_end"])}</div>
            </div>
            <div style="font-size:24px;color:#E8951D;">&#8594;</div>
            <div>
                <div style="font-size:10px;color:#64748B;text-transform:uppercase;
                            letter-spacing:1px;">Current</div>
                <div style="font-size:14px;font-weight:600;color:#CBD5E1;margin-top:2px;">
                    {curr_name}</div>
                <div style="font-size:11px;color:#475569;">
                    {format_date(m["curr_end"])}</div>
            </div>
            <div style="margin-left:auto;text-align:right;">
                <div style="font-size:10px;color:#64748B;text-transform:uppercase;
                            letter-spacing:1px;">Net Project Finish Movement</div>
                <div style="font-size:28px;font-weight:800;color:{pfm_colour};margin-top:2px;">
                    {pfm_str}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # ---- Summary metric cards ------------------------------------------------
    st.markdown(
        '<div style="font-size:12px;font-weight:700;color:#94A3B8;'
        'letter-spacing:1px;text-transform:uppercase;margin-bottom:10px;">'
        'Summary</div>',
        unsafe_allow_html=True,
    )

    r1 = st.columns(5)
    r1[0].markdown(_mi_card("Activities Added",    m["n_added"],    colour="#16a34a"), unsafe_allow_html=True)
    r1[1].markdown(_mi_card("Activities Removed",  m["n_removed"],  colour="#6b7280"), unsafe_allow_html=True)
    r1[2].markdown(_mi_card("Finish Date Moved",   m["n_delayed"] + m["n_improved"], colour="#0B1F33"), unsafe_allow_html=True)
    r1[3].markdown(_mi_card("Delayed",             m["n_delayed"],  colour="#dc2626"), unsafe_allow_html=True)
    r1[4].markdown(_mi_card("Improved",            m["n_improved"], colour="#16a34a"), unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    r2 = st.columns(5)
    r2[0].markdown(_mi_card("Became Critical",     m["n_became_crit"],    colour="#dc2626", bg="#fef2f2", border="#fca5a5"), unsafe_allow_html=True)
    r2[1].markdown(_mi_card("No Longer Critical",  m["n_no_longer_crit"], colour="#16a34a", bg="#f0fdf4", border="#86efac"), unsafe_allow_html=True)
    r2[2].markdown(_mi_card("Lost Float",          m["n_float_lost"],     colour="#d97706", bg="#fffbeb", border="#fcd34d"), unsafe_allow_html=True)
    r2[3].markdown(_mi_card("Gained Float",        m["n_float_gained"],   colour="#16a34a", bg="#f0fdf4", border="#86efac"), unsafe_allow_html=True)
    r2[4].markdown(_mi_card("Net Project Movement", pfm_str,              colour=pfm_colour), unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ---- Written commentary --------------------------------------------------
    st.markdown(
        '<div style="font-size:12px;font-weight:700;color:#94A3B8;'
        'letter-spacing:1px;text-transform:uppercase;margin-bottom:10px;">'
        'Intelligence Summary</div>',
        unsafe_allow_html=True,
    )

    commentary = _build_commentary(m, prev_name, curr_name)
    for heading, body, col, bg, border in commentary:
        st.markdown(_commentary_box(heading, body, col, bg, border), unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ---- Detailed sections (tabs) --------------------------------------------
    (tab_slips, tab_gains, tab_float_loss, tab_float_gain,
     tab_new_crit, tab_rem_crit, tab_added, tab_removed,
     tab_wbs, tab_export) = st.tabs([
        "Biggest Slips",
        "Biggest Gains",
        "Float Losses",
        "Float Gains",
        "New Critical",
        "No Longer Critical",
        "Added",
        "Removed",
        "WBS Movement",
        "Export",
    ])

    # Helper: common columns for merged common activities
    MOVE_COLS = {
        "task_code":           "Activity ID",
        "task_name_c":         "Activity Name",
        "wbs_path_c":          "WBS",
        "eff_finish_p":        "Previous Finish",
        "eff_finish_c":        "Current Finish",
        "finish_move":         "Finish Movement (d)",
        "total_float_days_p":  "Prev Float (d)",
        "total_float_days_c":  "Curr Float (d)",
        "float_move":          "Float Movement (d)",
    }
    ALT_NAME = {k.replace("_c", "_p"): v for k, v in MOVE_COLS.items()}

    def _safe_display(df, col_map=None):
        if df is None or df.empty:
            return pd.DataFrame(columns=["No data"])
        cols = col_map or MOVE_COLS
        avail = {k: v for k, v in cols.items() if k in df.columns}
        # Fallbacks for name column
        if "task_name_c" not in df.columns and "task_name_p" in df.columns:
            avail["task_name_p"] = "Activity Name"
        out = df[list(avail.keys())].copy().rename(columns=avail)
        for col in ["Previous Finish","Current Finish","Previous Start","Current Start"]:
            if col in out.columns:
                out[col] = out[col].apply(format_date)
        for col in ["Prev Float (d)","Curr Float (d)","Float Movement (d)","Finish Movement (d)"]:
            if col in out.columns:
                out[col] = out[col].apply(lambda x: round(float(x),1) if x is not None and str(x) not in ("","nan") else "-")
        return out

    # -- Biggest slips ---------------------------------------------------------
    with tab_slips:
        st.markdown("**Activities with the largest finish date slips** (most slipped first).")
        df_slips = _safe_display(m["delayed"].head(50))
        if df_slips.empty or "No data" in df_slips.columns:
            st.success("No activities slipped between these two revisions.")
        else:
            # Colour the movement column
            def _red_pos(val):
                try:
                    return "color:#dc2626;font-weight:600;" if float(val) > 0 else ""
                except Exception:
                    return ""
            if "Finish Movement (d)" in df_slips.columns:
                st.dataframe(
                    df_slips.style.map(_red_pos, subset=["Finish Movement (d)"]),
                    hide_index=True
                )
            else:
                st.dataframe(df_slips, hide_index=True)

            if not m["delayed"].empty:
                top20 = m["delayed"].head(20).copy()
                top20 = top20.dropna(subset=["finish_move"])
                label_col = "task_name_c" if "task_name_c" in top20.columns else "task_code"
                top20["Label"] = top20["task_code"].astype(str) + "  " + top20[label_col].astype(str).str[:35]
                fig = px.bar(
                    top20, x="finish_move", y="Label", orientation="h",
                    title="Top 20 Biggest Finish Date Slips (days)",
                    labels={"finish_move":"Slip (days)","Label":""},
                    color_discrete_sequence=["#dc2626"],
                )
                fig.update_yaxes(autorange="reversed")
                fig.update_layout(height=500, margin=dict(l=10,r=10,t=40,b=10),
                                  plot_bgcolor="#ffffff", paper_bgcolor="#ffffff")
                st.plotly_chart(fig)

    # -- Biggest gains ---------------------------------------------------------
    with tab_gains:
        st.markdown("**Activities with the largest finish date improvements** (most improved first).")
        df_gains = _safe_display(m["improved"].head(50))
        if df_gains.empty or "No data" in df_gains.columns:
            st.success("No activities improved their finish date between these two revisions.")
        else:
            st.dataframe(df_gains, hide_index=True)
            if not m["improved"].empty:
                top20g = m["improved"].head(20).copy().dropna(subset=["finish_move"])
                label_col = "task_name_c" if "task_name_c" in top20g.columns else "task_code"
                top20g["Label"] = top20g["task_code"].astype(str) + "  " + top20g[label_col].astype(str).str[:35]
                top20g["improvement"] = top20g["finish_move"].abs()
                fig2 = px.bar(
                    top20g, x="improvement", y="Label", orientation="h",
                    title="Top 20 Biggest Finish Date Improvements (days)",
                    labels={"improvement":"Improvement (days)","Label":""},
                    color_discrete_sequence=["#16a34a"],
                )
                fig2.update_yaxes(autorange="reversed")
                fig2.update_layout(height=500, margin=dict(l=10,r=10,t=40,b=10),
                                   plot_bgcolor="#ffffff", paper_bgcolor="#ffffff")
                st.plotly_chart(fig2)

    # -- Float losses ----------------------------------------------------------
    with tab_float_loss:
        st.markdown("**Activities that lost the most float** -- these are moving towards the critical path.")
        df_fl = _safe_display(m["float_lost"].head(50))
        if df_fl.empty or "No data" in df_fl.columns:
            st.success("No activities lost float between these two revisions.")
        else:
            st.dataframe(df_fl, hide_index=True)

    # -- Float gains -----------------------------------------------------------
    with tab_float_gain:
        st.markdown("**Activities that gained the most float** -- these have moved away from the critical path.")
        df_fg = _safe_display(m["float_gained"].head(50))
        if df_fg.empty or "No data" in df_fg.columns:
            st.success("No activities gained float between these two revisions.")
        else:
            st.dataframe(df_fg, hide_index=True)

    # -- New critical activities ------------------------------------------------
    with tab_new_crit:
        if m["became_crit"].empty:
            st.success("No activities became critical in the current revision. Good news.")
        else:
            st.warning(
                f"**{len(m['became_crit'])} activities became critical** in the current revision. "
                "These activities now have zero or negative float and require immediate attention."
            )
            df_nc = _safe_display(m["became_crit"])
            st.dataframe(df_nc, hide_index=True)

    # -- No longer critical ----------------------------------------------------
    with tab_rem_crit:
        if m["no_longer_crit"].empty:
            st.info("No activities moved off the critical path in the current revision.")
        else:
            st.success(
                f"**{len(m['no_longer_crit'])} activities are no longer critical.** "
                "This indicates recovery on those paths."
            )
            df_rc = _safe_display(m["no_longer_crit"])
            st.dataframe(df_rc, hide_index=True)

    # -- Added activities ------------------------------------------------------
    with tab_added:
        if m["added"].empty:
            st.info("No new activities were added in the current revision.")
        else:
            st.info(f"**{len(m['added'])} activities added** to the current programme.")
            added_cols = {
                "task_code":"Activity ID","task_name":"Activity Name",
                "wbs_path":"WBS","eff_start":"Start","eff_finish":"Finish",
                "total_float_days":"Float (d)","status":"Status",
            }
            st.dataframe(_safe_display(m["added"], added_cols), hide_index=True)

    # -- Removed activities ----------------------------------------------------
    with tab_removed:
        if m["removed"].empty:
            st.info("No activities were removed in the current revision.")
        else:
            st.warning(
                f"**{len(m['removed'])} activities removed.** "
                "Confirm these are intentional and not accidental deletions."
            )
            rem_cols = {
                "task_code":"Activity ID","task_name":"Activity Name",
                "wbs_path":"WBS","eff_start":"Start","eff_finish":"Finish",
                "total_float_days":"Float (d)","status":"Status",
            }
            st.dataframe(_safe_display(m["removed"], rem_cols), hide_index=True)

    # -- WBS-level movement ----------------------------------------------------
    with tab_wbs:
        st.markdown("**Finish date and critical path movement broken down by WBS area.**")
        if m["wbs_summary"].empty:
            st.info("WBS data not available for movement analysis.")
        else:
            wbs_disp = m["wbs_summary"].copy()
            wbs_disp = wbs_disp.rename(columns={
                "wbs_top":           "WBS Area",
                "total_activities":  "Total Activities",
                "delayed_count":     "Delayed",
                "improved_count":    "Improved",
                "avg_finish_move":   "Avg Finish Move (d)",
                "max_finish_slip":   "Max Slip (d)",
                "crit_gained":       "Became Critical",
                "crit_lost":         "No Longer Critical",
            })
            wbs_disp = wbs_disp.sort_values("Max Slip (d)", ascending=False)
            st.dataframe(wbs_disp, hide_index=True)

            # Chart: delayed count by WBS
            if not wbs_disp.empty and "Delayed" in wbs_disp.columns:
                fig_wbs = px.bar(
                    wbs_disp[wbs_disp["Delayed"] > 0].head(20),
                    x="Delayed", y="WBS Area", orientation="h",
                    title="Delayed Activities by WBS Area",
                    labels={"Delayed":"Activities Delayed","WBS Area":""},
                    color_discrete_sequence=["#dc2626"],
                )
                fig_wbs.update_yaxes(autorange="reversed")
                fig_wbs.update_layout(height=400, margin=dict(l=10,r=10,t=40,b=10),
                                      plot_bgcolor="#ffffff", paper_bgcolor="#ffffff")
                st.plotly_chart(fig_wbs)

            # Chart: average movement by WBS
            wbs_nonzero = wbs_disp[wbs_disp["Avg Finish Move (d)"] != 0].head(20)
            if not wbs_nonzero.empty:
                fig_avg = px.bar(
                    wbs_nonzero,
                    x="Avg Finish Move (d)", y="WBS Area", orientation="h",
                    title="Average Finish Movement by WBS (positive = slipped)",
                    labels={"Avg Finish Move (d)":"Average (days)","WBS Area":""},
                    color="Avg Finish Move (d)",
                    color_continuous_scale=["#16a34a","#f5f5f5","#dc2626"],
                    color_continuous_midpoint=0,
                )
                fig_avg.update_yaxes(autorange="reversed")
                fig_avg.update_layout(height=400, margin=dict(l=10,r=10,t=40,b=10),
                                      plot_bgcolor="#ffffff", paper_bgcolor="#ffffff",
                                      coloraxis_showscale=False)
                st.plotly_chart(fig_avg)

    # -- Export ----------------------------------------------------------------
    with tab_export:
        st.markdown("**Download the full movement analysis as a formatted Excel workbook.**")

        summary_sheet = pd.DataFrame({
            "Metric": [
                "Previous Programme", "Current Programme",
                "Previous Finish", "Current Finish",
                "Net Project Finish Movement",
                "Activities Added", "Activities Removed",
                "Activities Delayed", "Activities Improved",
                "Became Critical", "No Longer Critical",
                "Lost Float", "Gained Float",
            ],
            "Value": [
                prev_name, curr_name,
                format_date(m["prev_end"]), format_date(m["curr_end"]),
                pfm_str,
                m["n_added"], m["n_removed"],
                m["n_delayed"], m["n_improved"],
                m["n_became_crit"], m["n_no_longer_crit"],
                m["n_float_lost"], m["n_float_gained"],
            ],
        })

        commentary_sheet = pd.DataFrame({
            "Section": [h for h, *_ in commentary],
            "Commentary": [b.replace("<strong>","").replace("</strong>","") for _,b,*_ in commentary],
        })

        export_sheets = {
            "Summary":          summary_sheet,
            "Commentary":       commentary_sheet,
            "Biggest Slips":    _safe_display(m["delayed"].head(100)),
            "Biggest Gains":    _safe_display(m["improved"].head(100)),
            "Float Losses":     _safe_display(m["float_lost"].head(100)),
            "Float Gains":      _safe_display(m["float_gained"].head(100)),
            "Became Critical":  _safe_display(m["became_crit"]),
            "No Longer Critical": _safe_display(m["no_longer_crit"]),
            "Added Activities": _safe_display(
                m["added"],
                {"task_code":"Activity ID","task_name":"Activity Name",
                 "wbs_path":"WBS","eff_start":"Start","eff_finish":"Finish",
                 "total_float_days":"Float (d)","status":"Status"},
            ),
            "Removed Activities": _safe_display(
                m["removed"],
                {"task_code":"Activity ID","task_name":"Activity Name",
                 "wbs_path":"WBS","eff_start":"Start","eff_finish":"Finish",
                 "total_float_days":"Float (d)","status":"Status"},
            ),
        }
        if not m["wbs_summary"].empty:
            export_sheets["WBS Movement"] = m["wbs_summary"].rename(columns={
                "wbs_top":"WBS Area","total_activities":"Total Activities",
                "delayed_count":"Delayed","improved_count":"Improved",
                "avg_finish_move":"Avg Finish Move (d)","max_finish_slip":"Max Slip (d)",
                "crit_gained":"Became Critical","crit_lost":"No Longer Critical",
            })

        xls_bytes = export_df_to_excel(export_sheets)

        st.download_button(
            label="📥  Download Movement Intelligence Report",
            data=xls_bytes,
            file_name=f"movement_intelligence_{datetime.now().strftime('%Y%m%d')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            help="Exports all movement data across multiple sheets including commentary.",
            key="dl_007",
        )

        st.markdown(
            f"""
            <div style="background:#f8fafc;border:1px solid #E2E8F0;border-radius:8px;
                        padding:14px 18px;margin-top:12px;">
                <div style="font-size:12px;font-weight:700;color:#0B1F33;
                            margin-bottom:8px;">Workbook contents</div>
                <div style="font-size:12px;color:#64748B;line-height:2;">
                    {"".join(f"<strong>{k}</strong> &nbsp;|&nbsp; " for k in export_sheets.keys())}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )




# -----------------------------------------------------------------------------
# PAGE: PM ACTIONS
# -----------------------------------------------------------------------------

# Risk keywords for notes scanning
_RISK_WORDS = [
    "risk","delay","delayed","delays","blocked","block","constraint","access",
    "design","procurement","client","instruction","CE","EWN","change",
    "hold","pending","late","overrun","issue","dispute","claim",
]

def _generate_actions(
    tasks: pd.DataFrame,
    rels: pd.DataFrame,
    near_crit_days: float,
    notes_text: str = "",
) -> pd.DataFrame:
    """
    Analyse the programme and generate a prioritised PM action list.
    Returns a DataFrame with one row per action.
    """
    rows = []
    now  = datetime.now()
    eight_weeks = now + timedelta(weeks=8)
    four_weeks  = now + timedelta(weeks=4)

    tasks = get_critical_threshold(tasks, near_crit_days)

    # Build predecessor/successor sets from rels
    tasks_with_pred = set()
    tasks_with_succ = set()
    if not rels.empty:
        if "succ_task_id" in rels.columns:
            tasks_with_pred = set(rels["succ_task_id"].dropna())
        if "pred_task_id" in rels.columns:
            tasks_with_succ = set(rels["pred_task_id"].dropna())

    def _wbs(row):
        wbs = row.get("wbs_path", "") if "wbs_path" in row.index else ""
        return str(wbs).split(" > ")[0] if wbs and str(wbs).strip() not in ("","nan") else "-"

    def _row(priority, category, task_code, task_name, wbs, issue, why, action):
        return {
            "Priority":         priority,
            "Category":         category,
            "Activity ID":      str(task_code),
            "Activity Name":    str(task_name),
            "WBS":              str(wbs),
            "Issue":            str(issue),
            "Why It Matters":   str(why),
            "Suggested Action": str(action),
            "Owner":            "",
            "Due Date":         "",
            "Status":           "Open",
        }

    # -- 1. Negative float -----------------------------------------------------
    neg = tasks[tasks["total_float_days"].apply(lambda f: safe_float(f,0) < 0)] \
        if "total_float_days" in tasks.columns else pd.DataFrame()
    for _, t in neg.iterrows():
        tf = round(safe_float(t.get("total_float_days"), 0), 1)
        rows.append(_row(
            "High", "Negative Float",
            t.get("task_code",""), t.get("task_name",""), _wbs(t),
            f"Activity has {tf} days negative float.",
            "Negative float means the current schedule cannot meet its target date on this path. "
            "Every day of inaction increases the delay.",
            f"Investigate the cause of the {abs(tf)}-day overrun. "
            "Agree a recovery plan with the planner and update the programme.",
        ))

    # -- 2. Critical activities not started ------------------------------------
    crit_ns = tasks[
        tasks["is_critical"] &
        tasks["status"].apply(lambda s: str(s) in ("TK_NotStart","Not Started"))
    ] if "status" in tasks.columns and "is_critical" in tasks.columns else pd.DataFrame()
    for _, t in crit_ns.iterrows():
        finish = t.get("eff_finish")
        finish_str = format_date(finish) if finish else "unknown"
        rows.append(_row(
            "High", "Critical Not Started",
            t.get("task_code",""), t.get("task_name",""), _wbs(t),
            f"Critical activity not yet started. Target finish: {finish_str}.",
            "Any delay to a critical activity directly delays the project finish date. "
            "There is no float to absorb slippage.",
            "Confirm start date with the responsible party. "
            "If start is at risk, escalate immediately and review recovery options.",
        ))

    # -- 3. Near-critical due within 4 weeks -----------------------------------
    nc_due = tasks[
        tasks["is_near_critical"] &
        tasks["eff_finish"].apply(
            lambda d: d is not None and hasattr(d,"date") and d <= four_weeks
        )
    ] if "eff_finish" in tasks.columns else pd.DataFrame()
    for _, t in nc_due.iterrows():
        tf = round(safe_float(t.get("total_float_days"),0), 1)
        rows.append(_row(
            "High", "Near-Critical Due Soon",
            t.get("task_code",""), t.get("task_name",""), _wbs(t),
            f"Near-critical activity due within 4 weeks. Float: {tf} days.",
            "With only a small float buffer and a near-term finish, any disruption "
            "will push this activity onto the critical path.",
            "Confirm the activity is progressing on schedule. "
            "If at risk, treat as critical and raise with the team now.",
        ))

    # -- 4. Near-critical due within 4-8 weeks --------------------------------
    nc_due_med = tasks[
        tasks["is_near_critical"] &
        tasks["eff_finish"].apply(
            lambda d: d is not None and hasattr(d,"date") and four_weeks < d <= eight_weeks
        )
    ] if "eff_finish" in tasks.columns else pd.DataFrame()
    for _, t in nc_due_med.iterrows():
        tf = round(safe_float(t.get("total_float_days"),0), 1)
        rows.append(_row(
            "Medium", "Near-Critical Due Soon",
            t.get("task_code",""), t.get("task_name",""), _wbs(t),
            f"Near-critical activity due in 4-8 weeks. Float: {tf} days.",
            "Limited float with an upcoming finish date. Monitor closely.",
            "Review progress and confirm no blockers. Add to weekly look-ahead.",
        ))

    # -- 5. No predecessor (open start) ----------------------------------------
    no_pred = tasks[~tasks["task_id"].isin(tasks_with_pred)] \
        if "task_id" in tasks.columns else pd.DataFrame()
    # Exclude milestones with no predecessors intentionally (start milestones)
    no_pred = no_pred[no_pred.get("task_type","").apply(
        lambda t: "Milestone" not in str(t) and "LOE" not in str(t)
    )] if "task_type" in no_pred.columns else no_pred
    for _, t in no_pred.head(20).iterrows():
        rows.append(_row(
            "Medium", "Open Logic - No Predecessor",
            t.get("task_code",""), t.get("task_name",""), _wbs(t),
            "Activity has no predecessor. Logic is open at the start.",
            "Open-start activities are not driven by schedule logic. "
            "Float calculations may be unreliable for this activity.",
            "Review with the planner. Add a predecessor or confirm the open start is intentional. "
            "If intentional, add a Start On or After constraint.",
        ))

    # -- 6. No successor (open finish) -----------------------------------------
    no_succ = tasks[~tasks["task_id"].isin(tasks_with_succ)] \
        if "task_id" in tasks.columns else pd.DataFrame()
    no_succ = no_succ[no_succ.get("task_type","").apply(
        lambda t: "Finish Milestone" not in str(t) and "LOE" not in str(t)
    )] if "task_type" in no_succ.columns else no_succ
    for _, t in no_succ.head(20).iterrows():
        rows.append(_row(
            "Medium", "Open Logic - No Successor",
            t.get("task_code",""), t.get("task_name",""), _wbs(t),
            "Activity has no successor. Logic is open at the finish.",
            "Open-finish activities may have artificially high float and can mask "
            "schedule risk downstream.",
            "Review with the planner. Add a successor or confirm the open finish "
            "is a deliberate programme end point.",
        ))

    # -- 7. Excessive lag (> 10 days) ------------------------------------------
    if not rels.empty and "lag_days" in rels.columns:
        big_lag = rels[rels["lag_days"].apply(lambda l: safe_float(l,0) > 10)]
        for _, r in big_lag.head(15).iterrows():
            code = r.get("succ_task_code", r.get("succ_task_id",""))
            name = r.get("succ_task_name","")
            lag  = safe_float(r.get("lag_days",0),0)
            pred = r.get("pred_task_code", r.get("pred_task_id",""))
            # Get WBS from tasks
            match = tasks[tasks["task_code"] == code]
            wbs   = _wbs(match.iloc[0]) if not match.empty else "-"
            rows.append(_row(
                "Medium", "Excessive Lag",
                code, name, wbs,
                f"Relationship from {pred} has {int(lag)} days lag.",
                "Excessive lag hides schedule risk. It should be replaced with properly "
                "sequenced activities so the plan reflects real work.",
                f"Challenge the {int(lag)}-day lag with the planner. "
                "Replace with an intermediate activity or reduce the lag to the minimum justified.",
            ))

    # -- 8. Long duration activities (> 60 working days) -----------------------
    long_dur = tasks[tasks["orig_dur_days"].apply(lambda d: safe_float(d,0) > 60)] \
        if "orig_dur_days" in tasks.columns else pd.DataFrame()
    for _, t in long_dur.head(15).iterrows():
        dur = int(safe_float(t.get("orig_dur_days",0),0))
        rows.append(_row(
            "Low", "Long Duration",
            t.get("task_code",""), t.get("task_name",""), _wbs(t),
            f"Activity duration is {dur} working days.",
            "Activities longer than 60 days are difficult to manage and monitor. "
            "Progress is hard to measure and problems can go undetected for weeks.",
            f"Review whether the {dur}-day activity can be broken into smaller "
            "work packages of 20-30 days. Discuss with the planner.",
        ))

    # -- 9. Activities flagged in planning notes --------------------------------
    if notes_text and "task_code" in tasks.columns:
        for word in _RISK_WORDS:
            pattern = re.compile(r"\b" + re.escape(word) + r"\b", re.IGNORECASE)
            if not pattern.search(notes_text):
                continue
            # Find activity IDs mentioned near this risk word in the notes
            for _, t in tasks.iterrows():
                code = str(t.get("task_code",""))
                if code and code in notes_text:
                    # Check the risk word appears within 200 chars of the activity ID
                    idx = notes_text.find(code)
                    snippet = notes_text[max(0,idx-200):idx+200]
                    if pattern.search(snippet):
                        rows.append(_row(
                            "High", "Planning Notes Risk",
                            code, t.get("task_name",""), _wbs(t),
                            f"Activity mentioned in planning notes alongside the word '{word}'.",
                            "Planning notes flag a potential issue against this activity that "
                            "may not be visible in the programme alone.",
                            f"Review the planning notes for {code} and confirm the '{word}' "
                            "item has been actioned or is being tracked.",
                        ))
                        break   # one action per risk word per activity

    # -- 10. Activities with constraints ---------------------------------------
    if "cstr_type" in tasks.columns:
        constrained = tasks[tasks["cstr_type"].apply(
            lambda x: bool(x) and str(x).strip() not in ("","None","nan")
        )]
        for _, t in constrained.head(20).iterrows():
            cstr = str(t.get("cstr_type",""))
            rows.append(_row(
                "Low", "Constraint",
                t.get("task_code",""), t.get("task_name",""), _wbs(t),
                f"Activity has a constraint: {cstr}.",
                "Constraints override schedule logic and can cause artificial float "
                "or negative float. Each constraint should be justified.",
                f"Confirm the {cstr} constraint is still valid. "
                "If the constraint is no longer required, ask the planner to remove it.",
            ))

    if not rows:
        return pd.DataFrame()

    df = pd.DataFrame(rows)

    # Deduplicate: same Activity ID + Category
    df = df.drop_duplicates(subset=["Activity ID","Category"]).reset_index(drop=True)

    # Sort: High first, then Medium, then Low; within each by Category
    priority_order = {"High": 0, "Medium": 1, "Low": 2}
    df["_sort"] = df["Priority"].map(priority_order)
    df = df.sort_values(["_sort","Category","Activity ID"]).drop(columns=["_sort"]).reset_index(drop=True)

    return df


def page_pm_actions(data: dict, near_crit_days: float):
    """
    PM Action Dashboard.
    Automatically generates a prioritised action list from the uploaded programme.
    Each action has priority, category, issue, why it matters, suggested action,
    and editable owner / due date / status fields.
    """
    st.title("📋 PM Actions")
    st.caption(
        "Automatically generated from your programme. "
        "Each action is prioritised and includes a suggested next step. "
        "Update the Owner, Due Date and Status columns to track progress."
    )

    tasks = data["tasks_df"]
    rels  = data["relationships_df"]

    if tasks.empty:
        st.warning("No activities found. Please upload a programme first.")
        return

    # Retrieve any planning notes text from session state (set by Planning Notes page)
    notes_text = st.session_state.get("_notes_text", "")

    # ---- Generate / cache actions -------------------------------------------
    prog_key = st.session_state.get("_xer_cache_key", "")
    cache_key = f"_pm_actions_{prog_key}_{near_crit_days}"

    if st.session_state.get("_pm_actions_key") != cache_key:
        with st.spinner("Analysing programme..."):
            actions_df = _generate_actions(tasks, rels, near_crit_days, notes_text)
        st.session_state["_pm_actions_df"]  = actions_df
        st.session_state["_pm_actions_key"] = cache_key
    else:
        actions_df = st.session_state["_pm_actions_df"]

    if actions_df.empty:
        st.success(
            "No actions generated. The programme has no obvious issues detected by "
            "the automated checks. Review the Health Check page for more detail."
        )
        return

    # ---- Summary banner ------------------------------------------------------
    n_high   = int((actions_df["Priority"] == "High").sum())
    n_med    = int((actions_df["Priority"] == "Medium").sum())
    n_low    = int((actions_df["Priority"] == "Low").sum())
    n_total  = len(actions_df)

    st.markdown(
        f"""
        <div style="background:#0B1F33;border-radius:12px;padding:20px 24px;
                    margin-bottom:20px;display:flex;gap:24px;flex-wrap:wrap;
                    align-items:center;">
            <div>
                <div style="font-size:11px;color:#64748B;text-transform:uppercase;
                            letter-spacing:1px;">Total Actions</div>
                <div style="font-size:32px;font-weight:800;color:#E8951D;
                            line-height:1;margin-top:4px;">{n_total}</div>
            </div>
            <div style="width:1px;background:#1e3a5f;align-self:stretch;"></div>
            <div>
                <div style="font-size:11px;color:#64748B;text-transform:uppercase;
                            letter-spacing:1px;">High Priority</div>
                <div style="font-size:28px;font-weight:800;color:#dc2626;
                            line-height:1;margin-top:4px;">{n_high}</div>
            </div>
            <div>
                <div style="font-size:11px;color:#64748B;text-transform:uppercase;
                            letter-spacing:1px;">Medium</div>
                <div style="font-size:28px;font-weight:800;color:#d97706;
                            line-height:1;margin-top:4px;">{n_med}</div>
            </div>
            <div>
                <div style="font-size:11px;color:#64748B;text-transform:uppercase;
                            letter-spacing:1px;">Low</div>
                <div style="font-size:28px;font-weight:800;color:#64748B;
                            line-height:1;margin-top:4px;">{n_low}</div>
            </div>
            <div style="margin-left:auto;">
                <div style="font-size:11px;color:#475569;line-height:1.6;">
                    Generated from negative float, critical path,<br>
                    open logic, lag, duration and planning notes.
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # ---- Category breakdown cards -------------------------------------------
    cats = actions_df.groupby("Category").size().reset_index(name="n")
    cols = st.columns(min(len(cats), 5))
    cat_colours = {
        "Negative Float":           "#dc2626",
        "Critical Not Started":     "#b91c1c",
        "Near-Critical Due Soon":   "#d97706",
        "Open Logic - No Predecessor": "#7c3aed",
        "Open Logic - No Successor":   "#6d28d9",
        "Excessive Lag":            "#0369a1",
        "Long Duration":            "#0891b2",
        "Planning Notes Risk":      "#dc2626",
        "Constraint":               "#475569",
    }
    for i, (_, cat_row) in enumerate(cats.iterrows()):
        col = cols[i % len(cols)]
        colour = cat_colours.get(cat_row["Category"], "#374151")
        col.markdown(
            f'<div style="background:#ffffff;border:1px solid #E2E8F0;border-radius:8px;'
            f'padding:12px 14px;border-top:3px solid {colour};margin-bottom:6px;">'
            f'<div style="font-size:11px;color:#94A3B8;text-transform:uppercase;'
            f'letter-spacing:0.8px;margin-bottom:4px;">{cat_row["Category"]}</div>'
            f'<div style="font-size:22px;font-weight:800;color:{colour};">{cat_row["n"]}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    st.markdown("<br>", unsafe_allow_html=True)

    # ---- Filters ------------------------------------------------------------
    with st.expander("Filter actions", expanded=False):
        fc1, fc2, fc3, fc4, fc5 = st.columns(5)

        all_priorities = ["All"] + sorted(actions_df["Priority"].unique().tolist(),
                          key=lambda p: {"High":0,"Medium":1,"Low":2}.get(p,9))
        f_priority = fc1.selectbox("Priority",  all_priorities, key="pma_f_pri")

        all_cats = ["All"] + sorted(actions_df["Category"].unique().tolist())
        f_cat    = fc2.selectbox("Category",  all_cats, key="pma_f_cat")

        all_wbs = ["All"] + sorted(actions_df["WBS"].unique().tolist())
        f_wbs   = fc3.selectbox("WBS",        all_wbs, key="pma_f_wbs")

        all_owners = ["All"] + sorted([o for o in actions_df["Owner"].unique() if o])
        f_owner  = fc4.selectbox("Owner",      all_owners, key="pma_f_own")

        all_status = ["All", "Open", "In Progress", "Closed"]
        f_status = fc5.selectbox("Status",     all_status, key="pma_f_sta")

    # Apply filters
    display_df = actions_df.copy()
    if f_priority != "All":
        display_df = display_df[display_df["Priority"] == f_priority]
    if f_cat != "All":
        display_df = display_df[display_df["Category"] == f_cat]
    if f_wbs != "All":
        display_df = display_df[display_df["WBS"] == f_wbs]
    if f_owner != "All":
        display_df = display_df[display_df["Owner"] == f_owner]
    if f_status != "All":
        display_df = display_df[display_df["Status"] == f_status]

    n_shown = len(display_df)
    st.caption(f"Showing {n_shown} of {n_total} actions.")

    if display_df.empty:
        st.info("No actions match the current filters.")
        return

    # ---- Editable action table ----------------------------------------------
    st.markdown(
        '<div style="font-size:12px;font-weight:700;color:#94A3B8;letter-spacing:1px;'
        'text-transform:uppercase;margin-bottom:8px;">Action List</div>',
        unsafe_allow_html=True,
    )
    st.caption(
        "Edit the Owner, Due Date and Status columns directly in the table below. "
        "Changes are saved while you are on this page."
    )

    # Colour-code Priority column
    def _style_priority(row):
        p = row.get("Priority","")
        base = [""] * len(row)
        idx  = list(row.index)
        try:
            pi = idx.index("Priority")
        except ValueError:
            return base
        colours = {"High":"background-color:#fef2f2;color:#991b1b;font-weight:700;",
                   "Medium":"background-color:#fffbeb;color:#92400e;font-weight:600;",
                   "Low":"background-color:#f8fafc;color:#475569;"}
        base[pi] = colours.get(p,"")
        return base

    # Use st.data_editor for the editable fields
    edit_cols = [
        "Priority","Category","Activity ID","Activity Name","WBS",
        "Issue","Why It Matters","Suggested Action",
        "Owner","Due Date","Status",
    ]
    avail_edit = [c for c in edit_cols if c in display_df.columns]

    _edit_data = display_df[avail_edit].copy()
    _edit_data = _edit_data.loc[:, ~_edit_data.columns.duplicated()]
    edited_df = st.data_editor(
        _edit_data,
        hide_index=True,
        num_rows="fixed",
        column_config={
            "Priority": st.column_config.SelectboxColumn(
                "Priority",
                options=["High","Medium","Low"],
                width="small",
            ),
            "Status": st.column_config.SelectboxColumn(
                "Status",
                options=["Open","In Progress","Closed"],
                width="small",
            ),
            "Owner": st.column_config.TextColumn("Owner", width="small"),
            "Due Date": st.column_config.TextColumn("Due Date", width="small",
                                                     help="e.g. 01/06/2025"),
            "Issue":            st.column_config.TextColumn("Issue",            width="medium"),
            "Why It Matters":   st.column_config.TextColumn("Why It Matters",   width="large"),
            "Suggested Action": st.column_config.TextColumn("Suggested Action", width="large"),
            "Activity ID":      st.column_config.TextColumn("Activity ID",      width="small"),
            "Activity Name":    st.column_config.TextColumn("Activity Name",    width="medium"),
            "WBS":              st.column_config.TextColumn("WBS",              width="medium"),
            "Category":         st.column_config.TextColumn("Category",         width="medium"),
        },
        key="pma_editor",
    )

    # Persist edits back to session state
    if edited_df is not None and not edited_df.empty:
        for col in ["Owner","Due Date","Status"]:
            if col in edited_df.columns:
                actions_df.loc[display_df.index, col] = edited_df[col].values
        st.session_state["_pm_actions_df"] = actions_df

    # ---- Highlighted high-priority section ----------------------------------
    high_df = display_df[display_df["Priority"] == "High"]
    if not high_df.empty:
        st.divider()
        st.markdown(
            f'<div style="font-size:12px;font-weight:700;color:#dc2626;letter-spacing:1px;'
            f'text-transform:uppercase;margin-bottom:12px;">'
            f'High Priority Actions ({len(high_df)})</div>',
            unsafe_allow_html=True,
        )
        for _, act in high_df.iterrows():
            st.markdown(
                f"""
                <div style="background:#ffffff;border:1px solid #fca5a5;border-left:4px solid #dc2626;
                            border-radius:8px;padding:14px 18px;margin-bottom:8px;">
                    <div style="display:flex;justify-content:space-between;align-items:flex-start;
                                flex-wrap:wrap;gap:8px;">
                        <div style="flex:1;min-width:200px;">
                            <span style="background:#dc2626;color:white;padding:2px 8px;
                                         border-radius:4px;font-size:10px;font-weight:700;
                                         letter-spacing:0.5px;text-transform:uppercase;">
                                {act.get("Category","")}
                            </span>
                            <div style="font-weight:700;color:#0B1F33;font-size:14px;
                                        margin-top:6px;">
                                {act.get("Activity ID","")} - {act.get("Activity Name","")}
                            </div>
                            <div style="font-size:12px;color:#64748B;margin-top:2px;">
                                {act.get("WBS","")}
                            </div>
                        </div>
                    </div>
                    <div style="margin-top:10px;display:grid;grid-template-columns:1fr 1fr;
                                gap:10px;">
                        <div>
                            <div style="font-size:10px;font-weight:700;color:#94A3B8;
                                        text-transform:uppercase;letter-spacing:0.8px;">Issue</div>
                            <div style="font-size:13px;color:#334155;margin-top:3px;">
                                {act.get("Issue","")}</div>
                        </div>
                        <div>
                            <div style="font-size:10px;font-weight:700;color:#94A3B8;
                                        text-transform:uppercase;letter-spacing:0.8px;">
                                Suggested Action</div>
                            <div style="font-size:13px;color:#334155;margin-top:3px;">
                                {act.get("Suggested Action","")}</div>
                        </div>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    # ---- Export -------------------------------------------------------------
    st.divider()

    # Merge any edits back for export
    export_df = st.session_state.get("_pm_actions_df", actions_df).copy()

    # Add summary sheet
    summary_rows = {
        "Metric": ["Total Actions","High Priority","Medium Priority","Low Priority",
                   "Open","In Progress","Closed"],
        "Count":  [
            len(export_df),
            int((export_df["Priority"]=="High").sum()),
            int((export_df["Priority"]=="Medium").sum()),
            int((export_df["Priority"]=="Low").sum()),
            int((export_df["Status"]=="Open").sum()),
            int((export_df["Status"]=="In Progress").sum()),
            int((export_df["Status"]=="Closed").sum()),
        ],
    }

    export_sheets = {
        "Summary":           pd.DataFrame(summary_rows),
        "All Actions":       export_df,
        "High Priority":     export_df[export_df["Priority"]=="High"],
        "Medium Priority":   export_df[export_df["Priority"]=="Medium"],
        "Low Priority":      export_df[export_df["Priority"]=="Low"],
    }
    # Category sheets
    for cat in export_df["Category"].unique():
        safe_name = cat[:31]
        export_sheets[safe_name] = export_df[export_df["Category"]==cat]

    xls_bytes = export_df_to_excel(export_sheets)

    dl_col, _ = st.columns([1,3])
    dl_col.download_button(
        label="📥  Export Action List to Excel",
        data=xls_bytes,
        file_name=f"pm_actions_{datetime.now().strftime('%Y%m%d')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        help="Exports all actions with Summary, High/Medium/Low priority sheets and one sheet per category.",
        key="dl_008",
    )

    st.caption(
        "Tip: Export the action list and share it with your delivery team. "
        "Update Owner and Status as actions are completed."
    )



# -----------------------------------------------------------------------------
# PAGE: LOOKAHEAD PLANNER
# -----------------------------------------------------------------------------

def _lookahead_card(label: str, value, sublabel: str = "",
                    colour: str = "#0B1F33", bg: str = "#ffffff",
                    border_top: str = "#0B1F33") -> str:
    """Compact metric card matching PlanTrace brand."""
    return (
        f'<div style="background:{bg};border:1px solid #E2E8F0;border-radius:10px;'
        f'padding:16px 18px;border-top:3px solid {border_top};'
        f'box-shadow:0 1px 4px rgba(11,31,51,0.06);">'
        f'<div style="font-size:10px;font-weight:700;color:#94A3B8;letter-spacing:1px;'
        f'text-transform:uppercase;margin-bottom:6px;">{label}</div>'
        f'<div style="font-size:26px;font-weight:800;color:{colour};line-height:1;">{value}</div>'
        f'{"" if not sublabel else f"<div style=font-size:11px;color:#64748B;margin-top:4px;>{sublabel}</div>"}'
        f'</div>'
    )


def _lookahead_section_header(title: str, count: int, colour: str = "#0B1F33") -> str:
    return (
        f'<div style="display:flex;align-items:center;gap:10px;margin:16px 0 8px 0;">'
        f'<div style="font-size:14px;font-weight:700;color:{colour};">{title}</div>'
        f'<div style="background:{colour};color:white;border-radius:12px;'
        f'padding:1px 9px;font-size:11px;font-weight:700;">{count}</div>'
        f'</div>'
    )


def _week_labour(task_res_df: pd.DataFrame, tasks_df: pd.DataFrame,
                 window_start: datetime, window_end: datetime) -> pd.DataFrame:
    """
    Expand resource loading into a weekly series within the lookahead window.
    Returns a DataFrame with columns: week, rsrc_name, qty.
    """
    if task_res_df.empty:
        return pd.DataFrame()

    rows = []
    for _, r in task_res_df.iterrows():
        s = pd.to_datetime(r.get("target_start") or r.get("target_start_date"), errors="coerce")
        e = pd.to_datetime(r.get("target_finish") or r.get("target_end_date"), errors="coerce")
        if pd.isna(s) or pd.isna(e) or s > e:
            continue

        # Clip to window
        s_clip = max(s, pd.Timestamp(window_start))
        e_clip = min(e, pd.Timestamp(window_end))
        if s_clip > e_clip:
            continue

        qty = safe_float(r.get("target_qty", 0), 0)
        if qty == 0:
            continue

        total_days = max(1, (e - s).days)
        window_days = max(1, (e_clip - s_clip).days)
        qty_in_window = qty * (window_days / total_days)

        weeks = max(1, math.ceil(window_days / 7))
        qty_per_week = qty_in_window / weeks

        current = s_clip
        for _ in range(weeks):
            rows.append({
                "week":      current.to_period("W").start_time,
                "rsrc_name": str(r.get("rsrc_name", r.get("rsrc_id", "Unknown"))),
                "qty":       qty_per_week,
            })
            current += timedelta(weeks=1)

    return pd.DataFrame(rows) if rows else pd.DataFrame()


def page_lookahead(data: dict, near_crit_days: float):
    """
    Lookahead Planner page.
    Shows all activities starting, finishing, or at risk within a user-defined
    lookahead window, with labour demand and export.
    """
    st.title("📅 Lookahead Planner")
    st.caption(
        "A short-term delivery view from your programme. "
        "Select a lookahead window to see what is starting, finishing and at risk."
    )

    tasks    = data["tasks_df"]
    rels     = data["relationships_df"]
    task_res = data.get("task_resources_df", pd.DataFrame())

    if tasks.empty:
        st.warning("No activities found. Please upload a programme first.")
        return

    tasks = get_critical_threshold(tasks, near_crit_days)

    # ---- Lookahead window selector ------------------------------------------
    st.markdown(
        '<div style="font-size:12px;font-weight:700;color:#94A3B8;letter-spacing:1px;'
        'text-transform:uppercase;margin-bottom:8px;">Lookahead Window</div>',
        unsafe_allow_html=True,
    )

    wcol1, wcol2 = st.columns([2, 3])

    with wcol1:
        window_opt = st.selectbox(
            "Window",
            ["2 Weeks", "4 Weeks", "6 Weeks", "12 Weeks", "Custom Date Range"],
            index=1,
            label_visibility="collapsed",
            key="la_window_opt",
        )

    now = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)

    window_map = {
        "2 Weeks":  14,
        "4 Weeks":  28,
        "6 Weeks":  42,
        "12 Weeks": 84,
    }

    with wcol2:
        if window_opt == "Custom Date Range":
            dc1, dc2 = st.columns(2)
            window_start = dc1.date_input(
                "From", value=now.date(), key="la_custom_start"
            )
            window_end = dc2.date_input(
                "To", value=(now + timedelta(weeks=4)).date(), key="la_custom_end"
            )
            window_start = datetime.combine(window_start, datetime.min.time())
            window_end   = datetime.combine(window_end,   datetime.max.time().replace(microsecond=0))
        else:
            days = window_map[window_opt]
            window_start = now
            window_end   = now + timedelta(days=days)
            st.markdown(
                f'<div style="padding:8px 14px;background:#0B1F33;border-radius:8px;'
                f'font-size:13px;color:#CBD5E1;display:inline-block;">'
                f'{window_start.strftime("%d %b %Y")} &rarr; '
                f'{window_end.strftime("%d %b %Y")}'
                f'</div>',
                unsafe_allow_html=True,
            )

    st.markdown("<br>", unsafe_allow_html=True)

    # (tasks and rels already extracted above for window computation)
    # Quick-access SAP panel in expander
    with st.expander("Selected Activity Summary", expanded=False, key="sap_exp_lookahead"):
        render_selected_activity_panel(
            tasks_df=tasks,
            rels_df=rels,
            near_crit_days=near_crit_days,
            context_key="lookahead_sap",
        )

    # ---- Filters ------------------------------------------------------------
    with st.expander("Filters", expanded=False):
        fl1, fl2, fl3, fl4 = st.columns(4)

        wbs_opts = ["All"]
        if "wbs_path" in tasks.columns:
            tops = tasks["wbs_path"].dropna().apply(
                lambda x: str(x).split(" > ")[0] if x and str(x).strip() not in ("","nan") else "Unknown"
            ).unique().tolist()
            wbs_opts += sorted(tops)
        f_wbs = fl1.selectbox("WBS / Area", wbs_opts, key="la_f_wbs")

        status_opts = ["All"] + sorted([
            s for s in tasks["status"].dropna().unique() if str(s).strip()
        ]) if "status" in tasks.columns else ["All"]
        f_status = fl2.selectbox("Activity Status", status_opts, key="la_f_stat")

        f_crit = fl3.selectbox(
            "Float Filter",
            ["All", "Critical only", "Near-critical only", "Negative float only"],
            key="la_f_crit",
        )

        f_type = fl4.selectbox(
            "Activity Type",
            ["All"] + (sorted([t for t in tasks["task_type"].dropna().unique() if str(t).strip()])
                       if "task_type" in tasks.columns else []),
            key="la_f_type",
        )

    # ---- Build lookahead activity sets --------------------------------------
    def _in_window(d, strict_start=False):
        """True if datetime d falls within the lookahead window."""
        if d is None or not hasattr(d, "date"):
            return False
        if strict_start:
            return window_start <= d <= window_end
        return d <= window_end

    # Activities starting in window
    starting = tasks[
        tasks["eff_start"].apply(lambda d: _in_window(d, strict_start=True))
    ].copy() if "eff_start" in tasks.columns else pd.DataFrame()

    # Activities finishing in window
    finishing = tasks[
        tasks["eff_finish"].apply(lambda d: _in_window(d, strict_start=True))
    ].copy() if "eff_finish" in tasks.columns else pd.DataFrame()

    # Activities spanning the window (started before, finish after window start)
    spanning = tasks[
        tasks["eff_start"].apply(lambda d: d is not None and hasattr(d,"date") and d < window_start) &
        tasks["eff_finish"].apply(lambda d: d is not None and hasattr(d,"date") and d >= window_start)
    ].copy() if ("eff_start" in tasks.columns and "eff_finish" in tasks.columns) else pd.DataFrame()

    # All activities in window = union
    in_window_ids = (
        set(starting["task_id"].tolist() if not starting.empty else []) |
        set(finishing["task_id"].tolist() if not finishing.empty else []) |
        set(spanning["task_id"].tolist() if not spanning.empty else [])
    )
    all_window = tasks[tasks["task_id"].isin(in_window_ids)].copy()

    # Apply filters to all_window
    def _apply_filters(df):
        if df.empty:
            return df
        if f_wbs != "All" and "wbs_path" in df.columns:
            df = df[df["wbs_path"].astype(str).str.startswith(f_wbs)]
        if f_status != "All" and "status" in df.columns:
            df = df[df["status"] == f_status]
        if f_crit == "Critical only" and "is_critical" in df.columns:
            df = df[df["is_critical"] == True]
        elif f_crit == "Near-critical only" and "is_near_critical" in df.columns:
            df = df[df["is_near_critical"] == True]
        elif f_crit == "Negative float only" and "total_float_days" in df.columns:
            df = df[df["total_float_days"].apply(lambda f: safe_float(f, 0) < 0)]
        if f_type != "All" and "task_type" in df.columns:
            df = df[df["task_type"] == f_type]
        return df

    starting  = _apply_filters(starting)
    finishing = _apply_filters(finishing)
    spanning  = _apply_filters(spanning)
    all_window = _apply_filters(all_window)

    # Open logic in window
    tasks_with_pred = set(rels["succ_task_id"].dropna()) if not rels.empty and "succ_task_id" in rels.columns else set()
    tasks_with_succ = set(rels["pred_task_id"].dropna()) if not rels.empty and "pred_task_id" in rels.columns else set()

    open_start_win  = all_window[~all_window["task_id"].isin(tasks_with_pred)] if "task_id" in all_window.columns else pd.DataFrame()
    open_finish_win = all_window[~all_window["task_id"].isin(tasks_with_succ)] if "task_id" in all_window.columns else pd.DataFrame()

    # Critical & near-critical in window
    crit_win  = all_window[all_window["is_critical"] == True] if "is_critical" in all_window.columns else pd.DataFrame()
    nc_win    = all_window[all_window["is_near_critical"] == True] if "is_near_critical" in all_window.columns else pd.DataFrame()
    neg_win   = all_window[all_window["total_float_days"].apply(lambda f: safe_float(f,0) < 0)] if "total_float_days" in all_window.columns else pd.DataFrame()

    # Notes matching
    notes_text = st.session_state.get("_notes_text", "")
    notes_ids  = set()
    if notes_text and "task_code" in all_window.columns:
        for code in all_window["task_code"].dropna():
            if str(code) in notes_text:
                notes_ids.add(str(code))
    notes_win = all_window[all_window["task_code"].astype(str).isin(notes_ids)] if notes_ids else pd.DataFrame()

    risk_ids = set()
    if notes_text and "task_code" in all_window.columns:
        for _, t in all_window.iterrows():
            code = str(t.get("task_code",""))
            if not code:
                continue
            idx = notes_text.find(code)
            if idx < 0:
                continue
            snippet = notes_text[max(0,idx-300):idx+300]
            for word in _RISK_WORDS:
                if re.search(r'\b' + re.escape(word) + r'\b', snippet, re.IGNORECASE):
                    risk_ids.add(code)
                    break
    risk_win = all_window[all_window["task_code"].astype(str).isin(risk_ids)] if risk_ids else pd.DataFrame()

    # Labour in window
    labour_weekly = _week_labour(task_res, tasks, window_start, window_end) if not task_res.empty else pd.DataFrame()
    peak_labour   = int(labour_weekly.groupby("week")["qty"].sum().max()) if not labour_weekly.empty else 0

    # ---- Summary metric cards -----------------------------------------------
    open_logic_count = len(open_start_win) + len(open_finish_win)

    st.markdown(
        '<div style="font-size:12px;font-weight:700;color:#94A3B8;letter-spacing:1px;'
        'text-transform:uppercase;margin-bottom:10px;">Window Summary</div>',
        unsafe_allow_html=True,
    )

    r1 = st.columns(6)
    r1[0].markdown(_lookahead_card(
        "Starting",       len(starting),
        f"in {window_opt.lower()}",
        "#2563eb", "#eff6ff", "#2563eb"
    ), unsafe_allow_html=True)
    r1[1].markdown(_lookahead_card(
        "Finishing",      len(finishing),
        f"in {window_opt.lower()}",
        "#0B1F33", "#f0f9ff", "#0B1F33"
    ), unsafe_allow_html=True)
    r1[2].markdown(_lookahead_card(
        "Critical",       len(crit_win),
        "in window",
        "#dc2626", "#fef2f2" if crit_win.empty else "#fef2f2", "#dc2626"
    ), unsafe_allow_html=True)
    r1[3].markdown(_lookahead_card(
        "Negative Float", len(neg_win),
        "in window",
        "#dc2626" if not neg_win.empty else "#64748B",
        "#fef2f2" if not neg_win.empty else "#f8fafc",
        "#dc2626" if not neg_win.empty else "#E2E8F0"
    ), unsafe_allow_html=True)
    r1[4].markdown(_lookahead_card(
        "Labour Peak",
        f"{peak_labour:,}" if peak_labour else "N/A",
        "hrs/week in window",
        "#7c3aed", "#faf5ff", "#7c3aed"
    ), unsafe_allow_html=True)
    r1[5].markdown(_lookahead_card(
        "Open Logic",     open_logic_count,
        "in window",
        "#d97706" if open_logic_count else "#64748B",
        "#fffbeb" if open_logic_count else "#f8fafc",
        "#d97706" if open_logic_count else "#E2E8F0"
    ), unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    if all_window.empty:
        st.info(
            f"No activities found in the selected window "
            f"({window_start.strftime('%d %b %Y')} to {window_end.strftime('%d %b %Y')}). "
            "Try extending the lookahead window or removing filters."
        )
        return

    # ---- Display columns helper ---------------------------------------------
    BASE_COLS = {
        "task_code":        "Activity ID",
        "task_name":        "Activity Name",
        "wbs_path":         "WBS",
        "eff_start":        "Start",
        "eff_finish":       "Finish",
        "orig_dur_days":    "Orig Dur (d)",
        "rem_dur_days":     "Rem Dur (d)",
        "total_float_days": "Float (d)",
        "status":           "Status",
        "is_critical":      "Critical",
    }

    def _fmt_table(df, extra_cols=None):
        cols = dict(BASE_COLS)
        if extra_cols:
            cols.update(extra_cols)
        avail = {k: v for k, v in cols.items() if k in df.columns}
        out = df[list(avail.keys())].copy().rename(columns=avail)
        for col in ["Start","Finish"]:
            if col in out.columns:
                out[col] = out[col].apply(format_date)
        if "Status" in out.columns:
            out["Status"] = out["Status"].apply(_status_label)
        if "Critical" in out.columns:
            out["Critical"] = out["Critical"].apply(lambda x: "Yes" if x else "")
        if "Float (d)" in out.columns:
            out["Float (d)"] = out["Float (d)"].apply(
                lambda x: round(float(x),1) if x is not None and str(x) not in ("","nan") else "-"
            )
        return out

    def _crit_style(df):
        """Apply row colour based on float."""
        def _row_style(row):
            flag = _crit_flag(row.get("Float (d)", None))
            colour_map = {
                "Negative Float": "background-color:#fecaca;",
                "Critical":       "background-color:#fee2e2;",
                "Near-Critical":  "background-color:#fef3c7;",
            }
            style = colour_map.get(flag, "")
            return [style] * len(row)
        try:
            return df.style.apply(_row_style, axis=1)
        except Exception:
            return df

    # ---- Main tabs ----------------------------------------------------------
    (tab_start, tab_finish, tab_crit, tab_nc, tab_neg,
     tab_open, tab_notes, tab_labour, tab_gantt, tab_export) = st.tabs([
        f"Starting ({len(starting)})",
        f"Finishing ({len(finishing)})",
        f"Critical ({len(crit_win)})",
        f"Near-Critical ({len(nc_win)})",
        f"Negative Float ({len(neg_win)})",
        f"Open Logic ({open_logic_count})",
        f"Planning Notes ({len(notes_win) + len(risk_win)})",
        "Labour Demand",
        "Gantt",
        "Export",
    ])

    # -- Starting -------------------------------------------------------------
    with tab_start:
        st.markdown(
            f"Activities with a scheduled start date between "
            f"**{window_start.strftime('%d %b %Y')}** and "
            f"**{window_end.strftime('%d %b %Y')}**."
        )
        if starting.empty:
            st.info("No activities are scheduled to start in this window.")
        else:
            st.dataframe(_crit_style(_fmt_table(starting.sort_values("eff_start"))),
                         hide_index=True)

    # -- Finishing -------------------------------------------------------------
    with tab_finish:
        st.markdown(
            f"Activities with a scheduled finish date between "
            f"**{window_start.strftime('%d %b %Y')}** and "
            f"**{window_end.strftime('%d %b %Y')}**."
        )
        if finishing.empty:
            st.info("No activities are scheduled to finish in this window.")
        else:
            st.dataframe(_crit_style(_fmt_table(finishing.sort_values("eff_finish"))),
                         hide_index=True)

    # -- Critical --------------------------------------------------------------
    with tab_crit:
        if crit_win.empty:
            st.success("No critical activities fall within this window.")
        else:
            st.warning(
                f"**{len(crit_win)} critical activities** are active or due in this window. "
                "Any delay to these directly impacts the project finish date."
            )
            st.dataframe(_fmt_table(crit_win.sort_values("eff_finish")),
                         hide_index=True)

    # -- Near-Critical ---------------------------------------------------------
    with tab_nc:
        if nc_win.empty:
            st.success("No near-critical activities fall within this window.")
        else:
            st.info(
                f"**{len(nc_win)} near-critical activities** are active or due in this window. "
                f"Each has between 0 and {near_crit_days} days float."
            )
            st.dataframe(_fmt_table(nc_win.sort_values("total_float_days")),
                         hide_index=True)

    # -- Negative Float --------------------------------------------------------
    with tab_neg:
        if neg_win.empty:
            st.success("No activities with negative float are active in this window.")
        else:
            st.error(
                f"**{len(neg_win)} activities have negative float** in this window. "
                "These activities are already beyond their target dates and need immediate attention."
            )
            st.dataframe(_fmt_table(neg_win.sort_values("total_float_days")),
                         hide_index=True)

    # -- Open Logic ------------------------------------------------------------
    with tab_open:
        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown("**Open Start** -- no predecessors")
            if open_start_win.empty:
                st.success("No open-start activities in this window.")
            else:
                st.warning(f"{len(open_start_win)} activities have no predecessor in this window.")
                st.dataframe(_fmt_table(open_start_win), hide_index=True)
        with col_b:
            st.markdown("**Open Finish** -- no successors")
            if open_finish_win.empty:
                st.success("No open-finish activities in this window.")
            else:
                st.warning(f"{len(open_finish_win)} activities have no successor in this window.")
                st.dataframe(_fmt_table(open_finish_win), hide_index=True)

    # -- Planning Notes --------------------------------------------------------
    with tab_notes:
        if not notes_text:
            st.info(
                "No planning notes loaded. Upload notes on the Planning Notes page "
                "and return here to see which activities in this window are mentioned."
            )
        else:
            col_n1, col_n2 = st.columns(2)
            with col_n1:
                st.markdown("**Activities mentioned in planning notes**")
                if notes_win.empty:
                    st.info("No activities in this window are mentioned in the planning notes.")
                else:
                    st.dataframe(_fmt_table(notes_win), hide_index=True)
            with col_n2:
                st.markdown("**Activities with risk keywords in notes**")
                if risk_win.empty:
                    st.info("No risk keywords found against activities in this window.")
                else:
                    st.warning(
                        f"{len(risk_win)} activities in this window appear alongside "
                        "risk keywords (delay, blocked, CE, EWN, constraint etc.) in the planning notes."
                    )
                    st.dataframe(_fmt_table(risk_win), hide_index=True)

    # -- Labour Demand ---------------------------------------------------------
    with tab_labour:
        if labour_weekly.empty:
            st.info(
                "No resource loading data found in this XER. "
                "If the programme was resourced in P6, check the export included resources. "
                "You can also upload a resource CSV on the Labour Histogram page."
            )
        else:
            weekly_total = labour_weekly.groupby("week")["qty"].sum().reset_index()
            weekly_total["week_str"] = weekly_total["week"].dt.strftime("%d %b")

            peak_wk = weekly_total.loc[weekly_total["qty"].idxmax()]
            avg_hrs  = weekly_total["qty"].mean()
            total_hrs = weekly_total["qty"].sum()

            lm1, lm2, lm3 = st.columns(3)
            lm1.metric("Total Hours in Window", f"{total_hrs:,.0f}")
            lm2.metric("Peak Week",             f"{peak_wk['qty']:,.0f} hrs ({peak_wk['week_str']})")
            lm3.metric("Average per Week",      f"{avg_hrs:,.0f} hrs")

            fig_lab = px.bar(
                weekly_total, x="week_str", y="qty",
                title=f"Labour Demand by Week ({window_opt})",
                labels={"week_str": "Week", "qty": "Hours"},
                color_discrete_sequence=["#0B1F33"],
            )
            fig_lab.update_layout(
                height=320, margin=dict(l=10,r=10,t=40,b=10),
                plot_bgcolor="#ffffff", paper_bgcolor="#ffffff",
            )
            st.plotly_chart(fig_lab)

            # By resource
            if "rsrc_name" in labour_weekly.columns:
                by_res = labour_weekly.groupby("rsrc_name")["qty"].sum().reset_index() \
                             .sort_values("qty", ascending=False).head(15)
                fig_res = px.bar(
                    by_res, x="qty", y="rsrc_name", orientation="h",
                    title="Hours by Resource in Window",
                    labels={"qty":"Hours","rsrc_name":""},
                    color_discrete_sequence=["#7c3aed"],
                )
                fig_res.update_yaxes(autorange="reversed")
                fig_res.update_layout(
                    height=350, margin=dict(l=10,r=10,t=40,b=10),
                    plot_bgcolor="#ffffff", paper_bgcolor="#ffffff",
                )
                st.plotly_chart(fig_res)

    # -- Gantt -----------------------------------------------------------------
    with tab_gantt:
        st.markdown(
            f"All activities active in the **{window_opt.lower()}** window, "
            "colour-coded by float status."
        )
        gantt_src = all_window.dropna(subset=["eff_start","eff_finish"]).copy() \
            if "eff_start" in all_window.columns and "eff_finish" in all_window.columns \
            else pd.DataFrame()

        if gantt_src.empty:
            st.info("No date data available for the Gantt chart.")
        else:
            gantt_src["Label"] = (
                gantt_src["task_code"].astype(str) + "  " +
                gantt_src["task_name"].astype(str).str[:40]
            )
            gantt_src["Float Status"] = gantt_src["total_float_days"].apply(
                lambda f: (
                    "Negative Float" if safe_float(f,0) < 0 else
                    "Critical"       if safe_float(f,0) == 0 else
                    "Near-Critical"  if safe_float(f,0) <= near_crit_days else
                    "Has Float"
                )
            )
            gantt_src = gantt_src.sort_values("eff_start")
            max_gantt = 100

            fig_g = px.timeline(
                gantt_src.head(max_gantt),
                x_start="eff_start", x_end="eff_finish", y="Label",
                color="Float Status",
                color_discrete_map={
                    "Negative Float": "#7f1d1d",
                    "Critical":       "#dc2626",
                    "Near-Critical":  "#d97706",
                    "Has Float":      "#2563eb",
                },
                title=f"Lookahead Gantt: {window_opt} ({window_start.strftime('%d %b')} - {window_end.strftime('%d %b %Y')})",
                labels={"Label":""},
            )
            fig_g.update_yaxes(autorange="reversed")

            # Shade the window
            _ws_str = pd.Timestamp(window_start).strftime("%Y-%m-%dT%H:%M:%S")
            _we_str = pd.Timestamp(window_end).strftime("%Y-%m-%dT%H:%M:%S")
            fig_g.add_vrect(
                x0=_ws_str, x1=_we_str,
                fillcolor="#E8951D", opacity=0.06,
                line_width=0, annotation_text="Lookahead Window",
                annotation_position="top left",
                annotation=dict(font_color="#E8951D", font_size=11),
            )
            fig_g.add_vline(
                x=str(now), line_dash="dash", line_color="#0B1F33",
                annotation_text="Today", annotation_position="top right",
                annotation=dict(font_color="#0B1F33", font_size=11),
            )
            fig_g.update_layout(
                height=max(350, min(900, 55 + len(gantt_src.head(max_gantt)) * 26)),
                margin=dict(l=10,r=10,t=50,b=10),
                plot_bgcolor="#ffffff", paper_bgcolor="#ffffff",
                legend_title_text="Float Status",
            )
            st.plotly_chart(fig_g)
            if len(gantt_src) > max_gantt:
                st.caption(f"Showing first {max_gantt} of {len(gantt_src)} activities. Apply filters to narrow the view.")

    # -- Export ----------------------------------------------------------------
    with tab_export:
        st.markdown("**Download the full lookahead as a formatted Excel workbook.**")

        window_label = (
            f"{window_opt}: {window_start.strftime('%d %b %Y')} to {window_end.strftime('%d %b %Y')}"
        )

        summary_sheet = pd.DataFrame({
            "Metric": [
                "Lookahead Window", "From", "To",
                "Activities Starting", "Activities Finishing",
                "Activities in Window (total)",
                "Critical", "Near-Critical", "Negative Float",
                "Open Start", "Open Finish",
                "In Planning Notes", "Risk Keywords",
                "Labour Peak (hrs/week)",
            ],
            "Value": [
                window_label,
                window_start.strftime("%d %b %Y"),
                window_end.strftime("%d %b %Y"),
                len(starting), len(finishing), len(all_window),
                len(crit_win), len(nc_win), len(neg_win),
                len(open_start_win), len(open_finish_win),
                len(notes_win), len(risk_win),
                peak_labour if peak_labour else "N/A",
            ],
        })

        def _exp(df):
            if df is None or df.empty:
                return pd.DataFrame(columns=["No data"])
            return _fmt_table(df)

        export_sheets = {
            "Summary":          summary_sheet,
            "Starting":         _exp(starting.sort_values("eff_start")    if not starting.empty  else starting),
            "Finishing":        _exp(finishing.sort_values("eff_finish")   if not finishing.empty else finishing),
            "All in Window":    _exp(all_window.sort_values("eff_start")   if not all_window.empty else all_window),
            "Critical":         _exp(crit_win),
            "Near-Critical":    _exp(nc_win),
            "Negative Float":   _exp(neg_win),
            "Open Start":       _exp(open_start_win),
            "Open Finish":      _exp(open_finish_win),
        }
        if not notes_win.empty:
            export_sheets["In Planning Notes"] = _exp(notes_win)
        if not risk_win.empty:
            export_sheets["Risk Keywords"]     = _exp(risk_win)
        if not labour_weekly.empty:
            lw_exp = labour_weekly.copy()
            lw_exp["week"] = lw_exp["week"].dt.strftime("%d %b %Y")
            export_sheets["Labour by Week"] = lw_exp.groupby(["week","rsrc_name"])["qty"] \
                .sum().reset_index().rename(columns={"week":"Week","rsrc_name":"Resource","qty":"Hours"})

        xls_bytes = export_df_to_excel(export_sheets)

        st.download_button(
            label="📥  Download Lookahead Report",
            data=xls_bytes,
            file_name=f"lookahead_{window_opt.replace(' ','_')}_{datetime.now().strftime('%Y%m%d')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            help="Exports summary, starting, finishing, critical, open logic and labour sheets.",
            key="dl_009",
        )

        st.markdown(
            '<div style="background:#f8fafc;border:1px solid #E2E8F0;border-radius:8px;'
            'padding:14px 18px;margin-top:12px;">'
            '<div style="font-size:12px;font-weight:700;color:#0B1F33;margin-bottom:6px;">'
            'Workbook sheets</div>'
            '<div style="font-size:12px;color:#64748B;line-height:2;">'
            + "".join(f"<strong>{k}</strong> &nbsp;|&nbsp; " for k in export_sheets.keys()) +
            '</div></div>',
            unsafe_allow_html=True,
        )




# -----------------------------------------------------------------------------
# PAGE: MILESTONE TRACKER
# -----------------------------------------------------------------------------

# Keywords that indicate a milestone activity by name
_MILESTONE_KEYWORDS = [
    "milestone", "complete", "completion", "handover", "hand over",
    "energisation", "energization", "energise", "energize",
    "access", "installation start", "install start",
    "commissioning", "commission", "delivery", "deliver",
    "approval", "approve", "sign off", "sign-off", "signoff",
    "ready for", "available", "issued", "award", "award of contract",
    "start on site", "mobilisation", "mobilization", "end", "finish",
    "practical completion", "pc", "substantial completion",
    "first fix", "second fix", "inspection", "witness",
    "notice to proceed", "ntp", "key date",
]


def _detect_milestones(tasks: pd.DataFrame) -> pd.DataFrame:
    """
    Auto-detect milestone activities from the tasks DataFrame.
    Returns a copy with a boolean 'is_milestone_detected' column.
    """
    df = tasks.copy()

    flags = pd.Series(False, index=df.index)

    # 1. Activity type contains "milestone" (covers TT_Mile and TT_FinMile)
    if "task_type" in df.columns:
        flags |= df["task_type"].astype(str).str.lower().str.contains(
            "milestone", na=False
        )

    # 2. Zero original duration
    if "orig_dur_days" in df.columns:
        flags |= df["orig_dur_days"].apply(lambda d: safe_float(d, 1) == 0)

    # 3. Name contains a milestone keyword
    if "task_name" in df.columns:
        kw_pattern = "|".join(re.escape(k) for k in _MILESTONE_KEYWORDS)
        flags |= df["task_name"].astype(str).str.lower().str.contains(
            kw_pattern, na=False
        )

    df["is_milestone_detected"] = flags
    return df


def _risk_rating(tf, has_constraint: bool = False, in_notes_risk: bool = False) -> str:
    """Return a simple risk rating string for a milestone."""
    f = safe_float(tf, 9999)
    if f < 0 or in_notes_risk:
        return "High"
    if f == 0 or has_constraint:
        return "High"
    if f <= 10:
        return "Medium"
    return "Low"


def _risk_colour(rating: str) -> str:
    return {"High": "#dc2626", "Medium": "#d97706", "Low": "#16a34a"}.get(rating, "#6b7280")


def _risk_bg(rating: str) -> str:
    return {"High": "#fef2f2", "Medium": "#fffbeb", "Low": "#f0fdf4"}.get(rating, "#f8fafc")


def _milestone_header_card(row: pd.Series, movement_days=None) -> str:
    """Render a single milestone summary card as HTML."""
    tf        = safe_float(row.get("total_float_days"), None)
    f_col     = _float_color(tf)
    is_crit   = bool(row.get("is_critical", False))
    status    = _status_label(str(row.get("status", "")))
    s_col     = _status_colour(str(row.get("status", "")))
    code      = str(row.get("task_code", "-"))
    name      = str(row.get("task_name", "-"))
    wbs       = str(row.get("wbs_path", "-"))
    finish    = format_date(row.get("eff_finish"))
    start     = format_date(row.get("eff_start"))
    rating    = _risk_rating(
        tf,
        has_constraint=bool(row.get("cstr_type","")) and str(row.get("cstr_type","")).strip() not in ("","None","nan"),
    )
    r_col     = _risk_colour(rating)

    crit_pill = (
        '<span style="background:#dc2626;color:white;padding:1px 8px;'
        'border-radius:10px;font-size:10px;font-weight:700;margin-left:6px;">CRITICAL</span>'
        if is_crit else ""
    )

    move_html = ""
    if movement_days is not None:
        m_col  = "#dc2626" if movement_days > 0 else "#16a34a" if movement_days < 0 else "#6b7280"
        m_sign = "+" if movement_days > 0 else ""
        move_html = (
            f'<span style="background:{m_col};color:white;padding:2px 8px;'
            f'border-radius:10px;font-size:10px;font-weight:700;margin-left:6px;">'
            f'{m_sign}{movement_days}d movement</span>'
        )

    return f"""
    <div style="background:#ffffff;border:1px solid #E2E8F0;border-radius:12px;
                padding:18px 22px;margin-bottom:10px;
                border-left:5px solid {r_col};
                box-shadow:0 2px 6px rgba(11,31,51,0.07);">
        <div style="display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:8px;">
            <div style="flex:1;min-width:0;">
                <div style="font-size:10px;font-weight:700;color:#94A3B8;letter-spacing:1px;
                            text-transform:uppercase;margin-bottom:4px;">Milestone</div>
                <div style="font-size:16px;font-weight:800;color:#0B1F33;">
                    {code}{crit_pill}{move_html}
                </div>
                <div style="font-size:14px;color:#334155;margin-top:3px;
                            white-space:nowrap;overflow:hidden;text-overflow:ellipsis;"
                     title="{name}">{name}</div>
                <div style="font-size:11px;color:#94A3B8;margin-top:2px;">{wbs}</div>
            </div>
            <div style="display:grid;grid-template-columns:repeat(4,auto);gap:10px;align-items:start;">
                <div style="text-align:center;">
                    <div style="font-size:9px;color:#94A3B8;text-transform:uppercase;
                                letter-spacing:0.8px;">Start</div>
                    <div style="font-size:12px;font-weight:600;color:#0B1F33;margin-top:2px;">{start}</div>
                </div>
                <div style="text-align:center;">
                    <div style="font-size:9px;color:#94A3B8;text-transform:uppercase;
                                letter-spacing:0.8px;">Finish</div>
                    <div style="font-size:12px;font-weight:700;color:#0B1F33;margin-top:2px;">{finish}</div>
                </div>
                <div style="text-align:center;">
                    <div style="font-size:9px;color:#94A3B8;text-transform:uppercase;
                                letter-spacing:0.8px;">Float</div>
                    <div style="font-size:14px;font-weight:800;color:{f_col};margin-top:2px;">
                        {(str(tf) + "d") if tf is not None else "-"}</div>
                </div>
                <div style="text-align:center;">
                    <div style="font-size:9px;color:#94A3B8;text-transform:uppercase;
                                letter-spacing:0.8px;">Risk</div>
                    <div style="background:{r_col};color:white;border-radius:8px;
                                padding:2px 10px;font-size:11px;font-weight:700;margin-top:2px;">
                        {rating}</div>
                </div>
            </div>
        </div>
        <div style="margin-top:10px;">
            <span style="background:{s_col};color:white;padding:2px 9px;
                         border-radius:10px;font-size:10px;">{status}</span>
        </div>
    </div>"""


def page_milestone_tracker(data: dict, near_crit_days: float):
    """
    Milestone Tracker page.
    Auto-detects milestones from the XER, allows manual additions,
    and shows driving path, successors, movement, notes and risk per milestone.
    """
    st.title("🏁 Milestone Tracker")
    st.caption(
        "Key programme milestones auto-detected from your XER. "
        "Select any milestone to see what is driving it, what it drives, and its current risk."
    )

    tasks = data["tasks_df"]
    rels  = data["relationships_df"]

    if tasks.empty:
        st.warning("No activities found. Please upload a programme first.")
        return

    tasks = get_critical_threshold(tasks, near_crit_days)
    tasks_with_milestones = _detect_milestones(tasks)

    # Build graph
    G           = build_graph(tasks, rels)
    task_lookup = tasks.set_index("task_id").to_dict("index")

    # ---- Comparison data (finish movement) ----------------------------------
    movement_map = {}
    if "_mi_prev" in st.session_state and "_mi_curr" in st.session_state:
        try:
            prev_t = st.session_state["_mi_prev"]["tasks_df"]
            curr_t = st.session_state["_mi_curr"]["tasks_df"]
            if not prev_t.empty and not curr_t.empty:
                merged = prev_t[["task_code","eff_finish"]].merge(
                    curr_t[["task_code","eff_finish"]], on="task_code",
                    suffixes=("_p","_c"), how="inner"
                )
                for _, r in merged.iterrows():
                    try:
                        p = pd.Timestamp(r["eff_finish_p"])
                        c = pd.Timestamp(r["eff_finish_c"])
                        movement_map[str(r["task_code"])] = int((c - p).days)
                    except Exception:
                        pass
        except Exception:
            pass

    # ---- Notes data ---------------------------------------------------------
    notes_text = st.session_state.get("_notes_text", "")
    notes_ids  = set()
    risk_ids   = set()
    if notes_text and "task_code" in tasks.columns:
        for _, t in tasks.iterrows():
            code = str(t.get("task_code",""))
            if not code or code not in notes_text:
                continue
            notes_ids.add(code)
            idx = notes_text.find(code)
            snippet = notes_text[max(0,idx-300):idx+300]
            for word in _RISK_WORDS:
                if re.search(r'\b' + re.escape(word) + r'\b', snippet, re.IGNORECASE):
                    risk_ids.add(code)
                    break

    # ---- Milestone selection -------------------------------------------------
    auto_milestones = tasks_with_milestones[
        tasks_with_milestones["is_milestone_detected"]
    ].copy()

    # Session state for manually added milestones
    if "ms_manual_ids" not in st.session_state:
        st.session_state["ms_manual_ids"] = set()

    # ---- Sidebar-style controls panel ---------------------------------------
    with st.expander("Configure Milestones", expanded=True):
        cfg1, cfg2 = st.columns([1, 2])

        with cfg1:
            st.markdown(
                '<div style="font-size:11px;font-weight:700;color:#94A3B8;'
                'letter-spacing:1px;text-transform:uppercase;margin-bottom:6px;">'
                'Auto-Detected</div>',
                unsafe_allow_html=True,
            )
            st.caption(
                f"{len(auto_milestones)} milestones detected from activity type, "
                "zero duration, or name keywords."
            )

            show_auto = st.checkbox("Include auto-detected milestones", value=True, key="ms_show_auto")

        with cfg2:
            st.markdown(
                '<div style="font-size:11px;font-weight:700;color:#94A3B8;'
                'letter-spacing:1px;text-transform:uppercase;margin-bottom:6px;">'
                'Add Key Activities Manually</div>',
                unsafe_allow_html=True,
            )
            manual_search = st.text_input(
                "Search to add",
                placeholder="Type Activity ID or name",
                key="ms_manual_search",
                label_visibility="collapsed",
            )
            if manual_search.strip():
                mask = (
                    tasks["task_code"].astype(str).str.contains(manual_search.strip(), case=False, na=False) |
                    tasks["task_name"].astype(str).str.contains(manual_search.strip(), case=False, na=False)
                )
                search_results = tasks[mask]
                if not search_results.empty:
                    add_labels = search_results.apply(
                        lambda r: f"{r.get('task_code','')}  --  {r.get('task_name','')}", axis=1
                    ).tolist()
                    add_sel = st.selectbox("Select to add", add_labels, key="ms_add_sel", label_visibility="collapsed")
                    if st.button("Add as milestone", key="ms_add_btn"):
                        sel_code = add_sel.split("  --  ")[0].strip()
                        match = tasks[tasks["task_code"] == sel_code]
                        if not match.empty:
                            st.session_state["ms_manual_ids"].add(match.iloc[0]["task_id"])
                            st.success(f"Added {sel_code}")
                else:
                    st.caption("No activities match.")

            if st.session_state["ms_manual_ids"]:
                if st.button("Clear manual selections", key="ms_clear"):
                    st.session_state["ms_manual_ids"] = set()

        # Filters
        st.markdown("<hr style='border:none;border-top:1px solid #E2E8F0;margin:10px 0;'>", unsafe_allow_html=True)
        fa, fb, fc = st.columns(3)
        f_risk = fa.selectbox("Risk filter", ["All","High","Medium","Low"], key="ms_f_risk")
        f_crit = fb.selectbox("Float filter", ["All","Critical","Near-Critical","Negative Float"], key="ms_f_crit")
        f_wbs  = fc.text_input("WBS contains", placeholder="e.g. Civil", key="ms_f_wbs")

    # ---- Build final milestone list -----------------------------------------
    milestone_ids = set()
    if show_auto:
        milestone_ids |= set(auto_milestones["task_id"].tolist())
    milestone_ids |= st.session_state["ms_manual_ids"]

    if not milestone_ids:
        st.info(
            "No milestones detected or selected. "
            "Either enable auto-detection above or search for and add activities manually."
        )
        return

    milestones = tasks[tasks["task_id"].isin(milestone_ids)].copy()

    # Apply risk rating
    def _rate(row):
        tf   = safe_float(row.get("total_float_days"), 9999)
        cstr = str(row.get("cstr_type","")).strip() not in ("","None","nan")
        risk = str(row.get("task_code","")) in risk_ids
        return _risk_rating(tf, cstr, risk)

    milestones["risk_rating"] = milestones.apply(_rate, axis=1)

    # Apply filters
    if f_risk != "All":
        milestones = milestones[milestones["risk_rating"] == f_risk]
    if f_crit == "Critical":
        milestones = milestones[milestones["is_critical"] == True]
    elif f_crit == "Near-Critical":
        milestones = milestones[milestones["is_near_critical"] == True]
    elif f_crit == "Negative Float":
        milestones = milestones[milestones["total_float_days"].apply(lambda f: safe_float(f,0) < 0)]
    if f_wbs.strip() and "wbs_path" in milestones.columns:
        milestones = milestones[milestones["wbs_path"].astype(str).str.contains(f_wbs.strip(), case=False, na=False)]

    if milestones.empty:
        st.info("No milestones match the current filters.")
        return

    milestones = milestones.sort_values("eff_finish" if "eff_finish" in milestones.columns else "task_code")

    # ---- Summary metrics ----------------------------------------------------
    n_total  = len(milestones)
    n_high   = int((milestones["risk_rating"] == "High").sum())
    n_med    = int((milestones["risk_rating"] == "Medium").sum())
    n_low    = int((milestones["risk_rating"] == "Low").sum())
    n_crit   = int(milestones["is_critical"].sum()) if "is_critical" in milestones.columns else 0
    n_neg    = int(milestones["total_float_days"].apply(lambda f: safe_float(f,0) < 0).sum()) if "total_float_days" in milestones.columns else 0

    st.markdown(
        f"""
        <div style="background:#0B1F33;border-radius:12px;padding:16px 22px;
                    margin-bottom:18px;display:flex;gap:20px;flex-wrap:wrap;align-items:center;">
            <div>
                <div style="font-size:10px;color:#64748B;text-transform:uppercase;
                            letter-spacing:1px;">Milestones</div>
                <div style="font-size:28px;font-weight:800;color:#E8951D;line-height:1;margin-top:3px;">{n_total}</div>
            </div>
            <div style="width:1px;background:#1e3a5f;align-self:stretch;"></div>
            <div>
                <div style="font-size:10px;color:#64748B;text-transform:uppercase;letter-spacing:1px;">High Risk</div>
                <div style="font-size:24px;font-weight:800;color:#dc2626;line-height:1;margin-top:3px;">{n_high}</div>
            </div>
            <div>
                <div style="font-size:10px;color:#64748B;text-transform:uppercase;letter-spacing:1px;">Medium</div>
                <div style="font-size:24px;font-weight:800;color:#d97706;line-height:1;margin-top:3px;">{n_med}</div>
            </div>
            <div>
                <div style="font-size:10px;color:#64748B;text-transform:uppercase;letter-spacing:1px;">Low</div>
                <div style="font-size:24px;font-weight:800;color:#16a34a;line-height:1;margin-top:3px;">{n_low}</div>
            </div>
            <div style="width:1px;background:#1e3a5f;align-self:stretch;"></div>
            <div>
                <div style="font-size:10px;color:#64748B;text-transform:uppercase;letter-spacing:1px;">Critical</div>
                <div style="font-size:24px;font-weight:800;color:#dc2626;line-height:1;margin-top:3px;">{n_crit}</div>
            </div>
            <div>
                <div style="font-size:10px;color:#64748B;text-transform:uppercase;letter-spacing:1px;">Neg Float</div>
                <div style="font-size:24px;font-weight:800;color:#7f1d1d;line-height:1;margin-top:3px;">{n_neg}</div>
            </div>
            {"" if not movement_map else '<div style="margin-left:auto;font-size:11px;color:#E8951D;">Comparison data loaded</div>'}
        </div>
        """,
        unsafe_allow_html=True,
    )

    # ---- Summary table tab + individual milestone tabs ----------------------
    tab_names = ["Summary Table"] + [
        f"{row.get('task_code','?')}" for _, row in milestones.head(10).iterrows()
    ]
    if len(milestones) > 10:
        tab_names[-1] = f"{tab_names[-1]}..."
    tabs = st.tabs(tab_names)

    # =========================================================================
    # TAB 0: Summary Table
    # =========================================================================
    with tabs[0]:
        st.markdown(
            '<div style="font-size:12px;font-weight:700;color:#94A3B8;letter-spacing:1px;'
            'text-transform:uppercase;margin-bottom:8px;">All Milestones</div>',
            unsafe_allow_html=True,
        )

        table_rows = []
        for _, ms in milestones.iterrows():
            code = str(ms.get("task_code",""))
            tf   = safe_float(ms.get("total_float_days"), None)
            move = movement_map.get(code)
            table_rows.append({
                "Activity ID":      code,
                "Activity Name":    str(ms.get("task_name","")),
                "WBS":              str(ms.get("wbs_path","")).split(" > ")[0],
                "Forecast Start":   format_date(ms.get("eff_start")),
                "Forecast Finish":  format_date(ms.get("eff_finish")),
                "Total Float (d)":  round(tf, 1) if tf is not None else "-",
                "Critical":         "Yes" if ms.get("is_critical") else "",
                "Movement (d)":     move if move is not None else "-",
                "Status":           _status_label(str(ms.get("status",""))),
                "Risk":             ms.get("risk_rating",""),
                "In Notes":         "Yes" if code in notes_ids else "",
                "Risk Keywords":    "Yes" if code in risk_ids else "",
            })

        summary_df = pd.DataFrame(table_rows)

        def _style_summary(row):
            risk = row.get("Risk","")
            colour_map = {
                "High":   "background-color:#fef2f2;",
                "Medium": "background-color:#fffbeb;",
                "Low":    "",
            }
            return [colour_map.get(risk,"")] * len(row)

        st.dataframe(
            summary_df.style.apply(_style_summary, axis=1),
            hide_index=True,
        )

        # Timeline chart of all milestones
        gantt_src = milestones.dropna(subset=["eff_finish"]).copy() \
            if "eff_finish" in milestones.columns else pd.DataFrame()
        if not gantt_src.empty:
            gantt_src["Label"] = gantt_src["task_code"].astype(str) + "  " + gantt_src["task_name"].astype(str).str[:40]
            gantt_src["Risk"]  = gantt_src["risk_rating"]
            start_col  = "eff_start"  if "eff_start"  in gantt_src.columns else "eff_finish"
            finish_col = "eff_finish"

            # For zero-duration milestones, give a 1-day bar for visibility
            gantt_src["_plot_start"]  = gantt_src[start_col]
            gantt_src["_plot_finish"] = gantt_src.apply(
                lambda r: r[finish_col] + timedelta(days=1)
                if r.get(start_col) == r.get(finish_col) or pd.isna(r.get(start_col))
                else r[finish_col],
                axis=1,
            )

            fig_ms = px.timeline(
                gantt_src,
                x_start="_plot_start", x_end="_plot_finish", y="Label",
                color="Risk",
                color_discrete_map={
                    "High":   "#dc2626",
                    "Medium": "#d97706",
                    "Low":    "#16a34a",
                },
                title="Milestone Timeline",
                labels={"Label":""},
            )
            fig_ms.update_yaxes(autorange="reversed")
            fig_ms.add_vline(
                x=datetime.now().strftime("%Y-%m-%dT%H:%M:%S"), line_dash="dash", line_color="#0B1F33",
                annotation_text="Today", annotation_position="top right",
                annotation=dict(font_color="#0B1F33"),
            )
            fig_ms.update_layout(
                height=max(300, min(800, 50 + len(gantt_src) * 28)),
                margin=dict(l=10,r=10,t=50,b=10),
                plot_bgcolor="#ffffff", paper_bgcolor="#ffffff",
                legend_title_text="Risk",
            )
            st.plotly_chart(fig_ms)

    # =========================================================================
    # TABS 1-N: Individual milestone detail
    # =========================================================================
    for tab_idx, (_, ms) in enumerate(milestones.head(10).iterrows(), start=1):
        if tab_idx >= len(tabs):
            break
        with tabs[tab_idx]:
            ms_code = str(ms.get("task_code",""))
            ms_id   = ms["task_id"]
            move    = movement_map.get(ms_code)

            # Header card
            st.markdown(
                _milestone_header_card(ms, movement_days=move),
                unsafe_allow_html=True,
            )

            # Constraint warning
            if "cstr_type" in ms.index:
                cstr = str(ms.get("cstr_type","")).strip()
                if cstr not in ("","None","nan"):
                    cdate = format_date(ms.get("cstr_date") if "cstr_date" in ms.index else None)
                    st.markdown(
                        f'<div style="background:#fffbeb;border-left:4px solid #E8951D;'
                        f'border-radius:6px;padding:10px 16px;margin-bottom:10px;">'
                        f'<strong>Constraint:</strong> {cstr} &nbsp;|&nbsp; '
                        f'<strong>Date:</strong> {cdate}</div>',
                        unsafe_allow_html=True,
                    )

            # Notes against this milestone
            if ms_code in notes_ids:
                note_flag = "Risk keywords found in notes." if ms_code in risk_ids else "Mentioned in planning notes."
                note_col  = "#dc2626" if ms_code in risk_ids else "#2563eb"
                st.markdown(
                    f'<div style="background:#eff6ff;border-left:4px solid {note_col};'
                    f'border-radius:6px;padding:10px 16px;margin-bottom:10px;">'
                    f'<strong>Planning Notes:</strong> {note_flag}</div>',
                    unsafe_allow_html=True,
                )

            # ---- Action buttons ---------------------------------------------
            ba, bb = st.columns(2)
            btn_drive = ba.button("Find Driving Path",  key=f"ms_drive_{ms_id}", use_container_width=True)
            btn_succ  = bb.button("Show Successors",    key=f"ms_succ_{ms_id}",  use_container_width=True)

            # ---- Driving Path ------------------------------------------------
            if btn_drive:
                with st.spinner("Tracing driving path..."):
                    direct_preds = list(G.predecessors(ms_id))
                    if not direct_preds:
                        st.warning(f"{ms_code} has no predecessors. No driving path can be identified.")
                    else:
                        path = driving_path_to_activity(G, tasks, rels, ms_id)
                        path_rows = []
                        for i, tid in enumerate(path):
                            t  = task_lookup.get(tid, {})
                            tf = t.get("total_float_days")
                            is_tgt = (tid == ms_id)
                            # Relationship to next step
                            rl, lg = "-", 0
                            if i < len(path) - 1:
                                next_id = path[i+1]
                                if not rels.empty:
                                    rel = rels[
                                        (rels.get("pred_task_id", pd.Series(dtype=str)) == tid) &
                                        (rels.get("succ_task_id", pd.Series(dtype=str)) == next_id)
                                    ]
                                    if not rel.empty:
                                        rl = _rel_label(rel["rel_type"].iloc[0] if "rel_type" in rel.columns else "FS")
                                        lg = safe_float(rel["lag_days"].iloc[0] if "lag_days" in rel.columns else 0, 0)
                            path_rows.append({
                                "Step":            i + 1,
                                "Activity ID":     t.get("task_code", tid),
                                "Activity Name":   t.get("task_name",""),
                                "Start":           format_date(t.get("eff_start")),
                                "Finish":          format_date(t.get("eff_finish")),
                                "Float (d)":       round(float(tf),1) if tf is not None else "-",
                                "Link":            rl if not is_tgt else "-",
                                "Lag (d)":         lg if not is_tgt else "-",
                                "Flag":            _crit_flag(tf),
                                "Milestone":       "TARGET" if is_tgt else "",
                            })

                        path_df = pd.DataFrame(path_rows)
                        st.markdown("**Driving Path**")
                        st.caption("Activities ordered from chain start to the milestone target.")

                        def _path_row_style(row):
                            flag = row.get("Flag","")
                            is_t = row.get("Milestone","") == "TARGET"
                            if is_t:
                                return ["background-color:#0B1F33;color:white;font-weight:700;"] * len(row)
                            cm = {"Critical":"background-color:#fee2e2;",
                                  "Negative Float":"background-color:#fecaca;",
                                  "Near-Critical":"background-color:#fef3c7;"}
                            return [cm.get(flag,"")] * len(row)

                        st.dataframe(path_df.style.apply(_path_row_style, axis=1),
                                     hide_index=True)

                        # Mini Gantt for path
                        path_task_ids = [p for p in path if p != ms_id]
                        path_tasks    = tasks[tasks["task_id"].isin(path)].copy()
                        gantt_p = path_tasks.dropna(subset=["eff_start","eff_finish"]).copy() \
                            if "eff_start" in path_tasks.columns else pd.DataFrame()
                        if not gantt_p.empty:
                            gantt_p["Label"] = gantt_p["task_code"].astype(str) + "  " + gantt_p["task_name"].astype(str).str[:35]
                            gantt_p["Type"]  = gantt_p["task_id"].apply(lambda t: "Milestone" if t == ms_id else "Driving Path")
                            fig_dp = px.timeline(
                                gantt_p, x_start="eff_start", x_end="eff_finish", y="Label",
                                color="Type",
                                color_discrete_map={"Milestone":"#0B1F33","Driving Path":"#dc2626"},
                                title=f"Driving Path to {ms_code}",
                            )
                            fig_dp.update_yaxes(autorange="reversed")
                            fig_dp.add_vline(x=datetime.now().strftime("%Y-%m-%dT%H:%M:%S").strftime("%Y-%m-%dT%H:%M:%S"), line_dash="dot", line_color="#94A3B8",
                                             annotation_text="Today")
                            fig_dp.update_layout(height=max(250, 50+len(gantt_p)*28),
                                                 margin=dict(l=10,r=10,t=40,b=10),
                                                 plot_bgcolor="#ffffff", paper_bgcolor="#ffffff")
                            st.plotly_chart(fig_dp)

            # ---- Successors --------------------------------------------------
            if btn_succ:
                direct_succs = list(G.successors(ms_id))
                if not direct_succs:
                    st.warning(f"{ms_code} has no successors.")
                else:
                    all_succs = trace_successors(G, ms_id)
                    succ_df   = _build_full_trace_df(G, rels, task_lookup, ms_id, all_succs, "succ")
                    # Add WBS
                    if not succ_df.empty and "Activity ID" in succ_df.columns:
                        code_to_wbs = tasks.set_index("task_code")["wbs_path"].to_dict() if "wbs_path" in tasks.columns else {}
                        succ_df.insert(succ_df.columns.get_loc("Activity Name")+1, "WBS",
                                       succ_df["Activity ID"].map(code_to_wbs).fillna("-"))

                    st.markdown(f"**All Successors of {ms_code}** ({len(succ_df)} activities)")

                    def _succ_style(val):
                        return {"Critical":"background-color:#fee2e2;color:#991b1b;font-weight:600;",
                                "Negative Float":"background-color:#fecaca;color:#7f1d1d;font-weight:700;",
                                "Near-Critical":"background-color:#fef3c7;color:#92400e;font-weight:600;",
                                "Float":"background-color:#dcfce7;color:#166534;"}.get(val,"")

                    st.dataframe(
                        succ_df.style.map(_succ_style, subset=["Critical Flag"]),
                        hide_index=True, height=min(400, 45+len(succ_df)*35),
                    )

            # ---- Individual export -------------------------------------------
            st.divider()
            exp_col, _ = st.columns([1,3])

            exp_detail = pd.DataFrame([{
                "Activity ID":    ms_code,
                "Activity Name":  str(ms.get("task_name","")),
                "WBS":            str(ms.get("wbs_path","")),
                "Forecast Start": format_date(ms.get("eff_start")),
                "Forecast Finish":format_date(ms.get("eff_finish")),
                "Total Float (d)":safe_float(ms.get("total_float_days"), None),
                "Critical":       "Yes" if ms.get("is_critical") else "No",
                "Movement (d)":   move if move is not None else "N/A",
                "Status":         _status_label(str(ms.get("status",""))),
                "Risk Rating":    ms.get("risk_rating",""),
                "In Notes":       "Yes" if ms_code in notes_ids else "No",
                "Risk Keywords":  "Yes" if ms_code in risk_ids else "No",
                "Constraint":     str(ms.get("cstr_type","")).strip() or "None",
            }])

            xls_ms = export_df_to_excel({"Milestone Detail": exp_detail})
            exp_col.download_button(
                label=f"📥 Export {ms_code} Report",
                data=xls_ms,
                file_name=f"milestone_{ms_code}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key=f"ms_exp_{ms_id}",
            )

    # =========================================================================
    # FULL EXPORT
    # =========================================================================
    st.divider()

    full_exp_rows = []
    for _, ms in milestones.iterrows():
        code = str(ms.get("task_code",""))
        tf   = safe_float(ms.get("total_float_days"), None)
        full_exp_rows.append({
            "Activity ID":      code,
            "Activity Name":    str(ms.get("task_name","")),
            "WBS":              str(ms.get("wbs_path","")),
            "Forecast Start":   format_date(ms.get("eff_start")),
            "Forecast Finish":  format_date(ms.get("eff_finish")),
            "Total Float (d)":  round(tf,1) if tf is not None else "-",
            "Critical":         "Yes" if ms.get("is_critical") else "",
            "Risk Rating":      ms.get("risk_rating",""),
            "Movement (d)":     movement_map.get(code, "-"),
            "Status":           _status_label(str(ms.get("status",""))),
            "In Notes":         "Yes" if code in notes_ids else "",
            "Risk Keywords":    "Yes" if code in risk_ids else "",
            "Constraint":       str(ms.get("cstr_type","")).strip() or "",
        })

    full_df = pd.DataFrame(full_exp_rows)
    high_df = full_df[full_df["Risk Rating"] == "High"]
    med_df  = full_df[full_df["Risk Rating"] == "Medium"]

    xls_full = export_df_to_excel({
        "All Milestones": full_df,
        "High Risk":      high_df if not high_df.empty else pd.DataFrame(columns=["No data"]),
        "Medium Risk":    med_df  if not med_df.empty  else pd.DataFrame(columns=["No data"]),
    })

    dl_col, _ = st.columns([1,3])
    dl_col.download_button(
        label="📥  Export All Milestones to Excel",
        data=xls_full,
        file_name=f"milestones_{datetime.now().strftime('%Y%m%d')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        help="Exports All Milestones, High Risk and Medium Risk sheets.",
        key="dl_010",
    )




# -----------------------------------------------------------------------------
# PAGE: RISK & OPPORTUNITY REGISTER
# -----------------------------------------------------------------------------

# Probability and impact scoring for auto-generated items
_PROB_MAP   = {"High": "High", "Medium": "Medium", "Low": "Low"}
_IMPACT_MAP = {"High": "High", "Medium": "Medium", "Low": "Low"}

# RAG colour helpers reused from rest of app
def _rag_colour(priority: str) -> tuple:
    """Return (text_colour, bg_colour, border_colour) for a priority level."""
    return {
        "High":   ("#991b1b", "#fef2f2", "#fca5a5"),
        "Medium": ("#92400e", "#fffbeb", "#fcd34d"),
        "Low":    ("#166534", "#f0fdf4", "#86efac"),
    }.get(priority, ("#374151", "#f9fafb", "#E2E8F0"))


def _ro_card_pill(item_type: str) -> str:
    """Coloured pill for Risk vs Opportunity."""
    if item_type == "Risk":
        return (
            '<span style="background:#dc2626;color:white;padding:2px 10px;'
            'border-radius:10px;font-size:10px;font-weight:700;letter-spacing:0.5px;">'
            'RISK</span>'
        )
    return (
        '<span style="background:#16a34a;color:white;padding:2px 10px;'
        'border-radius:10px;font-size:10px;font-weight:700;letter-spacing:0.5px;">'
        'OPPORTUNITY</span>'
    )


def _ro_row(item_type, priority, act_code, act_name, wbs,
            description, cause, effect, mitigation):
    """Build a single register row dict."""
    return {
        "Type":            item_type,
        "Priority":        priority,
        "Activity ID":     str(act_code),
        "Activity Name":   str(act_name),
        "WBS":             str(wbs),
        "Description":     str(description),
        "Cause":           str(cause),
        "Effect":          str(effect),
        "Mitigation / Action": str(mitigation),
        "Owner":           "",
        "Due Date":        "",
        "Status":          "Open",
    }


def _generate_register(
    tasks: pd.DataFrame,
    rels:  pd.DataFrame,
    near_crit_days: float,
    notes_text: str = "",
) -> pd.DataFrame:
    """
    Generate a draft Risk & Opportunity register from programme data.
    Returns a DataFrame with one row per item.
    """
    rows = []
    now  = datetime.now()
    eight_weeks = now + timedelta(weeks=8)
    tasks = get_critical_threshold(tasks, near_crit_days)

    def _wbs(row):
        w = row.get("wbs_path","") if "wbs_path" in row.index else ""
        return str(w).split(" > ")[0] if w and str(w).strip() not in ("","nan") else "-"

    # -- Predecessor / successor lookup ----------------------------------------
    tasks_with_pred = set()
    tasks_with_succ = set()
    if not rels.empty:
        if "succ_task_id" in rels.columns:
            tasks_with_pred = set(rels["succ_task_id"].dropna())
        if "pred_task_id" in rels.columns:
            tasks_with_succ = set(rels["pred_task_id"].dropna())

    # =========================================================================
    # RISKS
    # =========================================================================

    # R1: Negative float
    if "total_float_days" in tasks.columns:
        neg = tasks[tasks["total_float_days"].apply(lambda f: safe_float(f,0) < 0)]
        for _, t in neg.iterrows():
            tf = round(safe_float(t.get("total_float_days"),0), 1)
            rows.append(_ro_row(
                "Risk", "High",
                t.get("task_code",""), t.get("task_name",""), _wbs(t),
                f"Activity has {abs(tf)} days negative float.",
                f"The current schedule logic and constraints result in a {abs(tf)}-day overrun on this activity.",
                "The project completion date cannot be achieved on this path without intervention. "
                "Every day of further delay worsens the position.",
                f"Raise with the planner immediately. Develop a recovery plan to recover the {abs(tf)} days. "
                "Consider acceleration, scope reduction, or parallel working.",
            ))

    # R2: Critical activities not started
    if "status" in tasks.columns and "is_critical" in tasks.columns:
        crit_ns = tasks[
            tasks["is_critical"] &
            tasks["status"].apply(lambda s: str(s) in ("TK_NotStart","Not Started"))
        ]
        for _, t in crit_ns.iterrows():
            finish = format_date(t.get("eff_finish"))
            rows.append(_ro_row(
                "Risk", "High",
                t.get("task_code",""), t.get("task_name",""), _wbs(t),
                f"Critical activity not yet started. Target finish: {finish}.",
                "The activity is on the critical path but mobilisation or enabling works have not commenced.",
                "Any further delay to this activity will directly delay the project finish date.",
                "Confirm the start date and mobilisation plan. If at risk, escalate to the delivery team and PM.",
            ))

    # R3: Near-critical due within 8 weeks
    if "eff_finish" in tasks.columns and "is_near_critical" in tasks.columns:
        nc_soon = tasks[
            tasks["is_near_critical"] &
            tasks["eff_finish"].apply(
                lambda d: d is not None and hasattr(d,"date") and d <= eight_weeks
            )
        ]
        for _, t in nc_soon.iterrows():
            tf = round(safe_float(t.get("total_float_days"),0), 1)
            rows.append(_ro_row(
                "Risk", "Medium",
                t.get("task_code",""), t.get("task_name",""), _wbs(t),
                f"Near-critical activity due within 8 weeks with only {tf} days float.",
                "Limited schedule buffer combined with imminent completion requirements.",
                "If this activity is delayed, it will move onto the critical path and threaten the project finish.",
                f"Monitor weekly. If float drops below 5 days, treat as critical and escalate.",
            ))

    # R4: No predecessor (open start)
    if "task_id" in tasks.columns and "task_type" in tasks.columns:
        no_pred = tasks[
            ~tasks["task_id"].isin(tasks_with_pred) &
            ~tasks["task_type"].astype(str).str.contains("Milestone|LOE|WBS", na=False)
        ]
        for _, t in no_pred.head(15).iterrows():
            rows.append(_ro_row(
                "Risk", "Medium",
                t.get("task_code",""), t.get("task_name",""), _wbs(t),
                "Activity has no predecessor. Programme logic is open at the start.",
                "Missing or incomplete logic in the schedule. The activity may start earlier or later than intended.",
                "Float calculations for this activity may be unreliable. "
                "The activity could be delayed without any schedule warning.",
                "Review with the planner. Add a logical predecessor or add a constraint if the start date is fixed.",
            ))

    # R5: No successor (open finish)
    if "task_id" in tasks.columns and "task_type" in tasks.columns:
        no_succ = tasks[
            ~tasks["task_id"].isin(tasks_with_succ) &
            ~tasks["task_type"].astype(str).str.contains("Finish Milestone|LOE|WBS", na=False)
        ]
        for _, t in no_succ.head(15).iterrows():
            rows.append(_ro_row(
                "Risk", "Medium",
                t.get("task_code",""), t.get("task_name",""), _wbs(t),
                "Activity has no successor. Programme logic is open at the finish.",
                "Missing logic means this activity does not drive any subsequent work in the schedule.",
                "The activity may show artificially high float. "
                "Delays may not cascade correctly through the programme.",
                "Review with the planner. Add a logical successor or confirm the activity is a deliberate end point.",
            ))

    # R6: Excessive lag (> 10 days)
    if not rels.empty and "lag_days" in rels.columns:
        big_lag = rels[rels["lag_days"].apply(lambda l: safe_float(l,0) > 10)]
        for _, r in big_lag.head(10).iterrows():
            code = r.get("succ_task_code", r.get("succ_task_id",""))
            name = r.get("succ_task_name","")
            pred = r.get("pred_task_code", r.get("pred_task_id",""))
            lag  = int(safe_float(r.get("lag_days",0),0))
            match = tasks[tasks["task_code"] == str(code)]
            wbs   = _wbs(match.iloc[0]) if not match.empty else "-"
            rows.append(_ro_row(
                "Risk", "Low",
                code, name, wbs,
                f"Relationship from {pred} has {lag} days lag.",
                f"A {lag}-day lag has been applied instead of modelling the actual work sequence.",
                "Excessive lag disguises schedule risk and inflates float on successor activities. "
                "It may also mask negative float.",
                f"Challenge the {lag}-day lag with the planner. Replace with a properly sequenced "
                "activity or reduce the lag to the minimum justified by contract or site conditions.",
            ))

    # R7: Long duration activities (> 60 days)
    if "orig_dur_days" in tasks.columns:
        long_dur = tasks[tasks["orig_dur_days"].apply(lambda d: safe_float(d,0) > 60)]
        for _, t in long_dur.head(10).iterrows():
            dur = int(safe_float(t.get("orig_dur_days",0),0))
            rows.append(_ro_row(
                "Risk", "Low",
                t.get("task_code",""), t.get("task_name",""), _wbs(t),
                f"Activity duration is {dur} working days.",
                "Activity is too long to manage or monitor effectively as a single work package.",
                "Problems can go undetected for weeks. Float burn and delays may not surface until it is too late.",
                f"Break the {dur}-day activity into smaller work packages of 20-30 days. "
                "Discuss with the planner to improve schedule resolution.",
            ))

    # R8: Planning notes risk words
    if notes_text and "task_code" in tasks.columns:
        for word in _RISK_WORDS:
            pattern = re.compile(r'\b' + re.escape(word) + r'\b', re.IGNORECASE)
            if not pattern.search(notes_text):
                continue
            for _, t in tasks.iterrows():
                code = str(t.get("task_code",""))
                if not code or code not in notes_text:
                    continue
                idx = notes_text.find(code)
                snippet = notes_text[max(0,idx-300):idx+300]
                if pattern.search(snippet):
                    rows.append(_ro_row(
                        "Risk", "High",
                        code, t.get("task_name",""), _wbs(t),
                        f"Planning notes reference '{word}' against this activity.",
                        f"The planning note indicates a potential {word} issue that is not fully visible in the programme.",
                        "This risk may not be reflected in the current schedule float or critical path.",
                        f"Review the planning note for {code}. Confirm whether the '{word}' item has been "
                        "resolved, raise a formal risk if not, and update the programme if dates are affected.",
                    ))
                    break

    # R9: Activities delayed in comparison (> 10 days slip)
    if "_mi_prev" in st.session_state and "_mi_curr" in st.session_state:
        try:
            prev_t = st.session_state["_mi_prev"]["tasks_df"]
            curr_t = st.session_state["_mi_curr"]["tasks_df"]
            if not prev_t.empty and not curr_t.empty:
                merged_comp = prev_t[["task_code","eff_finish","task_name"]].merge(
                    curr_t[["task_code","eff_finish"]], on="task_code", suffixes=("_p","_c"), how="inner"
                )
                for _, r in merged_comp.iterrows():
                    try:
                        slip = int((pd.Timestamp(r["eff_finish_c"]) - pd.Timestamp(r["eff_finish_p"])).days)
                        if slip > 10:
                            match = tasks[tasks["task_code"] == str(r["task_code"])]
                            wbs   = _wbs(match.iloc[0]) if not match.empty else "-"
                            rows.append(_ro_row(
                                "Risk", "High" if slip > 30 else "Medium",
                                str(r["task_code"]), str(r.get("task_name","")), wbs,
                                f"Activity finish date has slipped {slip} days since the previous programme revision.",
                                f"The activity's forecast finish moved {slip} days later between the two programme versions.",
                                f"A {slip}-day slip on this activity may delay downstream work and impact the project finish date.",
                                f"Investigate the reason for the {slip}-day slip. Agree a recovery programme with the delivery team. "
                                "Update the programme and issue a revised schedule.",
                            ))
                    except Exception:
                        pass
        except Exception:
            pass

    # =========================================================================
    # OPPORTUNITIES
    # =========================================================================

    # O1: Activities pulled earlier in comparison (> 5 days improvement)
    if "_mi_prev" in st.session_state and "_mi_curr" in st.session_state:
        try:
            prev_t = st.session_state["_mi_prev"]["tasks_df"]
            curr_t = st.session_state["_mi_curr"]["tasks_df"]
            if not prev_t.empty and not curr_t.empty:
                merged_opp = prev_t[["task_code","eff_finish","task_name"]].merge(
                    curr_t[["task_code","eff_finish"]], on="task_code", suffixes=("_p","_c"), how="inner"
                )
                for _, r in merged_opp.iterrows():
                    try:
                        gain = int((pd.Timestamp(r["eff_finish_p"]) - pd.Timestamp(r["eff_finish_c"])).days)
                        if gain > 5:
                            match = tasks[tasks["task_code"] == str(r["task_code"])]
                            wbs   = _wbs(match.iloc[0]) if not match.empty else "-"
                            rows.append(_ro_row(
                                "Opportunity", "Medium",
                                str(r["task_code"]), str(r.get("task_name","")), wbs,
                                f"Activity finish date has improved by {gain} days since the previous revision.",
                                f"The activity's forecast finish moved {gain} days earlier between programme versions.",
                                f"This improvement may allow successor activities to start earlier and could reduce the overall project duration.",
                                f"Review whether the {gain}-day improvement can be formally recognised in the programme. "
                                "Check if successor activities can be brought forward accordingly.",
                            ))
                    except Exception:
                        pass
        except Exception:
            pass

    # O2: High float activities that could be resequenced (float > 30 days)
    if "total_float_days" in tasks.columns:
        high_float = tasks[tasks["total_float_days"].apply(lambda f: safe_float(f,0) > 30)]
        for _, t in high_float.head(10).iterrows():
            tf = int(safe_float(t.get("total_float_days"),0))
            rows.append(_ro_row(
                "Opportunity", "Low",
                t.get("task_code",""), t.get("task_name",""), _wbs(t),
                f"Activity has {tf} days total float. There may be scope to resequence.",
                f"Significant float of {tf} days suggests this activity's resources or timing could be optimised.",
                "The float window could be used to smooth resource demand, resolve clashes, or accelerate predecessor activities.",
                f"Review with the planner whether the {tf} days of float can be used to resource-level, "
                "front-load critical activities, or release resource to other packages.",
            ))

    # O3: Float gained in comparison
    if "_mi_prev" in st.session_state and "_mi_curr" in st.session_state:
        try:
            prev_t = st.session_state["_mi_prev"]["tasks_df"]
            curr_t = st.session_state["_mi_curr"]["tasks_df"]
            if not prev_t.empty and not curr_t.empty and \
               "total_float_days" in prev_t.columns and "total_float_days" in curr_t.columns:
                mf = prev_t[["task_code","total_float_days","task_name"]].merge(
                    curr_t[["task_code","total_float_days"]], on="task_code", suffixes=("_p","_c"), how="inner"
                )
                for _, r in mf.iterrows():
                    fp = safe_float(r.get("total_float_days_p"), None)
                    fc = safe_float(r.get("total_float_days_c"), None)
                    if fp is not None and fc is not None and fc - fp > 10:
                        match = tasks[tasks["task_code"] == str(r["task_code"])]
                        wbs   = _wbs(match.iloc[0]) if not match.empty else "-"
                        gain  = round(fc - fp, 1)
                        rows.append(_ro_row(
                            "Opportunity", "Low",
                            str(r["task_code"]), str(r.get("task_name","")), wbs,
                            f"Activity gained {gain} days float since the previous revision.",
                            "Schedule logic or date changes in the current revision have created additional float.",
                            "This float could be used for resource smoothing, risk buffer, or to accommodate other priorities.",
                            f"Review whether the additional {gain} days of float has been deliberately created or is an unintended consequence. "
                            "Confirm with the planner that this aligns with programme strategy.",
                        ))
        except Exception:
            pass

    # O4: Near-critical activities where predecessors have high float
    if "total_float_days" in tasks.columns and not rels.empty:
        nc_acts = tasks[tasks["is_near_critical"]] if "is_near_critical" in tasks.columns else pd.DataFrame()
        if not nc_acts.empty and "pred_task_id" in rels.columns:
            task_float_map = tasks.set_index("task_id")["total_float_days"].to_dict() if "task_id" in tasks.columns else {}
            for _, t in nc_acts.head(10).iterrows():
                tid = t.get("task_id","")
                pred_ids = rels[rels["succ_task_id"] == tid]["pred_task_id"].tolist() if "succ_task_id" in rels.columns else []
                high_float_preds = [
                    p for p in pred_ids
                    if safe_float(task_float_map.get(p), 0) > 20
                ]
                if high_float_preds:
                    tf = round(safe_float(t.get("total_float_days"),0), 1)
                    rows.append(_ro_row(
                        "Opportunity", "Medium",
                        t.get("task_code",""), t.get("task_name",""), _wbs(t),
                        f"Near-critical activity ({tf}d float) has predecessors with high float.",
                        f"One or more predecessor activities have significant float, meaning they could potentially "
                        "be accelerated without impacting the overall programme.",
                        f"Accelerating high-float predecessors could create additional buffer on this near-critical path.",
                        f"Review whether the high-float predecessors can be started earlier or completed faster "
                        "to increase the float buffer on this activity.",
                    ))

    if not rows:
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    df = df.drop_duplicates(subset=["Type","Activity ID","Description"]).reset_index(drop=True)

    # Sort: Risks first, then Opportunities; within each by Priority
    priority_order = {"High": 0, "Medium": 1, "Low": 2}
    type_order     = {"Risk": 0, "Opportunity": 1}
    df["_ts"] = df["Type"].map(type_order)
    df["_ps"] = df["Priority"].map(priority_order)
    df = df.sort_values(["_ts","_ps","Activity ID"]).drop(columns=["_ts","_ps"]).reset_index(drop=True)

    return df


def page_risk_register(data: dict, near_crit_days: float):
    """
    Risk & Opportunity Register page.
    Auto-generates a draft register from programme data with
    editable owner, due date and status fields.
    """
    st.title("⚠️ Risk & Opportunity Register")
    st.caption(
        "Auto-generated draft register based on the uploaded programme. "
        "Edit Owner, Due Date and Status inline. Export when complete."
    )

    tasks = data["tasks_df"]
    rels  = data["relationships_df"]

    if tasks.empty:
        st.warning("No activities found. Please upload a programme first.")
        return

    notes_text = st.session_state.get("_notes_text", "")

    # ---- Generate / cache register ------------------------------------------
    prog_key  = st.session_state.get("_xer_cache_key", "")
    cache_key = f"_ro_register_{prog_key}_{near_crit_days}"

    if st.session_state.get("_ro_register_key") != cache_key:
        with st.spinner("Analysing programme for risks and opportunities..."):
            register_df = _generate_register(tasks, rels, near_crit_days, notes_text)
        st.session_state["_ro_register_df"]  = register_df
        st.session_state["_ro_register_key"] = cache_key
    else:
        register_df = st.session_state["_ro_register_df"]

    if register_df.empty:
        st.success("No risks or opportunities generated from the current programme data.")
        return

    # ---- Summary banner -----------------------------------------------------
    n_total = len(register_df)
    n_risk  = int((register_df["Type"] == "Risk").sum())
    n_opp   = int((register_df["Type"] == "Opportunity").sum())
    n_high  = int((register_df["Priority"] == "High").sum())
    n_med   = int((register_df["Priority"] == "Medium").sum())
    n_low   = int((register_df["Priority"] == "Low").sum())

    # Risk counts by priority
    r_high = int(((register_df["Type"]=="Risk") & (register_df["Priority"]=="High")).sum())
    r_med  = int(((register_df["Type"]=="Risk") & (register_df["Priority"]=="Medium")).sum())
    r_low  = int(((register_df["Type"]=="Risk") & (register_df["Priority"]=="Low")).sum())
    o_high = int(((register_df["Type"]=="Opportunity") & (register_df["Priority"]=="High")).sum())
    o_med  = int(((register_df["Type"]=="Opportunity") & (register_df["Priority"]=="Medium")).sum())
    o_low  = int(((register_df["Type"]=="Opportunity") & (register_df["Priority"]=="Low")).sum())

    st.markdown(
        f"""
        <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-bottom:20px;">
            <div style="background:#0B1F33;border-radius:12px;padding:16px 20px;">
                <div style="font-size:11px;font-weight:700;color:#64748B;text-transform:uppercase;
                            letter-spacing:1px;margin-bottom:8px;">Risks  ({n_risk})</div>
                <div style="display:flex;gap:12px;">
                    <div style="text-align:center;">
                        <div style="font-size:9px;color:#64748B;text-transform:uppercase;letter-spacing:0.8px;">High</div>
                        <div style="font-size:24px;font-weight:800;color:#dc2626;line-height:1.1;">{r_high}</div>
                    </div>
                    <div style="text-align:center;">
                        <div style="font-size:9px;color:#64748B;text-transform:uppercase;letter-spacing:0.8px;">Medium</div>
                        <div style="font-size:24px;font-weight:800;color:#d97706;line-height:1.1;">{r_med}</div>
                    </div>
                    <div style="text-align:center;">
                        <div style="font-size:9px;color:#64748B;text-transform:uppercase;letter-spacing:0.8px;">Low</div>
                        <div style="font-size:24px;font-weight:800;color:#94A3B8;line-height:1.1;">{r_low}</div>
                    </div>
                </div>
            </div>
            <div style="background:#0B1F33;border-radius:12px;padding:16px 20px;">
                <div style="font-size:11px;font-weight:700;color:#64748B;text-transform:uppercase;
                            letter-spacing:1px;margin-bottom:8px;">Opportunities  ({n_opp})</div>
                <div style="display:flex;gap:12px;">
                    <div style="text-align:center;">
                        <div style="font-size:9px;color:#64748B;text-transform:uppercase;letter-spacing:0.8px;">High</div>
                        <div style="font-size:24px;font-weight:800;color:#16a34a;line-height:1.1;">{o_high}</div>
                    </div>
                    <div style="text-align:center;">
                        <div style="font-size:9px;color:#64748B;text-transform:uppercase;letter-spacing:0.8px;">Medium</div>
                        <div style="font-size:24px;font-weight:800;color:#16a34a;line-height:1.1;">{o_med}</div>
                    </div>
                    <div style="text-align:center;">
                        <div style="font-size:9px;color:#64748B;text-transform:uppercase;letter-spacing:0.8px;">Low</div>
                        <div style="font-size:24px;font-weight:800;color:#94A3B8;line-height:1.1;">{o_low}</div>
                    </div>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # ---- Filters ------------------------------------------------------------
    with st.expander("Filter register", expanded=False):
        f1, f2, f3, f4 = st.columns(4)

        f_type   = f1.selectbox("Type",     ["All","Risk","Opportunity"], key="ro_f_type")
        f_pri    = f2.selectbox("Priority", ["All","High","Medium","Low"],  key="ro_f_pri")
        f_status = f3.selectbox("Status",   ["All","Open","In Progress","Closed"], key="ro_f_status")

        all_wbs  = ["All"] + sorted(register_df["WBS"].unique().tolist())
        f_wbs    = f4.selectbox("WBS",      all_wbs, key="ro_f_wbs")

    filtered = register_df.copy()
    if f_type   != "All": filtered = filtered[filtered["Type"]     == f_type]
    if f_pri    != "All": filtered = filtered[filtered["Priority"] == f_pri]
    if f_status != "All": filtered = filtered[filtered["Status"]   == f_status]
    if f_wbs    != "All": filtered = filtered[filtered["WBS"]      == f_wbs]

    st.caption(f"Showing {len(filtered)} of {n_total} items.")

    # ---- Tabs: Risks | Opportunities | Full Register | Export ----------------
    tab_risks, tab_opps, tab_full, tab_export = st.tabs([
        f"Risks ({n_risk})",
        f"Opportunities ({n_opp})",
        "Full Register",
        "Export",
    ])

    EDIT_COLS = [
        "Type","Priority","Activity ID","Activity Name","WBS",
        "Description","Cause","Effect","Mitigation / Action",
        "Owner","Due Date","Status",
    ]

    COL_CONFIG = {
        "Type": st.column_config.SelectboxColumn(
            "Type", options=["Risk","Opportunity"], width="small"
        ),
        "Priority": st.column_config.SelectboxColumn(
            "Priority", options=["High","Medium","Low"], width="small"
        ),
        "Status": st.column_config.SelectboxColumn(
            "Status", options=["Open","In Progress","Closed"], width="small"
        ),
        "Owner":              st.column_config.TextColumn("Owner",    width="small"),
        "Due Date":           st.column_config.TextColumn("Due Date", width="small"),
        "Activity ID":        st.column_config.TextColumn("Activity ID",   width="small"),
        "Activity Name":      st.column_config.TextColumn("Activity Name", width="medium"),
        "WBS":                st.column_config.TextColumn("WBS",           width="medium"),
        "Description":        st.column_config.TextColumn("Description",   width="large"),
        "Cause":              st.column_config.TextColumn("Cause",         width="large"),
        "Effect":             st.column_config.TextColumn("Effect",        width="large"),
        "Mitigation / Action":st.column_config.TextColumn("Mitigation / Action", width="large"),
    }

    def _style_row(row):
        """Row background by Type and Priority."""
        if row.get("Type","") == "Opportunity":
            return ["background-color:#f0fdf4;"] * len(row)
        priority = row.get("Priority","")
        cm = {
            "High":   "background-color:#fef2f2;",
            "Medium": "background-color:#fffbeb;",
        }
        return [cm.get(priority,"")] * len(row)

    def _show_editable(df, key_suffix):
        avail = [c for c in EDIT_COLS if c in df.columns]
        _df_edit = df[avail].copy()
        _df_edit = _df_edit.loc[:, ~_df_edit.columns.duplicated()]
        edited = st.data_editor(
            _df_edit.style.apply(_style_row, axis=1),
            hide_index=True,
            num_rows="fixed",
            column_config=COL_CONFIG,
            key=f"ro_editor_{key_suffix}",
        )
        return edited

    # -- Risks tab --------------------------------------------------------------
    with tab_risks:
        risks_df = filtered[filtered["Type"] == "Risk"].copy()
        if risks_df.empty:
            st.success("No risks match the current filters.")
        else:
            st.markdown(
                '<div style="font-size:12px;font-weight:700;color:#94A3B8;letter-spacing:1px;'
                'text-transform:uppercase;margin-bottom:8px;">Risks</div>',
                unsafe_allow_html=True,
            )
            st.caption("Edit Owner, Due Date and Status directly in the table.")
            edited_risks = _show_editable(risks_df, "risks")

            # Persist edits
            if edited_risks is not None and not edited_risks.empty:
                for col in ["Owner","Due Date","Status","Priority"]:
                    if col in edited_risks.columns:
                        register_df.loc[risks_df.index, col] = edited_risks[col].values
                st.session_state["_ro_register_df"] = register_df

            # High priority cards
            high_risks = risks_df[risks_df["Priority"] == "High"]
            if not high_risks.empty:
                st.markdown("---")
                st.markdown(
                    f'<div style="font-size:12px;font-weight:700;color:#dc2626;letter-spacing:1px;'
                    f'text-transform:uppercase;margin-bottom:10px;">'
                    f'High Priority Risks ({len(high_risks)})</div>',
                    unsafe_allow_html=True,
                )
                for _, item in high_risks.iterrows():
                    tc, bc, brc = _rag_colour("High")
                    st.markdown(
                        f"""
                        <div style="background:{bc};border:1px solid {brc};border-left:5px solid #dc2626;
                                    border-radius:8px;padding:14px 18px;margin-bottom:8px;">
                            <div style="display:flex;align-items:flex-start;justify-content:space-between;
                                        gap:12px;flex-wrap:wrap;">
                                <div style="flex:1;min-width:0;">
                                    {_ro_card_pill("Risk")}
                                    <div style="font-weight:700;color:#0B1F33;font-size:14px;margin-top:6px;">
                                        {item.get("Activity ID","")} - {item.get("Activity Name","")}
                                    </div>
                                    <div style="font-size:12px;color:#64748B;margin-top:2px;">
                                        {item.get("WBS","")}
                                    </div>
                                </div>
                            </div>
                            <div style="margin-top:10px;display:grid;grid-template-columns:1fr 1fr 1fr;gap:10px;">
                                <div>
                                    <div style="font-size:10px;font-weight:700;color:#94A3B8;
                                                text-transform:uppercase;letter-spacing:0.8px;">Description</div>
                                    <div style="font-size:13px;color:#334155;margin-top:3px;">{item.get("Description","")}</div>
                                </div>
                                <div>
                                    <div style="font-size:10px;font-weight:700;color:#94A3B8;
                                                text-transform:uppercase;letter-spacing:0.8px;">Effect</div>
                                    <div style="font-size:13px;color:#334155;margin-top:3px;">{item.get("Effect","")}</div>
                                </div>
                                <div>
                                    <div style="font-size:10px;font-weight:700;color:#94A3B8;
                                                text-transform:uppercase;letter-spacing:0.8px;">Mitigation</div>
                                    <div style="font-size:13px;color:#334155;margin-top:3px;">{item.get("Mitigation / Action","")}</div>
                                </div>
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

    # -- Opportunities tab ------------------------------------------------------
    with tab_opps:
        opps_df = filtered[filtered["Type"] == "Opportunity"].copy()
        if opps_df.empty:
            st.info("No opportunities detected from the current programme data.")
            if "_mi_prev" not in st.session_state:
                st.caption(
                    "Note: Comparison-based opportunities (pulled-forward activities, float gained) "
                    "require both previous and current XER files to be uploaded on the "
                    "Programme Comparison page."
                )
        else:
            st.markdown(
                '<div style="font-size:12px;font-weight:700;color:#16a34a;letter-spacing:1px;'
                'text-transform:uppercase;margin-bottom:8px;">Opportunities</div>',
                unsafe_allow_html=True,
            )
            st.caption("Review these opportunities with the planner to see if they can be captured.")
            edited_opps = _show_editable(opps_df, "opps")

            if edited_opps is not None and not edited_opps.empty:
                for col in ["Owner","Due Date","Status","Priority"]:
                    if col in edited_opps.columns:
                        register_df.loc[opps_df.index, col] = edited_opps[col].values
                st.session_state["_ro_register_df"] = register_df

    # -- Full register tab ------------------------------------------------------
    with tab_full:
        st.caption("Complete register -- all risks and opportunities combined.")
        edited_full = _show_editable(filtered, "full")
        if edited_full is not None and not edited_full.empty:
            for col in ["Owner","Due Date","Status","Priority"]:
                if col in edited_full.columns:
                    register_df.loc[filtered.index, col] = edited_full[col].values
            st.session_state["_ro_register_df"] = register_df

    # -- Export tab -------------------------------------------------------------
    with tab_export:
        st.markdown("Download the full Risk & Opportunity Register as a formatted Excel workbook.")

        final_df = st.session_state.get("_ro_register_df", register_df)

        # Summary sheet
        status_counts = final_df.groupby(["Type","Priority"]).size().reset_index(name="Count")
        owner_counts  = final_df[final_df["Owner"] != ""].groupby("Owner").size().reset_index(name="Assigned Items")

        summary_data  = pd.DataFrame({
            "Metric": [
                "Total Items","Total Risks","Total Opportunities",
                "High Priority Risks","Medium Priority Risks","Low Priority Risks",
                "High Priority Opportunities","Medium Priority Opportunities","Low Priority Opportunities",
                "Open Items","In Progress","Closed",
            ],
            "Count": [
                len(final_df),
                int((final_df["Type"]=="Risk").sum()),
                int((final_df["Type"]=="Opportunity").sum()),
                int(((final_df["Type"]=="Risk")&(final_df["Priority"]=="High")).sum()),
                int(((final_df["Type"]=="Risk")&(final_df["Priority"]=="Medium")).sum()),
                int(((final_df["Type"]=="Risk")&(final_df["Priority"]=="Low")).sum()),
                int(((final_df["Type"]=="Opportunity")&(final_df["Priority"]=="High")).sum()),
                int(((final_df["Type"]=="Opportunity")&(final_df["Priority"]=="Medium")).sum()),
                int(((final_df["Type"]=="Opportunity")&(final_df["Priority"]=="Low")).sum()),
                int((final_df["Status"]=="Open").sum()),
                int((final_df["Status"]=="In Progress").sum()),
                int((final_df["Status"]=="Closed").sum()),
            ],
        })

        export_sheets = {
            "Summary":          summary_data,
            "Full Register":    final_df,
            "Risks":            final_df[final_df["Type"]=="Risk"],
            "Opportunities":    final_df[final_df["Type"]=="Opportunity"],
            "High Priority":    final_df[final_df["Priority"]=="High"],
            "Open Items":       final_df[final_df["Status"]=="Open"],
        }
        if not owner_counts.empty:
            export_sheets["By Owner"] = owner_counts

        xls_bytes = export_df_to_excel(export_sheets)

        dl_col, _ = st.columns([1,3])
        dl_col.download_button(
            label="📥  Download Risk & Opportunity Register",
            data=xls_bytes,
            file_name=f"risk_register_{datetime.now().strftime('%Y%m%d')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            help="Exports Summary, Full Register, Risks, Opportunities, High Priority and Open Items sheets.",
            key="dl_011",
        )

        st.markdown(
            '<div style="background:#f8fafc;border:1px solid #E2E8F0;border-radius:8px;'
            'padding:14px 18px;margin-top:12px;">'
            '<div style="font-size:12px;font-weight:700;color:#0B1F33;margin-bottom:6px;">'
            'Workbook sheets</div>'
            '<div style="font-size:12px;color:#64748B;line-height:2;">'
            + "".join(f"<strong>{k}</strong> &nbsp;|&nbsp; " for k in export_sheets.keys()) +
            '</div></div>',
            unsafe_allow_html=True,
        )
        st.caption(
            "Note: This is an auto-generated draft register. "
            "Review all items with the project team before issuing formally."
        )



# -----------------------------------------------------------------------------
# PAGE: EXPORT REPORTS
# -----------------------------------------------------------------------------

def page_export_reports(data: dict, near_crit_days: float):
    st.title("📥 Export Reports")
    st.markdown("> Download all schedule data as formatted Excel reports.")

    tasks = data["tasks_df"]
    rels = data["relationships_df"]
    wbs = data["wbs_df"]
    resources = data["resources_df"]

    if tasks.empty:
        st.warning("No data loaded to export.")
        return

    tasks = get_critical_threshold(tasks, near_crit_days)
    critical = tasks[tasks["is_critical"]]
    neg_float = tasks[tasks["total_float_days"].apply(lambda f: f is not None and f < 0)]

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Single-Sheet Exports")

        # All activities
        avail = [c for c in ["task_code","task_name","wbs_path","eff_start","eff_finish",
                              "orig_dur_days","rem_dur_days","total_float_days","free_float_days",
                              "status","task_type","is_critical","cstr_type"] if c in tasks.columns]
        xls = export_df_to_excel({"All Activities": tasks[avail]})
        st.download_button("📄 All Activities", xls, "all_activities.xlsx",
        key="dl_012",
                           )

        # Critical path
        avail_c = [c for c in avail if c in critical.columns]
        xls2 = export_df_to_excel({"Critical Path": critical[avail_c]})
        st.download_button("🔴 Critical Path Activities", xls2, "critical_path.xlsx",
        key="dl_013",
                           )

        # Relationships
        if not rels.empty:
            xls3 = export_df_to_excel({"Relationships": rels})
            st.download_button("🔗 All Relationships", xls3, "relationships.xlsx",
        key="dl_014",
                               )

    with col2:
        st.subheader("Multi-Sheet Reports")

        # Full schedule pack
        sheets = {"All Activities": tasks[avail]}
        if not critical.empty:
            sheets["Critical Path"] = critical[avail_c]
        if not neg_float.empty:
            sheets["Negative Float"] = neg_float[[c for c in avail if c in neg_float.columns]]
        if not rels.empty:
            sheets["Relationships"] = rels
        if not wbs.empty:
            sheets["WBS"] = wbs
        if not resources.empty:
            sheets["Resources"] = resources

        xls_full = export_df_to_excel(sheets)
        st.download_button("📦 Full Schedule Data Pack", xls_full, "schedule_data_pack.xlsx",
        key="dl_015",
                           )

        # WBS summary
        if "wbs_path" in tasks.columns:
            tasks["wbs_top"] = tasks["wbs_path"].apply(
                lambda x: str(x).split(" > ")[0] if pd.notna(x) and x else "Unknown"
            )
            wbs_summary = tasks.groupby("wbs_top").agg(
                total=("task_id","count"),
                critical=("is_critical","sum"),
                near_critical=("is_near_critical","sum"),
            ).reset_index()
            xls_wbs = export_df_to_excel({"WBS Summary": wbs_summary})
            st.download_button("🌲 WBS Summary", xls_wbs, "wbs_summary.xlsx",
        key="dl_016",
                               )


# -----------------------------------------------------------------------------
# PAGE: HOME  (PlanTrace branded landing page)
# -----------------------------------------------------------------------------

def _page_home():
    """
    PlanTrace branded homepage.
    Shown when no XER is loaded, or when the user navigates to Home.
    """

    # ---- Hero section -------------------------------------------------------
    st.markdown(
        """
        <div style="padding: 48px 0 24px 0;">
            <div style="font-size: 13px; font-weight: 700; color: #F5A623;
                        letter-spacing: 2px; text-transform: uppercase;
                        margin-bottom: 10px;">
                Project Programme Intelligence
            </div>
            <div style="font-size: 52px; font-weight: 900; color: #0B1F33;
                        line-height: 1.1; letter-spacing: -1px;">
                PlanTrace
            </div>
            <div style="width: 56px; height: 4px; background: #F5A623;
                        border-radius: 2px; margin: 14px 0 18px 0;"></div>
            <div style="font-size: 20px; font-weight: 400; color: #334155;
                        margin-bottom: 10px;">
                Trace logic. Expose risk. Drive delivery.
            </div>
            <div style="font-size: 15px; color: #64748B; max-width: 680px;
                        line-height: 1.7; margin-bottom: 32px;">
                Project planning intelligence for delivery teams. Upload an XER programme,
                trace predecessors and successors, review critical paths, check programme
                health and understand labour demand &mdash; without opening P6.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # ---- Upload prompt -------------------------------------------------------
    st.markdown(
        """
        <div style="background:#0B1F33; border-radius:12px; padding:22px 28px;
                    display:flex; align-items:center; gap:20px; margin-bottom:36px;
                    max-width:560px;">
            <div style="font-size:28px;">📂</div>
            <div>
                <div style="color:#E8951D;font-weight:700;font-size:15px;
                            margin-bottom:4px;">Ready to start</div>
                <div style="color:#CBD5E1;font-size:13px;line-height:1.5;">
                    Upload your <strong style="color:#fff;">.xer file</strong>
                    using the panel on the left to begin analysis.
                    <br>Export from P6 via
                    <strong style="color:#E8951D;">File &rarr; Export &rarr; Primavera P6 XER</strong>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # ---- Feature cards -------------------------------------------------------
    st.markdown(
        '<div style="font-size:13px;font-weight:700;color:#94A3B8;'
        'letter-spacing:1.5px;text-transform:uppercase;margin-bottom:16px;">'
        'What PlanTrace does</div>',
        unsafe_allow_html=True,
    )

    CARDS = [
        {
            "icon": "🔗",
            "title": "Logic Trace",
            "body": (
                "See what drives an activity and what it impacts. "
                "Trace full predecessor and successor chains across the network, "
                "with depth levels and relationship types shown at every step."
            ),
        },
        {
            "icon": "🚨",
            "title": "Critical Path",
            "body": (
                "Review the full critical path, near-critical work and negative float. "
                "Identify which activity or milestone is at risk and understand "
                "exactly what is driving it."
            ),
        },
        {
            "icon": "👷",
            "title": "Labour Demand",
            "body": (
                "View labour histograms by week, month, WBS and resource. "
                "Identify peak demand periods and understand resource loading "
                "across the programme."
            ),
        },
        {
            "icon": "🩺",
            "title": "Programme Health",
            "body": (
                "Find missing logic, open ends, constraints, excessive lag and "
                "planning risk before they cause problems. "
                "Eleven automated quality checks with export."
            ),
        },
    ]

    cols = st.columns(4, gap="medium")
    for col, card in zip(cols, CARDS):
        with col:
            st.markdown(
                f"""
                <div class="pt-card">
                    <div class="pt-card-icon">{card["icon"]}</div>
                    <div class="pt-card-accent"></div>
                    <div class="pt-card-title">{card["title"]}</div>
                    <div class="pt-card-body">{card["body"]}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    # ---- What's in the tool -------------------------------------------------
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown(
        '<div style="font-size:13px;font-weight:700;color:#94A3B8;'
        'letter-spacing:1.5px;text-transform:uppercase;margin-bottom:16px;">'
        'All pages</div>',
        unsafe_allow_html=True,
    )

    PAGE_LIST = [
        ("📊", "Project Summary",          "Activity counts, float distribution, WBS breakdown and schedule span."),
        ("🔍", "Activity Search",           "Search and filter activities. View full detail, dates, float and logic."),
        ("🔗", "Logic Trace",               "Trace predecessors and successors through the network with depth levels."),
        ("🚨", "Critical Path Analysis",    "Full critical path, near-critical and negative float by WBS."),
        ("🎯", "Critical Path to Activity", "Identify the driving chain into any selected activity or milestone."),
        ("👷", "Labour Histogram",          "Weekly and monthly labour demand by resource, WBS and package."),
        ("🩺", "Schedule Health Check",     "Eleven automated quality checks with counts, tables and export."),
        ("📝", "Planning Notes",            "Upload notes, link to activities, keyword search and highlighting."),
        ("📅", "Programme Comparison",      "Compare two XER revisions. See what moved, changed or became critical."),
        ("📥", "Export Reports",            "Download all data as formatted Excel workbooks."),
    ]

    left, right = st.columns(2, gap="large")
    for i, (icon, title, desc) in enumerate(PAGE_LIST):
        col = left if i % 2 == 0 else right
        with col:
            st.markdown(
                f"""
                <div style="display:flex;gap:14px;align-items:flex-start;
                            padding:14px 0;border-bottom:1px solid #E2E8F0;">
                    <div style="font-size:22px;min-width:30px;margin-top:2px;">{icon}</div>
                    <div>
                        <div style="font-weight:700;color:#0B1F33;
                                    font-size:14px;margin-bottom:3px;">{title}</div>
                        <div style="color:#64748B;font-size:13px;
                                    line-height:1.5;">{desc}</div>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    # ---- Footer -------------------------------------------------------------
    st.markdown("<br><br>", unsafe_allow_html=True)
    st.markdown(
        """
        <div style="border-top:1px solid #E2E8F0;padding-top:18px;
                    display:flex;justify-content:space-between;align-items:center;">
            <div style="font-size:13px;color:#94A3B8;">
                <strong style="color:#0B1F33;">PlanTrace</strong>
                &nbsp;&nbsp;|&nbsp;&nbsp;
                Built for Primavera P6 XER programmes
                &nbsp;&nbsp;|&nbsp;&nbsp;
                No P6 licence required
            </div>
            <div style="font-size:12px;color:#CBD5E1;">
                Upload a .xer file to begin
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


# -----------------------------------------------------------------------------
# -----------------------------------------------------------------------------
# UI COMPONENTS & NAVIGATION  --  PlanTrace v3
# -----------------------------------------------------------------------------

import base64 as _b64

# -- Palette constants ---------------------------------------------------------
_NAVY    = "#0B1929"
_NAVY2   = "#0d2035"
_AMBER   = "#E8951D"
_RED     = "#C0392B"
_GREEN   = "#1E7A4E"
_TEXT    = "#1C2B3A"
_MUTED   = "#6B7C8E"

_MODE_KEY = "plantrace_mode"

def is_pm_mode():
    return st.session_state.get(_MODE_KEY, "PM Mode") == "PM Mode"

def mode_label():
    return st.session_state.get(_MODE_KEY, "PM Mode")

def pm_note(text):
    if is_pm_mode():
        st.info("What this means: " + text)

def planner_note(text):
    if not is_pm_mode():
        st.caption("Planner note: " + text)

def mode_badge():
    if is_pm_mode():
        return '<span style="font-size:11px;font-weight:700;color:#E8951D;padding:2px 8px;border:1px solid #E8951D;border-radius:4px;">PM Mode</span>'
    return '<span style="font-size:11px;font-weight:700;color:#2DD4BF;padding:2px 8px;border:1px solid #2DD4BF;border-radius:4px;">Planner Mode</span>'

def mode_toggle_bar():
    st.markdown('<div style="display:flex;justify-content:flex-end;margin-bottom:8px;">' + mode_badge() + '</div>', unsafe_allow_html=True)

_PM_TASK_COLS   = {"task_code":"Activity ID","task_name":"Activity Name","wbs_path":"WBS","eff_start":"Start","eff_finish":"Finish","total_float_days":"Float (d)","status":"Status","is_critical":"Critical"}
_PLANNER_TASK_COLS = {"task_code":"Activity ID","task_name":"Activity Name","wbs_path":"WBS","task_type":"Type","calendar":"Calendar","eff_start":"Start","eff_finish":"Finish","early_start":"Early Start","early_finish":"Early Finish","late_start":"Late Start","late_finish":"Late Finish","orig_dur_days":"Orig Dur (d)","rem_dur_days":"Rem Dur (d)","total_float_days":"Total Float (d)","free_float_days":"Free Float (d)","status":"Status","is_critical":"Critical","cstr_type":"Constraint","phys_pct":"% Complete"}
_PM_REL_COLS    = {"pred_task_code":"Predecessor ID","pred_task_name":"Predecessor Name","succ_task_code":"Successor ID","succ_task_name":"Successor Name","lag_days":"Lag (d)"}
_PLANNER_REL_COLS = {"pred_task_code":"Pred ID","pred_task_name":"Pred Name","succ_task_code":"Succ ID","succ_task_name":"Succ Name","rel_type":"Link Type","lag_days":"Lag (d)"}

def mode_cols(df, pm_cols=None, planner_cols=None):
    pm_cols = pm_cols or _PM_TASK_COLS
    planner_cols = planner_cols or _PLANNER_TASK_COLS
    cols  = pm_cols if is_pm_mode() else planner_cols
    avail = {k: v for k, v in cols.items() if k in df.columns}
    out   = df[list(avail.keys())].copy().rename(columns=avail)
    for col in out.columns:
        if any(w in col.lower() for w in ("start","finish","date")):
            try: out[col] = out[col].apply(format_date)
            except Exception: pass
    if "Critical" in out.columns: out["Critical"] = out["Critical"].apply(lambda x: "Yes" if x else "")
    if "Status"   in out.columns: out["Status"]   = out["Status"].apply(lambda x: _status_label(str(x)) if x else "")
    return out



# -- Nav definition (no emoji -- professional text icons) -----------------------
_NAV_GROUPS = [
    ("Dashboard", [
        ("overview",       "Overview"),
        ("executive",      "Executive Summary"),
        ("multiproject",   "Multi-Project"),
    ]),
    ("Programme Analysis", [
        ("programme",      "Programme"),
        ("logic",          "Logic"),
        ("critical",       "Critical Path"),
        ("labour",         "Labour"),
        ("health",         "Health Check"),
        ("dcma",           "DCMA Assessment"),
        ("scurve",         "S-Curve"),
        ("comparison",     "Comparison"),
    ]),
    ("Forensic & Planning", [
        ("forensic",       "Forensic Planner"),
        ("acceleration",   "Acceleration Analysis"),
        ("resources",      "Resource Levelling"),
        ("milestones",     "Contract Milestones"),
    ]),
    ("Commercial", [
        ("ce_tracker",     "Compensation Events"),
        ("client_review",  "Client Comment Review"),
        ("pm_actions",     "PM Actions"),
        ("risk",           "Risk Register"),
    ]),
    ("Reports & Settings", [
        ("reports",        "Reports"),
        ("settings",       "Settings"),
    ]),
]

_NAV = [(k, l) for _, items in _NAV_GROUPS for k, l in items]

_NEEDS_PROG = {
    "overview","executive","programme","logic","critical",
    "labour","health","pm_actions","risk","reports",
    "forensic","scurve","resources","milestones","multiproject","dcma",
    "client_review","acceleration","ce_tracker",
}


# -----------------------------------------------------------------------------
# Logo helper
# -----------------------------------------------------------------------------
def _logo_b64(width: int = 80) -> str:
    """Return an <img> tag with the PlanTrace logo embedded as base64.
    The image data is baked in so the logo works on any deployment.
    """
    _LOGO_DATA = "/9j/4AAQSkZJRgABAQAAAQABAAD/4gHYSUNDX1BST0ZJTEUAAQEAAAHIAAAAAAQwAABtbnRyUkdCIFhZWiAH4AABAAEAAAAAAABhY3NwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAQAA9tYAAQAAAADTLQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAlkZXNjAAAA8AAAACRyWFlaAAABFAAAABRnWFlaAAABKAAAABRiWFlaAAABPAAAABR3dHB0AAABUAAAABRyVFJDAAABZAAAAChnVFJDAAABZAAAAChiVFJDAAABZAAAAChjcHJ0AAABjAAAADxtbHVjAAAAAAAAAAEAAAAMZW5VUwAAAAgAAAAcAHMAUgBHAEJYWVogAAAAAAAAb6IAADj1AAADkFhZWiAAAAAAAABimQAAt4UAABjaWFlaIAAAAAAAACSgAAAPhAAAts9YWVogAAAAAAAA9tYAAQAAAADTLXBhcmEAAAAAAAQAAAACZmYAAPKnAAANWQAAE9AAAApbAAAAAAAAAABtbHVjAAAAAAAAAAEAAAAMZW5VUwAAACAAAAAcAEcAbwBvAGcAbABlACAASQBuAGMALgAgADIAMAAxADb/2wBDAAUDBAQEAwUEBAQFBQUGBwwIBwcHBw8LCwkMEQ8SEhEPERETFhwXExQaFRERGCEYGh0dHx8fExciJCIeJBweHx7/2wBDAQUFBQcGBw4ICA4eFBEUHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh7/wAARCASQAxADASIAAhEBAxEB/8QAHQABAQACAgMBAAAAAAAAAAAAAAECCAYHBAUJA//EAGMQAAIAAwYDBQEJBQ8SAwkBAAABAgMRBAUGITFBB1GBCBJhcZHREyIyobGys8HwFEJSdJIVFyMzNjdDU2JydYKT0uEWJic1RVRVY2Rlc4OElKKjwtMkNEQYJUZWhZXD4vHj/8QAGwEBAAEFAQAAAAAAAAAAAAAAAAEDBAUGBwL/xAA9EQEAAQMCAgUICQMEAwEAAAAAAQIDBAURBjESIUFRcRMyMzRhgaGxFBYiI1JykcHRFVPhJDVC8CVD8WL/2gAMAwEAAhEDEQA/ANvBR1KD08pQpcwA6AABTJijA5AKEoUAShaeAYAU8CNMtA/IDGgSz0MqeAVQJTwKk/EcggCRegIE7qQBhCOoSdCsgDoOgAAFQAApADHqAAVS0IqgBRkKOgCj5DbQABzG4VagCko6gmdQLmkAKgMyZlCAlGXMcxuAFAxmBANy+0CItAg6gBnQewZgM6gcgA5ZMZhPMZgNtB0AYGLXgKFYzp0AepdiAC5jcZ5BAAKMZ5ACAoESBQBH5MIOoVQBegGYDmGsxnmAIRLIoAFXkQoAAAH5E30LnyIALyyJ0KqgEVpkFQAFeQqAJrUFQEzqHUuYAADcAAEAAzKBKjMbEAq8huTMqbAZgDMCkZSZgUDMICOoDABMVYAFIVEAMBkVQKNwACG4QABNh1IBa+AqQAVVoN9yIvkAr4BMgAozJmAKhmRF3ADcDfcBmKse0AFUBIbAEBmUCMmZSPkBMzJGO5kgMkRlIwIMxmQCupKgZgVagKoAAZj1AAbAAAuoAZhAIAMxmADDI2VgANgq+IBkRcyICjMIAFUBV2GYDMmYADMu3QjKAzAG4AhSKoAuYIBRVimYADMBVAZgmYzAqD0JuUCk3DG4CugQ5AC10FSAAPviPQqpUANy9SAAMvswAqAgBQT7aioF6AgXQC7ggApKgMBUIMIBUVFBQCkHIAK/aoTFCAGy1IAC0KF0CAIbAAQoFACBPYXmAG4JTPYC+0vImwQFQIAAA5kQmQpAiUFCvTqQMC1FTH0HoBWwgAG4qQAZEY3FftUAwAAYIwBkmRvIhQHUVGQAAlC0Amxa/apGGBUNgOgBhMjCAqLUhQIhuAAGwAAvsJsNgCG4AAAgF3Fcx7ABVqH9ZNygCdA/IgFqSpGAKitkDAMZ1LlmNwFAB0AMhWGBjmNwXcCgDcCeuhQ+oAU8CULkHoBGgV0J6gRhfWUIAhXIqpmGAIUgTAVJh/WEwhSUZUH5ASgKTIBmQoX1ASjFGV66ACBalQAZhVGwQDPmKAZAQUKqDcBQm5RuAaZCvrqQAhRhFAmYCoAncaYzK6eI6BCZ1Jn8ZQ/LcJgQC0KwhPYM6alb+Qi0CdhjMu5MggAqqlyAUZKZl2DAjFPCgY6ARFGXiUCKpUFpoEA2FBXItUBjQNMofkApoShQBKArFV4gTMuZCtgF9RGUMDFlVQEAzBeg26ARVoMyrQZAKEKAGYHqNgG4egQAmZFUyqOgGNBQoyAi1DKBApKlJuBfAEroUAGGQCMu5Ni0zAqACYBgMACMpAAQAEKthsEgKAOQTIgwRsIGVEYQGSIAAqAAAWgIBaipABaivgQoDYiKRAVgABzG5B4gULUg3ApAQCoVyIUC8yLUIbgUAAQvtA9oAhfQjAPUlRQAXcvsMSgEGNwBW8hUgAPQIhUBUAGA2FR7SAKl5GOxQK+YfMj3AFCBNgGw3AQDcVHMANsy1ItOgYAIhV0AvQD2jYAgQAVkBAKWpABUwENtQBC76kAtSAASpSADKpK5lG+oE6FqQq8wHQdAxoAypoMq6IF31AlRvoV+ZNwD8gGACfgARgWvgQpUBHpoFtluWmWqIAT8A34Co9AJVciGRGAr4FRGVeYFROgKwImC18dyVAV8CVKAJXwHQo6gF5FItC7gCFRAFQGQBUVzIZABvoBvqAfluKjqAJ0HQo9AICk31AtSkdOYy5gOgZN9TL2gQPyAAj8h0KNgJvoh0LuAJXMFQaAlQgUCF6BgB6jLkOofmAqSuegr4hPyAZci9CF0Am2hehCvcAhUVFSA6BBvxCEB0HQAkF5E6FWmqD11AhU/AhVoAqKjIAEAgA6DmUiAb6ApAFcxX5R1J7QLUnQoQEy5FIXmBNw/IbiviABEAMgQoBkKxTMBohUUyI9QKEAABGAFPAAAVFRiUCkKToAA5kYAANARmSIwgKg/qCAEQKT1ABMmwQGWXIbBB76ABuQAUECArRAAA3INwMgiFAbAjABFMehQL4k3CYAqBAgL5oe0lPAP6wKQACkA2AoJuUBuAAA6BkAFIUAR/GUjAgQFALt0BCgCMMMCoIIIJgKgRA3UMgCBBgm+wFrqEQoFIOg9gAvQg5gUnMEAyYIVANyAAAmABSEKEpQMbhhKjclfAVzDyvQpKlAPyA6CoDYbkrkG8wKNxXzG4E6FoQoE2BaojYAIVGwAdAAL0J0CYqBehC1JXcA/Iq8iMqAdB0FSVAqIE0K+YTARb5FFUEBTGuZagBQIANtCIy9hAA6AVQAb6AuTABeQG4EfkN9A2MgIvItPAm5QHQCormAKiVAFIy1WQr8oEAAABk2AtSvyMa0LUBuOgWoYALyFQA6FJsUAgF5DICdBQtSVzAbaFGwAjXgGvANqgYAqWehCphO5TwCFUAhOgAAq00I14FTI35gToEGFQC76F6EFckJBDcJ5MbgOgoXoEBOgKSoAaiodAKOhEyrqBKDnkUlQ9Aemg3DYBLcm5dwHkAAFeoIwApkGsxsNwLQUzAACgY6AKEZfYRgACgQFehPaAXmNgAAAArQ6kYAE6lIAQC6gJKCmXQo26BBQAUAqBEUBTIlCkWgBoUG4AItDH2FAo3DJXMB1BK5ACgiKAA9SAUEqUAX2mKMvaAp4koUnqAoKAbARrMpGUBuKDfkABaEKgDWQoNsgADIR6AZbdSLzJ7ShJtqX0J7ChCBoepWBAl4j1AF3JyAQEKQoBfUSniBUAwgEgKqlpkABEKZkRUBaEAAMAASlWUi16lAJFAQAhSATcPcbiLQJ3ZIlcy0I9QgRegp4ACNivyB0CpUAnloWoohuArmWvgTyKgJ0LUEdAFc9CNl9SMBXwLXwJQq0AreRK+G4y8RUBnyFch4sdAHQN5FyIwI2VMkWoVAkT8B0L0GWWTCBDoFQdAFabBPLoOjGVH5AOgqMhkArloKkyzKA6BPIbAA34DPkNxkBBXMoyqBKiuZSbgTYtRQtAIn8g3LTwGQDveBOgKBAmWgyAxqZVIVUABsZB0AVfxEr4DcZUAVzL0G4y5AKkb8CkYDoWoyCAtRUMAToSpSOgBN5BPPQFy5AK5BsbACN+AbK9NCAWpK5aFKgJ0C8hkEA6EqWiDoBK/IK5jKnQU8AFSpkCoBU/AVIVUAJ+ACABvwJUoAlQigAteo2G/UIC1FQTIA3noOeRSZBMG5HoXcOmeQQAAABUVAMB6jcBsKgMAggUBUlQygToRlAAIAAKlIBAXmNgCBSAGRFYAIVAYBFIgwHQVAAMVyD8gBBuXYAKgAAwAAW4CKBAWpK5gOZSFqBAEGBAtSkAoqABEUb6D2gAFoGAqBUlQLuBuAA5AMBUVGwyAVAKBAUgFCIvIqYELUACBhh5gEEEABEysAEwAAXUcghuAIZEAgRWACAQ3AApOgABgAgyrUjAVIUgAJgAKlb1IVgPjHQZEqqgATIvgBcxTzJkXIABlsNwGY3AyAAhcgHQdBsMgBehMgqAB0AAdAMiAUPQEdABURlVAAYyI6AAFQZB6UE8hlQPKscyOgy6gAEXIBsQZBACjKgyAiL0JkALXkTcAJ2Ni8jEqoEKNiFyAAZEAtBtoMhVAKAVGVAkINSZBCgZE2ApSVQ3ApAqAClJl4DIA/IEZQHRhjIlQAXkTIuQF2AyoMgDD8iOgdALsERUKqAAPQioBQSqLVAOgqTKhcgA6DIgFJ0DarqMqdACKRUoXKoAdBkTLMCvyAbJkBegJuMgKShciKgF6BAZeAAOpK5htBOykpmUbhCUKAAAoVAQr1GxGAC1AWoApCgQAbgFoyrqRFQDIBj2hMJQMoCEXgOYIBQtA0ABGUMCIBAAANgALQlAGRWRB6gUiAQFAABB6jIPUBQUzApmBNiimQAm5eRABUOYWgAAACFAAUyIUnQANi+wbATfxKQoBDkAA8QAAYX1BgAgwHoBPaEWnykAoRCoB6hhoMJEEgirwCEIVkAi6lJQoFXqAqUIBSdSkAFG42AiKAAKgyLcAwVkoBNy7CmZaATcKgC5gACBMBGUMIUZAm4AoHxgH0KQUAbB9BsR6gUE6lSCd16AhQhAWhGAMjEqAbDkNV/QSgFQIi0AcibAeQDYBhAEAGAAQYALQgWgFazA3AERfQgAdAtCkSAoAADchUBV6ha7AUz2APQgf1gCAUAFQ32HMLUCgMgAempC7BPUB0AoBPYVLIewBAAAHoAgBaeQAAdCBhgB6AABTMUyCAZABAAwGAWwRCoCsgCAbbACgBAIAWg9CU+1CgERAUAIBIUzABD7aBAZAgAb9QNx6gCfbQACCg9S0Cd03D6ArXyBANxUm4FBK6ZlqAYDIBakepK5MoFCHqFqBSmJagCMBgKhEqXqAAeg3AqHIlQ3kAAI+oB15Mq8iMAXoHUlRUCoMiZQDInqUxAyzGZEyoAASoFCImAKyPyZWyMCKpkjEyWoGSBENwIM66CpGABEUAKipAMhsQAN9CoiKAQFQAIyk2AFMSgVCpKlqAFSNgAykKgCADAuZEPaKgANtwwAZGX1AIKoQABAIAUgAdAF9RGBasbaGJUBS7EFQKg6kqKgCVDIBkV+RAA30GxBUCgVFQnYAqOYQbh6EqRvIQLkN9QAGQFCAV+Y6kZQJkVUFMtBTMC5EyKTcBkXInQoDLIOgIwDoOQfkALsMuZC+0BlmTLmCAUF6E20Cdx0GXMMIIMiZFFACoVUCAAnUABkXKmpOgApAKARUoMqDkALlUZcyULQCKhVSpC7gVUC1IVagNtSU8SvTQcgJkXIiFAIBQLUBkUUABULkRF6ACPzL0IwAyoBsAy5lyJuKAMiF6EaAOhciUKkAoXIDoBSMBgTLmNx7SUAvUEL0AbFdOZiUC5cxlzIEBcgqAIAqDqQAVUIFoOgEyKGER1vQFQEJeVyGREWgDbUZAgFy5jIhQHUjKSmQDIIUz3CAuRKjoKeADqGKZhrwAyJTMpNwKPULzGXMCFI/MV8QLlQbiuWormACHqOoBlMX5lQFRGF5jqAAfmK6AAK+Ir5gQBACjYiGwBhCviF5gZLqPUi8w2ABF5gJhSewpOoQDpUnqVdQKtAwtB1AhC5cyeFQKB1HUAvIbhMq1AchuBXPqBfUhK+JUAXkCJ+JQmAm5eo31CAIlfEqYAr8iIV+UAGggAoB66CviA3Arn1AANCor4gKAdQAfkCdS18QAFQAXkQtfMlfEACVKAfUMPR5hhMKETqF5hDIiHqE/EAEAgCXmQqHUCFWg6sbAPUgrmQCrTQpFoXnmAJ0L6kr4gUcidSrzAdAK5j1AU8CU8GUifiBSDqVeYEDLkHTmA6jcDcBmKAARgrG4EWhN+pQwG+xQAJn8QzBcgAegRHqAHoAgKwtiAAvMAAVaB6AMCMqDCYDMOoIwAGw3AVC0BV0AlCgVAIIegbAE5l2IALmAwINwVoAqjOvUbjcDF9CgAEMwigEQpNwHUqCAB+YzIX4gnYDCAQDOgqTYBv1LmQAXMhURgUD0CAjKRlAhQg2AzMdy1IAKCoBsRlD8gIFoEVAAGOQEKvMhQKiP0CfkGA6jOgIgHULQCgFWgCABkKRAAqlHQCbgoAdSFIAKiBbgXcPzAYEoNyhagKAdABGCigEQY2FfACkHQIBuKAoBEZQwIX4gEBGCgCFIi7AAykYGLKiRFhAChUGBEgVFAlMgUnPICAryABCgFQJTIIpEBRQAAi7kRQIANwMWUAAi0CAAm5SAVAbCoEHtBegBIBaACewhk/qJsAAKBAVEYAAoEY9g5l2AEY3KBiEXbb1CAFIVAA0AwI0EgVACF3AEAKAQC00HQBzHMABQU8NgUCAoYEJQyIgAD8h0AbigTzLsBEsyFIAoECoAlmKDcrYGKG5RlXUCdCj1ABgNj1Cdk2JuZbah0qDZjQbmQ3CGPQyQC8wHQjKgwICgA+piZbE6hMIioFWniAD00KTqEMWvlC8jJoKgERfUBgEAgAG3QpFuBHroC5DICLQepaqg6gRBaaFQWgAAAQvQACDctSVz1AU8AOoqARV5EXUoEBSAQIpUBEX2gBIgwg6eIQjIsjL2E6gBQvVgCBloRgQGVPFigGJQ/NjYACoZAQi1L1CAhQAG2gBWE9ieoRUOoQnqUBBKUBQEItwXqAIVEKgFPMIoAiBUEBNuhEZE9QBOhQBNwUPqAREUL7ZAQIoQSIABAxXMdCZgWoqTcAGXcjL0CYNg3mXbQjCAJkfkNwFeRepCgKhgMAB0ACo67joEAHIdB0ApAHXkAYDAFI2CAEy1IggKWviQdAFSVAaAJipKMAWoQXkVIBXLUVD10IAqGEUCVJ7ShVAVCBaeABAZ8hmE7FSMpAgYTAQBPTMoSYX1gXqSpdyMBXPUlfEBJ00Au4r4jcMBXMgzqOgFAGfICMVI6lAqYqQrAbAE6ACkz5FpkA2DYehGAqKjPkOgBlWhAgLUVIVeQBMBB+QEKiMqqAqKheQ20AJipNtC+oBslQAFS1MfYUCoEHLzAtQmCdAKE9R0IE7LuNhRloEINwOQEzBSgYsoYIlMFWKgj1CAbhBEgi7EZkBMyFAEqygAHoNwOpAZgDYQA2AJTA3qA0KBCVYLQlABVWgQCYMxUAIORCgCAtBTxAiehVoKBAMwgAACQAgWpSbgNihAAh4AASpKupaCmbAIoADMV+UAJAwPAIQbF3AABalW4EBkYsBmSpaD0AgzDHMAGVEaAVCBeoDYEKAzoTMuQe4EzKtAUCEqUIBsMwACYbIVAABsAqyZ0DLQCBVCLQCApEAzAyADcudCLUqIgQIu5CUwqqAAgqMxuGAGVQggLkMqDkXmBHQbh1AESQZlsRgY5FFPMJAGPXQeooABcyAC5AATIZF9SAMqk2LQAMgFsADLQhUBAUjAKgCAFJkCrQCbgAJgBPUoBaALzCCBgAAgEABNxyC11AuQHUbgReRQkAmBDIvUm4JMgAEGQyG4AAoAgGfxDqAyGQAFyqR0GdRmAdBkBsAdBkQq+oAgF1ADIZD11CWYAAIBkGPUMAqAIq0Agy5hjqABC0AZDIKtAEwZBaALzCAUyA2AAdQAIikQFyCoC5gRFy+MIdQIAQCjICgEWpWTcBO6iuZSbhBUpAgLUpiVAXYN5kqRvMC7k3FQq1AdCkLUCkFSMAUgAvQEFdAKsiPT+kCoDoBmHzCYC5U0IwtAhftqRlqYtgVAJkAvQL6iVKtGA33H21AqAH21BHUCoJ5aDYgFH21AAIbkKwHQm/wDSUAAtvaSormBUXkRMtcgBNwAA6AbMAPtqTcuwFIVEAP6gtP6QGA3/AKQAA3DKtSMAAAIy+wjAFQCIwKEKgBX7VAqADAYYDb+kq8iJlAjBWRaABWn/APQGAQCFQH21GwAAbf0hj2AToXciKACAQAqBALXMj06gAKkqXMgBfbMpCoBuRl3IwKN9A6krmBegFdAABakdQFciVGY6gPUdCgCdC9CFAVoPUehAL0C8gPQJ3ByDIEKmSoqAL0HQcgAYzDqEBKupHUoAiHQqG4AL7ZlIAL6kAFJ0YLUCBeRSIC9ACAA/IcxuBSb6blJuEoOQfmAgRfUiKgKh0CG4DoNtAQJ3GAgDdV5Aegb8QgfkBUVyAP7ZjoNwADC1ADoOgKBi68gHsUCLyYZfQgBMdAAGwRCgH5BgrAi8i9BmEAKgSoDcbAAECIoDmSu5SZgOhdgEAXkAABCkAtQQqAIFIACKRACZlJmA3DJUMDIhWTcCgAAAwmAG4DAAFTzAgDYr8gAAgFAqUCAMAAKhAEAAKyBsIAAGBUQqIwAFQmApmCkqAVaAbCoAIBAKBlIwIULcABuBXMJ2QUL7RkEMUjInQoBFWpAnmABeRNgCAFcgAAYAEr8grkBdwRvMtQACDAoJUtdQIxQFAgegDYBAIeIEAqEAKQMCoIIKgFZBUIAAAACYYAcx0GwDcbAAARMoAhkyICFRGXYAgTRlr4AGBUABnQABQjWQrmVsATcormAArkSoBgMAXYMlRv1AoFRUCMqJUoDoHqAwAFRUAycitk9oDoCpiuQAPQCoBkRWACAJUCrQjKmR6hMARCrToEKAMwIV+QI2AKiJhPICigqKgEV6kqKgVmO5akzqEgJUtQgKiJlqAIiioDoOgqKgTcABMLsRgMIR7ApABRv1AFRBXMgFAABgABUAVApNy1IEgQrkK+AQEYFQKioxMkAehEUlQG5SFTAIUCYqAAqK6gGRlqAIi9CJlTzAAtSVAPyDFQBjuMi7gJAAEAAqADJzD0AyGXIOhALUlfAtCUAOgyAAZU0FVXTcUyFAFUKrkKF3AmRaoUADLkHTkEADGQoOgDKhMvjDIBkqZjKmhPIbAV0GVAKAMuQyqGAHQOldBQgBeRciAJ2VUCa+IU5D2BBlUOgADKmgyAQEVOQTVNBQIC5VIGAFUWq5GJlzAKnIZVA3AmVBkWg3IEVOWwFC0yJCqJVAbgMuRVQUAEqh7QPaBcuQdORAHpXTw0MU0XcjDytc9BkSmZQGQLuSgFy5DIUADLlsKrkGicwLkMqaBB6AFTkMq6D2gBlTQOhNigToOhWiMAqciqlNCUKgLlQmQASmXIqpyAaBuZfEHQLRigQZchVchzADLkKoCgETRU0RIAXIZVIALXw2FVyG4AZctw6CnyhAMssgqfECAXImWZQBMiOhSNAUAu4EHIo6gQFY1Amw3MvUgEBkiaAQuwexQMQUMABUVAjJyK2KgRAoAAqABkKZJA3YkMmmtiNUY2GKCMibhO4ggKhA2KgVCYAF9siglF9QKEEIyGTAGJQNwA3LtkGwJ7SobCoBAVKBiNyhgQqAAg9peQ9oGIMgBiNihaATqUvqAIQoAVL6gAQoCAVDYDYEJXMoABCuRQJsH5lAShUEEEA6gJgQBalAi+oVKAIBUqYEC06FAE9QXYeoTsgWRakTCBFAAIbCuYbyAhEWugQE3BkYgEVhPyGwAblIAqAEgBSAC7AbbACkBV0AxLVh+RAKRgAKkqUARhEiahhbiiUKSq28kkdD8Xe0ZceHY5104Olyb8vOGsEdpcX/AIWRF5rOY/BUXiXOLh3suvoWad5U7t6i1G9Uu8bfbrHd9kmWy32qz2SzSl3pk6dMUEEC5tuiR1HjPtHYAuOsm547RiK0p0assPckr/WRLP8AipmpGNMaYoxnbvuzEt8Wi3RVrBKb7smX+9gXvV6V8T0NTc8HhK3ERVk1bz3R/LC3tVqnqtxs79xJ2o8YWqKOC5Lnuu7Jbr3XH3p0xdXRfEcAvXjTxQvJv3XF9skwv72zQQSkvRVOv2ya7M2Czo+FZjam3Hz+ayqy71fOqXJJvEDHc2LvTMZ3+3+PRr6zyrFxP4iWNpyMaX1lp37Q4/lqcPbXMzh8C4nDxpjboR+kKflrvfLta5e0LxSu2OH3S+bPeEC1gtVmhir1VGdi4c7V81RQwYkwlBHDVd6dd9opElz7keT/ACkayNmFSyyNCwb3O3EeHUq282/Ryl9AsCcYOH+M3BJum/pMm2x/+jti9wnV5JRZRfxWznmfI+YcNKrwO1+F/HTGeC45VktNpivy6IXR2S2RtxwQ/wCLm5uHydV4I1vO4SqpjpY1W/slkbGqxM7XI97eWoOGcMuJeFeIFg91uS2921wQ1n2Ge1DPk+cP3y/dKqOZGoXbNyzXNFcbTHZLL0V01xvTO8LXyC0HqEU3pSkQazAVCIVAVgnIABuEeBiO9rNcVw2++bZDMis9hs8domqWqxOGFVdE98iaYmqYiCZiI3l5/oKnQsXanwHSsFz4iiqv73lL/wDIYLtUYHrncWIv5GT/ANwyf9Fz+fkpWv02x+J38NDoRdqjAdM7lxH/ACEr/uF/9qnAb/uLiP8AkZP/AHCP6Lnf2pPptj8TvrcHQcXaowN97cWI4v8AVSV/+Qxh7VGC287gxAlz7kn+eT/RM/8AtSfTbH4nfyB0vdXaY4aWuJQ2qZe9313n2JxJdYHEdj4TxzhDFkCiw7iK77wiaq5UualNXnA6RL0LW9g5NiN7lExHgq0ZFuvzaochJUrRC1VUqVMh6HiPe9sw/wAPsQ37d6lu13fds+0yFMh70LjggcSqt1VHqiia6opjtJmIjdyEj8maYQdqHiT3VWy4eiy/vOP/ALhJvag4ktZWfD8P+xRv5ZhsH1Xz+e0fqx/9Tsb7dbc5hHoeHN622/sAYfvu8vcvuy33dJtM73KDuwd+OFROiq6LM9/Q16qmaKppnsZCJiY3UAECegKGBAugAFA5DmA9CMu5AJ7QikSApSUKBH0Kw9AwkHoT0HQINiohUBCkAFIxmKeAEKgVAAAwIiMoAgRaBLUBUoAEqGwQAECroA2JzMuRAIUgYGQo+QqhuAo+QFUAI/IvQg5gFpoXcmwYFCFRUA0GRsNgMyFIwGx4173jYLouu0Xpelrk2SxWaBzJ06bFSGCFbv2bn62q0SLHZZtqtU6XJkSYHMmzJkVIYIUqtt8kjSHtC8XLZxCveK7bujjs+GrJMf3PKrR2mJfs0a+bDsvEymlaXc1C70KeqmOc9y1ysqnHp3nm83j3xyvLHE2dceHZk+7sOQtwx0fdm23xjp8GD9x68jpdaUWRGe3whhu+8W37JuXD93zbdbZuahgVIYId4o4nlDCubOl4+Nj6fZ6NH2YjnLXLly5kVdfXL1aZ2Dw94Q46xxBDaLqur7msET/89bW5UmnOHLvR/wAVNeJshwh7PmGsJwSryxJDJv8AvpUi/RIK2azxcoIH8J/uouiR3XDSGBQpKFJUSSySNa1Dizo70Yse+f2hksfSt+u5Pua94S7LeG7JLlzcTX3br0n09/Ls1LPKr4axP1R2JdnBXhddyhUnB13zol99anFPb/LbR2CRs1W/quZfneu5P67fJlLeLao5Uw49JwHgaVAoJeC8NwwpaK7JP808K8uGHDq8IHDacEXA67y7FBLfrBRnLnTMVLSnIu0zvFU7+Mqk26J7HTuI+zhw1vSGOKxWS33RMaydltUUUK/ix946lxl2W8S2JTJ+F75sd7QKrVntK9wmvwTzhb82jbsKhksfXc2xPVXM+PWt7mDZr/47eD5qYhuG+sN3jFd1/XVa7ttcP7FaJbgbXNbRLxVUevrVH0kxdhi4MW3PHdWIbrs94WWLNQzFnA/woYlnC/FM1M419n29cJy59+YVin3vcsNY5slqtpssPNpfDhX4SzW63Ny0riWzkzFu99mr4T/DE5Om1W/tUdcOlrrvG33TeMm8bstk+x2yzxKOVPkxuGKB+DRuF2f+Otjxl7jhzFMcmxYip3ZM7KGVbvL8GZ+50e3JaZJ18iqKOCOGZLjigjhaihihdHC1mmmtGZLVNJs6hb2q6quyf+9i3xsmuxV1cu59Pa+QOr+zFijE+LuGMi8MT2WNTpM1yLPbY8nbpUP7I1zTrC3pE1Xmdo02OVZFmqxdqt1c4nZs9uuK6Yq72MUUMELijiUMKVW26JI65tvHXhRZL0d3TcY2VzVH3IpkqROmSYX4zYYXBTxrQ6W7VXGdXjMtGA8J2ytigbgvS2yov06JayYGvvV9893lpWutiyNl0zhv6Ra8rfmY35QxuRqPQq6NEbvp9YrVZrdZJVrsdok2mzzoFHKmyo1FBHC9Gmsmj9Ua39ha/LfargxFh60TY5lju6bJn2VROvufuvfUcK5KsCipzb5myJr+diTiZFVmZ32X9i75WiK47TMZkLUtNlUXkcU4x1/OoxV/BNo+jZytanFOMn60+K6f4JtP0bK+L6ajxj5vF30c+D53S0u5D+9QYl/Ah8kZNHaY64abPN+bWREdrYK4DY3xhhix4iumO64bFa4YopXu9pcMVFE4XVd10zTPcPsw8SU/024/98f80xVesYVFU0zcjePauqcS9VG8UulFoU7tXZi4ir/1Fx/74/5plN7MfEWGX3oJ1yzIvwVa2n8cIjW8D+7H6onCv/hl0ikZypkyTOgnSZkcubA6wRwROGKF801mjnGL+EXELC1mmWq88NWuOyy/hT7NSfBCub7jbS8WjgidVVOpf2cixkU726oqj9VGu3Xbn7UbO5+F/aHxdhedJseIY48Q3SqKJTov/Ey4ecMx/C8oq+aNu8GYnuTGGH5F+3BbYbXY52VVlFBEtYI4dYYlumfNxnYXAfiRbOHWMpVpjmzIrltkcMu8rOs04NFNS/Dh18VVbmsa3w9avUTdx42qjsjlP+WSwtQqpmKbk7w35ocR42frO4x/gS1fRs5bJmyp9nl2iRMhmSpsCjlxwuqihaqmnumjiPGz9Z7GP8C2r6Nmh40ffUR7Y+bPXZ+xPg+eEOi8iRaFWgeh2bsadHN9FuEcPc4U4Sh5XLZfooTlWxxrhY1+dhhT+BrJ9FCckqqHF7/pavGW4W/NgA3BSexB+QQbECMCoEhmX2EegQDPkB0ABfbMb6BEqBltoDGviZIB6h1D0IwAFfMKgEdSoMLYSAAALyHQL6huAoyqtNCBAUmbRKlQBVLuRFAUJQtfMgFYRK/IUB0I/Iu4yEB0Yp4DICQIWpKgTOuhX5DcgGVSbgm4TstS9CUKwgqT7ahkWoF20DeYIwMqkrmABG/AvQlBQJ3UA4lxdxlZ8B4AvLEc5QxT5UHudklP9lnx5QQ+Vc34Qs92rdV2uKKY656oeaqopiapdDdsHibNmWx8PLltHdkSlDHe8yCL4cesMjySpFFzbS2ZrS2z9bda7Tb7bPt1snxT7TaJkU2dNidXHHE6tvzbJZLPPttrkWOySY51onzIZUqVAqxRxxOihXi20db03Ct4GNFuOznLVMi9Vfuby97w7wZfWPMT2e4LkkpzY/fzp8Sfudnlp5zI3yXLVvJG+HC/AOH+HuHYLpuSzr3SJKK1WuNL3W0zKfCifLlDoker4F8OLJw5wbLsLUEy97UoZ15WiFfCmUygT/AhrRc83udgcjQdc1irOudCifsR8fbLO4WHFmnpTzky5IVG4pka+vwbAMDFlXkGKAOmYeugDQEG7CAGsXab4IyYbPasb4NsagcCc287vlQ5NaxTpaWlNYoVtmt0+rez3wutPEfEyjtcMcrD9hihit0+HL3R6qTA/wAKLd7LPVo3ufqeDhq47pw7diu247vs9gsimRzPcZMNIe9FE4on1b+RaI2GxxHkWsScfnPKJ7o/7yWFen26rvT7O551gslmsFhk2KxyJdns0iXDLlSpcNIZcCVFClskjXztW8YYbjsU/BGGLX/72tEHdvC0yos7LLa/S01+yRLX8FeLVORdpTjDJwHdbuK450EzEtsl1hpRqxS3+yRL8J/ew9Xks9JbVOnWmfMtFomxzZs2NxzJkcTcUcTdW23q29y84f0SciYyb8fZ7I71HOzIo+7o5vHoeVdF322970s113bZplqtlqmwypEmWqxRxN0SR48MMUcaglwxRRRNQwwwqrbeyW5un2ZeDkGCbugxNiCRDHiS1yveS4lX7hlxfeL9218J7aLeu0arqVvAtdKec8oY7Gx6r9W0cnMeBHDqz8N8Ey7sijgnXpaolPvGfDpHNpRQQ/uYVkueb3Of5Ahy69ervXJuVzvMtkooiimKY5QdC1IUpvQvI4rxhz4VYqVP7kWn6NnKeZxfi6q8LcU/wTafo2Vsb01HjDxd8yfB87Jf6XD5ItSQL9Dh8kU7TTyabPNvj2XI+9wKw4vwZc1f82I7LiOq+ynH3uB1xr8GKcv+ZEdpt1ZxzUo2y7n5p+bcMad7VPghamJdyzVWSfI107VPCG77TcVqx1hmwQWa8LL+i3lZ5ENIbTK++mqFZKOHV0+Eq1zSNiluY2iRLtVmm2SfCopU6CKXHC91EqP5S7wcy5h3ou255fGFG/Zpu0TTU+YuoelD2GJbvd1YivG7H/6S1TJP5MTX1HrqHYKK/KUxV3tRqjozs3d7ImJpl/cJJFhtM1zLTc1oisUTbz9zoo5f/DFT+Kcw43OnBzGVF/cS1fRs6O7CdoiU/Ftj+8cFknLzTmw/I0d48bF/Yexh/Atq+jZy/UbEWNUmiOXSif12ls+PX08aJnufPOFZLyD0YWiD0OqbfZax2vovwrdeF+FP4Fsn0UJyU4xwo/Wvwp/A1l+ihOTHFb/pavGW5W/NhalIjiPEziPhfh9dsNqv62P3ecn9zWOSlHPn0/Bh2XOJ0SItWq7tcUW43meyCqumiN5lzAxieRpvjrtMY1vebHJw5Is2H7I8oYoUp1oa8Y4l3V0h6nVN9Y2xlfMxx3piq+7U3rDHbZih/JTSNlxuE8u5G9yYp+Msdc1S1TO1Mbvo37pBp34fUzTVMmfMd223uLvO322vP7oj9p7a5sZYvuabDMuzFN9WZp1Shtkbh9G2i5q4PubfZuR+inGrU9sPpIyGnmBu09iy6o5UnE9jkX7ZFlHHDCpNoS5qJe9i8ms+aNmuHHEDC/EC6XeGHLf7q5dFaLNMXcn2dvaODbwaqnszX87SMrBn7ynq745L+xl273mz1uUkehQzGrhj0FVXY687ROK75wVwxtN/XDPlybdLtUiXDHHLUxUiio8nloayvtHcU3/dawrysEv2GXwNEyM635S1ttvt1ytL+bbsVdGpu8ZLyNR+EXHTiHiDiZh+473vOzTbBbbWpU6GGyS4G4XC8qpVWdDZ/FmI7mwpcM++7/t8qxWGQvfRx6xRbQwrWKJ7JZlDN0y/h3YtV9czy263uxlUXqJrjlD3O2hjG0tfjNPeI/aXxPfE+bZcISVcV31cMM+JKO1TFzbdYYPJVfidOXtivFF7TnNvLEd8WqOLX3S2zGvStDNYnCeVdp6VyqKfZzlaXdUt0ztTG76RqKF6NMyTPmjZr7vuyxqOzX3esiJaRS7bMTX/ABHYGCOOvETDM+D3S+I75sifv7NeH6J3lyUfwofUq3+EMiine3XE/B5o1W3M7VRs3ur4A4Fwi4o4c4j3ZFNuyOKy3jIhTtV3zol7pK/dL8KD90utGc9WZqt6zcsVzRcjaYZOiumuN6Z6hDoF1DKT0ewVHsKA6CuRAAC0AAIpCgCVLQgD2CqLQgFWpH9sxuEAKn4EIgMqk5gAA2CMCsCo3AoAAjQKyANgKlrmE7HsC1DCCB0IVj0AhqN218XR2/FN34Os82tmuyWrTaYU8nPjXva/vYPns24mRwy4Io42lDCm4m9ktT5vY8vyPE2Nb5xBHMcat1tmzZbf7X3qQL8lI2jhXEi9lTcq5Ux8ZY3VLs0WopjteizNguxhgWG98UWvGt4Se/ZrofuViUSyitESzj/iQv1iT2NfY6QwNvRZs+hfAzC39R/Cy47nmSlLtf3OrRa1TP3aZ7+JPyqoehsXE+bOPi+Tpnrr6vd2sfptnylzpTyhzfYxK2SpzVsR6jkSuYrkBRUiY2CYVgNioQDxHoGwAoK6BdAI0dace+Kth4bYdSkKXab+tsLVhszeUOzmx/uU9t3lzpybidja58A4UtF+3vMT7qcFms6ipHaJtMoIfreyzPn9jfE17YwxNbMQ31P91tdqjq0vgy4fvYIVtClkvazYdB0ac655S5H2I+PsY/OzPI09GnnLwL4vK33xetqvW9LVMtdttUxzZ86Y6xRxPV/0baHhvxI+aNheytwejxBa5OOMUWNO5pEXesFlmw/+bmJ/pjX7XC1l+E1yWe+5uZZwLHTq6ojlHf7GEs2qr9e0OT9lXg1DYZVnx7iux1tkxKZdVjmw/pMO0+JP75/ep6LPVqmytXUUByzNzbmZem7c/wDjZrNmm1T0aRaBBMJlpCqvQCuYqSIcX4utQ8LcUP8AzTafo2cpWZxTjD+tVin+CbR9Gytjemo8YeLvmT4PnfA/ew/vUCQfBh8kXc7PS02ebebsoJw8Drm8Y57/AOZEdqHV3ZVp+cbcT8Z30sR2jTyOP6n65d/NPzbdi+hp8ELuECxV2SJMmy5EqOfNiUMuXC44m9klVsqyOiu1PxWsOHsOWvB9zWuGbf14S/cp/ucVfuOTEvfOJrSOJZKHXOvKt1h4lzLvU2rcc1K9dpt0TVLUvFtvV64pvS81pa7ZNnLyijbXynrCVGuR2Si3FuiKY7GoVT0p3bQdhKyRNYutzXvf/CyIXzf6JE/q9Tu3jcq8HMY/wLavo2cX7KWGZmGuEVjm2iBw2m95sV4zE1moY0oZa/IhhfU5XxlXe4RYvh/zLa/oojlmo5EX9Uqrjl0o+G0NosW+hixE9z53Q1oZMqWQeh1X/i1bfrfRThT+tdhT+BrJ9FCclOL8JG3wrwm/8zWX6JHKVqcUyPS1eM/NuVvzYcY4o4xsOBMFW3EVthU2KUlBZpFaOfOiygg8q5t7JNmgeLcRXtinEFqv2+7XFabdaYu9HE9IVtBCvvYVokd6duDEUydia5MLS5v6BZLK7bOhT1mTInDDXyhgi/KNcnEdE4X06izjRkTH2qvhDX9Sv1V3OhHKGb5nvMMYMxZiiGKPDuHLzvOXDF3YpsiQ3LT5d90hr4VOVdnTAknH/EKVYLfDFFddilu1W5J078CaSgrt3oml5VN7rDZbLYbHJsVis0mzWaRAoJUmVAoYIIVoklkkTrXEP0GvyNqnert35QYWB5eOlVO0NB5vBTirKg78WCLxappDHKifoo6nDr9uO+rhtn3Hfl0W+7J70gtUiKW4vKqz6H0vdGtD1mKcPXNii5Z9z37YJNusU6GkUuZDnC/woXrDEtms0YXH4vvRXHlaImPYva9Ko2+zPW+adD3WC8T3zg7Edlv+4rU5Frs7zVfeTYN5ca3he666o87ihhK04Jx3emG7RFFMhss2sia1T3WTFnBF50yfimcYN2mbWXZiedNUfNhd6rVfth9HuHWK7BjbBt3Ylu73sq1y6xy26uTMTpHLfiok14qj3OQM1f7DV/zVFiDDEybWUlBbpEDfwYvgR080oPQ2fbOUanifRMmu12Ry8G0413y1uKnTPbIipwVnw/hXjZ18bZpOjdXtl58HP/qciv8AxGlhu/Ckf6OfH+GE1Wfvvc5VwhvCzXRxOw3elsnQSLNZbwlzJ0yN0hggVat9D2nHTiNb+I2MJtqcyZLuayROXdtlbyhg/bIl+HFq+SotjgcDplsdm8B+E9s4kX3Mn2iZMslw2ONK12mFe+ji1UqXX75rNv71PxRl8ujGx6/pt7/jG3/z2rWxVcrjyVHa6xlwRTJily4Io43pDCqt9EfvPu68ZEv3Wfd9slS1rFHIjS9Wj6K4UwhhjCtjl2W4LksVhhlw93vwSk5sXjFG/fRPzZ71ttNRe+TWjzTNbr4y2q+xa6vbP+GSp0nq+1U+YkLharC014MG83GXgrhnGtzWi0XZd9kurEMELjs9rkS1LhmxfgTVDRRJ6d7Va809HrZZrRYrZPsVrkxSLTZ5kUqdKj1gjhdGn1RsWl6xZ1GiZpjaY5wsMrEqx56+uJedhe/71wxf9kv25bVFZbdZI+/LjWjW8MS3haya3R9CuG+KrFjXBV2YlscKghtkqsyVWvuU1ZRwdGn0ofOHc2l7DF/TZthxHhmbMbgs0yVbZCb0UdYYqdYYfUwvFOFTcseXiPtU/Jd6ZemK+hPKWzQIVM57DPCKRPUNkgBUIAwBUCLqUCoAgqEBdwCoCbgVzDAEFdAgAqAEwVzI9BXMMIBuWgp5AKCnUF2AjJTMtBQCUdBuKF3AUFCkAAhQOL8Wryjubhfie85cXdmyLrtDlP8AduBqH42j51qFQQwwrSFURvr2mJjlcDcTRJ0cUmVB6zoEaFvM3/g+3EWLlffO3wYHV6vvKafY5Bw0uhX/AMQ8PXNHD3pdqvGTDMXOBRKKL4kz6N1rmaJdliyq1ccrjbhqpEM6d5Ulv2m9pjOLrnSyqaO6FzpNO1uZ9o/MhcyGpsqCg3FPIAkKFoNtgncoAwEFBQFoBEj1GMcS3RhHDlrv+/LUrPY7NDWJ/fRxbQQreJvJI9hedusV13daLxvG0yrNZLPLcydOmRUhghWrZot2geJ9p4j4npZopsm4LDE4bBZ4su+950a/Ce3JZczLaRpVeoXujypjnK0y8qmxR7XouLnEC9uImKpt8Xi4pNml1gsVjUVYLPKrkvGJ6t7vwSOGaGTr6nO+CnDW9OJWKYbBIcdnuuzNR3hbaZSoNoYdnHFsvNvJHTKqrGBj9fVTTDXY6eRc75l73s58I5/EW/PzTvWCOVhmwzF90R6O1RrP3GB/Oeyy1eW8lmk2ey2WVZrLJlyLPJgUuXLlwpQwQpUSSWiSPDw1cd14cuGyXHctlgstgsktS5MuHZc2923m3q2zqjtJcX5GBrqjuC4p8EeJbXLyao1Ypb/ZIv3b+9hf755LPnGXlX9Zy4pojq7I7obBat0YdreXOYOIFwzuJsOALLN+6Lzhsky02iKBruSHD3aS3zjairTZLPU5ZQ0L7O18TrHxzw9bJ06ZHHa7THJnRxxNxRubDEm23q22nV7m+i0Kes6ZGn3abcTvvET73rDyJv0zVPelCUMqZEpkYhdmZKP4i7gAkcT4yZcKMVU/wTaPo2csOKcZE3wnxUl/gi0/Rsr4vpqPGPm8XfMnwfO6B+8h8kVswl/pcH71GTOz08mnTzd68K+0MsCYFsGF1g780fuRxt2h3l7l3u9G4vg+5xU1pqcpXa1hevD5ryvj/wDxNX4nsTQwl3h/Au1zXXR1zO89c/yvaM69TTERLaSHtaSd+H8a/wDq6/7J+do7WUTlxfc2AYYZn3rm3tWFeaUrP1NX6+JW/Epxw3p34PjP8vX9QyO93FjPtF8Q7+s8yy2O02S47PMh7rVgltTKf6SJuJdKHT8yZMmzY506ZHNmzInFHHHE3FE3q23qz8+8i13ZlMXDx8SNrVMQt7l25d86d2fkdgcB+Hlr4iY2k2KKXHDdFkiU68p6yUMuuUtP8KPReFXsedwi4LYrx9OlWqKzzLpuNusd4WmBrvw/4mF5xvx+D47G6WBcI3FgnDsi4sP2NWeyy/fRxPOZOj3jji++if8AQqIweu6/bsUTZsTvXPd2f5XuFgVV1RXXHU91Kly5MmCTJghlypcKggghVFDClRJeFDi3GN04SYvr/gW1fRRHLNziPGdtcIMYv/Mlr+iiOfY/Xepme+Pmzt2Pu5j2PnpsYxfBYTbRItGdp3+y03brfRfhXClwvwol/gayfRQnJDjXCp14X4Tf+ZbJ9FCcmOKX/S1eLc7cfZh1dxA4GYLxrim1YkviO9fu20wy4Y/crX3YEoIFCklTLJerZxyPswcO4llPvyHyti+uE70bOKY34hYLwZD3cR4gslkntVhs0L90nxf6uGsVPFpIvsbUM/aLVmur2RG6jcs2Y3qriHrOEfCvDfDabeM24ptvnTLwhlwzYrVNUfdUDiaUNEqfCdfJHPzXi/8AtU4Wss1y7lw9el4Jfsk6OCRC+nvmcRvHtXYgmRNXfhO7ZEOznWiON/EkXlWi6nlVeUrp6575hRjMx7UbRLbamRUabR9qXHr+Bc+H4POVNf8A1mD7UfEV/Bu7Di87LNf/AOQ9xwvqE9kfq8TqdiH7dteRBL4nWCcoUopt2Qd586RxJHQ9Dl3E7H19cQ78kXvflnsMm0SZCkQqyQRQwuFNurUUUWeZxShv2l41zHxKLVznEMBk3IuXaqqeUu6+xbMcHFu1S1pMuqbXpHAzcxmmXYyVOME3+Cp3zoDc5mh8Uxtnz4Qz+lzvY97pntjL+wvNfK8bP8rNKTdXtlNw8F41zvKz/KzSqpsvCfqU+P8ADHar6b3P0kSps+fLs8iHvTZscMuCHnE3RL1Z9FOGmFbHgzA114cscMNLLJTnRpZzZzzjjfi4m+lFsfPvCD/ruuXKq/NGz/SQn0oXwUY7jC7V93b7OuVxpFMfaq7WLQYe4NIZoo2suRon2prrl3XxuvxSoVDBavcrTRc44E2/WpvbCaT9seGnGq0PnYbO/wDhNl4VqmM3bvif2Y3VI+597pqh332IImuJ98wJ5R3LE2vKfK9p0IjvjsRVfFW9eX5iTPppJuGvx/4+54fvDEYMz5eluLQoaFDlTaUSdGKZlp5CgAtCehdgIFoUmwAUHoPQBTIiLTyCAtBQb7AARooAlCJFAEoGmUuwGNAy7hgWpNwNwLXQq2McigGxkHQmQFrkK5kJuBlVBvNEqigRlqTIcgOsu1I6cCsQ+P3Ov+fAaIN5m9vaor+cVf1Pw7P9NAaIxHQ+EfVavzftDXtV9NHg7d7IVPz7rD+J2j5pvCjRvsiNLjdd9d7LPX/AbxowXFfr3uhkNL9D71y5EyEcUEEEUccShhhTcTboklqzV7iJ2orXZ75nWLBNzWGfY5MbgVut/fi93p99DBC4e7DybbbWdEYbC0+/m1TTZp325ru9kUWY3rltCtRsdG8Du0BYMa3nLw/iSxWe575mulmjlTG7Pan+Cu9nBHyhbddnXI7yKWVh3sS55O9G0vVq9Tdp6VMqSoGRbqithNEyCoBUSdMlyZUc6bHDLlwQuKOOJ0UKWbbb0RVrQ1P7UXGX81rRaMDYWtNbtlxdy8rZLiytMSecqBr7xP4T++apos7/AE7T7mfei1R757oUMnIpsUdKXH+0jxfmY4vCPD1wzooMOWWZnGsvu2Yvv3+4X3q31e1OlWZOjPZYWuC9cUYgstxXLZYrTbrVH3YIVolvFE9oUs2zqmPjWNOx+jT1Uxzlq9y5XkV7zzl5vDbBV849xVZ7gueXRxe/tFoiXvLNKTzji+pbvI35wBhO5sE4YsuH7jke52eSqxxte/nzH8KZG94n8WiyR6jhBw9urh1hWXdVh7s62TaTLfbHDSK0TKfFCtIVsvFs5ojnOuaxVn3OjT1URy9vtbDhYkWKd55y9Vjm8rXc2Cb8vewQwR2uxXfPtEmGP4LjggcSr4VR84b0t9tvS8rReV42qba7ZaZjmzp8x1imRPNt/bI+ivEyJS+GuKI3pDc9rb/koj5xJe9XkZvg6imYu1bdfUstXqnemHIOGM2KTxLwvNhbThviy6f6WFH0dh5Hzk4XyHaOJuF5UKzd72b4pkL+o+jKLfi/09vwVdJ8ypnUhCo09lgtSAdgpxXjAq8K8VL/ADTafo2cpOL8W/1rsUV/wTafo2Vsb01HjDxd8yfB86JP6XD+9RnmYS8pcPkjKp2enk0+rm7s4X9nu04+wRY8TSMXSruVpimQ/c8d3ub3e7E4fhKYq6cjkn/slXp/8+2Lrdcf/dO1OyO/7B10f6a0fSxHbPU5zna9nWsmuimvqiZiOqP4bHYwrNVumZp7GqP/ALJV61/V5Yf/ALXH/wBwsPZLvL77HlkS8Lri/wC4bWdSlp9YtQ/H8IVfoNj8LW65uyfcUqbDFe+LbytcKWcFns8ElV824mdnYO4KcNsLxS59kw7JtlqgacM+3v7oiT5pRe9T8UjsIpZ39VzL8bV3J2/T5KlGLZoneKVyWiSpkg2QGPXEBxPjNnwhxjRf3Etf0URyzc4pxjdOEWMH/mS1/RRFbH9LT4w8XfMnwfO9aIkfwX5BaIRfBZ2jsabHN9F+Fipwvwmv8yWP6GE5IszjfCn9a/Cen9pbJ9DCOKWIHhTh5fuIIGlNsljjjk1/bH72D/iaOMV25uX5ojnM7fFuMVRTb3nudJdpPjjarpt9owdgu1e5WuVWC8LxgzcmLeVL/dLeLbRZ1a1XnxzJ8+ZaJ8yObOmxOKZMjicUccT1bbzb8WJ02bPnRzp0cUybMiccccTq4onm2/Gpg2dW03TLODZimiOvtntlq2RkV3q95lhHSFVrQ/e77vvC8I1Bd9gtVrieikSYo38SNq+zzwHuD+p2xYrxpYYLyt1sghn2awzs5NnlvOFxQ/fxtZ55KqVK5mw9gsVjsEmGTYbJZ7LLhVFBJlwwJLySRgc7im3YuTRap6W3bv1L+xplVdMVVTtu+eFn4d8QJ0KilYHxJHC9GrsnU+aeVK4YcRotMC4jXnd8xfKj6Hd50Wb9QqVqY364ZPZRT8Vx/Sbf4pfOLFGEcS4Vis0OIrntN2RWmFxSYZ6ScSTo3ROq6npWjYDtv2yGZj25bGok4pN2d+JJ6d6ZFT5DX6vM3TTMuvLxaL1cbTLB5VqLV2aKex3T2NMuMEz+Cp3zoDc1s007GK73Fy0PldU758BuXEaFxT6/PhDYNK9A6U7ZzrwcS/zlZ/8AqNLdzdDtn5cHV/Cdn/6jS82XhX1KfGf2Y3VPTe57XBtHjC40/wDCVn+khPpQsoYfI+a2EX3cX3I+V42f6SE+lUPwV5GI4v8ASW/CV3pHm1I9AV0I0jTWYVGk/bKp+fPP/ELP81m7CNI+2NFXjVa1ysVn+YjZOFfX/dP7MdqnoPe6c2Z332H1Xihe8XK5I/ppR0DU7/7DTrxLvr+BYvppRuXEH+33PD94YjBj7+luHUVWxjVFOUtoVMMxFUBWxUgQGVRUgAJlrqRFyAV8NiZDIAWuYqQZAWoZMqjIC1CZABQwFQA9SMZcwwFSFJkBcgQoAuRCgTKhGXXclFUJ3CoUQWoQFovAhcgOsu1LDXgViHw+53/z4DQ6LI3z7UVPzicR+Uj6eA0Mj1Oh8Iz/AKWr837Q17VfTR4O1+yTV8b7r/F5/wAxm860NGOyOq8cLs/F5/zGbzowXFc/62PCGQ0v0M+L1GObFbbxwTfl33d/5203fPkyM6VjigaSPm3aZM6zWibZrRKjkzpMbgmS41SKCJZNNbNM+ntctToHtMcFIcVSp2LsKWaGG/pcHetdkgSSt0KWq5TUvytNaEcN6pbw7s27nVTV29xqGLVepiqnnDT+CKKCOGOCKKCKFqKGKF0cLWjTWjNvOzTxshxHKk4RxfbIIb7gSgsVrmOitq2hif7av+LzrXUJwxQxOGKGKGKFuGKGJUaa1TWzEEcUEcMyCOKCKFpwxQujha0aezN41PTLOo2ejVz7J/72MNj5FePXvHvh9PWiM6E7NHGuHFMmRhHFdoSv+XDSyWqN0VuhS0f+NSX8ZZ61O+mcszMO7h3ZtXY2mPi2azepu09KlRsQ6e7RvF2TgK6Xc1zTYJmJLZLrL0asct/ssS/C/Bh6vJZ+MXFuZV2LVuN5ku3abVM1VcnGu1Vxf/MayTsDYYtbV5z4e7eVqlRZ2aW1+lQtaRxLX8FeLy1KhdKU0P0tU+dabRMtE+bHOnTYnHMmRxVijibq229W2YSYJk6dLkSZccybMiUEEEENYoom6JJLVtnVNN061p1joRz7Zaxk5FWRXvLzbku28L6vay3TdVlmWu3WuYpUiTAs44n8i3b0STZvPwM4WXbw4w/SP3K1X7a4E7fbEt9fc4OUC+N5vZL0HZs4RQYDux35fkqXMxHbZdGtVYpb/Y0/wn98+iyWfcqNL4g1ucuvyFqfsR8Z/hmNPwvJR06+ZRFQyIawyjhHH63q7uC2Kp/eScd3xyIfFzKQf9R8/Xrkbh9tO/4LBw4sVxwx/o16W2F037kpd5v1cJp1Wp0bhKxNGJVXP/Kfk1zVa+ldiO6HYHZ2sKt3GzDEtrvQy7U5zy07kET+Whv5TI007Ft0u3cVrTeUUNZd3XdMjryijagX/V6G5hgOK70V5sUx2RDI6XRNNnfvYheZSI1hkVDAAhxbi+6cLMU/wTafo2cp31OKcYv1qsU0/wAE2n6NlbF9NR4x83i75k+D52QZwQ+SKSX+lw+SLTM7NTHU0+ebebslJLgbc3jMnv8A5sR2wdVdlCFLgbcmmcU/6WI7Vp4nIdT9cufmltuN6KnwhFQpEXqWKsDIE3AoGQ2AbnFOMdPzo8Xp/wCBLX9FEcry5nFeMn60eMP4Etf0URWx/S0+MPF3zJfOuHRCL4L8hB8FeQi+C/I7T/xaf2vo5wvhUPDLCq/zNZPoYThvavUx8C79Uv8ACkd7y91hqc04ZUfDTC1NPzGsn0MBnxFw7BivAt84fbSjttkjly29FMpWF/lJHHrF2LWXFdXKKt/i2yunpWZiO583UzLvQppxaJ1ZlabPOslqnWW0S4pU+TMilzIIlnDFC6NPyaaPz1Ou7xXT1drVNtpfS3Ctos9pwxddosrhikTLHKiltaOFwKh7I1K7O/Hmx4YuSRhLGSnqwWb3thvCVA5nuMDdfc5kKzcK2iVaLJrKpsDZeKfDu02dT5WMbncDVaxWmGF+jozlOdpWTj3qqZomY36piOptFjKtV0RO7mp496W6x3XdtovK8bTLstks0uKbOnTHSGCFKrbOtMUcfeGdxWaOOC/FetohXvbPd8tzYony72UK82zWbjRxmxBxGiVh9z/Mq4oIlFBYZcfec1rSKbF98+SWS8XmV9P0DKy643pmmntmep4yM+1ap6p3lxnixi6bjfiDe2I43HDJtE3uWWXFrLkQ+9gXg6Kr8Wzi1dyd2P3NzVDF7molC4qZJ8q88n6GPeOmWLVFmiLdHKI2a1cma6pqntd7dipJ8VLc6Zq6Zn0kBuO8zTbsURL89W3Q87pmfSQG5L6HOeJ/X58IbFpnoHSPbRX9h6DP+6cj/qNL1qbp9s5V4OrwvKz/APUaW0zNn4V9SnxljNT9N7nssKfqrub+ELP9JCfSuH4C8j5q4RVcX3Iud5Wf6SE+laXvV5GI4w9Jb8J/Zd6R5tQ/Mjz8RzGRprMKjSHtifr2278Us/0aN3sjSPtjL+zZbXzsdn+YjZeFPX/dP7Mdqnofe6ZaO/Ow22uJ19LncsX00o6Fod9dh9f2T73f+ZY/ppRufEEf+PueEfOGJwZ+/pbiIZEKcnbOKgyFPEZBJQpMioIMiqg3DoAogqVJ6FAELl4E9AGRciNjqAyrqUipXUuXMBREL1QSAeoAyAbhoZcw6DYSg3FQwJTzLTzAAMhWAJsXcmw3AyBBVgBsQqzQHWXak/WJxD/s/wBPAaHvNs3x7Uf6xWIF42f6eA0PazOh8Jeq1fm/aGvar6WPB232QoK8bbBF+DZLQ/8AhN4DSTsf/r1WX8Sn/NN29jA8V+u+6GQ0r0M+Kk9hCmtMls117T3BVXzKtONsI2T/AN6QpzLxsMqH/wA3CtZsC/bFuvvl+6WepadT6fZmtPaf4J/dStOOMHWL/wASqzLzsEmH9NWrnS0vvvwoVrqs613Lh/XvJ7Y2RPV2T3exiM/C6W9yjn2tXbLPn2W0SrTZp0yTPkxqOXMlxOGKCJOqiTWjTNzOzjxml44s0GHMRTJUjEkiD3keUMNugWsSW0a3h31W6Wl8DqjyLHabRYrZJtlktE2z2mRGpkqbKicMcuJOqiTWjRtGq6Xa1GztPVVHKf8AvYxeNlVY9fs7W9nHjihYeG+H0pSl2m/bZC/uGyt5Lb3WPlAnt988udNGr6vK33xetqvW87VMtdttUxzJ86Y6xRxP7abI8jFWIb6xRfk6+r+t8y3W6coYYpkdFRJUSSWSS5Lx5nqqlLRtIo06319dc85/hOZl1ZFXsYxUSNtuy5wchuOyyMbYpsn/AL3nw9+77LNh/wDKS2spkSekyJfkrxbpxvsr8HobdMkY8xVY1FZoWo7psc2HKY08p8af3q+9T1fvtEjalt1zNd4i1zpzOLYnq7Z/ZkdPwto8pX7hZioRKM0xmVLDmYnXXaBx/JwDgG0WmVNhV7W9RWa75dc++176Z5QJ186Lcr49iu/cpt0R1z1PFy5Fumap5Q1j7VWMIMU8U7RZbJOUy77lg+4pLhdYY5idZsa/je9/iHUuhnNiccTiibibdW282zzcN3PbcQX/AGG5LtluZa7dPhkyoVzb18lq/BHXcezbwsam3E9VMf8A1qdyuq9cme2W2fYnw5HdvDy34htEHdm3zbP0J0z9wk1hhfWNzH0R34epwlc1lw5hq7risS/QLBZ4LPA6a91Ub6ur6ntfE5Nn5M5ORXd75bVYt+TtxT3FCUKNi0VAjAAI4rxg/WqxU/8ANFp+jZylM8S+bvsd73Ta7qt8DmWW1yYpM6BROFxQRKjVVpkz3Zrii5TVPZMPNdPSpmHzNk/pcPkjOhvRB2fOE8MChWHp+S/v6b/OC7PvCmv9oJ/+/TfadBp4txI/41fD+WBnSr2++8HZRr+cbcn7+f8ASxHalMz1WEsPXRhW4ZFx3HZnZrBIcTly3McdO86vN56s9qaJmXov367lPKZmWdtUTRbimeyEKAWypCcwtRzAQooBUAcV4xfrSYv/AIEtf0MRyo8O+7usd83NbbovCCKZZLdZ47PPgUThcUEcLhiVVpkz3ar6FdNU9kvNcdKmYh8zofgpeBWsmbyQdnjhRDDT8xLW/O8JvtMY+ztwqelz22HyvCZ7Tof1uw9tujV8P5YGdKu777w5vwub/Oywr/Atk+hhORp0PFuiw2a67qsd2WKBy7LY5EEiTC4quGCCFQwqr1ySzPJ2OeV1RVXNUdss9RG1MQ1o7T3BK2XleNpxxg+yxWi0Tvf3ld8uGscyLedLW7a+FDq9VnU1cigigjigjhcMULaihao090+R9O9zhWOuFGA8aTYrTfdxylbYlnbLNE5M5+bhyi/jJm06TxLONRFq/G9McpjnDGZeneUnpUTtL59woy7i1aRtdevZXuCZMiiu3E942eHaGdJgmU6qh+Fj7K11wzE7Xi62zIN1LssML9W2bHHE2nzG/Sn9JY3+m5G/Jqz3Msjl/DDhpijiFecNmuSxuCxQRJWi8J0LUiSt6v76L9ys/LU2swx2e+HFzxwTrVYbVfEyGjSts+sFf3kNE/J1O1rHZ7NYrJLslis8my2aVD3ZcmTAoIIFyUKySMXncXU9Gacanr75/hd4+kzvvdn3NZO0tw6ubAnAK47uuSGKNWS+YI7VaI0vdLRMjlTIXMip4qFJbJJGr6Z9JMc4UuPGtwxXHiGzTLRYYp0E5wQTYpb70Dqs4czgq7PHCdf/AA/av/uM/wDnFnpPENvGszTf3mZmZ3VsrAquV70bRDoTsWTe5xhny3+yXROp0jls3SZwTBPCTAeDb+V+YdumdZbcpMUnvxWybMXcipVUibWyOdvQwus51vNyfK24mI2iOteYliqzb6NTpjtkwp8GY3TS8bP8rNKXqfR/HGE7kxncf5i3/Z5k+x+6wzu5BNctuKHTNeZwaHs88Kt7ktT87fN9pmNE13HwcebVyJmd9+pZ5mDcvXOlTs0uwUq43uBUf9tLN9LCfSnZHWF38BuF9gt1nt1muKfDaLPNhnSo3bZr7scLTT15o7ObLHXtUtahXRVaiY235q+Di12ImKu1GTmV6EZgF+bGlHbHp+fTaaPWw2ev5Juv4HAsa8IMCYzxBHfuILutNotscEMuKKC1xy13YVRKkLMvomfbwMry1yJmNpjq9y0zbFV+30ae9oClkd+diCF/nk30+VzP6aWd3y+z7wng1w7Ni/fW6d/OORYI4ZYKwTec+8sM3RFYrVPk+4TY3aZkxRQd5RUpFE1qlmbDqnEuNl4tdmimd579u/xWGNp1y1diuZjaHLdypBBGjs0JaihQBAigAKE3KAQCHQAQpOYCgoUgD2gDYChECAyQoSpagNyMVDAEyqUnQC5AmRQHQbgASuWhehMqF30AdAXIm4ABsAdZ9qN/2C8Q/wCo+mgND4jfDtRfrGYgX+g+mgND4tTovCPqtX5v2hr2q+mjwdv9kD9euyfiU/5pu6aQ9kL9eux/ik/5pu9kYDiv173Qv9L9D7wLyQFTWWSVAgrUgap9qPgx+Z8dpx3hKy/+DibmXpYZUP6S3rOgS+9f3y210rTXCF7n05jhhjgigjhUUMScLTVU09jVfjd2dLzhvabffDuyy7RZLRE4511e6QwRyInq5TiaTg/ctpraq03fQOIKaaYx8meXKZ+UsNn4E1T07ceMNcKI7o7NPCR41vb+qG/pDWHbFMooIlRW2avvF+4X3z3+Dzphw87PeOL5vyTBiS7Ztw3VBGnaJ06ZB7rFD+DLhhbzelXRLXPQ3GuO67Bct0WW6brssuy2KySlKkyYFlBCvtm926lzr2u0U2/I41W8zzmOyFHAwJmrp3I6oeXKgglwQwS4IYIIF3YYYVRJLRJbIyfkBkaDLPgyqD0+L8T3FhG5Jt84gvCVYrHL3ifvo4toYIdYonyRNFFVdUU0xvMomqKY3l+uKL+uvDOH7Zft82qGzWGxy3MmRvXwhS3ibyS3bNBOLePLy4h4xn37bVFJs8P6FYrL3qqzyU8ofN6t7t+CPeccOK968Sb3UChmWK4bNG3Y7D3s29PdJlNY2ukKyW7fWzVDo2gaJ9Cp8te8+fhDXs7N8tPRp5IbRdjLh7FKgn8Qr0s9HNhis91KNaQ6TJy8/gJ8lFzOmuB/De38RsXS7EoZkq6LLFDMvK0pU7kuvwIX+HFRpcs3sb9XfZLJd9hs9gsMiXZ7LZpcMqTJgVIYIIVRQpcki04o1WKaPotueuefh3Kum4u8+Uq5dj9loQpHQ0JnV20HQi+oq8gHREdDKiIwJ0I9SryJlUJgKmQq1CF2HxkyKAQIioJ2ToFroi5Eoq+ITsqLkTIZUCJSpakyLsJQBgdAIXYEAtR0A1Am5DLoToACYyFFQbi1LsY5FQFDYyAE6CoXkRgWuREydCoCh6AOgEyKiIqoE7rUVJlyCoEDpyHRFyqR0AIr6Ey5Fy5AQdBUZAPQdEMgARSLyLkAy5Ey5FdOROeQDLIvQjLkIE32GRVSugy5CRORCsmQFCZNtCqgDcPyCpmV7AR0BQBiUtBQCAFoBMqAtMiUVQDIUUAnQooAOsu1H+sXiD/UfTQGiERvf2pMuBOIf9n+ngNEHqdE4Q9Vq/N+0Ne1b00eDt7sgU/Pqslf7zn/ADTd5Gj3ZEipxtsC52Wf803hMDxX697oX+l+h950IV6DI1lkhE20Ki0VAJTwDLRHB+MvEm6OGmHZV6XlIm2u0Wqa5VkskppRzokqt1eUMKWr8VrUqWbVd6uKKI3mXmuuKI3nk5tRLYZUNYIO1pB9/gOPpea/7Z+Fs7WVpctqxYEkwx7OfebaXSGWvlMzHD2oT/6/jH8rT6fY/E2m1eh+NttVlsNmjtNttEqzSIFWOZOjUEMK8W8jTG/O0xxJvGFwWH8yLmh/Cstlccf5U1xL0SOscT4qxJiaf7tf9+W+8Yq5KfOcUK8odF0RfY3CeTXO92qKY/WVvd1W3T5kbtsOJnaRwpcMubYcKwrEN5KsPusLcNllPm49Y/KHLxRqrjnGOIsbX071xHeMy1zlVSpfwZUmH8GCFZQr43u2cfVEU23TtFxsGN6I3q75YrIzLl/qnktXU5bwtwBfnEPEcN1XRLUuTLpFa7ZMT9zs0uur5vlDq/KrXJODfBbEfECdKt86GO6sP9739umQe+mrlJhfwn+6+Cub0NzsFYTuHBtxSblw/YYLLZZfvot45se8ccWsUT5+lFkY/WeIbeLE2rM71/CFxh6fVdnpV9UPy4e4QubA+GLNcFySO5Ile+mTYl+iT5j+FMje7fxKiWSOQ5cg0MjnVy5Vcqmuqd5lsNNMUxtCdCGQpnkeEoqFRKFVALkMi7igGFMz8orRZoY3LitEmGJawuYqr4z96GtvENP+r6+XX/1L+RHu3R052YfWdVnTbUXIp33lsap9naynyvy0VTrPX9Plflo1ZUcaXw4l/GY78f7ZH+Uyr9H9rXPrnP8Aa+Laf3aQ/wBnlflr2j3WT+3S/wAtGrHfjX38Xqywzpn7ZH+Ux5D2n1zn+18f8NpvdZP7dL/LRPdpC/Z5X5aNXFOm/tsf5TJFOm/tsf5TH0f2p+uk/wBr4/4bSKfZ/wBvlflojtFn/viT/KL2mrLmzPw4vVmPukz8OL1Y+j+1H10n+18W1H3RZv75k/yi9o+6LN/fEn+UXtNWFHM/Di9WXvx/hxflMmMf2vM8az/a+P8AhtNBNkxxd2CbLifJRpn6dDoHhHHE8fXcnHE1WPJv9wzv/Yo3KOhOzZtF1T+p2Ju9HbadjIdAhRHhmEy5DIoWgQb6DLkBkBOgdCkAAo5gToQrFAC8gypBoCIF2GQGIKMgIV0I6FYDLIhUiqlAnqYlRRSgQmRSFAKnImXIqDQBDoMsxlQBkNgNkBANSgMuRDLImQEdB0K0EAXkHQLUMCMnQy5ESAgVMy0RaATcMpGAAG4BbaF9CVLVgAKioE2ArkNwKTctSVAhdhvuFUQOsu1J+sViH/Z/p4DRCLVm93amr+cXf9Odn+mgNEXqdE4S9Vq/N+0Ne1X00eDtnsjJ/n33d+LT/mG8RpB2RGlxtsFd7LPp+QbwVMBxX67HhC/0r0M+KPzHIpORrTJ7mRkvIxqVMGzJLwNXe3o338GQ7Utr+OQbRJmrnbzf6Ng395bflkmY4e69Qt+/5Ss8/wBBU1foZIiO3OAXB6zcULvvW1WjEE66vuCfLlKGXZlN90UULdauJU0Ol5eVbxLc3bs7RDXLVqq7V0aebqYVNu7p7K2DrPH3rzxBfl4fuIHLkQ/FC38Z2NhThDw6wy4Jl34Yscc6BZTrUnPj9Y6mAvcWYlEfdxNXwX1GlXZ5zs0zwFwuxvjScvzFuScrLWkdstK9ykQ/xovheUNWbLcMOzlhfD3uVvxRGsQXhC1EpUUPdsst/vNY/wCNl4HeEKUKUMKSSySSokipmtZ/EWXlRNNM9Gnuj+WRsadatdc9cpLgglwQwS4YYIIEoYYYVRQpaJLZGWY+MVMBLIQEK2QIFoUVAEKmKhAUpBUBua38Q3CseXwqpP7pfyI2PqfhMsNimzIpk2x2aZHFrFFKhbfWh7t19Cd2H1rSp1K1TRFW207tXVQNeBtErFYVpYrMvKVD7B9x2P8AvOz/AMlD7Cr5f2Nb+plX934f5asxuhiok4qVzNpnYrC61sVmfnKh9hxrifYrHLwFekcqx2eXGpcFIoJaT/TId0iYv7ztsoZPCNdi1Vc8pv0YmeXd73QSZdjBsV3LhpvakWT5lSrsbIYJsdkiwdc8UdlkRRRWKS24pabfvF4HtvuCwV/8jZv5KH2FH6RETybpa4PquW4r8rziJ5f5at0BtL9x2L+8rN/JQ+wjsVhazsVm/kofYR9Ijuep4Jqn/wBvw/y6D4TR04hXYucUfzIjYKuWx48uw2GVOU2VYrPLmQ6RwyoU11SP3WSKFyvpzu2fRNKnTLM2pq33ndmRiorkeGaASrKnUI2KsAlQhdwEAAqGxUCdSrQlcghuMgE8g2AIG2SuYAuw2FQI9B1KwBUFpuEEwncZORWyIIAtECgFoWgTI2wDfiE3zI2EyBahAEibBCoAIAKoF9gzJUVAqYehKj4wFAhXQICgVFWBNwy1zMW8hAyG4JuBR0HIAGQMgF2D1JoUCsbkG4DoXoTIuQ3HWXalX9gnEP8As/08Boc9TfDtUOnAq/1zis6/58Boi0dD4R9Vq/N+0Ne1Wfvo8Ha3ZKdON91+Miev+Bm8qNGuyYv7N11/6Cf8xm8iMHxZ67HhH7r7SvQz4qQpHp/SawykARAtAQzqatdvJ1tWDl/i7Z86UbSI1e7eKXumDot+7bF8ckzXDv8AuFv3/KVlqHoKmsKNruwrRYexO9/u2T9GzVJG1nYVX9b+J3/lkn6Nm68Tf7fV4x82G071iGyIaAOYNlKBFouQoBaCmXQIAKAMn21AtB0ImUQFPAiLQiApdiAAvIMZCmYEFC6EQSHGeKzpw+vX/Rw/SQHJ8jjHFan5317L/FQ/PhPVHnQstS9Uu/ln5Ndw9H5GLeYrkX/Y4r2tnMF/qQuf8Rk/MR7VnqMFuuD7m1/8jJ+Yj2+5j6ubt+J6CjwhAKBELgpmeHfV52C57vmW+8bRBIkQbvVvklu/A/DFF/3bhy64rdeE2i0ly4fhzYvwYV9kjoHF2JbwxNeP3VbY+7LhqpMiF+9lQ+HN83uVLdua59jAazrlvT6ejT11zyju8XZE/i7dUExqTdVsmQp5ROKFVPxfGG76f2ltf8pCdRMxZdeQoaPPFGozPnR+kO3vz4bDX+0lq/lIS/nxXelV3Lal/rITp8jHkKUfWjUfxx+kNpLjvCXe1zWO85UEUuC1SYZsMETzSarRnl7+Bx/hs+9gO5Kf3nB8h+mNsSWXDNzxWycvdJ8b7lnkp0cyP6ktWyy6M9LaHS7eVTTiU37s7dUTMvZXret33RZHarytcqzSlknG9XyS1b8jgd+cWLDJjcu6rvmWmn7JNi9zhfTNnVt+XxeF93hFbrytEU2a8oVpDLX4MK2R4DddS5psRHNomocWZFyqaceOjT39rnlp4r4iji/QbJdsqHxgji/6kflDxWxRC6xSLrjXL3KNf9RwWJpbmPoVIt09zCzrWoTO/lJdl2DjBboYkrwuSzzFvFInNP0iTOb4Z4gYdvyOCRDaXY7VG6QybR73vPknozX1oxihTVCKrFMwvsTinOs1R056Ud0tsEGdS8Icbz4rRKw7fE5zO/72yT43WKv7XE9/B9OR21kWldM0ztLounahaz7MXbfvhNipZgbnhf8AYUy0PQ4lxfcFwJy7ba1FaEqqRKXej6rbqcM4j8Q45cyZdOH56ThbhnWuHOj3hg/neh1PNjjmRxRxxxRRxNxRRROrb5tlxbs79ctN1fimnHqm1jRvMc57HZl7cX7Q24brueXCtorRNb+KH2np4+LGKm6wybrhXJSY3/1HBnQxZWi1THY1C7xBqFyd/KTHg7BsvF3EMES93u+7Z0O6hUcD+VnKbh4tXNa44ZV62Sfd8b+/T90l/Fn8R0qkWnMiq1TKrj8S6hZnfp7+yW1FitdlttlgtVjny58mYqwzJcVYX1R+yNbcJYovPDVt92sUxxSYmvdrPE33Ji+p+Kz+Q2Aw1fdiv+6JV42GNuXHlFA/hS4lrC/Ff0ltXbmlvuja7a1GnozG1cc4/h7IqIXfc8M8qPCvm9btueyO13nbJVlk6JxvOJ8klm34I9HxAxjZcLWFKGGG0XhOX6BIrkv3UXKH5Toe+L1vG+bfFbrztUdonxaNvKFcoVokVaLU1dctb1jiK3g/d246VfydpX1xdsUqOKC6rrm2lLSZOj7ifRVZxq0cW8TRP9Bsl1ylycuOP/qRwZ6H4zGk8y5ps0R2NJv8Rahdnfp7eDnsni1ilRe/s91TFy9xjh/6z390cX4Iu7DelzxQc47PM73xM6hgabyZ+8B6m1RMclO3xBqNqd/KTPi2Ww7iW5b+lt3bbII5iVYpMXvZkP8AFfyo9uasWe0TrNPgn2edHJmy3WCOCJpwvmmdycMcf/mzHBdF8RQQXhT9Cm6KeuVNovlLW5amnrhuOjcT0ZdUWsiOjVPKeyXYdCUKzEottZdAQAKZgD7agWgQIE7Ki08CAINyNFepGBSblJuAA6gCPQGRKAYlFMhTMCgFpmBCkLzA6w7U6b4FX/4Ozv8A58Bok9Te/tS5cCcQ/wCz/TwGiDZ0ThD1Sr837Q13VvTR4O1+ydT8++6n/iJ/zGbxo0b7JzX5990/6Gf8xm61/wB8XVh+6J97X1b5FhsNnh70ydNiol4eLeyWbMJxXTNWdTEdc7R+6+0qYixMz3vPFG1kmzVbiR2oLfOnTLFgO7pdmkptK8LbB35kXjDL0h/jV8kdN31xO4h3xMcdvxnfcdXXuy7VFKh/Jg7q+Ip4nC2Zfp6Ve1Pjz/R6u6naonaOt9C3lqmuhkkfOm7OImPrrnwzrBjO/pUULyUVtjmQvzhjbT6o7a4edp7El3T5dmxnYpN8WRukVps8EMmfCudF7yLyovMnJ4Vy7NPSomKvmW9UtVedGzb2hq928l+i4N/e235ZJsPgzFVwYwuWXfGHbxlW2yRukThyilxbwxwvOGLwZrv28X+i4OX7i2fLJLbQKKqNSopqjaY3+UqudVFWPMw1hRth2FV/WviaLnb5S/5ZqcmbY9hRp4XxMuVvlfRm5cTf7fV4x82H0708NjgWgOYNlEULYAB7AADMTJkoBEUBAUiKlkQBsUUATBzAI9QgIPQICrQ4vxZ/W+vb/RwfSQHKNji/Fdf2Pr1/0cH0kB6o86FnqXql38s/Jro3mK5Bkej8jIdjiva2ewZ+o+5vxGT8xHtWeqwYv60LnX+QyfmI9q9zHVc3bsT0FHhCgjCIXLorjlNmR46hlRzIopcuxS3BC3lD3nFWnnRehwZM5rxy/XAf4jJ+WM4SXlufsw49rczOfd373KMAYXeK7fabIrcrH7hKUzvOV361dKaqhy+Pg7MWmIpfWyP+ceFwA/VDeP4ovpEdzM8V3aonaJbToWh4WXhU3btG9U79suo3wetO1/yOtli/nGMXB217Ygs3+6xfzjt4Hjy9fezH1Y038Hxn+XrcK3XFcuHLBdUc+G0R2WSpbmQw91RNb0eh0nxYvuK9sX2iVDE3Z7C3Z5arlVP3z9cuh3860y1pkdEzOGOMZ0+ZPnQXf7pMjijiatLebbb+98RamN95WPEuPkVYtGPjUzMdu3dHJwutWc04bYK/qlmzLXbZkcq75MXdfcdIpsX4Key5syXC3Fa+9sH+8P8Amna+A7onXFhSx3daoZatEHeim9yLvJxOJvXfKhVuXIinqlgdD0C7cyYnKomKYjfr7Vu7CmHbvlQy7Nc9kSS1ilqNvzbq2ftasN3DaoHDPuixRp/4lL5D2tSqhaxVPe6HGHYinoxRG3g6j4i8O7PYLvm3vcajUuSu9PszfepDvFC9ct09jrBo2rmy4JsmOVHCooI4XDEnunkat3jJVmvC1WZaSp0cC8k2i6s3JmNpc+4p0u1iXKbtqNoq5w/GVHHKnQTZcTgmQRKKGJapp1T9TZrDd4q9rhsN5JJO0SYY4ktot/jqaxnfvBudFNwFY1E6+5xzIF5KJkX46t1Tg6/NORXa7Jjf9HL+R1/xlxNHdN1y7osUxw2u2wtxxQujlytHTk3p5VOwWnQ1y4iXk72xjeNpUXelwTHJlfvYPe/Km+pRs0dKpsfE2fViYnRonaqrq/l6FRPIbUW+iMNDsHgph+XeV9zb2tUtRybDRy01k5r0fRZ+dC8qq6Mbub4WHXm5FNmnnLysHcLZ1ukwWzEE2ZZJUSrDZpeUxr9038Hy18jm8jh1hCVLUH5kwzMtY5kTfynLHQjLKq5VVPN1PE0LCxqIpiiJnvnrcDvrhZcFqkxO7XNsE+nvaRuOBvxTz9GdR4luK8MP3lFYbwld2LWCOHOGZDzhe5s0jivFG4pd94VtLhgTtVkhc+RFTOqVXD5NfUe7d2YnaWK1vhzHu2ZuY9PRqjr6u1ryzmXCTEEVzYml2WbN7tjtzUqanpDF95F65eTOGd74xDMcEajhdIk6p+JcTG8OfYWRcxL9N2jnEtrzw77vGRdF02m8rU6SbPLccVNXyS8W6Lqflhe3q9cO2C8N58iGOLzpn8ZwfjxeTk3RYrpgip90TfdZirrDBovV/EWlNO9Wzredn02MKciO7q8Z5Oqr9vO1Xxek+8rbH3p06KrVcoVtCvBLI8B6mVGcg4eXHBf2K7NZJ8Pes0us6eucMP3vV0XqXUz0YcmtUXMzIinnVVL32BOHM++ZEq8b2mR2WwxpRS5cGUyauefwYfHV/Gdl3dgvC9hlqCTc1liovhTIO/E+sVT36SUKhhSSSokloUtqrlVXa6pgaHiYlERFMTPbMvR23CWGrXA4J1y2Jp7wylC/VHBMXcLoZcmO14emTIooVV2SY695fuYufg9eZ2uF5EU3KonmqZej4mVRNNdEeMc2q8acEcUEacMULo01RpmMEyZJnQTpMyKXNlxKKCKF0cLWjRzjjVdEF3YqgtsmBQybfL90aWimQukXrk+rOCl3TPSjdynNxasPIqtTzplsdgC/1iPDNnt8VFPhrKtEK2mQ6+uT6nvzpzgFeEUm+bwuuKL9DtMpToE3pHA6P1ha/JO40izrjarZ1PRMyczCouTz5T7jcpCnllk3KChOyAtCNBARMpEgLzI2xuGgKMqiuWpM66gUJZEr4lz5gB7B1HUCUVBQq8xuEyUKkOoCAhXUmfMJ2dY9qf8AWJxD/s/08BogzeztVtrgTf8A4x2Zf8+A0T3OicI+qVfm/aGvat6WPB2h2W7VZbBxfsV4W6fBZ7LZbJaZ06bG6KCCGW22+h4XG/ifefEjEcc5xTLNcdmjau+xt5KH9sjW8cXxLJb14BKmzJXf9zmRQKOFwRUdKwvVeRgZucC3Vl/SauudoiPYsvpFUWvJxyWHURRKHVpebOScM8G3rjzFtlw7dKhgmTaxzp0a95IlL4UyLyrkt20tzdbAfBnAOEbBLlWe5LNeVsUK91t1vlQzpsx7td5UhXhCl1LXVNdsafMUTHSq7o/dVxsGvI+1HVDQJRwxfBiT8mMj6FYw4VYBxTYYrNeeG7DBG01BabLKhkTpb5wxwJejqvA0z43cM7x4aYmgsM6c7ZdlrUUyw2vu0cyFPOCJbRw1VdmmmtaKnpnEFjPq8nt0au7vesjArsx0ucPXcLMdX1w/xPKvq6JjiltqG12WKL9DtMuucMXJ7qLVPqn2x2v8SXVi/DeAMR3LNcyyWuVbWlEqRQRVkqKCJbNNNPyNe06HkT7da5t2yLvmT44rJZ5kc2VKbygijUKia8+7D6F5f0yivKt5VPVVTz9sbKVvIqpt1W55S8Optf2EM8OYo/HpP0bNUGbXdhCqw5ih/wCXSfo2Y/iWf9BV4x81xp3p4bJ5Booz5nMmxogi9SPzCdwKgr4hfUEDApnqUJ7GJSkCFRC1fMlcgKwSueoCTIlCjcG6UG5RnUIEjjHFb9b29v8ARQ/PhOT5nGeK363t8f6GH58J6p86FnqPql38s/JrkSJ+8i8mNyRr3sXky/cV7W0GDf1JXR+JSfmI9o6HrcIKmE7oX+RSfmI9mzH1c3b8T0NHhCNCHUdQiFxLorjnDTHsL52CV86M4NQ57x0X9e8p/wCQS/nRnA6F5b82HHtb9fu+LsfgF+qK8fxNfPR3OdNcAf1QXl+KQ/PO5X5lvd850Dhf/b6Pf8yiBeo6lNsOwjIxrQ6xx/xRhsM6bdeHVLn2mBuCba4s5cuLdQr75rnovEmmmap2hZZufYwbflL07fu7LtU+RZpTm2mfLky1rFHEoV8Zxy8Md4TsUThm31Z4mtpVZj/4UzoC8r0vC9J7tF5W2dapj3mxt08louh422RdU48drS8njO7M/c29o9rvKfxTwrA6S5ltnfvLM/raPwi4t4eWlhvSL/VwL5YjpOjDyWZ7+j0MfPFmo1T1bfo7pi4v3FDpdV6v+LK/nnUF5WiG13jarVBDFDDOnRzEoqVSbbzoeE5kNad5ep+0GcNUKbdNM9TG6hq+VnxEX55cupDvXgl+oaX+MTfnHRlDvPgj+oWX+MzfnHi/5rK8I+vT4T+zl97Wj7juq2WtayJEcz0hbNX46t1ize78TZPGcTgwhfDWv3HM+aa1xPMjH5SvuNK5m5ap9ksIkd8cFrIrNgOzzaJRWqdMmvxXe7q+KFHQ0WhsVwwUKwDc3d0+5q/Gyb8/ZWvB9ETl1VT2R/DkbDQfmM+ZaOkiDUMULhiXehao09wh1Cdu96mHC+HFpcdg6yUZf1M4cf8AcK7v93hPaNlT8Sd5W/0Sx+CP0fnYrLZrFZoLLZJEuRIgygly1SGHOuSOk+N1pinY0hkVrDIssCSro2239R3itToDi9E/6v7cntBKp5dxFSzG9TXeLKuhgxTHKZhxRHL+G+JbrwxabZarfItU6OdBDBApEELok6utWvA4emYxuqLqqmKo2lzzFybmLdi7b5w7lfF3D6/uZe35Er+eR8Xbg/wZe35Mr+edLxeREeYs0MxPFepd8fo7n/PfuH/BN7+kn+eVcX7i2ui9vSV/POmKGXddCZs0I+tWpfij9HOuJWMrpxVYLJLsljt1nn2ea4k58MFHC1RqsMT8NtjgpjFqKimmKeqGIzMy9mXZu3fOlyrhZaHZsdXbFWnfjct9YWjYShrfw+beNboS/vmE2RrnuW9/zm/cHVT9Frie/wDZKFHUJ+JRbcLUBfWOu4FIK+IVQkSFEUdQhMqkaVDLfUjAE32KTcBUoAEY3DHMC7AAiBSFG5IjBChPY6w7VMNeBGIa7Ozv/nwGiFT6SY5wzdmMcLW3Dl7+7fcVsUPujkx9yNd2JRJp+aR1bD2Y+GiXvor8ifN23/8AU2zQdax8CxVRdid5nfqYnOwq79cVU9zS6uRGzvftN8H7h4e3Bc974aht8Umfa47Pa4rRP907rcHegpkqfBi+I6H2N1ws63m2ou2+TDXrFVmro1Nqewld1nV04pvpwwu0x2uTY4Yms4YIYO+0vNxr8lGzFTTfsa44sOHcWXhhq9Z8Fns99+5uzTI4qQw2mCqULe3ehip5wpbm47WZzriK1XRn1zX27THhs2LAqpqsxt2DZ0j2zrvstq4P/d06GH3aw3hJjkRbpxNwNdU/iO7XoavdtzGsibKu3AdhnQTJ0E1W28e7FX3OialS34uripyUPMoaJaruZ1uKOyd/c9ZlVNNmrdrEmHoRHbXZm4c3VxExTednv6C0xXbYbEpkXuEzuP3WKNKBV8lG6eB1PMyqMWzN65yhrVq1N2uKaecuo3obYdhLLC+Jn/l8r6M5fL7NvC+BUiu+8Jn763R/Uc34dcP8NYBslrsuG7LNs8u1zIZk5TJ0UyrSoqV0NI1jX8bNxZtW4nedufizOJgXLVyKqnK0UgNOZZSFRAAAQFQYABBhACZkVS0AApMygT5Q9S0IwAQG4FRxniqq8P73/wBCvnwnJTjXFN/1gXsv8VD8+EmnnCz1H1S7+Wfk1yazES96/Jn6NIkS96/Iv3Fp5tnMHuuE7of+RSvmI9m9T1WDf1IXP+JSfmI9tuWFXN23E9BR4QxCLQqIXDo3jn+rWV+Iy/nRHA2c945/q2l/iUv50RwJ6l7a82HHtc9fu+LsjgD/AG/vP8Vh+edynTXAFf1wXl+KQ/PO5WW17znQeFp/8dT7/mVJUFRS3bFLh3F2/ZtyYRmqyzHLtdti+55USdHAmvfRLxS+U17ghoqLY7Z7QUcbtdzya+87k2OnjWFHVqgL2xTtTu5ZxRlVXc6aJ5U8mMEMcUcMEEMUUUTShhhVW29kdr4R4UuZZ5dqxDaY5cUar9zSaVh8IoufgvU4twmskq0Y8sCnQqJS+/MhT/CUORsG9TzeuzE7Qv8AhjRrGVRVfvRvtO0Q43Y8CYTs0KUNzSJj5zm42/VnsJeGcOwKkNxXYv8AZoPYe09Sot+nVPa3ijBxqOqm3Ee561YeuFOquS7V4/csHsNd8TqBYjvJS4YYYFaZihhhVEkonokbG4gvOz3Pc9pvO1RKGVIgcX757JeLdEayWmdHaLRMtEz4c2NxxebdSvZ35y03jCbVFNu3TERPN+TeZ3lwQ/UNB+NTfnHRp3nwR/UNL/GZvzib3msdwj6/7p/ZyrEch2nDt5WaBVjm2WZDCub7roaxN7m1i0z6ms+L7rjubEtvu+KGkMudE5fjA84fiZ5x55wyvGdiZpt3Y5dcPUM7+4N2pWjANil1rFZ45kmLpE2viaOgWdlcCr8l2W8rTcloj7sNrpMkVf36VGuq+QqXqd6WD4Xy6cfOiKuVUbO5CkKWTqoloHRLkVLI41xNvyC4sIWyd3oVaJ8DkWeFvNxxKleiq+h6iJmdlHIv02LVV2rlEPfK1WaLJWiS/wDWQ+0rtEhZufKS/fr2mp0MqFKiRfcXG+5DB3onklTVlfyHtaV9cqt9otfH/DbOVHBMgUcuKGOF6OFpp9TozjbIjk45c5r3s+yy4l5pxQv6juDCF3fmRhm7rtok7PZ4IGlzpn8dTg/Hq63Ouyw3vLhq7PG5U1r8GLNP1XxlO1O1bK8Q2q8nTOnttMbTs6gqc44N2a7bdiG0WO8rFZ7UorO45anQKKjTzpXzOCKp7rBd7/mHiew3lG37lLmd2b+8iyi+J16FzXG8dTnul3qLOXRXXy3693fDwnhhurw/dv8Au8I/qUwxp/U/dv8Au8J7mVHBMlwzJcSigiScMSdU09ylp0qo7XXoxMaY3iiP0h6T+pDC7/uBd/8AIoLB2Fn/AHAu/wDkUe8Rdh057z6Djf24/SHo4cG4U/8Al27etnhK8G4V/wDl67f5CE9ZxAx3ZsJTbLZ/uN220z04/c4Zqg7kCy7zdHq8l5M4tDxnT1w2/wDfF/NPVNNyrrhiMnP0jGuTauRTEx/+f8OwbJhXDljtMu1WW5LBJny33oJkElJwvmme3OE4Fx9Fiq95lgguZ2WCXJc2Ob90d+maSVO6ta/Ec1PNUTE9bKYF7GvWvKY23RnujYBCnleG4FMygQFIBQRlQE3DAdcwA3FCUAoG46AGKEY3AuxXqTYPUgUEJ0JF5BEKEwvUgoRhDinF7CMGOeHt6YdcUME+dLUyyxvSGdA+9A34Nqj8Gz55W2y2iw22fYrXJjkWmzzIpU2VGqRQRwujT8U0fTg6J7RfA2HGMc3FGFoZcm/1D/4mztqGC2pLJ10hmUyrpFvnmbPw5q9OJXNm9O1NXb3SxuoYk3o6dPOGnB2zgftBcRMMXfLu6O02S+7JKShlq8YIopkEOyUyFqJr99U6uva77fdF5TruvWxT7FbJEXdmSJ8DgjhfimePWuhv1/Fx82iPKUxVHYwlFy5Zn7M7S7txL2meIN6WGOy3dZrruVxqkU+zwRTJq/euNtLzodL2y02i2WqbarXPm2i0To3MmzZsbijjibq4m3m2+Z+VaCVDHNmQy5cEUccTpDDCqtvkkecbBxsOJ8lTFP8A3vLl65e86d0bob09lvAk7BXDaXNvKQ5N63vGrXaYIlSKXDSkuW/FQ5tbOJnV3Zz4EW2G87JjDHNj9wlSIlNsN2TYffxx6wzJq+9S1ULzbzdEqPahxNo0viXWKcj/AE1md4jnPtZfTsSaPvK+fYj6Agp4GoMuoINwhkuoIgEwrQWhKAGwA0RkSQyRCFJQtARAC0BABUGRFAMm4p4E3Apxfiu6cP71/eQL/mQnKKHFuLK/se3p+9l/SwHqjnCy1L1O7+Wfk16qR6PyJUN5MvZcX7WzeDP1IXP+IyfmI9semwQ64NuZ/wCQyfmI9wWNXN23D9BR4QPYqIKELnsdGcc3/XxB+JSvljOCM51xz/VvB+Iy/nRHBHoXtrzYcd1v1+74uy+z/wD2/vP8Vg+ezuVnTXZ+/t9ef4rB887kZb3vOdA4X/26n3/MoESmZUqf/wAKTYt3W3Hm7o5t13feUELas82KXG6aKJZP1XxnUCVDZ++bus17XVaLutcNZM+Bwumq5NeKefQ12xXcNuw9ekditsDpWsqal72ZDzX1rYvMeuJp6Lm3FunV28j6TTH2aufsl+WHb0m3LfdlvOTD34pEdXDWneh0a6qpsRcV83dfdggtl3WmCbLi1SfvoH+DEtUzWVsykTZsmapsibMlTFpHLjcL9UerlqK2P0XXbmm709HemW0/U9Rf2J7juOW3eN4SYJiWUmGJRTIvBQrM18nX1e82DuTL1t0cNKUdoja+U9XMp3nF983Vt7lOnH75Zy/xnVNO1q3tPtlyriBjK2YptUMvuuz3fKirKkJ1q/wonu/iRxhaHu8EYWvDFF6QyLPDFLskEX/iLQ1lLXJc4uS9cjyeJt22a6MX2iw2OUpVnly5fucK5dxZ+LrXMqb0xPRhrOVaysi1Obe5TO3X/wB5ONneXBH9Q0v8Zm/OOjdjvDge64HhXK0zflKd7zWU4Sn/AF0+E/s54jrfjZhyK2WCC/rJL706yw920KFZuX+F0fxPwOxixQwxQuGJKKFqjTWTRb01TTO7oWoYVGbYqs19vzaqLMzkzJkifLnyY4pc2XEooI4XRwtaNHYXEbh7Pu6dMvO45EU2wxVijkQKsUjyW8PyeR121mXtNUVRvDkOdg38C9NFyNp7JdxYR4n2CfZ4LPf/AHrNPhVHPhgblx+LSzT+I5jKxRhqbL78F/XY4aVztUC+Jupra0YRZ65lKqzEyz2Jxbl2KIorpirbtd/35xGwxdsiJybZDb5yWUuze+q/GLRHTOMsS3hii81a7Y1LlS01IkQP3stfW3uz0pD3RbilYalr+VqFPQq6qe6BI5twjw474xFDbZ8utjsLUyNtZRR/ew/W/LxPS4RwveeJbepFjluCRC/0a0RL3ktfW+SNgMPXPYriumTd1hg7sqWs29Y4t4n4s83bm0bQv+HdGryrsX7kfYp+L2J4d93bZ74ui1XZalWVaZbgdNVya8U6PoeWiotIl0qu3TXRNFUdUtX76uu13Nelou23Qd2dJi7re0S2iXg1meGbC8QcHWbFFhUcDhkXhJT9xmtZNfgReHyHQt83Zbrot8dhvCzRyJ8DzhiWq5p7rxReW7kVQ5PrOjXdPuzMRvRPKXNeHnEOZcsmC672gmWiwQ5SpkGcclcqbw+Gq+I7Zui/7lvaBR3deVln84VMSiXnC6NehrMGlE/fJPlVCqzFS607ifJw6It1x0ojlvzbVTJsmVLcc2dLggWsUUSSXqcRxTxFuG6JEcFjnwXla0qQy5MVYE/3UWnpVnQdNmqrxMnmqHmLERzXWTxjfuU9G1RFM9/N+t+Xlbb5vafeVvmuZaJzq3ooVtClskeHQ/SlWdhcO+HlqvK0yrxvqRFJu+GkcMqNUin8qraH5fjKk1RRDX8bFyNRvdGmN5nnP8uXcErhjuzDsd52iBwT7wajhTWalL4PrVvqjnxIYVAlDCkklkkgiyqqmqd3W8DEpw8emzT2QrKtCUG5C7VLwDC1DAEQIBQtyU8ABluGjFalAAvUmVQAGXMOgEqXcjpzKqcyJTA2R6mSSoGlUjZLGpUGlzFEenkKTIuQDoR7FVBkBiy1FEMgOPY1wThTGdlUjEtyWW39xUgmxw92bB+9jVIl6nUV79ljBVqnOZd1+X5d0L/Y1FLnQr8qGvxnfzpnmMuZe4+o5WNG1quYjxUq7Fu5O9VO7Xy7+yng6TOUVtxFf1rgWsCcqUn1ULZ2lgXhfgXBThm3Dh+zS7UlT7qnVmzvy4qtdKHMlQNIX9Sy8iNrlyZhFGPaonemlHz3JVloiULJWUqMepkgIC5DLmBEXoRDKmoFJ7C5cyZBPUE5ldA6BAgMiqgETCeoy5hb5gAH5jLmALUmReoEG6LkMgImeqxddMV+4dtd1QT4ZEU+GFKZFB3lDSJRaVVdD2mxkhEzEvN23Tdom3VymNpdQrg7PfwsQSl5WR/zxFwcn0f9cUv/AHJ/zzt7IjKnlqmC+rOnfg+M/wAvCuKxfmbctiu5zPdXZbPBJcahp3u6kq02PMqXInsKc9bO0URRTFNPKFqKk6lVA9OB484fTMT37DeUN7Q2RKRBK7js7j0bda95c/iPQ/nNzN8Rw/7n/wDudtpl9p7i7VEbRLDX+H8G/cm5XRvM8+uXC+H2BYsJ2+02p3p92e7ylL7v3P7n3aRVrXvOpzNly5mL8zzVVNU7yyOJiWsW3Fq1G1MBakdPiCpQhcskzwr8um775sMVjvKzQT5T0rrC+aeqZ5mVdRkImYeLlqm5TNNcbxLqS++Ek73WKO570luW3lKtMLqv40PsPQTeGeLZcVIbNZZq5wWmGnx0O+XqGirF+qGvXuFsC5O8RMeEui7PwwxRNipNhsVnXOO0d74oUzktw8JrHKmwzr6t8Vro6+4yU4IH4N6v4js6i1FEKr1UvePwzgWKul0d59r8rBZbLd9kgstis8qzyIFSGCXDRI4Vjjh7/VNfv5qQ3qrJWTDLcDs/fq4a517y5r0OdGSKdNUxO8MrlYNjKteSuU/Z7nVK4PU/+IV/uf8A+5zfA2HHhi5YrtdtVrTnRTFGpXcpWmVKs991GXM9VXKquqVtiaLh4dzylmnafGSoJlzKqHhldlTOLYmwHcN+Rxz4pLslqi1nSKLvPm1ozlOVC5UJiZiepb5GLZyaehdpiY9rpu8uE16S4ond942W0QbKanLf1o9PO4aYuhdIbFImeMNpg+uh35lQjp1qVIvVQwF3hPArneN48JdFWThbimdF+jfcNmh3cc/vP0hTOV4f4T3dZ44Zt8W2O2xLP3KUvc4Or1fxHZWVSKhE3apVsfhnAsTv0d59r8bFZLLYbLBZbFZ5dnkwKkMuXDRI/UuXMKhT3Z+mmKY2hEqCpRQJVM8C/bkuq/LL9z3pY5dogWcLeUUPims0eeqcw34jrjk83LdF2maa43ieyXVt88I5UUUUd1Xq4FtLtEFUv4y9hxq18MMVyY2pUiy2hbOXaEviiod75CvjuVab1UNfv8L6fdneKZp8Ja/rhvjFxUd1KHxdol0+ce3u3hNfk5wu3W2xWSHdQtzYviovjO6VQKnxEzfqlSt8JYFE7zvPv/hxLC3D+4bjihtDlO22qHNTZ9H3X4Q6I5YwwqFKZmebP42JZxaOhapiIAi5BUIXKVBciKnMIUBU5jKmoE3IjKiIqACVMsqEaVXmBFqWuRMiulAFCFaAEAABoqWYFKMC0yD1D0D1CdyhKF5gIQFY2Agp4gmwGVCFIgI0F1KRgVIbEqUJkoShUGEI0VIjKgLQUACUQp5FRAgSFMugGzAUJQoAJZCgSACgpkCoCMULTwJQCItAioCUG+pdibhMCKkQZBC01J6FRAJQUzLzIAFCkAIyMTIAYtGRKAQhaDYJ3FqUlMyoIWhKFQYEoKFGQEaLQjKhuJQNFI0BKBFouRMgLy1KjHkUC7BojD1AtCAKgAIBAC0AAlAUdAJ0Ii0QAIInIuwEoWhEi0QCgoUbgKEKKICJFIi+0AQviyABQBAKeIaoXcPQBQjKQCMFYQAo5AAGsyACgBbAMgNigQlCkAooAgHUj2K9CAQJAAVFYRADQX1hlQACgYBE3FQBaCioBsAaJlQrIAQYQApEABQxUAAEALsR66gjCYCIAICkRdgKiUzA3AtCJFCoBBsXcgAFoTkAp8gpkAA3KkRlTAqIwmGBHuAEAZUQoANAeYEp8pEVABQUGxUAaDDDAIAbAKBIMICFIUAtBQIAQhWAJuWhCoAipBIoEIUgAvIjFQKhT5QgAFCkQCgQADcMlcw2BSCgzCdhigoUIQpCgQDbcjAoIVLMCgmZcwBGWnmQC7AJAAQpAJQoHUCogABoqXMNBeYAMBgY+gRRQAEAtwKPQACIAZiBAi08yUAFFBmARTEu4FJuUm+oEAYoEqgEUIRF3FAAAYeQEGwFAL6BjMMAT0KxsBHqBqWgBEZUKAAKDYAwGPYA5EKQAhQoQEFSsgD0BGXqBUCcswuoBhBhAKApACDqEUCAU1AAIACoEXUvUAQrIqgVgACLUoXmAH21AIBQAtAIHoXcjAo3A3CesCBQhGC/bQjAj0IZbDcCdSlIBChl9nIARlJ9tABSFAgK/tkQAQo5ewAgPtoNgLsCP7ZFX2yAELoQAhuFUAEFoUIJgZGV6kf2yAxKKCghAEFp/QVAAVh9QMSgbgCZ1KN/6AIUntKgCKQoDmTcv20Jv/QBSD7aDYCblAABj7aACfGXYnsC0AblH20ABagAC7AACMD7aDwCdgfKPtoPtoEBSVCYF2IABGGV6f0B/bICBF+2gQDYdQ/tkPUAKchQJAT2FA9QAL9tCbf0ANwPtoFp05AEAXcCMIv20IBC7oBAEBv/AED2gAikQAoAE3Iy0zDQgBnXQE3I2TuvIVAJQNsN59B0GQEqAQDKoIipgCv6iABUhSAUJk9CoBmMwAAAAEKRhMLmETIqoQSDcAkEOYVAAYWmg3FQg6DUCoTuLQMDIIAhsEAYHQZABuMqAAM6gVzAAOgqgAAYEKq1CC1AAACblf1kLkAzGYDoBM/iAy+IICuozAyAIMEdALUjBVQDF1KV0GQTuiHqFQZBACigEq6DMdC5VAgZQ6eASIIi2CoEKwgEAKQVAVHQlUALmQtUSqAoTyIXIJkXkAgEHsCHQnRAVkDCAudR8tQnmAHQIVCYGWZNijICBjcNgQhaEAoIUAOgHPyAPyIXYjABBoAUEAFIwAKCMvkAZC7D2gAEGAIykegSMJka1KqhCggAqAoNgBC0JRgOgA5gEXoQIDJEQC0AdBuBmAQIigCblzqQCogoAC8iogAqG5BUJUEFAhDJMxp8hl7QnZSdAGEHsINQBQAwIToWgAq8giFAMD2CgEHQpKAVfWXoRVGYAEZQAY2AECFBQChD2BIACkoBB0AYCoDIBS7EQWiAqHIAAECAVkAALyKiFYDcAICghQJuGxzIwL1IZCgGPUqAQAV+QUFAJV01D1K14CmYEY3K0AIAxQC+pGy9CUAoFMgAfmFUNBAF5jYeg2AEehaeAaIlO7FheZWgSgQ6igpkBUNgkXoAIZU8CUAx5Eehk0RoAghQoECZaES8ALnzGdQKeAEz+IvUJACVG5RTMCbDqWngKeAEHUU8ABM+YqWmpKagUdQUCF9pKFCRdQyolPAEp1GZlTwJQIQZl3FABGWgaAnUZgqWWgD2CoaGwABAASuZSASpSFoA2eZeZKFayI2AIChIZ8whQIAUgYEQz5igoBK6lFPAqAiqEWgoBFWgzFCpAQiMqESAAtBTwAi1FfEbingAqEKeAS8AKK6gcwJmR+ZaZhgZEA3APyBC+oDlkKBgBsCbal3APIE9QE7DHqT2AIXoCDqEwyqDGpQhX5EQFAKikVABehGOpNgD8hyyDQSQDox0Ii5AVeQqReZAllrsymKCCGTISpQCAVCZAKheTIFuBl0BABeg3MTJagNi76E9RvqA6AnUAKAIAVAi8ygRheRcsiATcyXluY5Fp8oGSHQiDApOg0JsBXroDF69Q2BlUjIvMewC9C9CFAj8ggEBV5DbQg9QL0IQZVAdC9BsABTF6F2AqKvJkRMgKERhAWoMfUqCZF5CngRF33CFp4E20KTYChGJVoBRuFoNwKQdSc8wKCBAUnPIbkCWXQgCoEL0JzyGVRlmEwb6BjfcUAo3IAhQQoBgPUgFWSG42AAchUlQBVqQICohfUjAdSog5AUehGAMiBMgFAqABV5ke4qBTF6lIwCYIigAmAgKVMxAFDFCAUImwQFDFQBAA9QKN9SAB1L1JsACZSIICobjmSoGXIbEqNgIymJfaBktARACkeSKYvqAeoI9SgF5gIMClRAAAGwAjFQwHVjqOXmAKUxLyAcysxfkGwKCDYCsiDCADcACrzHUiFdALUEbAAuxCoB8o3IAKQAAwOgAu49pCeuoFqOpCoCggQArGYYAhUTcACehQFCjMZgNiMKtA9QAQG4ApCqoAMZhgAiP6hUAB0CAAIegBdQCPQDJk2I2VVyAoYQAIAAKAIIAAwAGozAAhdibAUAgFQZF6FAUG4z5DfQACMZiAAQYFJuMyb7AZAhVoBC+0g9oFQCACviBn8RFoAeoAAqDRNygAH5AAxsH0ADcgzD6AULUgTzAuwJsADDI9AE7KvMIIi02CFCIVAUjAAIMINvwAnMqJmVAUUQTY6AA0BmBHqEXOmxFuBaBjMZ8gMd9R7Qw9wKtSoi6BVAoHQBMAAAbaE30KQIB0BUAp5gZDcBTLQPXcbDcB6igGQBrwFAAAYQegEp4DoUAToOhQBEvAvQIZAKEZQEwjRVtkHQBB0FAAKSgQAUIvLYuQAChSZACMpAARciICtEoUARLwL0CoMgHQm5ciZVAdB0BQIvIu2hFQoE6CmZSbsJg5ZF+MZAEm4p8oWpQSnQUKQIGvDYlCvUjAUBNygVLwHQAAx0AAnQvsI6F2AIlCh/bMCU8CLXRlJQBt0L0C0LkBKeAaKGE7iXgKeACoAIV0CCEoAUAtCdB6+oAlPAdAPUB0KhlULYCrTQU8AtCoDF+WwXkXIm4F6DoAApnoKBa7gCJBIoX2zAeug6ABOxuGK5kYOwJqZZE3CE31L6DIARlD8hUANwvIu4EKHQbgAOg6AOpGXLkTmBeor4jLkADC+sMJgAEAA2AypogADCYDqTcuVNCMJgTKRNCqBspOY6FVOQQdSMuXIjpyAEKggCBU1yC8gBOpcqheQEKRMoEG5cidAHUDLkMuQBAItUBBuMhkEwVAy5DIINy+0xRar4wKGEAD9CPMoyoBAXfQAQrJyGQFAdCdAA3DAFDCoMuQE9oKqchlyAnIvIZACB7lemhH5AKioyGVAHUIhVyAdShUCAiD8xlyGQECL0CALQqCpyGVACA2GXIAQr8icwBRuMgC+sCq5bivgEoVBUGXIIAKkqA3I2XKpH5AZUZKGRKASngWgzLTIDFkK6jMBlQMudA6gAAgAGYzAgKGgAYAEfIFZKAVDYKozAoYzAEYQdQgBH5GXoSgGPtCLnUUAAu5EBQwMwkQoAEBChVAUFAGBEUIUYAm5cx1AxG5aCjAItCKpQJQblJuABQ06AYloKZlowCDKiUYE9hGZUIBCijFALTMBVqHUCUHRloxRgRoFaGYBaAIOoADMZgRlGwABoB1AlBQtGAJQIrTIkwBdyJFowJ7BTwLT5ABjzKhnmMwL0AGdAACABgpAAXkMytAYl12G4dQAQzABkKFoBNw0XcUAtajcgrmE7KNiFCEY6F32JkAG42G4SE3Ll9mHrqED8ghkAAb8whQAC7D0AnIIZAAgNw3kAIXIbAGEGF9swKRsqIARCoAAgF9swAG4qvAAvIpFQACFqFSgAAoAm4LUCMbj0AD2ingHp1AEKFqMstABNy+0ZAAC7AQAemoFAJlkA9gL7CLQBvoCioAgqigCjYVAjAFQAYK+gGPQhkTcJhOhRyDoEBWSuoAIFVCLQACsmX2YApCoCBl+2pHqAG2gY2APyA9AtACARQJ7B0A+2oFBPZzHoA3ArmHoATCFfL1Cp9mAKRPyKE7G46CuZGELQm5kzFgGAwAG4AAj1LsNwkIihAk9gBUEBChgCAVAhSOgXkBakegQAAACsIBAUxZSAEUi8gtQKNgEBGSpk+piEqmGEAhSAICgAAhuQoCoTzDJuBaipC8gINhkALsTcDcCgciPyAF2IALsQAC+whSaoCsAgBalIAMhUgAN6AjKgAqAwAAQCuQBADKyB+QTuqCCKghGEGFTUAtSsgAIbgdABC8yAXcAgFTKYooADcARhlZAC1DCD0AjeYTKSgF2LUiQCdyuZGwGghnkR0qSpa5gMhkQoBoAAY7F3KTcClVCBAXIKhC1AZU0IymLAZBUA2CYMuRMsihBCAyAGPQZFSDQTBkAwggyGXIo3AmWQouRUKAOg20AAOhAUJYpFA3CDKgXkUi1AuXIZcgA9IN9AX1Dyg30KTcCZFA5AAReRQncRMuRkTcIMhQFAlA6F3yHtAmROhSbgMgXfoTYCghQCFAnmAGVC0QCAOnImXIoAg5ArAmVBkOQ3AcsiF2AEyDKACoVUCKEoyFZAgyAAFp4DfQIAToQyMQGXIZfENwgCpQyyMV4GQChCkAuQy2IVATcDcoEyIZEAKgyATAUQogisCCuZaEAFQpkAKGQoE2EQDAgWoKkABGUAQtCUAtACgR6BCgoBRsTOoAqRGUMCNBBoAMgKCgFWxAUAyLQrJTzAMAUAIBABsCoiAAtBQCBrMACkKTcBQAUAiKEKAAAAGw9RsAAoAAYACi+IxLT5BTICPUpGi7MAtQVIgFKQIACMqAIFIwAWpABRTIbCgE5lfQBgVAiFADIisgAEKBUQJCgFIOYAjC0MqEoARSIoBkAQBlQAAZUZC0Ab6kLQgAIUyKkA3K0QrAhH5ig3AoAa8QIK/IGvEnUC1y1G4zpqKZgXLwFVXUlPEoEdAC0AEehaEaAVyKiU8SpACimRGgmCvkTqN9RTIhK5BvxFBsSjcdAiNeJUvEIEBTxDAbAiKgBEVBKqAnIZFaFAIqDLmAl4gAipBICgUJTxADcCgCornqyU8QlnqBVpqFqSmRUmAVAEvkFM9QLtuQNahoCMCniKACkoPaE7L1JkKBoIMuY21HULTUCZVKKZhoABTPUtAICtEoAKqczFoqQFDFA0BAKeISz1AbFyJQUANrMN+QZGgLkVEp4lSy1AERWupEgAASALfyGXMU+QNAMgtBTxKkABKCgAdQkKAAqCgSArIy0FAMfaUU8VqKAUmRaEoAKmY0zFAMt9QyblayAg3K2TcBQCpagYsFAE2KBXMJAGwEIUdA/qAgKABRUAHoQPyHtCQbCoBIWhKlqEI9wVkrkADLuRgRfWVEXkXcPSoewILToAoRmVSMPLHUoFQKiFqRAABUAitEFQAWorUVzCdihSFTCBEKgwnZNiVKQABXQBAVfWQtdAFCMoCU9g0RWQICk3MkBKZhl3AEFC1AEehPYZbECYEAAg9oSzFQn9qgXYCuQqBGiMyqRsCBAqAESKFsACAWoTAtGGip5dA2DZiUVANghQECAqK5gCItSICgVLUCChagByIUgEoKF6CuoERdiVK9ALREepRuBKaFoTqXqAI0GAmEpkC7AARalAQFoiFQCiDAYEA2CAbD2hgJgSATQCFJ6lDoAaCQfmEE9RTzFARsAki+pE/EqYQIewahaagX1IAE7lFmQIbhBTzCp4l2CAUIXqQCDcB6gVEpmCgAiFAFIigEkSmZQBKChQBKZhIF21CdygoCMIH9RC18SbAXcAAKAbkAyG4qEHooKCor4gKDIEYeTIiL1ADYa8wQJhQ0QrCAJAqoAoKALzAmQA2AIciV8RUPS0Q2IVB5MvEAABTMJgBQLUAC8hTqOoqA33GXiTfUtQBKF5EAABAKBoBgVk3KyAB4DYvIDFmW7IUCUyG42G4AF6iuYEY6gvsAhCkyAoAAEKWiyAxLsVDkAIUgTAUhVoCRkfmZbmLCEQRUACLUAA3oR6CqFQItC+oAFIi7E2ApCkYEDGQAqAC6AANhuAQArkAC1A3AvUAfKBAAAqRlJuAIi0QQAFADcha5gAK+JSB6CBjmEdSpgIBABlkFQCCjLQATYDmVgEVEKqAGEwwgIAXUDEFRHqA5haDqFoBdxsC7dAMQUAQcy0IBQAABR1WoEQLllmAICjmAQa6grYEG5CAUpC8gAAf1ATmNxsTcDIIlKlpmA6AUACnkQUAFBC08wADJ7QKCJACmLMqGLpQJhSoxZUEKRlRKAF0AQAo5+QIvqANeBCtEp4gEUiAGS0IiFQFYJ6jIAhQZFALoPaCdQmApPjzFEECAFEAQLREArAoKAQpB7QKGCOgAB+YyoA3D8huAJuPQtCU0ApfQlC0AxZSULQAisgYAInVBAXYbE9hfUBzD3AAIIEoBk3nsERoICggAu3QjWeiBGkAoEWgAq10AoSmQBF3JQKgFIWnyESAoAyAD2k3L6AB7B1IgA9BkPQBXMNkpmQDIhlQlAGZaCgAAACDcrACg3KTcBuAwBPQjKRgVFIioA9ABQAAKAORGsioMJhAisIJPQlPIoDyCgQISgRSUJAFFAIgEHqEIChAAWgoBEWhCgQb9SsJZgQtSD1ADcFWwBAbAAxtQDyAhafKQoAxZlsR+QCuYWgAAbgoEoCrUAQoAEKGFuBAykYECKAJsUhQJQpaBhMoCgIRgBAEAAIUbAAUiKgA5FFMgMR8ZQAoAABNyvUAQAMBUiG5UABaeYAm5GvMyoIkBaIlEBuAdAAABGABVqYooTKjICoQjGQADLYBEApUTRBVAooNgAAz5EYFyI9BmTYJhXQKhGVBCqhBmGAVC8yIJhMKFv5DUiCVZGM/EjAAECJZIIiCCGW4yJVlqAVAQN5gUUzIEAYdA8iAAQoBehakVRnUC5DIhQBcvjJQP6wKg6EzDqAoNiBAUE3KABNw6gXIZeIqQClMc+RQL6kAAZABVqwnZOVCpDMBChkqHUJ2UmQzCBsESKOQIAMwEGwyJ7C5gCrQgAo2JmEwKgQbgUgz8SAV0HqAEm4oTcVaQQqSCIF9QTsyGRBmBQyVdQ603CFqNzFGS1AAKgAj29oKKATb+kuVQtBvqAqQvUAPtqAAHoGX1IAFAUAPaXYIAYun2ZkiOgGLIZZEy8QmAqFAkgAZSBCBFAAVKiLQJ3B50G4ByQFoMgMftqVFpkEtQgA3AEQev9JkhTUCBPMtCbhMICvQZA2YooSKggWjFCoASgAQDIFD+sCehMigCewbFa+QLQCZAu4YE3AoABQUCAMUAegFAA2BURagMvAAtAMSsbBgPtqRdChAAmKCgBD7agqAns5j7alDAgFAvEAT7amVESmQEAKl5gCFAEKKCgE6l9o3DAfbUhkiAQIpAnsNysiKwhiXOpKACiozLmAzAADYb7jOg6gEM6hBARlGYAgAAoGfMAAHUZ8wHMjCqKgBmUjAOpURlQFzI6lJuARCoMATMtfEKoADcZgEMxmMwGwzAQB1AdQAzG7AdQBNy5kVagAGhmAKiIoBABAQqHqEAGYzqNtwACAEGZafIR6AK5gZ1KgIEUgFGYzAEYSZWACIyh+YCpMxmEBdtwAADKRgQqIVAUhcyIBmAFUAgKjkAHMZhVAEKNgIN2CoAQoQDMBkAo1GdQgFWQpMwBC5kCYNw6lzqxSoSAEq67hDIEqy18QgY6irJUBtqXfUxT8WWrAo3IVMAx1FWK+IDqRotfEmYTsvUEqwgbL1AzFWEJQDcmYGSFCVY2CYVheZAmwMvUj8x1I2+YQo6siY3AepUSr5hNgUU8R1FXmBOo6sVy1J1CdmSCJVlTCCgAz5gKChMy76gAtR1ACgGeQ2ADYVddRXzAdRuOZNwKPUxqWrAdS9SeoCVQIqlq+YQAEQDfUpKvmK+ISu5H5iviw2EL1CJUJgVkFSNhLJCniYp+Ja+JECkFWNyUHUpKjPmBX5h+ZjVlq/EJEVETCYQvUdRVkzADqSor4gXMepE3zHUDLPmEQVdALQbakr4kqBUCVYTfMDIBeZK+I7BX5k9RUCQ3KE8xVgUlMxVhV5gQJAIBQNZFDbA//Z"
    return (
        f'<img src="data:image/png;base64,{_LOGO_DATA}" width="{width}" '
        f'style="display:block;" '
        f'alt="PlanTrace">'
    )


# -----------------------------------------------------------------------------
# Core component helpers
# -----------------------------------------------------------------------------

def chip(label: str, style: str = "grey") -> str:
    """Inline HTML status chip. style: red|amber|green|blue|grey|navy"""
    return f'<span class="chip chip-{style}">{label}</span>'


def float_chip(tf) -> str:
    if tf is None: return chip("-","grey")
    try: f = float(tf)
    except Exception: return chip(str(tf),"grey")
    if f < 0: return chip(f"{f}d","red")
    if f == 0: return chip("Critical","red")
    if f <= 10: return chip(f"{f}d","amber")
    return chip(f"{f}d","green")


def status_chip(s: str) -> str:
    m = {
        "TK_NotStart":"grey","Not Started":"grey",
        "TK_Active":"blue","In Progress":"blue",
        "TK_Complete":"green","Complete":"green",
    }
    labels = {
        "TK_NotStart":"Not Started","Not Started":"Not Started",
        "TK_Active":"In Progress","In Progress":"In Progress",
        "TK_Complete":"Complete","Complete":"Complete",
    }
    st_key = str(s).strip()
    return chip(labels.get(st_key, st_key or "-"), m.get(st_key,"grey"))


def kpi_card(label: str, value, sub: str = "",
             style: str = "navy") -> str:
    """
    Returns HTML for a premium KPI card.
    style: navy | red | amber | green | blue
    """
    border_cls = f"kpi-border-top-{style}"
    num_cls    = f"kpi-{style}"
    sub_html   = f'<div class="kpi-sub">{sub}</div>' if sub else ""
    return (
        f'<div class="kpi {num_cls} {border_cls}">'
        f'<div class="kpi-label">{label}</div>'
        f'<div class="kpi-num">{value}</div>'
        f'{sub_html}</div>'
    )


def kpi_row(items: list):
    """
    items = list of (label, value, sub, style)
    sub and style are optional.
    """
    cols = st.columns(len(items))
    for col, item in zip(cols, items):
        if len(item) == 2:
            lbl, val = item; sub, sty = "", "navy"
        elif len(item) == 3:
            lbl, val, sub = item; sty = "navy"
        else:
            lbl, val, sub, sty = item[:4]
        col.markdown(kpi_card(lbl, val, sub, sty), unsafe_allow_html=True)


# -----------------------------------------------------------------------------
# Control bar  (top of every page)
# -----------------------------------------------------------------------------

def ctrl_bar(title: str, description: str = ""):
    """Dark navy control bar at the top of the main content area."""
    prog_loaded = "programme" in st.session_state

    if prog_loaded:
        prog   = st.session_state["programme"]
        pname  = prog.get("project_info",{}).get("name","")
        ddate  = prog.get("project_info",{}).get("data_date")
        ntasks = len(prog.get("tasks_df",[]))
        nrels  = len(prog.get("relationships_df",[]))
        has_res = not prog.get("task_resources_df", pd.DataFrame()).empty

        dd_str = format_date(ddate) if ddate else "N/A"

        def meta(label, value, loaded=True):
            cls = "loaded" if loaded else ""
            return (
                f'<div class="ctrl-meta-item">'
                f'<div class="ctrl-meta-label">{label}</div>'
                f'<div class="ctrl-meta-value {cls}">{value}</div>'
                f'</div>'
            )

        meta_html = (
            meta("Programme",   pname[:28] if pname else "Loaded") +
            meta("Data Date",   dd_str) +
            meta("Activities",  f"{ntasks:,}") +
            meta("Rels",        f"{nrels:,}") +
            meta("Resources",   "Yes" if has_res else "None", has_res)
        )
    else:
        def meta_empty(label):
            return (f'<div class="ctrl-meta-item">'
                    f'<div class="ctrl-meta-label">{label}</div>'
                    f'<div class="ctrl-meta-value">-</div></div>')
        meta_html = (
            meta_empty("Programme") + meta_empty("Data Date") +
            meta_empty("Activities") + meta_empty("Rels") + meta_empty("Resources")
        )

    desc_html = (f'<div class="ctrl-bar-desc">{description}</div>'
                 if description else "")

    st.markdown(
        f'<div class="ctrl-bar">'
        f'<div class="ctrl-bar-left">'
        f'<div class="ctrl-bar-title">{title}</div>'
        f'{desc_html}</div>'
        f'<div class="ctrl-bar-meta">{meta_html}</div>'
        f'</div>',
        unsafe_allow_html=True,
    )


# -----------------------------------------------------------------------------
# Data quality card
# -----------------------------------------------------------------------------

def data_quality_card(data: dict):
    """Compact sidebar-style data quality summary card."""
    tasks = data.get("tasks_df", pd.DataFrame())
    rels  = data.get("relationships_df", pd.DataFrame())
    res   = data.get("task_resources_df", pd.DataFrame())
    cals  = data.get("calendars_df", pd.DataFrame())
    notes = bool(st.session_state.get("_notes_text", ""))
    comp  = "_mi_prev" in st.session_state

    items = [
        ("Activities",    len(tasks),                             not tasks.empty),
        ("Relationships", len(rels),                             not rels.empty),
        ("Resources",     len(res) if not res.empty else 0,     not res.empty),
        ("Calendars",     len(cals) if not cals.empty else 0,   not cals.empty),
        ("Notes",         "Yes" if notes else "No",              notes),
        ("Comparison",    "Yes" if comp else "No",               comp),
    ]
    ok_count = sum(1 for *_, ok in items if ok)
    overall_label, overall_style = (
        ("Good",    "green") if ok_count >= 5 else
        ("Partial", "amber") if ok_count >= 3 else
        ("Limited", "grey")
    )

    rows_html = ""
    for lbl, val, ok in items:
        dot_cls = "dot-green" if ok else "dot-grey"
        rows_html += (
            f'<div class="dq-row">'
            f'<span style="color:#374151;">{lbl}</span>'
            f'<span style="display:flex;align-items:center;gap:5px;">'
            f'<span style="font-weight:600;color:#071827;">{val}</span>'
            f'<span class="{dot_cls}"></span>'
            f'</span></div>'
        )

    st.markdown(
        f'<div class="card-flat" style="margin-top:0;">'
        f'<div style="display:flex;justify-content:space-between;'
        f'align-items:center;margin-bottom:10px;">'
        f'<span class="section-label" style="margin-bottom:0;">Data Quality</span>'
        f'{chip(overall_label, overall_style)}'
        f'</div>{rows_html}</div>',
        unsafe_allow_html=True,
    )


def programme_readiness_check(data: dict):
    """
    Programme Readiness Check panel.
    Shown on the Overview page immediately after upload.
    Checks what data was found in the XER and warns about limitations.
    """
    tasks = data.get("tasks_df", pd.DataFrame())
    rels  = data.get("relationships_df", pd.DataFrame())
    res   = data.get("task_resources_df", pd.DataFrame())
    cals  = data.get("calendars_df", pd.DataFrame())
    wbs   = data.get("wbs_df", pd.DataFrame())

    if tasks.empty:
        return

    # -- Evaluate each data element --------------------------------------------

    def _pct_with(col, tasks_df):
        """Percentage of tasks that have a non-null value for col."""
        if col not in tasks_df.columns:
            return 0.0
        n = len(tasks_df)
        if n == 0:
            return 0.0
        filled = tasks_df[col].apply(
            lambda v: v is not None and str(v).strip() not in ("", "None", "nan", "NaT")
        ).sum()
        return round(filled / n * 100, 0)

    n_tasks = len(tasks)
    n_rels  = len(rels)
    n_cals  = len(cals)
    n_wbs   = len(wbs)

    pct_start       = _pct_with("eff_start",        tasks)
    pct_finish      = _pct_with("eff_finish",        tasks)
    pct_total_float = _pct_with("total_float_days",  tasks)
    pct_free_float  = _pct_with("free_float_days",   tasks)
    pct_constraints = _pct_with("cstr_type",         tasks)
    has_res         = not res.empty
    has_labour      = has_res and "target_qty" in res.columns and res["target_qty"].apply(
        lambda v: safe_float(v, 0) > 0
    ).any()

    def _status(pct_or_bool, good_thresh=90, partial_thresh=50):
        """Return (label, style, dot_cls) for a percentage or boolean."""
        # Boolean values (e.g. has_res, has_labour) — must check first
        # because bool is a subclass of int in Python
        if isinstance(pct_or_bool, bool):
            if pct_or_bool:
                return "Found",   "green", "dot-green"
            return "Missing", "grey",  "dot-grey"
        # Count-based (good_thresh=None means: >0 is good, 0 is missing)
        if good_thresh is None:
            if pct_or_bool > 0:
                return f"{int(pct_or_bool):,}", "green", "dot-green"
            return "None found", "grey", "dot-grey"
        # Percentage-based
        try:
            val = float(pct_or_bool)
        except (TypeError, ValueError):
            return "-", "grey", "dot-grey"
        if val >= good_thresh:
            return f"{int(val)}%", "green", "dot-green"
        if val >= partial_thresh:
            return f"{int(val)}%", "amber", "dot-red"
        return f"{int(val)}%", "red", "dot-red"

    checks = [
        # (label, description, value, good_thresh, partial_thresh)
        ("Activities",      "programme activities extracted",       n_tasks,              None, None),
        ("Relationships",   "logic links between activities",       n_rels,               None, None),
        ("WBS",             "work breakdown structure nodes",       n_wbs,                None, None),
        ("Start Dates",     "activities with a start date",         pct_start,            90,   50),
        ("Finish Dates",    "activities with a finish date",        pct_finish,           90,   50),
        ("Total Float",     "activities with total float value",    pct_total_float,      85,   40),
        ("Free Float",      "activities with free float value",     pct_free_float,       85,   40),
        ("Calendars",       "calendars found in XER",               n_cals,               None, None),
        ("Constraints",     "constrained activities detected",      _pct_with("cstr_type", tasks), 0, 0),
        ("Resource Loading","resource assignments found",            has_res,             None, None),
        ("Labour Hours",    "quantified labour hours found",        has_labour,           None, None),
    ]

    # Compute statuses
    statuses = []
    for label, desc, val, good_t, partial_t in checks:
        if isinstance(val, bool):
            s_label, s_style, dot_cls = _status(val)
        elif isinstance(val, int):
            # Count-based: > 0 = found, 0 = missing
            if val > 0:
                s_label, s_style, dot_cls = f"{val:,}", "green", "dot-green"
            else:
                s_label, s_style, dot_cls = "None found", "grey", "dot-grey"
        else:
            # Percentage
            if label == "Constraints":
                # Constraints: any % is informational, not a quality issue
                s_label = f"{int(val)}%" if val > 0 else "None"
                s_style, dot_cls = "blue", "dot-green"
            else:
                s_label, s_style, dot_cls = _status(val, good_t, partial_t)
        statuses.append((label, desc, s_label, s_style, dot_cls))

    # -- Overall readiness score ------------------------------------------------
    core_ok = sum(1 for lbl, _, s_lbl, s_sty, _ in statuses
                  if lbl in ("Activities","Relationships","Start Dates","Finish Dates","Total Float")
                  and s_sty == "green")
    if core_ok >= 5:
        overall_label, overall_style, overall_desc = (
            "Good",
            "green",
            "All core schedule data found. PlanTrace is ready to use.",
        )
    elif core_ok >= 3:
        overall_label, overall_style, overall_desc = (
            "Partial",
            "amber",
            "Some core data is missing or incomplete. Some features may be limited.",
        )
    else:
        overall_label, overall_style, overall_desc = (
            "Poor",
            "red",
            "Critical programme data is missing. Check the XER export includes all required tables.",
        )

    # -- Render the panel ------------------------------------------------------
    st.markdown(
        '<div class="section-label">Programme Readiness Check</div>',
        unsafe_allow_html=True,
    )

    # Overall status strip
    strip_colour = {
        "green": "#F0FDF4", "amber": "#FFFBEB", "red": "#FEF2F2"
    }.get(overall_style, "#F9FAFB")
    strip_border = {
        "green": "#16A34A", "amber": "#F59E0B", "red": "#DC2626"
    }.get(overall_style, "#E5E7EB")

    st.markdown(
        f'<div style="background:{strip_colour};border:1px solid {strip_border};'
        f'border-radius:6px;padding:12px 16px;margin-bottom:14px;'
        f'display:flex;align-items:center;justify-content:space-between;gap:12px;">'
        f'<div>'
        f'<div style="font-size:13px;font-weight:700;color:#071827;margin-bottom:2px;">'
        f'Readiness: {chip(overall_label, overall_style)}'
        f'</div>'
        f'<div style="font-size:12px;color:#374151;">{overall_desc}</div>'
        f'</div>'
        f'<div style="font-size:11px;color:#6B7280;text-align:right;white-space:nowrap;">'
        f'{n_tasks:,} activities &nbsp;|&nbsp; {n_rels:,} relationships'
        f'</div>'
        f'</div>',
        unsafe_allow_html=True,
    )

    # Check rows in a 2-col grid
    col_left, col_right = st.columns(2, gap="medium")

    half = math.ceil(len(statuses) / 2)
    left_checks  = statuses[:half]
    right_checks = statuses[half:]

    def _render_checks(col, check_list):
        rows_html = ""
        for label, desc, s_label, s_style, dot_cls in check_list:
            chip_html = chip(s_label, s_style)
            rows_html += (
                f'<div style="display:flex;align-items:center;justify-content:space-between;'
                f'padding:8px 0;border-bottom:1px solid #F3F4F6;">'
                f'<div>'
                f'<div style="font-size:13px;font-weight:600;color:#111827;">{label}</div>'
                f'<div style="font-size:11px;color:#9CA3AF;">{desc}</div>'
                f'</div>'
                f'<div>{chip_html}</div>'
                f'</div>'
            )
        col.markdown(
            f'<div style="background:#FFFFFF;border:1px solid #E5E7EB;border-radius:8px;'
            f'padding:12px 16px;box-shadow:0 1px 3px rgba(7,24,39,0.05);">'
            f'{rows_html}</div>',
            unsafe_allow_html=True,
        )

    _render_checks(col_left,  left_checks)
    _render_checks(col_right, right_checks)

    # -- Targeted warnings ----------------------------------------------------
    warnings = []

    if n_rels == 0:
        warnings.append((
            "red",
            "No Relationships Found",
            "Logic tracing and critical path analysis require relationship data. "
            "Ensure the XER was exported with task relationships included. "
            "In P6: File > Export > ensure 'Relationships' is checked.",
        ))

    if pct_total_float < 50:
        warnings.append((
            "amber",
            "Total Float Data Limited",
            "Float values are missing for a significant number of activities. "
            "Critical path analysis, near-critical detection and risk ratings may be inaccurate. "
            "Ensure the programme has been scheduled in P6 before export.",
        ))

    if not has_res:
        warnings.append((
            "amber",
            "No Resource Loading Found",
            "Labour histograms require resource data. "
            "This XER does not contain resource assignments. "
            "You can upload a separate resource CSV on the Labour page, "
            "or re-export from P6 with resources included.",
        ))

    if n_wbs == 0:
        warnings.append((
            "amber",
            "WBS Data Not Found",
            "Work breakdown structure data is missing. "
            "WBS-level filtering, grouping and summary charts will not be available.",
        ))

    if pct_start < 50 or pct_finish < 50:
        warnings.append((
            "red",
            "Dates Missing for Many Activities",
            "A significant number of activities are missing start or finish dates. "
            "The programme may not be fully scheduled. "
            "Run a schedule calculation in P6 before exporting.",
        ))

    if warnings:
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown(
            '<div class="section-label">Readiness Warnings</div>',
            unsafe_allow_html=True,
        )
        for w_style, w_title, w_body in warnings:
            border_col = {"red": "#DC2626", "amber": "#F59E0B"}.get(w_style, "#E5E7EB")
            bg_col     = {"red": "#FEF2F2", "amber": "#FFFBEB"}.get(w_style, "#F9FAFB")
            icon       = {"red": "!", "amber": "!"}.get(w_style, "i")
            st.markdown(
                f'<div style="background:{bg_col};border:1px solid {border_col};'
                f'border-left:4px solid {border_col};border-radius:0 6px 6px 0;'
                f'padding:12px 16px;margin-bottom:8px;'
                f'display:flex;gap:12px;align-items:flex-start;">'
                f'<div style="font-size:14px;font-weight:800;color:{border_col};'
                f'flex-shrink:0;margin-top:1px;">{icon}</div>'
                f'<div>'
                f'<div style="font-size:13px;font-weight:700;color:#071827;margin-bottom:3px;">'
                f'{w_title}</div>'
                f'<div style="font-size:12px;color:#374151;line-height:1.6;">{w_body}</div>'
                f'</div></div>',
                unsafe_allow_html=True,
            )

    st.markdown("<br>", unsafe_allow_html=True)


def pm_attention_panel(data: dict, near_crit_days: float):
    """
    Recommended First Checks panel.
    Generates a numbered, prioritised list of specific actions for the PM,
    each with an exact count, a plain-English action verb and the relevant page.
    """
    tasks = data.get("tasks_df", pd.DataFrame())
    rels  = data.get("relationships_df", pd.DataFrame())
    res   = data.get("task_resources_df", pd.DataFrame())
    notes = st.session_state.get("_notes_text", "")

    if tasks.empty:
        return

    tasks_c = get_critical_threshold(tasks, near_crit_days)
    checks  = []   # list of (priority, count, action_text, why, page, style)
    # priority: 1=High 2=Medium 3=Low  (used for sorting)

    # 1. Negative float --------------------------------------------------
    if "total_float_days" in tasks_c.columns:
        neg = tasks_c[tasks_c["total_float_days"].apply(lambda f: safe_float(f, 0) < 0)]
        if not neg.empty:
            n = len(neg)
            worst = round(float(neg["total_float_days"].min()), 1)
            checks.append((
                1, n,
                f"Review {n} {'activity' if n==1 else 'activities'} with negative float",
                f"Worst case: {worst} days. These cannot meet their target dates without recovery action.",
                "Critical Path",
                "red",
            ))

    # 2. Critical activities not started --------------------------------
    if "is_critical" in tasks_c.columns and "status" in tasks_c.columns:
        cns = tasks_c[
            tasks_c["is_critical"] &
            tasks_c["status"].apply(lambda s: str(s) in ("TK_NotStart", "Not Started"))
        ]
        if not cns.empty:
            n = len(cns)
            checks.append((
                1, n,
                f"Check {n} critical {'activity' if n==1 else 'activities'} not yet started",
                "Any further delay to these will directly push out the project finish date.",
                "Critical Path",
                "red",
            ))

    # 3. Activities with no successor (open finish) ----------------------
    if not rels.empty and "task_id" in tasks_c.columns:
        with_succ = set(rels["pred_task_id"].dropna()) if "pred_task_id" in rels.columns else set()
        no_succ = tasks_c[~tasks_c["task_id"].isin(with_succ)]
        # Exclude finish milestones and LOE
        if "task_type" in no_succ.columns:
            no_succ = no_succ[~no_succ["task_type"].astype(str).str.contains(
                "Finish Milestone|LOE|WBS", na=False
            )]
        if not no_succ.empty:
            n = len(no_succ)
            checks.append((
                2, n,
                f"Check {n} {'activity' if n==1 else 'activities'} with missing successors",
                "Open-finish activities may carry artificially high float and can mask downstream risk.",
                "Health Check",
                "amber",
            ))

    # 4. Activities with no predecessor (open start) ---------------------
    if not rels.empty and "task_id" in tasks_c.columns:
        with_pred = set(rels["succ_task_id"].dropna()) if "succ_task_id" in rels.columns else set()
        no_pred = tasks_c[~tasks_c["task_id"].isin(with_pred)]
        if "task_type" in no_pred.columns:
            no_pred = no_pred[~no_pred["task_type"].astype(str).str.contains(
                "Start Milestone|LOE|WBS", na=False
            )]
        if not no_pred.empty:
            n = len(no_pred)
            checks.append((
                2, n,
                f"Check {n} {'activity' if n==1 else 'activities'} with missing predecessors",
                "Open-start activities are not driven by logic. Float values may be unreliable.",
                "Health Check",
                "amber",
            ))

    # 5. Excessive lag (> 10 days) ---------------------------------------
    if not rels.empty and "lag_days" in rels.columns:
        big_lag = rels[rels["lag_days"].apply(lambda l: safe_float(l, 0) > 10)]
        if not big_lag.empty:
            n = len(big_lag)
            max_lag = int(big_lag["lag_days"].max())
            checks.append((
                2, n,
                f"Review {n} {'relationship' if n==1 else 'relationships'} with excessive lag",
                f"Largest lag: {max_lag} days. Excessive lag can hide critical path issues and inflate float.",
                "Health Check",
                "amber",
            ))

    # 6. Near-critical finishing within 4 weeks -------------------------
    from datetime import timedelta
    four_wks = datetime.now() + timedelta(weeks=4)
    if "is_near_critical" in tasks_c.columns and "eff_finish" in tasks_c.columns:
        nc = tasks_c[
            tasks_c["is_near_critical"] &
            tasks_c["eff_finish"].apply(
                lambda d: d is not None and hasattr(d, "date") and d <= four_wks
            )
        ]
        if not nc.empty:
            n = len(nc)
            checks.append((
                1, n,
                f"Confirm {n} near-critical {'activity' if n==1 else 'activities'} finishing within 4 weeks",
                "Limited float and an imminent finish date -- any disruption will make these critical.",
                "Programme",
                "amber",
            ))

    # 7. Long duration activities (> 60 days) ----------------------------
    if "orig_dur_days" in tasks_c.columns:
        long_dur = tasks_c[tasks_c["orig_dur_days"].apply(lambda d: safe_float(d, 0) > 60)]
        if not long_dur.empty:
            n = len(long_dur)
            worst_d = int(tasks_c["orig_dur_days"].max())
            checks.append((
                3, n,
                f"Review {n} {'activity' if n==1 else 'activities'} with excessive duration",
                f"Longest: {worst_d} days. Activities over 60 days are hard to control and monitor.",
                "Health Check",
                "amber",
            ))

    # 8. Resource loading missing ----------------------------------------
    if res.empty:
        checks.append((
            2, 0,
            "Confirm resource loading is missing from this programme",
            "Labour histograms will not be available. Upload a resource CSV or re-export from P6 with resources.",
            "Labour",
            "amber",
        ))

    # 9. Constrained activities ------------------------------------------
    if "cstr_type" in tasks_c.columns:
        constrained = tasks_c[tasks_c["cstr_type"].apply(
            lambda x: bool(x) and str(x).strip() not in ("", "None", "nan")
        )]
        if not constrained.empty:
            n = len(constrained)
            checks.append((
                3, n,
                f"Confirm {n} constrained {'activity' if n==1 else 'activities'} are still valid",
                "Constraints override logic and can create artificial float or negative float.",
                "Health Check",
                "amber",
            ))

    # 10. Planning notes risk keywords ----------------------------------
    if notes and "task_code" in tasks_c.columns:
        import re as _re
        risk_acts = []
        for _, t in tasks_c.head(300).iterrows():
            code = str(t.get("task_code", ""))
            if not code or code not in notes:
                continue
            idx = notes.find(code)
            snippet = notes[max(0, idx-200):idx+200]
            for w in _RISK_WORDS:
                if _re.search(r"\b" + _re.escape(w) + r"\b", snippet, _re.IGNORECASE):
                    risk_acts.append(code)
                    break
        if risk_acts:
            n = len(risk_acts)
            checks.append((
                1, n,
                f"Investigate {n} {'activity' if n==1 else 'activities'} flagged in planning notes",
                "Risk keywords (delay, blocked, CE, EWN etc.) found in notes against these activities.",
                "Programme",
                "red",
            ))

    # 10b. Activities starting in the next 2-4 weeks ---------------------
    if "eff_start" in tasks_c.columns:
        from datetime import timedelta as _td
        now_dt     = datetime.now()
        two_wks    = now_dt + _td(weeks=2)
        four_wks   = now_dt + _td(weeks=4)
        starting   = tasks_c[
            tasks_c["eff_start"].apply(
                lambda d: d is not None and hasattr(d, "date")
                          and now_dt <= d <= four_wks
            )
        ]
        if not starting.empty:
            n = len(starting)
            crit_starting = int(starting["is_critical"].sum()) if "is_critical" in starting.columns else 0
            crit_note = f" ({crit_starting} critical)" if crit_starting else ""
            checks.append((
                2, n,
                f"Confirm {n} {'activity is' if n==1 else 'activities are'} due to start in the next 4 weeks",
                f"These activities should be mobilised now{crit_note}. Confirm readiness with the delivery team.",
                "Programme",
                "blue",
            ))

    # 11. Comparison: major slips ----------------------------------------
    if "_mi_prev" in st.session_state and "_mi_curr" in st.session_state:
        try:
            prev_t = st.session_state["_mi_prev"]["tasks_df"]
            curr_t = st.session_state["_mi_curr"]["tasks_df"]
            if not prev_t.empty and not curr_t.empty:
                mg = prev_t[["task_code", "eff_finish"]].merge(
                    curr_t[["task_code", "eff_finish"]], on="task_code",
                    suffixes=("_p", "_c"), how="inner"
                )
                big_slips = []
                for _, r in mg.iterrows():
                    try:
                        d = int((pd.Timestamp(r["eff_finish_c"]) -
                                 pd.Timestamp(r["eff_finish_p"])).days)
                        if d > 14:
                            big_slips.append(d)
                    except Exception:
                        pass
                if big_slips:
                    n = len(big_slips)
                    worst_slip = max(big_slips)
                    checks.append((
                        1, n,
                        f"Review {n} {'activity' if n==1 else 'activities'} with major date movement",
                        f"Worst slip: {worst_slip} days vs previous revision. Understand causes and agree recovery.",
                        "Comparison",
                        "red",
                    ))
        except Exception:
            pass

    # -- Sort by priority then count (highest count first within priority) -
    checks.sort(key=lambda x: (x[0], -x[1]))

    # -- Render ----------------------------------------------------------------
    if not checks:
        st.markdown(
            f'<div class="card-flat" style="padding:14px 18px;">'
            f'<div class="section-label" style="margin-bottom:8px;">'
            f'Recommended First Checks</div>'
            f'<div style="font-size:13px;color:#6B7280;display:flex;'
            f'align-items:center;gap:8px;">'
            f'{chip("No issues detected", "green")}'
            f'No immediate actions identified from this programme.'
            f'</div></div>',
            unsafe_allow_html=True,
        )
        return

    # Header row
    st.markdown(
        f'<div style="display:flex;align-items:baseline;'
        f'justify-content:space-between;margin-bottom:10px;">'
        f'<div class="section-label" style="margin-bottom:0;">'
        f'Recommended First Checks</div>'
        f'<div style="font-size:11px;color:#9CA3AF;">'
        f'{len(checks)} item{"s" if len(checks)!=1 else ""} identified'
        f'</div></div>',
        unsafe_allow_html=True,
    )

    # Numbered check rows
    rows_html = ""
    for i, (priority, count, action, why, page, style) in enumerate(checks, start=1):
        border_col = {"red": "#DC2626", "amber": "#F59E0B", "blue": "#1D4ED8"}.get(style, "#E5E7EB")
        bg_col     = {"red": "#FEF2F2", "amber": "#FFFBEB", "blue": "#EFF6FF"}.get(style, "#F9FAFB")
        num_col    = {"red": "#DC2626", "amber": "#B45309", "blue": "#1D4ED8"}.get(style, "#6B7280")
        pri_label  = {1: "High", 2: "Medium", 3: "Low"}.get(priority, "Low")
        pri_chip   = {1: "red",  2: "amber",  3: "grey"}.get(priority, "grey")

        # Count badge (only show if count > 0)
        count_badge = (
            f'<span style="background:{border_col};color:white;'
            f'border-radius:4px;padding:1px 7px;font-size:12px;font-weight:700;'
            f'margin-left:8px;">{count}</span>'
            if count > 0 else ""
        )

        rows_html += (
            f'<div style="display:flex;gap:14px;align-items:flex-start;'
            f'padding:11px 14px;border-bottom:1px solid #F3F4F6;'
            f'background:#FFFFFF;">'
            # Number circle
            f'<div style="width:24px;height:24px;border-radius:50%;'
            f'background:{border_col};color:white;font-size:12px;font-weight:800;'
            f'display:flex;align-items:center;justify-content:center;'
            f'flex-shrink:0;margin-top:1px;">{i}</div>'
            # Content
            f'<div style="flex:1;min-width:0;">'
            f'<div style="font-size:13px;font-weight:600;color:#071827;'
            f'margin-bottom:3px;">{action}{count_badge}</div>'
            f'<div style="font-size:12px;color:#6B7280;line-height:1.5;">{why}</div>'
            f'</div>'
            # Right side: priority chip + page label
            f'<div style="display:flex;flex-direction:column;align-items:flex-end;'
            f'gap:4px;flex-shrink:0;">'
            f'{chip(pri_label, pri_chip)}'
            f'<span style="font-size:10px;color:#9CA3AF;">-> {page}</span>'
            f'</div>'
            f'</div>'
        )

    st.markdown(
        f'<div style="background:#FFFFFF;border:1px solid #E5E7EB;'
        f'border-radius:8px;overflow:hidden;'
        f'box-shadow:0 1px 3px rgba(7,24,39,0.06);">'
        f'{rows_html}</div>',
        unsafe_allow_html=True,
    )
    st.markdown("<br>", unsafe_allow_html=True)


def empty_state(icon: str, title: str, body: str, cta: str = ""):
    cta_html = (
        f'<div style="margin-top:16px;display:inline-block;background:{_NAVY};'
        f'color:#FFFFFF;border-radius:6px;padding:8px 20px;font-size:13px;'
        f'font-weight:600;letter-spacing:0.1px;">{cta}</div>'
    ) if cta else ""
    st.markdown(
        f'<div class="empty-state">'
        f'<div style="font-size:32px;margin-bottom:12px;">{icon}</div>'
        f'<div style="font-size:16px;font-weight:700;color:#071827;margin-bottom:8px;">{title}</div>'
        f'<div style="font-size:13px;color:#6B7280;max-width:420px;margin:0 auto;'
        f'line-height:1.6;">{body}</div>{cta_html}</div>',
        unsafe_allow_html=True,
    )


# -----------------------------------------------------------------------------
# LANDING PAGE  --  control centre style
# -----------------------------------------------------------------------------

def _landing_page():
    """Professional control-centre landing page when no programme is loaded."""

    # Status strip
    def status_item(label, status, ok):
        dot_cls = "dot-green" if ok else "dot-grey"
        val_col = "#16A34A" if ok else "#6B7280"
        return (
            f'<div style="display:flex;align-items:center;gap:6px;padding:8px 14px;'
            f'background:#0B2438;border-radius:5px;white-space:nowrap;">'
            f'<span class="{dot_cls}"></span>'
            f'<div>'
            f'<div style="font-size:10px;color:#3D5268;text-transform:uppercase;'
            f'letter-spacing:0.8px;font-weight:700;">{label}</div>'
            f'<div style="font-size:12px;font-weight:700;color:{val_col};">{status}</div>'
            f'</div></div>'
        )

    status_html = (
        status_item("Programme",  "Not Loaded",  False) +
        status_item("Relationships","Waiting",   False) +
        status_item("Resources",  "Waiting",     False) +
        status_item("Notes",      "Not Loaded",  False) +
        status_item("Comparison", "Not Loaded",  False)
    )

    logo_html = _logo_b64(width=44)

    st.markdown(
        f"""
        <div style="background:{_NAVY};padding:18px 28px;margin:-28px -28px 0 -28px;
                    border-bottom:2px solid {_AMBER};">
            <div style="display:flex;align-items:center;gap:14px;margin-bottom:12px;">
                {logo_html}
                <div>
                    <div style="font-size:22px;font-weight:800;color:#FFFFFF;
                                letter-spacing:-0.3px;line-height:1;">PlanTrace</div>
                    <div style="font-size:12px;color:#4B6478;margin-top:2px;">
                        Planning intelligence for project delivery teams</div>
                </div>
            </div>
            <div style="display:flex;gap:8px;flex-wrap:wrap;">{status_html}</div>
        </div>
        <div style="height:28px;"></div>
        """,
        unsafe_allow_html=True,
    )

    # Left-aligned hero
    col_hero, col_upload = st.columns([3, 2], gap="large")

    with col_hero:
        st.markdown(
            f"""
            <div style="padding-top:4px;">
                <div style="font-size:11px;font-weight:700;color:{_AMBER};
                            text-transform:uppercase;letter-spacing:1.5px;
                            margin-bottom:10px;">PlanTrace Control Centre</div>
                <div style="font-size:32px;font-weight:900;color:{_NAVY};
                            line-height:1.1;letter-spacing:-0.5px;margin-bottom:14px;">
                    Programme intelligence<br>without opening P6.
                </div>
                <div style="font-size:14px;color:{_MUTED};line-height:1.7;
                            max-width:480px;margin-bottom:24px;">
                    Upload a Primavera P6 XER programme to interrogate logic, critical paths,
                    labour demand, schedule quality and programme movement. No P6 licence required.
                </div>
                <div style="display:flex;gap:10px;flex-wrap:wrap;margin-bottom:28px;">
                    {chip("Logic Trace","navy")}
                    {chip("Critical Path","navy")}
                    {chip("Labour Demand","navy")}
                    {chip("Health Check","navy")}
                    {chip("PM Actions","navy")}
                    {chip("Risk Register","navy")}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with col_upload:
        st.markdown(
            f"""
            <div style="background:{_NAVY};border-radius:8px;padding:24px;
                        border:1px solid #0B2438;">
                <div style="font-size:10px;font-weight:700;color:{_AMBER};
                            text-transform:uppercase;letter-spacing:1.2px;
                            margin-bottom:12px;">Upload Programme</div>
                <div style="font-size:13px;color:#4B6478;margin-bottom:16px;line-height:1.6;">
                    Export from Primavera P6 via<br>
                    <strong style="color:#94A3B8;">File &rarr; Export &rarr; Primavera P6 XER</strong>
                </div>
                <div style="font-size:12px;color:#4B6478;padding:10px 12px;
                            background:#0B2438;border-radius:5px;border:1px dashed #1e3a5f;">
                    Use the file uploader in the sidebar on the left to upload your .xer file.
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    # Module cards grid
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown(
        '<div class="section-label">Product Modules</div>',
        unsafe_allow_html=True,
    )

    modules = [
        (_NAVY,  "Logic",         "Logic Trace",
         "Trace predecessor and successor chains through the schedule network with depth, link type and lag.",
         ["Predecessor chains","Successor chains","Full logic chain","Network depth levels","Excel export"]),
        (_RED,   "Critical Path", "Critical Path Analysis",
         "Identify the critical path, near-critical work and negative float across the programme.",
         ["Full critical path","Near-critical work","Negative float","Driving path to activity","Gantt view"]),
        ("#5B21B6","Labour",       "Labour Demand",
         "View labour histograms by week, month, resource and WBS area.",
         ["Weekly histogram","Monthly histogram","Resource breakdown","WBS breakdown","Peak demand"]),
        (_GREEN, "Health",        "Programme Health Check",
         "Automated quality checks covering logic, constraints, durations and float.",
         ["Open logic detection","Constraint review","Long durations","Float issues","11 automated checks"]),
    ]

    cols = st.columns(4, gap="medium")
    for col, (colour, tag, title, desc, outputs) in zip(cols, modules):
        outputs_html = "".join(
            f'<div style="display:flex;align-items:center;gap:6px;padding:3px 0;'
            f'font-size:12px;color:#4B5563;border-bottom:1px solid #F3F4F6;">'
            f'<span style="color:{colour};font-size:9px;">&#9654;</span>{o}</div>'
            for o in outputs
        )
        col.markdown(
            f'<div style="background:#FFFFFF;border:1px solid #E5E7EB;border-radius:8px;'
            f'border-top:3px solid {colour};padding:0;overflow:hidden;'
            f'box-shadow:0 1px 4px rgba(7,24,39,0.07);height:100%;">'
            f'<div style="padding:16px 18px 14px 18px;">'
            f'<div style="font-size:10px;font-weight:700;color:{colour};text-transform:uppercase;'
            f'letter-spacing:1px;margin-bottom:6px;">{tag}</div>'
            f'<div style="font-size:15px;font-weight:700;color:#071827;margin-bottom:6px;">{title}</div>'
            f'<div style="font-size:12px;color:#6B7280;line-height:1.5;margin-bottom:12px;">{desc}</div>'
            f'<div style="font-size:10px;font-weight:700;color:#9CA3AF;text-transform:uppercase;'
            f'letter-spacing:0.8px;margin-bottom:6px;">Outputs</div>'
            f'<div>{outputs_html}</div>'
            f'</div></div>',
            unsafe_allow_html=True,
        )

    # Full module list at bottom
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="section-label">All Modules</div>', unsafe_allow_html=True)

    page_list = [
        ("Overview",       "KPI summary, PM attention panel and schedule health overview."),
        ("Programme",      "Activity search, lookahead planner, milestone tracker and planning notes."),
        ("Logic",          "Logic trace and critical path to selected activity."),
        ("Critical Path",  "Critical path, near-critical and negative float by WBS."),
        ("Labour",         "Labour histograms by week, month, resource and WBS."),
        ("Health Check",   "Eleven automated schedule quality checks."),
        ("Comparison",     "Programme movement and risk & opportunity register."),
        ("PM Actions",     "Auto-generated prioritised action list."),
        ("Risk Register",  "Auto-generated risk and opportunity register."),
        ("Reports",        "Export all data to formatted Excel workbooks."),
    ]

    lc, rc = st.columns(2, gap="large")
    for i, (mod_title, desc) in enumerate(page_list):
        col = lc if i % 2 == 0 else rc
        col.markdown(
            f'<div style="display:flex;gap:12px;padding:8px 0;'
            f'border-bottom:1px solid #F3F4F6;align-items:flex-start;">'
            f'<div style="width:4px;height:16px;background:{_NAVY};border-radius:2px;'
            f'margin-top:2px;flex-shrink:0;"></div>'
            f'<div><div style="font-weight:600;color:{_NAVY};font-size:13px;">{mod_title}</div>'
            f'<div style="font-size:12px;color:#6B7280;line-height:1.5;">{desc}</div>'
            f'</div></div>',
            unsafe_allow_html=True,
        )

    st.markdown(
        f'<div style="margin-top:32px;padding-top:16px;border-top:1px solid #E5E7EB;">'
        f'<span style="font-size:11px;color:#9CA3AF;">PlanTrace &nbsp;|&nbsp; '
        f'Planning intelligence for project delivery &nbsp;|&nbsp; '
        f'No Primavera P6 licence required</span></div>',
        unsafe_allow_html=True,
    )



# -----------------------------------------------------------------------------
# REUSABLE: SELECTED ACTIVITY PANEL
# -----------------------------------------------------------------------------

_SAP_KEY = "selected_activity_id"   # session state key for cross-page persistence


def _get_float_status(tf) -> tuple:
    """Return (label, chip_style) for a float value."""
    if tf is None:
        return "-", "grey"
    try:
        f = float(tf)
    except Exception:
        return str(tf), "grey"
    if f < 0:
        return f"{f}d  Negative Float", "red"
    if f == 0:
        return "0d  Critical", "red"
    if f <= 10:
        return f"{f}d  Near-Critical", "amber"
    return f"{f}d", "green"


def _sap_field(label: str, value, suffix: str = "") -> str:
    """Render a single labelled field row."""
    if value is None or str(value).strip() in ("", "None", "nan", "NaT", "-"):
        val_html = '<span style="color:#9CA3AF;font-size:12px;">Not available</span>'
    else:
        val_html = (
            f'<span style="font-weight:600;color:#111827;font-size:13px;">'
            f'{value}{suffix}</span>'
        )
    return (
        f'<div style="padding:6px 0;border-bottom:1px solid #F3F4F6;">'
        f'<div style="font-size:10px;font-weight:700;color:#9CA3AF;'
        f'text-transform:uppercase;letter-spacing:0.8px;margin-bottom:3px;">{label}</div>'
        f'{val_html}</div>'
    )


def render_selected_activity_panel(
    tasks_df: pd.DataFrame,
    rels_df:  pd.DataFrame,
    near_crit_days: float = 10.0,
    context_key: str = "default",
):
    """
    Reusable Selected Activity Summary Panel.

    Displays a consistent activity detail card with status chips and
    action buttons. Persists the selected activity in session state so
    it is available across pages.

    Parameters
    ----------
    tasks_df       : full tasks DataFrame from the loaded programme
    rels_df        : relationships DataFrame
    near_crit_days : near-critical float threshold
    context_key    : unique suffix to avoid widget key collisions across pages
    """
    if tasks_df.empty:
        empty_state(
            "",
            "No Activities Available",
            "Upload a programme to enable activity selection.",
        )
        return

    tasks_c = get_critical_threshold(tasks_df, near_crit_days)
    # Guard: deduplicate on task_id to prevent 'column label not unique' errors
    if 'task_id' in tasks_c.columns:
        tasks_c = tasks_c.drop_duplicates(subset='task_id').reset_index(drop=True)

    # -- Activity selector ----------------------------------------------------
    def _label(r):
        code = str(r.get("task_code", "?"))
        name = str(r.get("task_name", "?"))
        tf   = r.get("total_float_days")
        try:
            f = float(tf) if tf is not None else None
        except Exception:
            f = None
        flag = "  [CRITICAL]" if f is not None and f <= 0 else ""
        return f"{code}  -  {name}{flag}"

    act_labels = tasks_c.apply(_label, axis=1).tolist()

    # Restore previously selected activity if available
    saved_id = st.session_state.get(_SAP_KEY)
    default_idx = 0
    if saved_id is not None and "task_id" in tasks_c.columns:
        matches = tasks_c[tasks_c["task_id"] == saved_id]
        if not matches.empty:
            pos = tasks_c.index.get_loc(matches.index[0])
            default_idx = pos

    sel_label = st.selectbox(
        "Select activity",
        options=act_labels,
        index=default_idx,
        key=f"sap_selector_{context_key}",
        label_visibility="collapsed",
    )

    sel_idx  = act_labels.index(sel_label)
    row      = tasks_c.iloc[sel_idx]
    task_id  = row.get("task_id", "")

    # Persist to session state
    st.session_state[_SAP_KEY] = task_id

    # -- Extract all fields (safe) ---------------------------------------------
    def _get(col, default=None):
        if col in row.index:
            v = row[col]
            if v is None or (isinstance(v, float) and math.isnan(v)):
                return default
            return v
        return default

    task_code  = str(_get("task_code", "-"))
    task_name  = str(_get("task_name", "-"))
    wbs_path   = str(_get("wbs_path",  "-"))
    task_type  = str(_get("task_type", "-"))
    calendar   = str(_get("calendar",  "-"))
    status_raw = str(_get("status",    ""))
    pct        = _get("phys_pct")
    is_crit    = bool(_get("is_critical", False))
    is_nc      = bool(_get("is_near_critical", False))
    tf_raw     = _get("total_float_days")
    ff_raw     = _get("free_float_days")
    orig_dur   = _get("orig_dur_days")
    rem_dur    = _get("rem_dur_days")
    early_s    = format_date(_get("early_start"))
    early_f    = format_date(_get("early_finish"))
    late_s     = format_date(_get("late_start"))
    late_f     = format_date(_get("late_finish"))
    act_s      = format_date(_get("act_start"))
    act_f      = format_date(_get("act_finish"))
    eff_s      = format_date(_get("eff_start"))
    eff_f      = format_date(_get("eff_finish"))
    cstr_type  = str(_get("cstr_type", "") or "")
    cstr_date  = format_date(_get("cstr_date"))

    try:
        tf_num = float(tf_raw) if tf_raw is not None else None
    except Exception:
        tf_num = None
    try:
        ff_num = float(ff_raw) if ff_raw is not None else None
    except Exception:
        ff_num = None

    tf_label, tf_style = _get_float_status(tf_num)
    s_label    = _status_label(status_raw)
    s_style    = {
        "Not Started": "grey", "TK_NotStart": "grey",
        "In Progress": "blue", "TK_Active":   "blue",
        "Complete":    "green","TK_Complete":  "green",
    }.get(status_raw, "grey")

    # -- Summary banner --------------------------------------------------------
    crit_chip_html = (
        f'&nbsp;{chip("Critical", "red")}' if is_crit else
        (f'&nbsp;{chip("Near-Critical", "amber")}' if is_nc else "")
    )
    neg_chip_html = (
        f'&nbsp;{chip("Negative Float", "red")}' if tf_num is not None and tf_num < 0 else ""
    )

    st.markdown(
        f"""
        <div style="background:#071827;border-radius:8px;padding:16px 20px;
                    margin-bottom:14px;border-left:3px solid #E8951D;">
            <div style="font-size:11px;font-weight:700;color:#4B6478;
                        text-transform:uppercase;letter-spacing:1px;margin-bottom:5px;">
                Selected Activity
            </div>
            <div style="font-size:18px;font-weight:800;color:#FFFFFF;
                        letter-spacing:-0.2px;line-height:1.1;margin-bottom:4px;">
                {task_code}
            </div>
            <div style="font-size:13px;color:#94A3B8;margin-bottom:10px;
                        white-space:nowrap;overflow:hidden;text-overflow:ellipsis;"
                 title="{task_name}">{task_name}</div>
            <div style="display:flex;flex-wrap:wrap;gap:6px;align-items:center;">
                {chip(s_label, s_style)}
                {chip(tf_label, tf_style)}
                {crit_chip_html}
                {neg_chip_html}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # -- Two-column detail grid ------------------------------------------------
    col_left, col_right = st.columns(2, gap="medium")

    with col_left:
        st.markdown(
            '<div style="font-size:10px;font-weight:700;color:#9CA3AF;'
            'text-transform:uppercase;letter-spacing:1px;margin-bottom:8px;">'
            'Identity & Schedule</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            _sap_field("Activity ID",       task_code) +
            _sap_field("Activity Name",     task_name) +
            _sap_field("WBS",               wbs_path) +
            _sap_field("Activity Type",     task_type) +
            _sap_field("Calendar",          calendar) +
            _sap_field("Start",             eff_s) +
            _sap_field("Finish",            eff_f) +
            _sap_field("Early Start",       early_s) +
            _sap_field("Early Finish",      early_f) +
            _sap_field("Late Start",        late_s) +
            _sap_field("Late Finish",       late_f),
            unsafe_allow_html=True,
        )

    with col_right:
        st.markdown(
            '<div style="font-size:10px;font-weight:700;color:#9CA3AF;'
            'text-transform:uppercase;letter-spacing:1px;margin-bottom:8px;">'
            'Float & Duration</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            _sap_field("Status",             s_label) +
            _sap_field("% Complete",         pct,    suffix="%") +
            _sap_field("Total Float",        f"{tf_num}d" if tf_num is not None else None) +
            _sap_field("Free Float",         f"{ff_num}d" if ff_num is not None else None) +
            _sap_field("Original Duration",  orig_dur, suffix="d") +
            _sap_field("Remaining Duration", rem_dur,  suffix="d") +
            _sap_field("Actual Start",       act_s) +
            _sap_field("Actual Finish",      act_f),
            unsafe_allow_html=True,
        )

    # Constraint warning
    if cstr_type and cstr_type.strip() not in ("", "None", "nan"):
        st.markdown(
            f'<div style="background:#FFFBEB;border:1px solid #FDE68A;'
            f'border-left:3px solid #F59E0B;border-radius:0 6px 6px 0;'
            f'padding:9px 14px;margin-top:10px;">'
            f'<span style="font-size:11px;font-weight:700;color:#B45309;">Constraint:</span>'
            f'<span style="font-size:12px;color:#374151;margin-left:6px;">'
            f'{cstr_type}</span>'
            f'<span style="font-size:11px;color:#9CA3AF;margin-left:8px;">{cstr_date}</span>'
            f'</div>',
            unsafe_allow_html=True,
        )

    # -- Direct predecessors / successors summary ------------------------------
    if not rels_df.empty and task_id:
        n_preds = len(rels_df[rels_df["succ_task_id"] == task_id]) if "succ_task_id" in rels_df.columns else 0
        n_succs = len(rels_df[rels_df["pred_task_id"] == task_id]) if "pred_task_id" in rels_df.columns else 0

        warn_open = []
        if n_preds == 0:
            warn_open.append("no predecessors")
        if n_succs == 0:
            warn_open.append("no successors")

        st.markdown(
            f'<div style="display:flex;gap:10px;margin-top:10px;flex-wrap:wrap;">'
            f'<div style="background:#F9FAFB;border:1px solid #E5E7EB;border-radius:6px;'
            f'padding:8px 14px;flex:1;min-width:100px;text-align:center;">'
            f'<div style="font-size:10px;font-weight:700;color:#9CA3AF;'
            f'text-transform:uppercase;letter-spacing:0.8px;">Predecessors</div>'
            f'<div style="font-size:20px;font-weight:800;color:#071827;margin-top:3px;">{n_preds}</div>'
            f'</div>'
            f'<div style="background:#F9FAFB;border:1px solid #E5E7EB;border-radius:6px;'
            f'padding:8px 14px;flex:1;min-width:100px;text-align:center;">'
            f'<div style="font-size:10px;font-weight:700;color:#9CA3AF;'
            f'text-transform:uppercase;letter-spacing:0.8px;">Successors</div>'
            f'<div style="font-size:20px;font-weight:800;color:#071827;margin-top:3px;">{n_succs}</div>'
            f'</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

        if warn_open:
            st.markdown(
                f'<div style="background:#FFFBEB;border:1px solid #FDE68A;'
                f'border-radius:6px;padding:8px 12px;margin-top:8px;">'
                f'<span style="font-size:12px;color:#B45309;">Open logic: '
                f'This activity has {" and ".join(warn_open)}.</span></div>',
                unsafe_allow_html=True,
            )

    # -- Action buttons --------------------------------------------------------
    st.markdown(
        '<div style="font-size:10px;font-weight:700;color:#9CA3AF;'
        'text-transform:uppercase;letter-spacing:1px;margin:14px 0 8px 0;">Actions</div>',
        unsafe_allow_html=True,
    )

    btn_row1 = st.columns(2, gap="small")
    btn_row2 = st.columns(2, gap="small")
    btn_row3 = st.columns(2, gap="small")

    if btn_row1[0].button("Trace Predecessors", key=f"sap_pred_{context_key}",
                          ):
        st.session_state["nav_page"]           = "logic"
        st.session_state[_SAP_KEY]             = task_id
        st.session_state["logic_trace_selector"] = sel_label
        st.session_state["btn_ap"]             = True
        st.rerun()

    if btn_row1[1].button("Trace Successors", key=f"sap_succ_{context_key}",
                          ):
        st.session_state["nav_page"]           = "logic"
        st.session_state[_SAP_KEY]             = task_id
        st.session_state["logic_trace_selector"] = sel_label
        st.session_state["btn_as"]             = True
        st.rerun()

    if btn_row2[0].button("Show Driving Path", key=f"sap_drive_{context_key}",
                          ):
        st.session_state["nav_page"]   = "logic"
        st.session_state[_SAP_KEY]     = task_id
        st.session_state["cpta_selector"] = sel_label
        st.rerun()

    if btn_row2[1].button("Add to PM Actions", key=f"sap_pm_{context_key}",
                          ):
        # Inject a manual action into the PM Actions list
        existing = st.session_state.get("_pm_actions_df", pd.DataFrame())
        new_row = pd.DataFrame([{
            "Type":             "Risk",
            "Priority":         "High" if is_crit else "Medium",
            "Category":         "Manual Review",
            "Activity ID":      task_code,
            "Activity Name":    task_name,
            "WBS":              wbs_path.split(" > ")[0] if wbs_path else "-",
            "Issue":            f"Manually added for review from Activity Panel.",
            "Why It Matters":   f"Total float: {tf_num}d. Status: {s_label}.",
            "Mitigation / Action": "Review this activity with the delivery team.",
            "Suggested Action": "Review and update the programme accordingly.",
            "Owner":            "",
            "Due Date":         "",
            "Status":           "Open",
        }])
        st.session_state["_pm_actions_df"] = (
            pd.concat([existing, new_row], ignore_index=True)
            if not existing.empty else new_row
        )
        st.success(f"Added {task_code} to PM Actions.")

    # Export pack
    export_data = {
        "Activity ID":        task_code,
        "Activity Name":      task_name,
        "WBS":                wbs_path,
        "Activity Type":      task_type,
        "Calendar":           calendar,
        "Status":             s_label,
        "% Complete":         pct,
        "Start":              eff_s,
        "Finish":             eff_f,
        "Early Start":        early_s,
        "Early Finish":       early_f,
        "Late Start":         late_s,
        "Late Finish":        late_f,
        "Actual Start":       act_s,
        "Actual Finish":      act_f,
        "Original Duration":  f"{orig_dur}d" if orig_dur is not None else "-",
        "Remaining Duration": f"{rem_dur}d"  if rem_dur  is not None else "-",
        "Total Float":        f"{tf_num}d"   if tf_num   is not None else "-",
        "Free Float":         f"{ff_num}d"   if ff_num   is not None else "-",
        "Critical":           "Yes" if is_crit else "No",
        "Constraint":         cstr_type or "-",
        "Constraint Date":    cstr_date,
    }
    detail_df = pd.DataFrame([export_data])

    # Predecessors/successors for export
    export_sheets = {"Activity Detail": detail_df}
    if not rels_df.empty and task_id:
        if "succ_task_id" in rels_df.columns:
            preds = rels_df[rels_df["succ_task_id"] == task_id]
            if not preds.empty:
                pred_cols = [c for c in ["pred_task_code","pred_task_name","rel_type","lag_days"] if c in preds.columns]
                export_sheets["Predecessors"] = preds[pred_cols] if pred_cols else preds
        if "pred_task_id" in rels_df.columns:
            succs = rels_df[rels_df["pred_task_id"] == task_id]
            if not succs.empty:
                succ_cols = [c for c in ["succ_task_code","succ_task_name","rel_type","lag_days"] if c in succs.columns]
                export_sheets["Successors"] = succs[succ_cols] if succ_cols else succs

    xls = export_df_to_excel(export_sheets)

    btn_row3[0].download_button(
        label="Export Activity Pack",
        data=xls,
        file_name=f"activity_{task_code}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key=f"sap_export_{context_key}",
    )

    if btn_row3[1].button("View in Logic Trace", key=f"sap_logic_{context_key}",
                          ):
        st.session_state["nav_page"]             = "logic"
        st.session_state[_SAP_KEY]               = task_id
        st.session_state["logic_trace_selector"] = sel_label
        st.rerun()




# -----------------------------------------------------------------------------
# PAGE: EXECUTIVE SUMMARY
# -----------------------------------------------------------------------------

def _exec_narrative(data: dict, near_crit_days: float) -> dict:
    """
    Analyse programme data and return a dict of narrative components.
    All text is written in plain English for senior stakeholders.
    """
    tasks = data.get("tasks_df", pd.DataFrame())
    rels  = data.get("relationships_df", pd.DataFrame())
    res   = data.get("task_resources_df", pd.DataFrame())
    proj  = data.get("project_info", {})

    if tasks.empty:
        return {}

    tasks_c = get_critical_threshold(tasks, near_crit_days)
    now_dt  = datetime.now()

    n_total = len(tasks_c)
    n_rels  = len(rels)

    # -- Float metrics ---------------------------------------------------------
    neg_df   = tasks_c[tasks_c["total_float_days"].apply(lambda f: safe_float(f,0) < 0)] \
                if "total_float_days" in tasks_c.columns else pd.DataFrame()
    crit_df  = tasks_c[tasks_c["is_critical"] == True] \
                if "is_critical" in tasks_c.columns else pd.DataFrame()
    nc_df    = tasks_c[tasks_c["is_near_critical"] == True] \
                if "is_near_critical" in tasks_c.columns else pd.DataFrame()

    n_neg    = len(neg_df)
    n_crit   = len(crit_df)
    n_nc     = len(nc_df)
    worst_float = round(float(neg_df["total_float_days"].min()), 1) \
                  if not neg_df.empty and "total_float_days" in neg_df.columns else None

    # -- Logic metrics ---------------------------------------------------------
    n_open_start = n_open_end = 0
    if not rels.empty and "task_id" in tasks_c.columns:
        wp = set(rels["succ_task_id"].dropna()) if "succ_task_id" in rels.columns else set()
        ws = set(rels["pred_task_id"].dropna()) if "pred_task_id" in rels.columns else set()
        no_pred = tasks_c[~tasks_c["task_id"].isin(wp)]
        no_succ = tasks_c[~tasks_c["task_id"].isin(ws)]
        if "task_type" in tasks_c.columns:
            no_pred = no_pred[~no_pred["task_type"].astype(str).str.contains("Milestone|LOE|WBS", na=False)]
            no_succ = no_succ[~no_succ["task_type"].astype(str).str.contains("Milestone|LOE|WBS", na=False)]
        n_open_start = len(no_pred)
        n_open_end   = len(no_succ)

    # -- Critical not started --------------------------------------------------
    crit_ns_df = pd.DataFrame()
    if "is_critical" in tasks_c.columns and "status" in tasks_c.columns:
        crit_ns_df = tasks_c[
            tasks_c["is_critical"] &
            tasks_c["status"].apply(lambda s: str(s) in ("TK_NotStart","Not Started"))
        ]
    n_crit_ns = len(crit_ns_df)

    # -- Schedule dates --------------------------------------------------------
    pname    = proj.get("name","")
    ddate    = proj.get("data_date")
    dd_str   = format_date(ddate) if ddate else "not available"

    valid_ends = tasks_c["eff_finish"].dropna() if "eff_finish" in tasks_c.columns else pd.Series()
    proj_end   = valid_ends.max() if not valid_ends.empty else None
    proj_end_s = format_date(proj_end) if proj_end else "not available"

    # -- WBS areas of concern --------------------------------------------------
    concern_areas = []
    if "wbs_path" in tasks_c.columns and not neg_df.empty:
        neg_df2 = neg_df.copy()
        neg_df2["wbs_top"] = neg_df2["wbs_path"].apply(
            lambda x: str(x).split(" > ")[0].strip() if x and str(x).strip() not in ("","nan") else "Unknown"
        )
        wbs_neg = neg_df2.groupby("wbs_top").size().sort_values(ascending=False)
        concern_areas = wbs_neg.head(3).index.tolist()

    if not concern_areas and "wbs_path" in crit_df.columns and not crit_df.empty:
        crit_df2 = crit_df.copy()
        crit_df2["wbs_top"] = crit_df2["wbs_path"].apply(
            lambda x: str(x).split(" > ")[0].strip() if x and str(x).strip() not in ("","nan") else "Unknown"
        )
        wbs_crit = crit_df2.groupby("wbs_top").size().sort_values(ascending=False)
        concern_areas = wbs_crit.head(3).index.tolist()

    # -- Resources -------------------------------------------------------------
    has_res = not res.empty
    peak_hrs = 0
    if has_res and "target_qty" in res.columns:
        try:
            peak_hrs = int(res.groupby(
                res.get("target_start", pd.Series()).apply(
                    lambda d: pd.Timestamp(d).to_period("W").start_time
                    if d is not None else None
                )
            )["target_qty"].sum().max())
        except Exception:
            peak_hrs = int(res["target_qty"].sum())

    # -- Comparison ------------------------------------------------------------
    comparison_text = ""
    n_slips = n_improvements = 0
    worst_slip = 0
    if "_mi_prev" in st.session_state and "_mi_curr" in st.session_state:
        try:
            prev_t = st.session_state["_mi_prev"]["tasks_df"]
            curr_t = st.session_state["_mi_curr"]["tasks_df"]
            if not prev_t.empty and not curr_t.empty:
                mg = prev_t[["task_code","eff_finish"]].merge(
                    curr_t[["task_code","eff_finish"]], on="task_code",
                    suffixes=("_p","_c"), how="inner"
                )
                for _, r in mg.iterrows():
                    try:
                        d = int((pd.Timestamp(r["eff_finish_c"]) -
                                 pd.Timestamp(r["eff_finish_p"])).days)
                        if d > 0:
                            n_slips += 1
                            worst_slip = max(worst_slip, d)
                        elif d < 0:
                            n_improvements += 1
                    except Exception:
                        pass
        except Exception:
            pass

        if n_slips or n_improvements:
            parts = []
            if n_slips:
                parts.append(f"{n_slips} activities have later finish dates (worst slip: {worst_slip} days)")
            if n_improvements:
                parts.append(f"{n_improvements} activities have improved finish dates")
            comparison_text = "; ".join(parts) + "."

    # -- Planning notes risk ---------------------------------------------------
    notes_text = st.session_state.get("_notes_text","")
    n_risk_notes = 0
    if notes_text and "task_code" in tasks_c.columns:
        import re as _re
        for _, t in tasks_c.head(300).iterrows():
            code = str(t.get("task_code",""))
            if not code or code not in notes_text: continue
            idx = notes_text.find(code)
            snippet = notes_text[max(0,idx-200):idx+200]
            for w in _RISK_WORDS:
                if _re.search(r"\b"+_re.escape(w)+r"\b", snippet, _re.IGNORECASE):
                    n_risk_notes += 1
                    break

    # -- Overall risk assessment -----------------------------------------------
    risk_score = 0
    if n_neg > 0:           risk_score += 3
    if n_crit_ns > 3:       risk_score += 2
    if n_crit_ns > 0:       risk_score += 1
    if n_open_start > 10:   risk_score += 1
    if n_open_end > 10:     risk_score += 1
    if n_nc > 20:           risk_score += 1
    if n_risk_notes > 0:    risk_score += 2
    if worst_slip > 14:     risk_score += 1

    if risk_score >= 6:
        risk_level, risk_colour = "High Risk", "#DC2626"
        risk_rag = "red"
    elif risk_score >= 3:
        risk_level, risk_colour = "Medium Risk", "#F59E0B"
        risk_rag = "amber"
    else:
        risk_level, risk_colour = "Low Risk", "#16A34A"
        risk_rag = "green"

    return dict(
        pname=pname, dd_str=dd_str, proj_end_s=proj_end_s,
        n_total=n_total, n_rels=n_rels,
        n_neg=n_neg, n_crit=n_crit, n_nc=n_nc,
        n_crit_ns=n_crit_ns, worst_float=worst_float,
        n_open_start=n_open_start, n_open_end=n_open_end,
        concern_areas=concern_areas,
        has_res=has_res, peak_hrs=peak_hrs,
        comparison_text=comparison_text,
        n_slips=n_slips, n_improvements=n_improvements, worst_slip=worst_slip,
        n_risk_notes=n_risk_notes,
        risk_level=risk_level, risk_colour=risk_colour, risk_rag=risk_rag,
        near_crit_days=near_crit_days,
        neg_df=neg_df, crit_df=crit_df, nc_df=nc_df,
        crit_ns_df=crit_ns_df,
    )


def _exec_prose(n: dict) -> str:
    """
    Build the main executive narrative paragraph from metrics dict n.
    Returns plain-English HTML paragraph text.
    """
    paras = []

    # -- Programme position ----------------------------------------------------
    name_str = f"<strong>{n['pname']}</strong>" if n['pname'] else "The programme"
    paras.append(
        f"{name_str} contains <strong>{n['n_total']:,} activities</strong> and "
        f"<strong>{n['n_rels']:,} relationships</strong>. "
        f"The programme data date is {n['dd_str']} and the current forecast "
        f"project completion is <strong>{n['proj_end_s']}</strong>."
    )

    # -- Critical path ---------------------------------------------------------
    if n['n_crit'] > 0:
        crit_para = f"There are <strong>{n['n_crit']} critical activities</strong> in the programme."
        if n['n_crit_ns'] > 0:
            crit_para += (
                f" Of these, <strong>{n['n_crit_ns']} have not yet started</strong>. "
                "Any delay to a critical activity will directly push out the project "
                "completion date with no schedule buffer."
            )
        paras.append(crit_para)

    # -- Negative float --------------------------------------------------------
    if n['n_neg'] > 0:
        neg_para = (
            f"<strong>{n['n_neg']} {'activity' if n['n_neg']==1 else 'activities'} "
            f"{'has' if n['n_neg']==1 else 'have'} negative float</strong>, "
            f"meaning the current schedule cannot achieve its target dates on "
            f"{'this path' if n['n_neg']==1 else 'these paths'}."
        )
        if n['worst_float'] is not None:
            neg_para += (
                f" The worst case is <strong>{abs(n['worst_float'])} days negative float</strong>."
            )
        if n['concern_areas']:
            areas = ", ".join(n['concern_areas'])
            neg_para += f" Programme risk is concentrated in <strong>{areas}</strong>."
        paras.append(neg_para)
    else:
        paras.append("There are no activities with negative float. The programme is not currently behind its target dates.")

    # -- Near-critical ---------------------------------------------------------
    if n['n_nc'] > 0:
        paras.append(
            f"<strong>{n['n_nc']} activities are near-critical</strong>, with float between "
            f"0 and {n['near_crit_days']} days. "
            "These activities have limited contingency and should be monitored closely."
        )

    # -- Open logic ------------------------------------------------------------
    n_open = n['n_open_start'] + n['n_open_end']
    if n_open > 0:
        open_parts = []
        if n['n_open_start'] > 0:
            open_parts.append(f"{n['n_open_start']} with no predecessor")
        if n['n_open_end'] > 0:
            open_parts.append(f"{n['n_open_end']} with no successor")
        paras.append(
            f"<strong>{n_open} activities have open-ended logic</strong> "
            f"({' and '.join(open_parts)}). "
            "Activities without logic are not properly constrained in the schedule and "
            "their float values may be unreliable."
        )

    # -- Resources -------------------------------------------------------------
    if n['has_res']:
        paras.append(
            "The programme contains resource loading data. "
            f"Peak labour demand is approximately <strong>{n['peak_hrs']:,} hours per week</strong>."
        )
    else:
        paras.append(
            "The programme does not contain resource loading. "
            "Labour demand cannot be assessed from this XER."
        )

    # -- Comparison -----------------------------------------------------------
    if n['comparison_text']:
        paras.append(
            "Compared to the previous programme revision: " + n['comparison_text']
        )

    # -- Planning notes --------------------------------------------------------
    if n['n_risk_notes'] > 0:
        paras.append(
            f"Planning notes flag <strong>{n['n_risk_notes']} "
            f"{'activity' if n['n_risk_notes']==1 else 'activities'}</strong> with "
            "risk-related keywords (delay, blocked, CE, EWN, constraint etc.). "
            "These should be reviewed against the current programme position."
        )

    return "<br><br>".join(paras)


def _exec_actions(n: dict, data: dict, near_crit_days: float) -> list:
    """
    Generate recommended PM actions from programme metrics.
    Returns list of (priority, action_text, detail, page).
    """
    actions = []
    tasks_c = get_critical_threshold(data.get("tasks_df", pd.DataFrame()), near_crit_days)
    rels    = data.get("relationships_df", pd.DataFrame())

    if n['n_neg'] > 0:
        actions.append((
            1,
            f"Recover {n['n_neg']} activities with negative float",
            f"Worst case: {abs(n['worst_float'])}d. Develop a recovery programme "
            "with the delivery team. Do not accept a programme with negative float "
            "without a clear recovery plan in place.",
            "Critical Path",
        ))

    if n['n_crit_ns'] > 0:
        actions.append((
            1,
            f"Mobilise {n['n_crit_ns']} critical activities not yet started",
            "Any further delay to these activities will directly push out the "
            "project completion date. Confirm start dates and resource availability "
            "with the responsible party immediately.",
            "Critical Path",
        ))

    if n['n_nc'] > 0:
        from datetime import timedelta
        four_wks = datetime.now() + timedelta(weeks=4)
        nc_soon = n['nc_df']
        if "eff_finish" in nc_soon.columns:
            nc_soon = nc_soon[nc_soon["eff_finish"].apply(
                lambda d: d is not None and hasattr(d,"date") and d <= four_wks
            )]
        if not nc_soon.empty:
            actions.append((
                1,
                f"Confirm readiness for {len(nc_soon)} near-critical activities finishing within 4 weeks",
                "These activities have limited float and are finishing soon. "
                "Any disruption will move them onto the critical path.",
                "Programme",
            ))
        elif n['n_nc'] > 10:
            actions.append((
                2,
                f"Monitor {n['n_nc']} near-critical activities",
                f"Float buffer is between 0 and {near_crit_days} days. "
                "Weekly monitoring is recommended.",
                "Programme",
            ))

    n_open = n['n_open_start'] + n['n_open_end']
    if n_open > 5:
        actions.append((
            2,
            f"Resolve {n_open} open-ended logic issues with the planner",
            f"{n['n_open_start']} activities have no predecessor and "
            f"{n['n_open_end']} have no successor. "
            "Open logic reduces schedule reliability. Ask the planner to "
            "add logic or justify the open ends.",
            "Health Check",
        ))

    # Excessive lag
    if not rels.empty and "lag_days" in rels.columns:
        big_lag = rels[rels["lag_days"].apply(lambda l: safe_float(l,0) > 10)]
        if not big_lag.empty:
            n_lag = len(big_lag)
            max_lag = int(big_lag["lag_days"].max())
            actions.append((
                2,
                f"Review {n_lag} relationships with excessive lag (largest: {max_lag}d)",
                "Excessive lag can hide critical path risk and inflate float values. "
                "Each lag should be reviewed and replaced with properly sequenced "
                "activities wherever possible.",
                "Health Check",
            ))

    # Missing resources
    if not n['has_res']:
        actions.append((
            2,
            "Request resource-loaded programme from the planner",
            "The current programme does not contain resource loading. "
            "A resourced programme is required to assess labour demand, "
            "identify peaks and support cost reporting.",
            "Labour",
        ))

    # Planning notes
    if n['n_risk_notes'] > 0:
        actions.append((
            1,
            f"Investigate {n['n_risk_notes']} activities flagged in planning notes",
            "Risk-related keywords found in planning notes. "
            "Confirm these issues are being tracked and that the programme "
            "reflects any agreed mitigation or recovery.",
            "Programme",
        ))

    # Comparison slips
    if n['n_slips'] > 5 and n['worst_slip'] > 14:
        actions.append((
            1,
            f"Understand {n['n_slips']} activities that have slipped since the previous revision",
            f"Worst slip: {n['worst_slip']} days. "
            "The programme has moved significantly. Confirm the reasons for movement "
            "and whether recovery is achievable within the current plan.",
            "Comparison",
        ))

    # Long durations
    if "orig_dur_days" in tasks_c.columns:
        long_dur = tasks_c[tasks_c["orig_dur_days"].apply(lambda d: safe_float(d,0) > 60)]
        if not long_dur.empty:
            actions.append((
                3,
                f"Review {len(long_dur)} activities with duration over 60 days",
                "Long-duration activities are difficult to control and monitor. "
                "Consider breaking into smaller work packages with the planner.",
                "Health Check",
            ))

    # Sort: High first, then Medium, then Low; within each by count (numeric in text)
    pmap = {1:"High", 2:"Medium", 3:"Low"}
    actions.sort(key=lambda x: x[0])
    return [(pmap[p], txt, detail, page) for p, txt, detail, page in actions]


def page_executive_summary(data: dict, near_crit_days: float):
    """
    Executive Summary page.
    Generates a plain-English programme summary suitable for directors,
    project managers and project controls managers.
    """
    ctrl_bar(
        "Executive Summary",
        "Programme position, risk assessment and management actions in plain English.",
    )
    mode_toggle_bar()

    tasks = data.get("tasks_df", pd.DataFrame())
    if tasks.empty:
        empty_state("", "No programme data available",
                    "Upload a programme to generate an Executive Summary.", "")
        return

    # Generate all metrics
    n = _exec_narrative(data, near_crit_days)
    if not n:
        st.warning("Could not generate summary from this programme.")
        return

    # -- OVERALL STATUS HEADER -------------------------------------------------
    risk_bg  = {"red":"#FEF2F2","amber":"#FFFBEB","green":"#F0FDF4"}.get(n["risk_rag"],"#F9FAFB")
    risk_brd = {"red":"#DC2626","amber":"#F59E0B","green":"#16A34A"}.get(n["risk_rag"],"#E5E7EB")

    st.markdown(
        f"""
        <div style="background:{risk_brd};border-radius:10px;padding:20px 28px;
                    margin-bottom:20px;color:white;">
            <div style="display:flex;justify-content:space-between;
                        align-items:flex-start;flex-wrap:wrap;gap:16px;">
                <div>
                    <div style="font-size:11px;font-weight:700;opacity:0.85;
                                text-transform:uppercase;letter-spacing:1.2px;
                                margin-bottom:6px;">Overall Programme Risk</div>
                    <div style="font-size:30px;font-weight:900;line-height:1;
                                margin-bottom:8px;">{n["risk_level"]}</div>
                    <div style="font-size:13px;opacity:0.9;">
                        {n["pname"] or "Programme"} &nbsp;|&nbsp;
                        Data date: {n["dd_str"]} &nbsp;|&nbsp;
                        Forecast completion: {n["proj_end_s"]}
                    </div>
                </div>
                <div style="display:grid;grid-template-columns:repeat(4,auto);
                            gap:16px;align-self:center;text-align:center;">
                    <div>
                        <div style="font-size:9px;opacity:0.7;text-transform:uppercase;
                                    letter-spacing:0.8px;">Activities</div>
                        <div style="font-size:24px;font-weight:800;">{n["n_total"]:,}</div>
                    </div>
                    <div>
                        <div style="font-size:9px;opacity:0.7;text-transform:uppercase;
                                    letter-spacing:0.8px;">Critical</div>
                        <div style="font-size:24px;font-weight:800;">{n["n_crit"]}</div>
                    </div>
                    <div>
                        <div style="font-size:9px;opacity:0.7;text-transform:uppercase;
                                    letter-spacing:0.8px;">Neg Float</div>
                        <div style="font-size:24px;font-weight:800;">{n["n_neg"]}</div>
                    </div>
                    <div>
                        <div style="font-size:9px;opacity:0.7;text-transform:uppercase;
                                    letter-spacing:0.8px;">Near-Crit</div>
                        <div style="font-size:24px;font-weight:800;">{n["n_nc"]}</div>
                    </div>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # -- TOP 5 MANAGEMENT POINTS -----------------------------------------------
    st.markdown(
        '<div class="section-label">Top 5 Management Points</div>',
        unsafe_allow_html=True,
    )

    mgmt_points = []

    if n["n_neg"] > 0:
        mgmt_points.append((
            "red",
            f"{n['n_neg']} activities have negative float.",
            f"The schedule cannot currently achieve its target dates. "
            f"{'Worst: ' + str(abs(n['worst_float'])) + 'd.' if n['worst_float'] else ''} "
            f"{'Risk concentrated in: ' + ', '.join(n['concern_areas']) + '.' if n['concern_areas'] else ''}",
        ))

    if n["n_crit_ns"] > 0:
        mgmt_points.append((
            "red",
            f"{n['n_crit_ns']} critical activities have not started.",
            "Any further delay will directly push out the project completion date. "
            "Confirm mobilisation plans and resource availability.",
        ))

    n_open = n["n_open_start"] + n["n_open_end"]
    if n_open > 0:
        mgmt_points.append((
            "amber",
            f"{n_open} activities have open-ended logic.",
            f"{n['n_open_start']} with no predecessor, {n['n_open_end']} with no successor. "
            "Open logic reduces programme reliability. Review with the planner.",
        ))

    if n["n_nc"] > 0:
        mgmt_points.append((
            "amber",
            f"{n['n_nc']} near-critical activities with limited contingency.",
            f"Float between 0 and {near_crit_days} days. Monitor weekly and "
            "escalate immediately if float reduces further.",
        ))

    if n["comparison_text"]:
        mgmt_points.append((
            "amber",
            "Programme has moved since the previous revision.",
            n["comparison_text"],
        ))
    elif not n["has_res"]:
        mgmt_points.append((
            "amber",
            "No resource loading found in this programme.",
            "Labour demand, peak resource requirements and cost loading cannot be assessed.",
        ))
    elif n["n_risk_notes"] > 0:
        mgmt_points.append((
            "red",
            f"{n['n_risk_notes']} activities flagged in planning notes.",
            "Risk keywords detected. Confirm these issues are being actively managed.",
        ))

    # Show top 5
    for i, (style, headline, detail) in enumerate(mgmt_points[:5], start=1):
        border_col = {"red":"#DC2626","amber":"#F59E0B","green":"#16A34A"}.get(style,"#E5E7EB")
        bg_col     = {"red":"#FEF2F2","amber":"#FFFBEB","green":"#F0FDF4"}.get(style,"#F9FAFB")
        st.markdown(
            f'<div style="background:{bg_col};border:1px solid {border_col};'
            f'border-left:4px solid {border_col};border-radius:0 8px 8px 0;'
            f'padding:12px 16px;margin-bottom:8px;display:flex;gap:14px;">'
            f'<div style="font-size:14px;font-weight:800;color:{border_col};'
            f'flex-shrink:0;min-width:20px;">{i}</div>'
            f'<div>'
            f'<div style="font-weight:700;color:#071827;font-size:13px;'
            f'margin-bottom:3px;">{headline}</div>'
            f'<div style="font-size:12px;color:#374151;line-height:1.6;">'
            f'{detail}</div></div></div>',
            unsafe_allow_html=True,
        )

    st.markdown("<br>", unsafe_allow_html=True)

    # -- TWO COLUMN: NARRATIVE + METRICS ---------------------------------------
    col_narr, col_metrics = st.columns([3, 2], gap="large")

    # -- Programme Position narrative ------------------------------------------
    with col_narr:
        st.markdown(
            '<div class="section-label">Programme Position</div>',
            unsafe_allow_html=True,
        )
        prose = _exec_prose(n)
        st.markdown(
            f'<div style="background:#FFFFFF;border:1px solid #E5E7EB;'
            f'border-radius:8px;padding:20px 22px;line-height:1.8;'
            f'font-size:13px;color:#374151;">{prose}</div>',
            unsafe_allow_html=True,
        )

        # -- What this means ---------------------------------------------------
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown(
            '<div class="section-label">What This Means</div>',
            unsafe_allow_html=True,
        )

        if n["risk_rag"] == "red":
            meaning = (
                "The programme is currently in a <strong>high-risk position</strong>. "
                "There are activities that cannot achieve their target dates without "
                "recovery action. The project completion date is at risk. "
                "PM review is recommended before accepting the programme position. "
                "A recovery plan should be developed with the delivery team "
                "and reviewed by the Project Director."
            )
        elif n["risk_rag"] == "amber":
            meaning = (
                "The programme shows <strong>moderate schedule risk</strong>. "
                "There are areas of concern that should be addressed to prevent "
                "them from becoming critical issues. The project completion date "
                "is achievable but is dependent on the near-critical and open-logic "
                "items being resolved. PM attention is required."
            )
        else:
            meaning = (
                "The programme appears to be in a <strong>healthy position</strong>. "
                "No immediate critical schedule risks have been identified. "
                "Continue to monitor near-critical activities and maintain "
                "the current schedule performance."
            )

        st.markdown(
            f'<div style="background:#EFF6FF;border:1px solid #BFDBFE;'
            f'border-left:4px solid #1AADAA;border-radius:0 8px 8px 0;'
            f'padding:16px 20px;font-size:13px;color:#1E3A5F;line-height:1.7;">'
            f'{meaning}</div>',
            unsafe_allow_html=True,
        )

    # -- Metrics column --------------------------------------------------------
    with col_metrics:
        st.markdown(
            '<div class="section-label">Schedule Metrics</div>',
            unsafe_allow_html=True,
        )

        metrics = [
            ("Total Activities",    f"{n['n_total']:,}",  "navy"),
            ("Relationships",        f"{n['n_rels']:,}",   "navy"),
            ("Critical Activities",  str(n["n_crit"]),     "red"   if n["n_crit"] > 0 else "green"),
            ("Negative Float",       str(n["n_neg"]),      "red"   if n["n_neg"]  > 0 else "green"),
            ("Near-Critical",        str(n["n_nc"]),       "amber" if n["n_nc"]   > 0 else "green"),
            ("Not Started (Critical)", str(n["n_crit_ns"]),"red"   if n["n_crit_ns"] > 0 else "green"),
            ("Open Start Logic",     str(n["n_open_start"]),"amber" if n["n_open_start"] > 0 else "green"),
            ("Open Finish Logic",    str(n["n_open_end"]), "amber" if n["n_open_end"]   > 0 else "green"),
        ]

        for label, val, style in metrics:
            st.markdown(kpi_card(label, val, style=style), unsafe_allow_html=True)
            st.markdown("<div style='height:6px;'></div>", unsafe_allow_html=True)

        # WBS concerns
        if n["concern_areas"]:
            st.markdown(
                '<div class="section-label" style="margin-top:14px;">Areas of Concern</div>',
                unsafe_allow_html=True,
            )
            for area in n["concern_areas"]:
                st.markdown(
                    f'<div style="background:#FEF2F2;border:1px solid #FECACA;'
                    f'border-radius:6px;padding:8px 12px;margin-bottom:4px;'
                    f'font-size:12px;font-weight:600;color:#991B1B;">'
                    f'{area}</div>',
                    unsafe_allow_html=True,
                )

    # -- Recommended Actions ---------------------------------------------------
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown(
        '<div class="section-label">Recommended Actions</div>',
        unsafe_allow_html=True,
    )

    actions = _exec_actions(n, data, near_crit_days)

    if not actions:
        st.success("No immediate recommended actions. Programme appears healthy.")
    else:
        priority_colours = {
            "High":   ("#DC2626","#FEF2F2","#FECACA"),
            "Medium": ("#F59E0B","#FFFBEB","#FDE68A"),
            "Low":    ("#6B7280","#F9FAFB","#E5E7EB"),
        }
        for i, (priority, action, detail, page) in enumerate(actions, start=1):
            tc, bg, brd = priority_colours.get(priority, priority_colours["Low"])
            st.markdown(
                f'<div style="background:{bg};border:1px solid {brd};'
                f'border-left:4px solid {tc};border-radius:0 8px 8px 0;'
                f'padding:12px 16px;margin-bottom:8px;display:flex;gap:14px;">'
                f'<div style="width:22px;height:22px;border-radius:50%;'
                f'background:{tc};color:white;font-size:11px;font-weight:800;'
                f'display:flex;align-items:center;justify-content:center;'
                f'flex-shrink:0;">{i}</div>'
                f'<div style="flex:1;">'
                f'<div style="display:flex;justify-content:space-between;'
                f'align-items:center;gap:10px;flex-wrap:wrap;margin-bottom:3px;">'
                f'<div style="font-weight:700;color:#071827;font-size:13px;">'
                f'{action}</div>'
                f'<div style="display:flex;gap:6px;flex-shrink:0;">'
                f'{chip(priority, priority.lower() if priority=="High" else ("amber" if priority=="Medium" else "grey"))}'
                f'<span style="font-size:11px;color:#9CA3AF;">-> {page}</span>'
                f'</div></div>'
                f'<div style="font-size:12px;color:#374151;line-height:1.5;">'
                f'{detail}</div>'
                f'</div></div>',
                unsafe_allow_html=True,
            )

    # -- WBS Risk Heatmap ---------------------------------------------------
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown(
        '<div class="section-label">WBS Risk Heatmap</div>',
        unsafe_allow_html=True,
    )
    st.caption(
        "Risk score for each WBS area based on negative float, critical activities, "
        "near-critical work, open logic and other schedule quality factors."
    )
    render_wbs_heatmap(data, near_crit_days)
    st.markdown("<br>", unsafe_allow_html=True)

    # -- Export ----------------------------------------------------------------
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown(
        '<div class="section-label">Export</div>',
        unsafe_allow_html=True,
    )

    # Build export content
    summary_rows = {
        "Item":  [
            "Programme Name", "Data Date", "Forecast Completion", "Overall Risk",
            "Total Activities", "Relationships",
            "Critical Activities", "Critical Not Started",
            "Negative Float Activities", "Worst Float (days)",
            "Near-Critical Activities", "Open Logic (starts)",
            "Open Logic (finishes)", "Resource Loading",
            "Activities in Planning Notes (risk)", "Comparison Slips",
        ],
        "Value": [
            n["pname"], n["dd_str"], n["proj_end_s"], n["risk_level"],
            n["n_total"], n["n_rels"],
            n["n_crit"], n["n_crit_ns"],
            n["n_neg"], n["worst_float"] if n["worst_float"] else "None",
            n["n_nc"], n["n_open_start"],
            n["n_open_end"], "Yes" if n["has_res"] else "No",
            n["n_risk_notes"], n["n_slips"],
        ],
    }

    top5_rows = {
        "No.":     list(range(1, len(mgmt_points[:5])+1)),
        "Priority":["High" if s=="red" else "Medium" if s=="amber" else "Low"
                    for s, *_ in mgmt_points[:5]],
        "Headline":[h for _, h, _ in mgmt_points[:5]],
        "Detail":  [d for _, _, d in mgmt_points[:5]],
    }

    actions_rows = {
        "No.":     list(range(1, len(actions)+1)),
        "Priority":[a[0] for a in actions],
        "Action":  [a[1] for a in actions],
        "Detail":  [a[2] for a in actions],
        "Page":    [a[3] for a in actions],
    }

    prose_clean = _exec_prose(n).replace("<br><br>","\n\n").replace("<strong>","").replace("</strong>","")

    narrative_rows = {
        "Section":    ["Programme Position"],
        "Commentary": [prose_clean],
    }

    export_sheets = {
        "Summary":           pd.DataFrame(summary_rows),
        "Top 5 Points":      pd.DataFrame(top5_rows),
        "Recommended Actions": pd.DataFrame(actions_rows),
        "Programme Narrative": pd.DataFrame(narrative_rows),
    }

    # Include neg float and crit not started lists
    if not n["neg_df"].empty:
        neg_exp_cols = {k:v for k,v in {
            "task_code":"Activity ID","task_name":"Activity Name",
            "wbs_path":"WBS","eff_finish":"Finish","total_float_days":"Float (d)",
            "status":"Status"
        }.items() if k in n["neg_df"].columns}
        neg_exp = n["neg_df"][list(neg_exp_cols.keys())].rename(columns=neg_exp_cols).copy()
        if "Finish" in neg_exp.columns:
            neg_exp["Finish"] = neg_exp["Finish"].apply(format_date)
        export_sheets["Negative Float Activities"] = neg_exp

    if not n["crit_ns_df"].empty:
        cns_cols = {k:v for k,v in {
            "task_code":"Activity ID","task_name":"Activity Name",
            "wbs_path":"WBS","eff_start":"Start","eff_finish":"Finish",
        }.items() if k in n["crit_ns_df"].columns}
        cns_exp = n["crit_ns_df"][list(cns_cols.keys())].rename(columns=cns_cols).copy()
        for col in ["Start","Finish"]:
            if col in cns_exp.columns:
                cns_exp[col] = cns_exp[col].apply(format_date)
        export_sheets["Critical Not Started"] = cns_exp

    xls = export_df_to_excel(export_sheets)

    exp_col, _ = st.columns([1, 3])
    exp_col.download_button(
        label="Export Executive Summary to Excel",
        data=xls,
        file_name=f"executive_summary_{datetime.now().strftime('%Y%m%d')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        help="Exports Summary, Top 5 Points, Recommended Actions, Programme Narrative, Negative Float and Critical Not Started sheets.",
        key="dl_017",
    )

    # -- Word export -----------------------------------------------------------
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        def _make_word_doc(n: dict, actions: list, mgmt_points: list, prose: str) -> bytes:
            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)

            # Title
            title = doc.add_heading("PlanTrace - Executive Summary", 0)
            title.runs[0].font.color.rgb = RGBColor(7, 24, 39)

            # Programme line
            p = doc.add_paragraph()
            p.add_run(f"Programme: {n['pname'] or 'Programme'}").bold = True
            doc.add_paragraph(f"Data Date: {n['dd_str']}  |  Forecast Completion: {n['proj_end_s']}")
            doc.add_paragraph(f"Overall Risk: {n['risk_level']}")
            doc.add_paragraph("")

            # Metrics table
            doc.add_heading("Schedule Metrics", 2)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Metric"
            hdr[1].text = "Value"
            for label, val, _ in [
                ("Total Activities",    f"{n['n_total']:,}", ""),
                ("Critical Activities", str(n["n_crit"]),    ""),
                ("Negative Float",      str(n["n_neg"]),     ""),
                ("Near-Critical",       str(n["n_nc"]),      ""),
                ("Not Started (Critical)", str(n["n_crit_ns"]), ""),
                ("Open Logic Issues",   str(n["n_open_start"]+n["n_open_end"]), ""),
            ]:
                row = table.add_row().cells
                row[0].text = label
                row[1].text = val

            doc.add_paragraph("")

            # Top 5
            doc.add_heading("Top 5 Management Points", 2)
            for i, (_, headline, detail) in enumerate(mgmt_points[:5], 1):
                doc.add_paragraph(f"{i}. {headline}", style="List Number")
                doc.add_paragraph(detail)

            # Narrative
            doc.add_heading("Programme Position", 2)
            clean = prose.replace("<br><br>","\n").replace("<strong>","").replace("</strong>","")
            doc.add_paragraph(clean)

            # What this means
            doc.add_heading("What This Means", 2)
            if n["risk_rag"] == "red":
                doc.add_paragraph("The programme is currently in a high-risk position. PM review is recommended before accepting the programme position. A recovery plan should be developed.")
            elif n["risk_rag"] == "amber":
                doc.add_paragraph("The programme shows moderate schedule risk. PM attention is required on the items noted above.")
            else:
                doc.add_paragraph("The programme appears to be in a healthy position. Continue to monitor near-critical activities.")

            # Recommended Actions
            doc.add_heading("Recommended Actions", 2)
            for i, (priority, action, detail, page) in enumerate(actions, 1):
                p = doc.add_paragraph(style="List Number")
                run = p.add_run(f"[{priority}] {action}")
                run.bold = True
                doc.add_paragraph(f"{detail} (Page: {page})")

            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        word_bytes = _make_word_doc(n, actions, mgmt_points, _exec_prose(n))
        exp_col.download_button(
            label="Export Executive Summary to Word",
            data=word_bytes,
            file_name=f"executive_summary_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help="Exports formatted Word document with narrative, metrics and recommended actions.",
            key="dl_018",
        )
    except Exception:
        pass  # Word export optional - Excel always works




# -----------------------------------------------------------------------------
# WBS RISK HEATMAP
# -----------------------------------------------------------------------------

_WBS_RISK_WEIGHTS = {
    "neg_float":     10,   # Negative float activities
    "crit":           6,   # Critical activities
    "crit_ns":        8,   # Critical not started
    "near_crit":      4,   # Near-critical activities
    "open_start":     3,   # No predecessor
    "open_end":       3,   # No successor
    "long_dur":       2,   # Long duration (>60d)
    "constrained":    2,   # Constrained activities
    "date_slip":      8,   # Major finish date movement
    "labour_peak":    2,   # Labour peak weighting (normalised)
}


def _wbs_risk_rating(score: float) -> tuple:
    """
    Map a raw WBS risk score to a (label, colour, chip_style) tuple.
    Thresholds tuned so a clean WBS = Low and a heavily negative WBS = Severe.
    """
    if score >= 50:
        return "Severe", "#7F1D1D", "red"
    if score >= 25:
        return "High",   "#DC2626", "red"
    if score >= 10:
        return "Medium", "#F59E0B", "amber"
    return "Low",    "#16A34A", "green"


def _wbs_reason(row: dict) -> str:
    """Generate a plain-English primary reason for the WBS risk rating."""
    reasons = []
    if row["neg_float"] > 0:
        reasons.append(f"{row['neg_float']} negative float")
    if row["crit_ns"] > 0:
        reasons.append(f"{row['crit_ns']} critical not started")
    if row["crit"] > 0 and not reasons:
        reasons.append(f"{row['crit']} critical activities")
    if row["near_crit"] > 0 and not reasons:
        reasons.append(f"{row['near_crit']} near-critical")
    if row["open_start"] + row["open_end"] > 0 and not reasons:
        n = row["open_start"] + row["open_end"]
        reasons.append(f"{n} open logic")
    if row.get("date_slip", 0) > 0 and not reasons:
        reasons.append(f"{row['date_slip']} activities slipped")
    if not reasons:
        return "No immediate issues"
    return "; ".join(reasons[:2]).capitalize()


def _wbs_action(row: dict) -> str:
    """Return a short recommended action for the WBS."""
    if row["neg_float"] > 0:
        return "Develop recovery plan — schedule cannot meet target dates."
    if row["crit_ns"] > 0:
        return "Confirm mobilisation of critical activities immediately."
    if row["crit"] > 5:
        return "Review critical path and confirm resources are in place."
    if row["near_crit"] > 5:
        return "Monitor closely — limited contingency remaining."
    if row["open_start"] + row["open_end"] > 5:
        return "Resolve open logic with the planner to improve schedule reliability."
    if row.get("date_slip", 0) > 0:
        return "Investigate date slippage and agree recovery with the team."
    if row["constrained"] > 0:
        return "Review constraints — may be causing artificial float."
    return "Continue to monitor."


def compute_wbs_heatmap(
    tasks: pd.DataFrame,
    rels:  pd.DataFrame,
    near_crit_days: float = 10.0,
    task_res: pd.DataFrame = pd.DataFrame(),
) -> pd.DataFrame:
    """
    Compute the WBS risk heatmap DataFrame.
    Each row = one top-level WBS area with all risk metrics and a composite score.
    """
    if tasks.empty or "wbs_path" not in tasks.columns:
        return pd.DataFrame()

    tasks_c = get_critical_threshold(tasks, near_crit_days)

    # Assign top-level WBS to each activity
    tasks_c = tasks_c.copy()
    tasks_c["wbs_top"] = tasks_c["wbs_path"].apply(
        lambda x: (
            str(x).split(" > ")[0].strip()
            if x and str(x).strip() not in ("", "nan", "None")
            else "Unknown"
        )
    )

    # Build predecessor/successor sets
    tasks_with_pred = set(rels["succ_task_id"].dropna()) if not rels.empty and "succ_task_id" in rels.columns else set()
    tasks_with_succ = set(rels["pred_task_id"].dropna()) if not rels.empty and "pred_task_id" in rels.columns else set()

    # Comparison data: finish movement per task
    slip_map = {}
    if "_mi_prev" in st.session_state and "_mi_curr" in st.session_state:
        try:
            prev_t = st.session_state["_mi_prev"]["tasks_df"]
            curr_t = st.session_state["_mi_curr"]["tasks_df"]
            if not prev_t.empty and not curr_t.empty:
                mg = prev_t[["task_code","eff_finish"]].merge(
                    curr_t[["task_code","eff_finish"]], on="task_code",
                    suffixes=("_p","_c"), how="inner"
                )
                for _, r in mg.iterrows():
                    try:
                        d = int((pd.Timestamp(r["eff_finish_c"]) - pd.Timestamp(r["eff_finish_p"])).days)
                        if d > 14:
                            slip_map[str(r["task_code"])] = d
                    except Exception:
                        pass
        except Exception:
            pass

    # Labour hours by WBS top (if resource data available)
    labour_by_wbs = {}
    if not task_res.empty and "task_id" in task_res.columns:
        try:
            tr_merged = task_res.merge(
                tasks_c[["task_id","wbs_top"]], on="task_id", how="left"
            )
            if "target_qty" in tr_merged.columns and "wbs_top" in tr_merged.columns:
                labour_by_wbs = (
                    tr_merged.groupby("wbs_top")["target_qty"]
                    .sum()
                    .to_dict()
                )
        except Exception:
            pass

    # -- Build one row per WBS top ---------------------------------------------
    rows = []
    for wbs_name, grp in tasks_c.groupby("wbs_top"):
        n_total     = len(grp)
        n_crit      = int(grp["is_critical"].sum()) if "is_critical" in grp.columns else 0
        n_neg       = int(grp["total_float_days"].apply(lambda f: safe_float(f,0) < 0).sum()) \
                      if "total_float_days" in grp.columns else 0
        n_nc        = int(grp["is_near_critical"].sum()) if "is_near_critical" in grp.columns else 0

        # Worst float
        worst_float = None
        if "total_float_days" in grp.columns:
            neg_floats = grp["total_float_days"].apply(lambda f: safe_float(f, 0)).values
            mins = [v for v in neg_floats if v < 0]
            worst_float = round(min(mins), 1) if mins else None

        # Open logic
        n_open_start = int((~grp["task_id"].isin(tasks_with_pred)).sum()) \
                       if "task_id" in grp.columns else 0
        n_open_end   = int((~grp["task_id"].isin(tasks_with_succ)).sum()) \
                       if "task_id" in grp.columns else 0

        # Critical not started
        n_crit_ns = 0
        if "is_critical" in grp.columns and "status" in grp.columns:
            n_crit_ns = int((
                grp["is_critical"] &
                grp["status"].apply(lambda s: str(s) in ("TK_NotStart","Not Started"))
            ).sum())

        # Long duration (>60d)
        n_long_dur = int(grp["orig_dur_days"].apply(lambda d: safe_float(d,0) > 60).sum()) \
                     if "orig_dur_days" in grp.columns else 0

        # Constrained
        n_constrained = 0
        if "cstr_type" in grp.columns:
            n_constrained = int(grp["cstr_type"].apply(
                lambda x: bool(x) and str(x).strip() not in ("","None","nan")
            ).sum())

        # Date slips in this WBS
        n_slips = 0
        if slip_map and "task_code" in grp.columns:
            n_slips = int(grp["task_code"].apply(lambda c: str(c) in slip_map).sum())

        # Labour
        labour_hrs = int(labour_by_wbs.get(wbs_name, 0))

        # -- Composite risk score ----------------------------------------------
        w = _WBS_RISK_WEIGHTS
        score = (
            n_neg        * w["neg_float"] +
            n_crit       * w["crit"] +
            n_crit_ns    * w["crit_ns"] +
            n_nc         * w["near_crit"] +
            n_open_start * w["open_start"] +
            n_open_end   * w["open_end"] +
            n_long_dur   * w["long_dur"] +
            n_constrained * w["constrained"] +
            n_slips      * w["date_slip"] +
            (1 if labour_hrs > 1000 else 0) * w["labour_peak"]
        )

        # Normalise by activity count so large WBS doesn't dominate unfairly
        density_score = round(score / max(n_total, 1) * 10, 1)
        raw_score     = score

        rating_label, rating_col, rating_chip = _wbs_risk_rating(density_score)

        row_dict = {
            "WBS Area":         wbs_name,
            "Activities":       n_total,
            "Critical":         n_crit,
            "Negative Float":   n_neg,
            "Worst Float (d)":  worst_float if worst_float is not None else "-",
            "Near-Critical":    n_nc,
            "Crit Not Started": n_crit_ns,
            "Open Logic":       n_open_start + n_open_end,
            "Long Duration":    n_long_dur,
            "Constrained":      n_constrained,
            "Date Slips":       n_slips,
            "Labour (hrs)":     labour_hrs if labour_hrs > 0 else "-",
            "Risk Score":       density_score,
            "Raw Score":        raw_score,
            "Rating":           rating_label,
            "_colour":          rating_col,
            "_chip":            rating_chip,
            "_reason":          _wbs_reason({
                "neg_float": n_neg, "crit": n_crit, "crit_ns": n_crit_ns,
                "near_crit": n_nc, "open_start": n_open_start, "open_end": n_open_end,
                "date_slip": n_slips, "constrained": n_constrained,
            }),
            "_action": _wbs_action({
                "neg_float": n_neg, "crit": n_crit, "crit_ns": n_crit_ns,
                "near_crit": n_nc, "open_start": n_open_start, "open_end": n_open_end,
                "date_slip": n_slips, "constrained": n_constrained,
            }),
        }
        rows.append(row_dict)

    if not rows:
        return pd.DataFrame()

    df = pd.DataFrame(rows).sort_values("Risk Score", ascending=False).reset_index(drop=True)
    return df


def render_wbs_heatmap(data: dict, near_crit_days: float):
    """
    Render the WBS Risk Heatmap panel.
    Can be called from the Executive Summary or Overview page.
    """
    tasks   = data.get("tasks_df",          pd.DataFrame())
    rels    = data.get("relationships_df",   pd.DataFrame())
    task_res = data.get("task_resources_df", pd.DataFrame())

    if tasks.empty:
        return

    if "wbs_path" not in tasks.columns or tasks["wbs_path"].dropna().empty:
        st.info("No WBS data available. WBS structure is required to generate the heatmap.")
        return

    # Compute
    with st.spinner("Computing WBS risk scores..."):
        hmap_df = compute_wbs_heatmap(tasks, rels, near_crit_days, task_res)

    if hmap_df.empty:
        st.info("Not enough WBS data to generate the heatmap.")
        return

    n_wbs = len(hmap_df)

    # -- Summary strip ---------------------------------------------------------
    n_severe = int((hmap_df["Rating"] == "Severe").sum())
    n_high   = int((hmap_df["Rating"] == "High").sum())
    n_medium = int((hmap_df["Rating"] == "Medium").sum())
    n_low    = int((hmap_df["Rating"] == "Low").sum())

    st.markdown(
        f'<div style="display:flex;gap:8px;flex-wrap:wrap;margin-bottom:14px;">'
        f'<div style="background:#7F1D1D;color:white;border-radius:6px;padding:8px 16px;'
        f'min-width:80px;text-align:center;">'
        f'<div style="font-size:9px;font-weight:700;text-transform:uppercase;'
        f'letter-spacing:0.8px;opacity:0.8;margin-bottom:4px;">Severe</div>'
        f'<div style="font-size:22px;font-weight:800;">{n_severe}</div></div>'
        f'<div style="background:#DC2626;color:white;border-radius:6px;padding:8px 16px;'
        f'min-width:80px;text-align:center;">'
        f'<div style="font-size:9px;font-weight:700;text-transform:uppercase;'
        f'letter-spacing:0.8px;opacity:0.8;margin-bottom:4px;">High</div>'
        f'<div style="font-size:22px;font-weight:800;">{n_high}</div></div>'
        f'<div style="background:#F59E0B;color:white;border-radius:6px;padding:8px 16px;'
        f'min-width:80px;text-align:center;">'
        f'<div style="font-size:9px;font-weight:700;text-transform:uppercase;'
        f'letter-spacing:0.8px;opacity:0.8;margin-bottom:4px;">Medium</div>'
        f'<div style="font-size:22px;font-weight:800;">{n_medium}</div></div>'
        f'<div style="background:#16A34A;color:white;border-radius:6px;padding:8px 16px;'
        f'min-width:80px;text-align:center;">'
        f'<div style="font-size:9px;font-weight:700;text-transform:uppercase;'
        f'letter-spacing:0.8px;opacity:0.8;margin-bottom:4px;">Low</div>'
        f'<div style="font-size:22px;font-weight:800;">{n_low}</div></div>'
        f'<div style="margin-left:auto;align-self:center;font-size:12px;color:#6B7280;">'
        f'{n_wbs} WBS areas assessed</div>'
        f'</div>',
        unsafe_allow_html=True,
    )

    # -- Tabs -----------------------------------------------------------------
    tab_heat, tab_cards, tab_table, tab_export = st.tabs([
        "Heatmap Chart", "Risk Cards", "Full Table", "Export"
    ])

    # -- HEATMAP CHART ---------------------------------------------------------
    with tab_heat:
        st.caption(
            "Each bar represents one WBS area. Height = risk score. "
            "Colour = risk rating. Hover for details."
        )

        colour_map = {
            "Severe": "#7F1D1D",
            "High":   "#DC2626",
            "Medium": "#F59E0B",
            "Low":    "#16A34A",
        }
        plot_df = hmap_df.copy()
        plot_df["Colour"] = plot_df["Rating"].map(colour_map)
        plot_df["Hover"] = plot_df.apply(
            lambda r: (
                f"WBS: {r['WBS Area']}<br>"
                f"Risk Score: {r['Risk Score']}<br>"
                f"Rating: {r['Rating']}<br>"
                f"Activities: {r['Activities']}<br>"
                f"Critical: {r['Critical']}<br>"
                f"Negative Float: {r['Negative Float']}<br>"
                f"Near-Critical: {r['Near-Critical']}<br>"
                f"Open Logic: {r['Open Logic']}<br>"
                f"Primary Concern: {r['_reason']}"
            ),
            axis=1,
        )

        fig = go.Figure()
        for rating in ["Severe","High","Medium","Low"]:
            sub = plot_df[plot_df["Rating"] == rating]
            if sub.empty:
                continue
            fig.add_trace(go.Bar(
                x=sub["WBS Area"],
                y=sub["Risk Score"],
                name=rating,
                marker_color=colour_map[rating],
                hovertext=sub["Hover"],
                hoverinfo="text",
                text=sub["Rating"],
                textposition="outside",
                textfont=dict(size=10, color="#374151"),
            ))

        fig.update_layout(
            barmode="stack",
            title=dict(text="WBS Risk Heatmap — Risk Score by Area", font=dict(size=14, color="#071827")),
            xaxis=dict(title="", tickfont=dict(size=11), tickangle=-30),
            yaxis=dict(title="Risk Score", gridcolor="#F3F4F6"),
            legend=dict(orientation="h", y=1.08, x=0),
            plot_bgcolor="#FFFFFF",
            paper_bgcolor="#FFFFFF",
            height=420,
            margin=dict(l=10, r=10, t=60, b=80),
            showlegend=True,
        )
        st.plotly_chart(fig)

        # Treemap for WBS proportional view
        if n_wbs > 2:
            st.caption("Proportional view — tile size = number of activities, colour = risk rating.")
            plot_df["label_text"] = plot_df.apply(
                lambda r: f"{r['WBS Area']}<br>{r['Rating']} ({r['Risk Score']})",
                axis=1
            )
            fig2 = go.Figure(go.Treemap(
                labels=plot_df["WBS Area"],
                parents=[""] * len(plot_df),
                values=plot_df["Activities"],
                customdata=plot_df[["Risk Score","Rating","_reason","Critical","Negative Float"]].values,
                hovertemplate=(
                    "<b>%{label}</b><br>"
                    "Activities: %{value}<br>"
                    "Risk Score: %{customdata[0]}<br>"
                    "Rating: %{customdata[1]}<br>"
                    "Primary concern: %{customdata[2]}<br>"
                    "Critical: %{customdata[3]}<br>"
                    "Negative Float: %{customdata[4]}<extra></extra>"
                ),
                marker=dict(
                    colors=plot_df["Rating"].map(colour_map),
                    line=dict(color="#FFFFFF", width=2),
                ),
                textfont=dict(size=12, color="#FFFFFF"),
                texttemplate="<b>%{label}</b><br>%{customdata[1]}",
            ))
            fig2.update_layout(
                height=380,
                margin=dict(l=5, r=5, t=5, b=5),
                paper_bgcolor="#FFFFFF",
            )
            st.plotly_chart(fig2)

    # -- RISK CARDS ------------------------------------------------------------
    with tab_cards:
        st.caption("Each card shows a WBS area ranked by risk. Red = most at risk, Green = lowest risk.")

        # Top worst first, then improve
        show_df = hmap_df.head(min(n_wbs, 20))

        for _, row in show_df.iterrows():
            r_col  = row["_colour"]
            r_label= row["Rating"]
            r_chip_style = row["_chip"]
            r_bg   = {"Severe":"#FEF2F2","High":"#FEF2F2","Medium":"#FFFBEB","Low":"#F0FDF4"}.get(r_label,"#F9FAFB")
            r_brd  = {"Severe":"#FECACA","High":"#FECACA","Medium":"#FDE68A","Low":"#BBF7D0"}.get(r_label,"#E5E7EB")

            # Metric pills
            pills = ""
            if row["Negative Float"] > 0:
                pills += f'<span style="background:#FEE2E2;color:#991B1B;font-size:10px;font-weight:700;padding:2px 7px;border-radius:4px;margin-right:4px;">{row["Negative Float"]} neg float</span>'
            if row["Critical"] > 0:
                pills += f'<span style="background:#FECACA;color:#991B1B;font-size:10px;font-weight:700;padding:2px 7px;border-radius:4px;margin-right:4px;">{row["Critical"]} critical</span>'
            if row["Near-Critical"] > 0:
                pills += f'<span style="background:#FEF3C7;color:#B45309;font-size:10px;font-weight:700;padding:2px 7px;border-radius:4px;margin-right:4px;">{row["Near-Critical"]} near-crit</span>'
            if row["Open Logic"] > 0:
                pills += f'<span style="background:#DBEAFE;color:#1D4ED8;font-size:10px;font-weight:700;padding:2px 7px;border-radius:4px;margin-right:4px;">{row["Open Logic"]} open logic</span>'
            if row["Crit Not Started"] > 0:
                pills += f'<span style="background:#FEE2E2;color:#991B1B;font-size:10px;font-weight:700;padding:2px 7px;border-radius:4px;margin-right:4px;">{row["Crit Not Started"]} crit NS</span>'

            st.markdown(
                f'<div style="background:{r_bg};border:1px solid {r_brd};'
                f'border-left:4px solid {r_col};border-radius:0 8px 8px 0;'
                f'padding:12px 16px;margin-bottom:8px;">'
                f'<div style="display:flex;justify-content:space-between;'
                f'align-items:flex-start;flex-wrap:wrap;gap:8px;margin-bottom:8px;">'
                f'<div>'
                f'<div style="font-size:15px;font-weight:700;color:#071827;">'
                f'{row["WBS Area"]}</div>'
                f'<div style="font-size:11px;color:#6B7280;margin-top:2px;">'
                f'{row["Activities"]} activities</div>'
                f'</div>'
                f'<div style="display:flex;align-items:center;gap:8px;">'
                f'<div style="font-size:11px;color:#6B7280;">Score: '
                f'<strong style="color:#071827;">{row["Risk Score"]}</strong></div>'
                f'{chip(r_label, r_chip_style)}'
                f'</div></div>'
                f'<div style="margin-bottom:8px;display:flex;flex-wrap:wrap;gap:3px;">{pills}</div>'
                f'<div style="display:grid;grid-template-columns:1fr 1fr;gap:8px;">'
                f'<div>'
                f'<div style="font-size:10px;font-weight:700;color:#9CA3AF;'
                f'text-transform:uppercase;letter-spacing:0.8px;margin-bottom:2px;">Primary Concern</div>'
                f'<div style="font-size:12px;color:#374151;">{row["_reason"]}</div>'
                f'</div>'
                f'<div>'
                f'<div style="font-size:10px;font-weight:700;color:#9CA3AF;'
                f'text-transform:uppercase;letter-spacing:0.8px;margin-bottom:2px;">Recommended Action</div>'
                f'<div style="font-size:12px;color:#374151;">{row["_action"]}</div>'
                f'</div></div>'
                f'</div>',
                unsafe_allow_html=True,
            )

    # -- FULL TABLE ------------------------------------------------------------
    with tab_table:
        st.caption(f"All {n_wbs} WBS areas ranked by risk score (highest first).")

        # Display columns for table
        display_cols = [
            "WBS Area", "Rating", "Risk Score", "Activities",
            "Critical", "Negative Float", "Near-Critical",
            "Crit Not Started", "Open Logic", "Long Duration",
            "Constrained", "Date Slips", "Labour (hrs)",
            "Worst Float (d)",
        ]
        disp_df = hmap_df[[c for c in display_cols if c in hmap_df.columns]].copy()

        # Colour code the Rating column
        def _style_rating(val):
            colour_map_row = {
                "Severe": "background-color:#FEF2F2;color:#7F1D1D;font-weight:800;",
                "High":   "background-color:#FEF2F2;color:#B91C1C;font-weight:700;",
                "Medium": "background-color:#FFFBEB;color:#B45309;font-weight:600;",
                "Low":    "background-color:#F0FDF4;color:#15803D;font-weight:600;",
            }
            return colour_map_row.get(val, "")

        def _style_neg(val):
            try:
                if int(val) > 0:
                    return "background-color:#FEF2F2;color:#B91C1C;font-weight:700;"
            except Exception:
                pass
            return ""

        st.dataframe(
            disp_df.style
            .map(_style_rating, subset=["Rating"])
            .map(_style_neg, subset=["Negative Float"]),
            hide_index=True,
            height=min(600, 45 + n_wbs * 40),
        )

    # -- EXPORT ----------------------------------------------------------------
    with tab_export:
        st.caption("Download the full WBS risk heatmap data as a formatted Excel workbook.")

        # Build export
        export_cols = [
            "WBS Area","Rating","Risk Score","Activities",
            "Critical","Negative Float","Near-Critical",
            "Crit Not Started","Open Logic","Long Duration",
            "Constrained","Date Slips","Labour (hrs)","Worst Float (d)",
        ]
        export_df = hmap_df[[c for c in export_cols if c in hmap_df.columns]].copy()

        # Add reason + action
        export_df["Primary Concern"] = hmap_df["_reason"]
        export_df["Recommended Action"] = hmap_df["_action"]

        # Summary sheet
        summary_exp = pd.DataFrame({
            "Rating":         ["Severe","High","Medium","Low"],
            "WBS Areas":      [n_severe, n_high, n_medium, n_low],
            "Description": [
                "Immediate PM attention required — critical schedule issues",
                "High priority review — significant risk factors present",
                "Monitor closely — some risk factors present",
                "Healthy — no significant immediate issues",
            ],
        })

        xls = export_df_to_excel({
            "Summary":       summary_exp,
            "WBS Heatmap":   export_df,
            "Severe Areas":  export_df[export_df["Rating"]=="Severe"] if n_severe else pd.DataFrame(columns=["No data"]),
            "High Areas":    export_df[export_df["Rating"]=="High"]   if n_high   else pd.DataFrame(columns=["No data"]),
        })

        dl_col, _ = st.columns([1,3])
        dl_col.download_button(
            label="Download WBS Risk Heatmap",
            data=xls,
            file_name=f"wbs_risk_heatmap_{datetime.now().strftime('%Y%m%d')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="dl_019",
        )




# =============================================================================
# PAGE: FORENSIC PLANNER
# Six tools: driving path · delay analysis · constraint impact ·
#            float paths · logic integrity · delay narrative
# =============================================================================

# -- Helpers -------------------------------------------------------------------

def _fp_activity_label(row) -> str:
    code = str(row.get("task_code", "?"))
    name = str(row.get("task_name", "?"))
    return f"{code}  —  {name[:60]}"


def _fp_float_colour(tf) -> str:
    try:
        f = float(tf)
        if f < 0:   return "#C0392B"
        if f == 0:  return "#C0392B"
        if f <= 10: return "#E8951D"
        return "#1E7A4E"
    except Exception:
        return "#6B7280"


def _fp_row_style(val, col):
    """Apply cell colour based on value and column name."""
    if "float" in col.lower() or col == "Float (d)":
        try:
            f = float(val)
            if f < 0:   return "background-color:#FAEAEA;color:#9E2828;font-weight:700;"
            if f == 0:  return "background-color:#FAEAEA;color:#9E2828;"
            if f <= 10: return "background-color:#FDF3E0;color:#8A5A00;"
        except Exception:
            pass
    return ""


def _fmt_task_table(df: pd.DataFrame, extra_cols: dict = None) -> pd.DataFrame:
    """Return a clean display DataFrame with standard columns."""
    base = {
        "task_code":        "Activity ID",
        "task_name":        "Activity Name",
        "wbs_path":         "WBS",
        "eff_start":        "Start",
        "eff_finish":       "Finish",
        "orig_dur_days":    "Dur (d)",
        "total_float_days": "Float (d)",
        "status":           "Status",
        "is_critical":      "Critical",
    }
    if extra_cols:
        base.update(extra_cols)
    avail = {k: v for k, v in base.items() if k in df.columns}
    out = df[list(avail.keys())].copy().rename(columns=avail)
    for col in ["Start", "Finish"]:
        if col in out.columns:
            out[col] = out[col].apply(format_date)
    if "Status" in out.columns:
        out["Status"] = out["Status"].apply(lambda x: _status_label(str(x)))
    if "Critical" in out.columns:
        out["Critical"] = out["Critical"].apply(lambda x: "Yes" if x else "")
    if "Float (d)" in out.columns:
        out["Float (d)"] = out["Float (d)"].apply(
            lambda x: round(float(x), 1) if x is not None and str(x) not in ("", "nan") else "-"
        )
    return out


# -----------------------------------------------------------------------------
# FEATURE 1: Driving Path to Activity
# -----------------------------------------------------------------------------

def _fp_driving_path(tasks: pd.DataFrame, rels: pd.DataFrame,
                     near_crit_days: float):
    st.markdown(
        '<div class="section-label">Driving Path to Activity</div>',
        unsafe_allow_html=True,
    )
    st.caption(
        "Traces backwards from the selected activity through the schedule network "
        "to identify the chain of activities that is driving its current forecast date. "
        "Use this to understand the root cause of a date."
    )

    G           = build_graph(tasks, rels)
    task_lookup = tasks.set_index("task_id").to_dict("index")
    act_labels  = tasks.apply(_fp_activity_label, axis=1).tolist()

    sel = st.selectbox("Select target activity or milestone",
                       act_labels, key="fp_dp_sel", label_visibility="collapsed")
    sel_idx = act_labels.index(sel)
    target  = tasks.iloc[sel_idx]
    tid     = target["task_id"]

    if st.button("Trace Driving Path", key="fp_dp_btn"):
        with st.spinner("Tracing driving path..."):
            path = driving_path_to_activity(G, tasks, rels, tid)

        if not path:
            st.warning("No driving path found. Check the activity has predecessor logic.")
            return

        st.session_state["fp_dp_result"] = path
        st.session_state["fp_dp_target"] = tid

    path = st.session_state.get("fp_dp_result")
    if not path:
        return

    # Build display table
    rows = []
    for i, pid in enumerate(path):
        t  = task_lookup.get(pid, {})
        tf = t.get("total_float_days")
        is_target = (pid == st.session_state.get("fp_dp_target"))
        # Relationship to next
        link, lag = "-", 0
        if i < len(path) - 1:
            nxt = path[i + 1]
            rel_row = rels[(rels.get("pred_task_id", pd.Series(dtype=str)) == pid) &
                          (rels.get("succ_task_id", pd.Series(dtype=str)) == nxt)] \
                if not rels.empty else pd.DataFrame()
            if not rel_row.empty:
                link = _rel_label(rel_row["rel_type"].iloc[0]) if "rel_type" in rel_row.columns else "FS"
                lag  = safe_float(rel_row["lag_days"].iloc[0] if "lag_days" in rel_row.columns else 0, 0)
        rows.append({
            "Step":          i + 1,
            "Activity ID":   t.get("task_code", pid),
            "Activity Name": str(t.get("task_name", ""))[:55],
            "Start":         format_date(t.get("eff_start")),
            "Finish":        format_date(t.get("eff_finish")),
            "Float (d)":     round(float(tf), 1) if tf is not None else "-",
            "Link to Next":  link if not is_target else "TARGET",
            "Lag (d)":       lag  if not is_target else "-",
            "Note":          "◀ TARGET" if is_target else "",
        })

    path_df = pd.DataFrame(rows)

    def _style(row):
        if row.get("Note") == "◀ TARGET":
            return ["background:#0B1929;color:white;font-weight:700;"] * len(row)
        tf = row.get("Float (d)", 999)
        try:
            f = float(tf)
            if f < 0:  return ["background:#FAEAEA;"] * len(row)
            if f == 0: return ["background:#FDF3E0;"] * len(row)
        except Exception:
            pass
        return [""] * len(row)

    st.markdown(
        f'<div style="font-size:13px;margin-bottom:8px;">'
        f'Driving path has <strong>{len(path)}</strong> activities. '
        f'The chain starts at <strong>{task_lookup.get(path[0],{}).get("task_code","?")}</strong> '
        f'and ends at <strong>{task_lookup.get(path[-1],{}).get("task_code","?")}</strong>.'
        f'</div>',
        unsafe_allow_html=True,
    )
    st.dataframe(path_df.style.apply(_style, axis=1),
                 hide_index=True)

    # Gantt
    path_tasks = tasks[tasks["task_id"].isin(path)].copy()
    gantt = path_tasks.dropna(subset=["eff_start","eff_finish"]) if "eff_start" in path_tasks.columns else pd.DataFrame()
    if not gantt.empty:
        gantt["Label"] = gantt["task_code"].astype(str) + "  " + gantt["task_name"].astype(str).str[:40]
        gantt["Type"]  = gantt["task_id"].apply(
            lambda x: "Target" if x == st.session_state.get("fp_dp_target") else "Driving Chain"
        )
        fig = px.timeline(gantt, x_start="eff_start", x_end="eff_finish",
                          y="Label", color="Type",
                          color_discrete_map={"Target":"#0B1929","Driving Chain":"#C0392B"},
                          title="Driving Path Gantt")
        fig.update_yaxes(autorange="reversed")
        fig.add_vline(x=str(datetime.now()), line_dash="dot", line_color="#8A9DB0",
                      annotation_text="Today")
        fig.update_layout(height=max(280, 50+len(gantt)*28),
                          margin=dict(l=10,r=10,t=40,b=10),
                          plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF")
        st.plotly_chart(fig)

    # Export
    xls = export_df_to_excel({"Driving Path": path_df})
    st.download_button("Export Driving Path", xls,
                       f"driving_path_{task_lookup.get(path[-1],{}).get('task_code','activity')}.xlsx",
                       "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                       key="fp_dp_export")


# -----------------------------------------------------------------------------
# FEATURE 2: Delay Analysis (two XER revisions)
# -----------------------------------------------------------------------------

def _fp_delay_analysis(tasks: pd.DataFrame, near_crit_days: float):
    st.markdown('<div class="section-label">Delay Analysis — Programme Comparison</div>',
                unsafe_allow_html=True)
    st.caption(
        "Compare two programme revisions to identify which activities have moved, "
        "how much they have slipped or improved, and what the overall delay position is."
    )

    has_comp = "_mi_prev" in st.session_state and "_mi_curr" in st.session_state

    if not has_comp:
        st.info(
            "No comparison data loaded. Upload a previous and current XER on the "
            "**Comparison** page, then return here for delay analysis."
        )
        return

    prev_t = st.session_state["_mi_prev"]["tasks_df"]
    curr_t = st.session_state["_mi_curr"]["tasks_df"]

    if prev_t.empty or curr_t.empty:
        st.warning("Comparison data is incomplete.")
        return

    # Merge on task_code
    merged = prev_t[["task_code","task_name","eff_start","eff_finish","total_float_days",
                      "wbs_path"]].merge(
        curr_t[["task_code","eff_start","eff_finish","total_float_days"]],
        on="task_code", suffixes=("_prev","_curr"), how="inner"
    )

    def _days(prev, curr):
        try:
            return int((pd.Timestamp(curr) - pd.Timestamp(prev)).days)
        except Exception:
            return None

    merged["Finish Movement (d)"] = merged.apply(
        lambda r: _days(r["eff_finish_prev"], r["eff_finish_curr"]), axis=1)
    merged["Start Movement (d)"]  = merged.apply(
        lambda r: _days(r["eff_start_prev"],  r["eff_start_curr"]),  axis=1)
    merged["Float Change (d)"]    = merged.apply(
        lambda r: (safe_float(r["total_float_days_curr"],0) -
                   safe_float(r["total_float_days_prev"],0)), axis=1)

    merged["Status"] = merged["Finish Movement (d)"].apply(
        lambda d: "Slipped" if d and d > 0 else ("Improved" if d and d < 0 else "Unchanged")
    )

    # Summary KPIs
    n_slip    = int((merged["Finish Movement (d)"] > 0).sum())
    n_improve = int((merged["Finish Movement (d)"] < 0).sum())
    n_same    = int((merged["Finish Movement (d)"] == 0).sum())
    worst     = int(merged["Finish Movement (d)"].max()) if n_slip else 0
    best      = int(merged["Finish Movement (d)"].min()) if n_improve else 0

    c1, c2, c3, c4 = st.columns(4)
    c1.markdown(kpi_card("Slipped", n_slip,     "activities later",  "red"   if n_slip   else "green"), unsafe_allow_html=True)
    c2.markdown(kpi_card("Improved",n_improve,   "activities earlier","green" if n_improve else "grey"),  unsafe_allow_html=True)
    c3.markdown(kpi_card("Unchanged",n_same,     "no movement",       "navy"),                            unsafe_allow_html=True)
    c4.markdown(kpi_card("Worst Slip",f"{worst}d","days",             "red"   if worst > 0 else "green"), unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)

    # Display table
    disp = merged[[
        "task_code","task_name","wbs_path",
        "eff_finish_prev","eff_finish_curr",
        "Finish Movement (d)","Float Change (d)","Status"
    ]].copy().rename(columns={
        "task_code":"Activity ID","task_name":"Activity Name","wbs_path":"WBS",
        "eff_finish_prev":"Previous Finish","eff_finish_curr":"Current Finish",
    })
    for col in ["Previous Finish","Current Finish"]:
        disp[col] = disp[col].apply(format_date)

    filter_opt = st.selectbox("Filter",
                              ["All","Slipped only","Improved only","Critical slips (>14d)"],
                              key="fp_da_filter", label_visibility="collapsed")
    if filter_opt == "Slipped only":
        disp = disp[disp["Status"] == "Slipped"]
    elif filter_opt == "Improved only":
        disp = disp[disp["Status"] == "Improved"]
    elif filter_opt == "Critical slips (>14d)":
        disp = disp[disp["Finish Movement (d)"] > 14]

    def _da_style(row):
        m = row.get("Finish Movement (d)", 0)
        try:
            if m > 14:  return ["background:#FAEAEA;"] * len(row)
            if m > 0:   return ["background:#FDF3E0;"] * len(row)
            if m < 0:   return ["background:#E4F5EC;"] * len(row)
        except Exception:
            pass
        return [""] * len(row)

    st.markdown(f'<div class="section-label">{len(disp)} activities shown</div>',
                unsafe_allow_html=True)
    st.dataframe(disp.style.apply(_da_style, axis=1),
                 hide_index=True, height=400)

    # Movement chart
    chart_df = merged.copy()
    chart_df["WBS_top"] = chart_df["wbs_path"].apply(
        lambda x: str(x).split(" > ")[0][:30] if x and str(x).strip() not in ("","nan") else "Unknown"
    )
    wbs_mv = chart_df.groupby("WBS_top").agg(
        avg_slip=("Finish Movement (d)", "mean"),
        n=("task_code","count")
    ).reset_index().sort_values("avg_slip", ascending=False)
    if not wbs_mv.empty:
        fig_mv = px.bar(wbs_mv, x="WBS_top", y="avg_slip",
                        title="Average Finish Movement by WBS Area (days)",
                        color="avg_slip",
                        color_continuous_scale=["#1E7A4E","#F5F0E8","#C0392B"],
                        labels={"WBS_top":"WBS Area","avg_slip":"Avg Movement (d)"})
        fig_mv.add_hline(y=0, line_color="#0B1929", line_width=1)
        fig_mv.update_layout(height=320, margin=dict(l=10,r=10,t=40,b=80),
                             plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
                             coloraxis_showscale=False)
        st.plotly_chart(fig_mv)

    # Export
    xls = export_df_to_excel({
        "Delay Analysis":    disp,
        "All Movements":     merged[["task_code","task_name","eff_finish_prev","eff_finish_curr",
                                     "Finish Movement (d)","Float Change (d)","Status"]],
        "Slipped >14d":      merged[merged["Finish Movement (d)"] > 14],
        "WBS Summary":       wbs_mv,
    })
    st.download_button("Export Delay Analysis", xls,
                       f"delay_analysis_{datetime.now().strftime('%Y%m%d')}.xlsx",
                       "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                       key="fp_da_export")


# -----------------------------------------------------------------------------
# FEATURE 3: Constraint Impact Analysis
# -----------------------------------------------------------------------------

def _fp_constraint_impact(tasks: pd.DataFrame):
    st.markdown('<div class="section-label">Constraint Impact Analysis</div>',
                unsafe_allow_html=True)
    st.caption(
        "Identifies all constrained activities and estimates the impact each constraint "
        "has on the programme. Constraints override logic and can create artificial float "
        "or drive negative float."
    )

    if "cstr_type" not in tasks.columns:
        st.info("No constraint data found in this programme.")
        return

    constrained = tasks[
        tasks["cstr_type"].apply(
            lambda x: bool(x) and str(x).strip() not in ("","None","nan")
        )
    ].copy()

    if constrained.empty:
        st.success("No constraints found in this programme.")
        return

    n_total = len(constrained)
    n_neg   = int(constrained["total_float_days"].apply(
        lambda f: safe_float(f, 0) < 0).sum()) if "total_float_days" in constrained.columns else 0
    n_crit  = int(constrained["is_critical"].sum()) if "is_critical" in constrained.columns else 0

    c1, c2, c3 = st.columns(3)
    c1.markdown(kpi_card("Constrained Activities", n_total, "", "amber"),   unsafe_allow_html=True)
    c2.markdown(kpi_card("With Negative Float",    n_neg,   "constraint may be causing delay", "red" if n_neg else "green"), unsafe_allow_html=True)
    c3.markdown(kpi_card("Critical + Constrained", n_crit,  "high risk combination", "red" if n_crit else "green"),         unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)

    # Categorise constraints
    constrained["Impact"] = constrained.apply(lambda r: (
        "DRIVING — Negative Float"   if safe_float(r.get("total_float_days"), 0) < 0 else
        "HIGH — Critical + Constrained" if r.get("is_critical") else
        "MEDIUM — Near-Critical"     if r.get("is_near_critical") else
        "LOW — Has Float"
    ), axis=1)

    disp = constrained[["task_code","task_name","wbs_path","cstr_type","cstr_date",
                         "eff_finish","total_float_days","is_critical","Impact"]].copy()
    disp.rename(columns={
        "task_code":"Activity ID","task_name":"Activity Name","wbs_path":"WBS",
        "cstr_type":"Constraint Type","cstr_date":"Constraint Date",
        "eff_finish":"Forecast Finish","total_float_days":"Float (d)","is_critical":"Critical"
    }, inplace=True)
    for col in ["Constraint Date","Forecast Finish"]:
        if col in disp.columns:
            disp[col] = disp[col].apply(format_date)
    if "Critical" in disp.columns:
        disp["Critical"] = disp["Critical"].apply(lambda x: "Yes" if x else "")
    if "Float (d)" in disp.columns:
        disp["Float (d)"] = disp["Float (d)"].apply(
            lambda x: round(float(x),1) if x is not None and str(x) not in ("","nan") else "-")

    disp = disp.sort_values("Impact")

    def _cst_style(row):
        imp = row.get("Impact","")
        if "Negative" in imp: return ["background:#FAEAEA;font-weight:700;"] * len(row)
        if "HIGH"     in imp: return ["background:#FDF3E0;"] * len(row)
        if "MEDIUM"   in imp: return ["background:#FEFCE8;"] * len(row)
        return [""] * len(row)

    st.dataframe(disp.style.apply(_cst_style, axis=1),
                 hide_index=True)

    # Plain English summary
    neg_names = constrained[constrained["total_float_days"].apply(
        lambda f: safe_float(f,0) < 0)]["task_code"].tolist()[:5]
    if neg_names:
        st.markdown(
            f'<div style="background:#FAEAEA;border-left:3px solid #C0392B;'
            f'border-radius:0 4px 4px 0;padding:10px 14px;margin-top:10px;">'
            f'<strong>What this means:</strong> '
            f'Activities {", ".join(neg_names)} have constraints that are creating negative float. '
            f'These constraints are overriding the programme logic and making the schedule '
            f'show dates that cannot be achieved. Review whether these constraints are still '
            f'valid or whether they should be removed/revised.'
            f'</div>',
            unsafe_allow_html=True,
        )

    xls = export_df_to_excel({"Constraint Impact": disp})
    st.download_button("Export Constraint Analysis", xls,
                       f"constraint_analysis_{datetime.now().strftime('%Y%m%d')}.xlsx",
                       "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                       key="fp_ci_export")


# -----------------------------------------------------------------------------
# FEATURE 4: Float Path Analysis
# -----------------------------------------------------------------------------

def _fp_float_paths(tasks: pd.DataFrame, rels: pd.DataFrame, near_crit_days: float):
    st.markdown('<div class="section-label">Float Path Analysis</div>',
                unsafe_allow_html=True)
    st.caption(
        "Ranks every schedule path by total float, from most critical to most healthy. "
        "Identifies the second, third and fourth critical paths — the ones that will "
        "become critical next if the programme slips further."
    )

    if tasks.empty or rels.empty:
        st.info("Logic data required for float path analysis.")
        return

    tasks_c = get_critical_threshold(tasks, near_crit_days)
    G       = build_graph(tasks_c, rels)

    # Group by float band
    if "total_float_days" not in tasks_c.columns:
        st.warning("Total float data not available.")
        return

    tasks_c["float_band"] = tasks_c["total_float_days"].apply(lambda f: (
        "Negative Float"    if safe_float(f,0) < 0   else
        "Critical (0d)"     if safe_float(f,0) == 0  else
        f"0–{near_crit_days}d (Near-Critical)" if safe_float(f,0) <= near_crit_days else
        f"{near_crit_days+1}–30d"              if safe_float(f,0) <= 30 else
        "31–60d"            if safe_float(f,0) <= 60 else
        ">60d"
    ))

    band_order = [
        "Negative Float", "Critical (0d)",
        f"0–{near_crit_days}d (Near-Critical)",
        f"{near_crit_days+1}–30d", "31–60d", ">60d"
    ]
    band_colours = {
        "Negative Float": "#FAEAEA",
        "Critical (0d)": "#FDF3E0",
        f"0–{near_crit_days}d (Near-Critical)": "#FEFCE8",
        f"{near_crit_days+1}–30d": "#F0F9F4",
        "31–60d": "#E8F5EE", ">60d": "#F5F9FF",
    }
    band_chip = {
        "Negative Float": "red", "Critical (0d)": "red",
        f"0–{near_crit_days}d (Near-Critical)": "amber",
        f"{near_crit_days+1}–30d": "green",
        "31–60d": "green", ">60d": "blue",
    }

    summary_rows = []
    for band in band_order:
        grp = tasks_c[tasks_c["float_band"] == band]
        if grp.empty:
            continue
        min_f = grp["total_float_days"].apply(lambda x: safe_float(x,999)).min()
        max_f = grp["total_float_days"].apply(lambda x: safe_float(x,0)).max()
        summary_rows.append({
            "Float Band":  band,
            "Activities":  len(grp),
            "Min Float":   round(min_f,1),
            "Max Float":   round(max_f,1),
            "% of Total":  f"{len(grp)/len(tasks_c)*100:.0f}%",
        })

    if summary_rows:
        st.markdown('<div class="section-label" style="margin-bottom:8px;">Summary by Float Band</div>',
                    unsafe_allow_html=True)
        summary_df = pd.DataFrame(summary_rows)

        # Donut chart
        fig_pie = px.pie(summary_df, names="Float Band", values="Activities",
                         hole=0.55, title="Activities by Float Band",
                         color="Float Band",
                         color_discrete_map={
                             "Negative Float":"#C0392B", "Critical (0d)":"#E8951D",
                             f"0–{near_crit_days}d (Near-Critical)":"#F5C842",
                             f"{near_crit_days+1}–30d":"#5BA08A",
                             "31–60d":"#1E7A4E",">60d":"#8AAABF"
                         })
        fig_pie.update_traces(textposition="outside", textinfo="percent+label")
        fig_pie.update_layout(height=320, margin=dict(l=10,r=10,t=40,b=10),
                              paper_bgcolor="#FFFFFF", showlegend=False)

        col_chart, col_table = st.columns([2,3])
        with col_chart:
            st.plotly_chart(fig_pie)
        with col_table:
            st.dataframe(summary_df, hide_index=True)

    # Drill into selected band
    st.markdown("<br>", unsafe_allow_html=True)
    band_opts = [b for b in band_order if b in tasks_c["float_band"].values]
    selected_band = st.selectbox("Drill into float band", band_opts,
                                 key="fp_fp_band", label_visibility="collapsed")

    band_tasks = tasks_c[tasks_c["float_band"] == selected_band]
    band_tasks_sorted = band_tasks.sort_values("total_float_days")

    st.markdown(
        f'<div style="margin-bottom:8px;">{chip(selected_band, band_chip.get(selected_band,"grey"))} '
        f'&nbsp;{len(band_tasks)} activities</div>',
        unsafe_allow_html=True,
    )
    st.dataframe(_fmt_task_table(band_tasks_sorted), hide_index=True)

    xls = export_df_to_excel(
        {band: _fmt_task_table(tasks_c[tasks_c["float_band"] == band])
         for band in band_order if band in tasks_c["float_band"].values}
    )
    st.download_button("Export Float Path Analysis", xls,
                       f"float_paths_{datetime.now().strftime('%Y%m%d')}.xlsx",
                       "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                       key="fp_fp_export")


# -----------------------------------------------------------------------------
# FEATURE 5: Logic Integrity Check
# -----------------------------------------------------------------------------

def _fp_logic_integrity(tasks: pd.DataFrame, rels: pd.DataFrame):
    st.markdown('<div class="section-label">Logic Integrity Check</div>',
                unsafe_allow_html=True)
    st.caption(
        "Checks the programme for structural logic problems: open ends, circular dependencies, "
        "excessive lags, redundant logic and missing links. These issues reduce schedule "
        "reliability and should be resolved before relying on float calculations."
    )

    if tasks.empty:
        return

    checks = []

    # Check 1: Open starts (no predecessor)
    if not rels.empty and "task_id" in tasks.columns:
        with_pred = set(rels["succ_task_id"].dropna()) if "succ_task_id" in rels.columns else set()
        no_pred   = tasks[~tasks["task_id"].isin(with_pred)]
        if "task_type" in no_pred.columns:
            no_pred = no_pred[~no_pred["task_type"].astype(str).str.contains("Milestone|LOE|WBS",na=False)]
        checks.append(("OPEN START", "No Predecessor",
                        len(no_pred), "amber",
                        "Activities not driven by logic. Float values may be unreliable.",
                        no_pred))

    # Check 2: Open finishes (no successor)
    if not rels.empty and "task_id" in tasks.columns:
        with_succ = set(rels["pred_task_id"].dropna()) if "pred_task_id" in rels.columns else set()
        no_succ   = tasks[~tasks["task_id"].isin(with_succ)]
        if "task_type" in no_succ.columns:
            no_succ = no_succ[~no_succ["task_type"].astype(str).str.contains("Finish Milestone|LOE|WBS",na=False)]
        checks.append(("OPEN FINISH", "No Successor",
                        len(no_succ), "amber",
                        "Activities not driving any subsequent work. May inflate float.",
                        no_succ))

    # Check 3: Excessive lag (>10 days)
    if not rels.empty and "lag_days" in rels.columns:
        big_lag = rels[rels["lag_days"].apply(lambda l: safe_float(l,0) > 10)]
        n_lag = len(big_lag)
        max_lag = int(big_lag["lag_days"].max()) if n_lag > 0 else 0
        checks.append(("EXCESSIVE LAG", f"Lag >10d (max: {max_lag}d)",
                        n_lag, "amber" if n_lag > 0 else "green",
                        "Lags disguise missing activities and can hide critical path risk.",
                        None))

    # Check 4: Circular dependencies
    circ_count = 0
    circ_examples = []
    if not rels.empty:
        try:
            import networkx as nx
            G_check = build_graph(tasks, rels)
            cycles = list(nx.simple_cycles(G_check))
            circ_count = len(cycles)
            for c in cycles[:3]:
                names = [str(tasks[tasks["task_id"]==n]["task_code"].values[0])
                         if not tasks[tasks["task_id"]==n].empty else n for n in c[:3]]
                circ_examples.append(" → ".join(names))
        except Exception:
            pass
    checks.append(("CIRCULAR LOGIC", "Circular dependencies",
                   circ_count, "red" if circ_count > 0 else "green",
                   "Circular logic makes float calculation impossible and crashes scheduling software.",
                   None))

    # Check 5: Long durations (>60 days)
    if "orig_dur_days" in tasks.columns:
        long_dur = tasks[tasks["orig_dur_days"].apply(lambda d: safe_float(d,0) > 60)]
        max_dur  = int(tasks["orig_dur_days"].max()) if not tasks.empty else 0
        checks.append(("LONG DURATION", f"Duration >60d (max: {max_dur}d)",
                        len(long_dur), "amber" if len(long_dur) > 0 else "green",
                        "Long activities are hard to control and mask schedule problems.",
                        long_dur))

    # Check 6: Start-to-Start or Finish-to-Finish with zero lag (potential soft logic)
    if not rels.empty and "rel_type" in rels.columns:
        soft = rels[rels["rel_type"].isin(["SS","FF"]) &
                    rels["lag_days"].apply(lambda l: safe_float(l,0) == 0)]
        checks.append(("SOFT LOGIC", "SS/FF with zero lag",
                        len(soft), "grey",
                        "SS/FF relationships with no lag may indicate logic that needs review.",
                        None))

    # Render summary cards
    cols = st.columns(3)
    for i, (code, label, count, style, detail, _df) in enumerate(checks):
        col = cols[i % 3]
        col.markdown(kpi_card(label, count, detail[:40] + ("..." if len(detail) > 40 else ""), style),
                     unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # Circular dependency warning
    if circ_count > 0:
        st.markdown(
            f'<div class="attn">'
            f'<strong>Circular Dependencies Found: {circ_count}</strong><br>'
            f'{"<br>".join(circ_examples)}<br>'
            f'These must be resolved before the programme can be scheduled correctly.'
            f'</div>',
            unsafe_allow_html=True,
        )

    # Detail expanders
    for code, label, count, style, detail, df in checks:
        if df is not None and not df.empty and count > 0:
            with st.expander(f"{label}  —  {count} activities", expanded=False):
                st.dataframe(_fmt_task_table(df), hide_index=True)

    # Export
    export_sheets = {}
    for code, label, count, style, detail, df in checks:
        if df is not None and not df.empty:
            export_sheets[label[:30]] = _fmt_task_table(df)
    summary_data = pd.DataFrame([
        {"Check": label, "Count": count, "Severity": style.title(), "Description": detail}
        for _, label, count, style, detail, _ in checks
    ])
    export_sheets["Summary"] = summary_data

    xls = export_df_to_excel(export_sheets)
    st.download_button("Export Logic Integrity Report", xls,
                       f"logic_integrity_{datetime.now().strftime('%Y%m%d')}.xlsx",
                       "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                       key="fp_li_export")


# -----------------------------------------------------------------------------
# FEATURE 6: Auto-Generate Delay Narrative
# -----------------------------------------------------------------------------

def _fp_delay_narrative(data: dict, near_crit_days: float):
    st.markdown('<div class="section-label">Auto-Generate Delay Narrative</div>',
                unsafe_allow_html=True)
    st.caption(
        "Generates a plain-English narrative of the programme position and delay events, "
        "suitable for inclusion in a project report, contract notice or delay analysis."
    )

    tasks = data.get("tasks_df", pd.DataFrame())
    rels  = data.get("relationships_df", pd.DataFrame())
    proj  = data.get("project_info", {})

    if tasks.empty:
        return

    tasks_c = get_critical_threshold(tasks, near_crit_days)
    pname   = proj.get("name","")
    ddate   = proj.get("data_date")
    dd_str  = format_date(ddate) if ddate else "not available"

    # Gather metrics
    n_total  = len(tasks_c)
    n_crit   = int(tasks_c["is_critical"].sum()) if "is_critical" in tasks_c.columns else 0
    n_neg    = int(tasks_c["total_float_days"].apply(lambda f: safe_float(f,0)<0).sum()) if "total_float_days" in tasks_c.columns else 0
    n_nc     = int(tasks_c["is_near_critical"].sum()) if "is_near_critical" in tasks_c.columns else 0

    worst_float = None
    if n_neg > 0 and "total_float_days" in tasks_c.columns:
        worst_float = round(float(tasks_c["total_float_days"].min()), 1)

    # WBS areas
    concern_wbs = []
    if "wbs_path" in tasks_c.columns and n_neg > 0:
        neg_df = tasks_c[tasks_c["total_float_days"].apply(lambda f: safe_float(f,0)<0)].copy()
        neg_df["wbs_top"] = neg_df["wbs_path"].apply(
            lambda x: str(x).split(" > ")[0].strip() if x and str(x).strip() not in ("","nan") else "Unknown"
        )
        concern_wbs = neg_df.groupby("wbs_top").size().sort_values(ascending=False).head(3).index.tolist()

    # Open logic
    n_open = 0
    if not rels.empty and "task_id" in tasks_c.columns:
        wp = set(rels["succ_task_id"].dropna()) if "succ_task_id" in rels.columns else set()
        ws = set(rels["pred_task_id"].dropna()) if "pred_task_id" in rels.columns else set()
        n_open = len(tasks_c[~tasks_c["task_id"].isin(wp)]) + len(tasks_c[~tasks_c["task_id"].isin(ws)])

    # Comparison
    has_comp = "_mi_prev" in st.session_state and "_mi_curr" in st.session_state
    slip_count, worst_slip = 0, 0
    if has_comp:
        try:
            prev_t = st.session_state["_mi_prev"]["tasks_df"]
            curr_t = st.session_state["_mi_curr"]["tasks_df"]
            mg = prev_t[["task_code","eff_finish"]].merge(
                curr_t[["task_code","eff_finish"]], on="task_code", suffixes=("_p","_c"), how="inner")
            for _, r in mg.iterrows():
                try:
                    d = int((pd.Timestamp(r["eff_finish_c"])-pd.Timestamp(r["eff_finish_p"])).days)
                    if d > 0:
                        slip_count += 1
                        worst_slip = max(worst_slip, d)
                except Exception:
                    pass
        except Exception:
            pass

    # Build narrative
    paras = []

    # Para 1: Programme identity
    name_str = f"**{pname}**" if pname else "The programme"
    paras.append(
        f"{name_str} comprises **{n_total:,} activities** with a programme data date of **{dd_str}**."
    )

    # Para 2: Critical path
    if n_crit > 0:
        paras.append(
            f"The critical path currently contains **{n_crit} activities** with zero total float. "
            "Any delay to activities on the critical path will directly delay the project completion date "
            "with no available contingency."
        )

    # Para 3: Delay / negative float
    if n_neg > 0:
        wbs_str = f" Programme risk is concentrated in **{', '.join(concern_wbs)}**." if concern_wbs else ""
        paras.append(
            f"**{n_neg} {'activity' if n_neg==1 else 'activities'} currently "
            f"{'has' if n_neg==1 else 'have'} negative total float**, "
            f"with the worst position being **{abs(worst_float)} days** negative. "
            f"This means the programme as currently scheduled cannot achieve its target completion dates "
            f"on these paths without recovery action.{wbs_str}"
        )
    else:
        paras.append("There are no activities with negative float at the current data date.")

    # Para 4: Near-critical
    if n_nc > 0:
        paras.append(
            f"A further **{n_nc} activities are near-critical**, each with between 0 and "
            f"{near_crit_days} days of total float. These activities have limited contingency "
            "and are at risk of moving onto the critical path if any further delays are encountered."
        )

    # Para 5: Open logic
    if n_open > 10:
        paras.append(
            f"The programme contains **{n_open} activities with open-ended logic** "
            "(missing predecessor or successor relationships). "
            "This reduces the reliability of the float calculations and the critical path analysis, "
            "as these activities are not fully constrained within the programme network."
        )

    # Para 6: Comparison
    if has_comp and slip_count > 0:
        paras.append(
            f"Comparison against the previous programme revision shows that **{slip_count} activities "
            f"have later forecast finish dates**, with the worst individual slippage being **{worst_slip} days**. "
            "This movement should be investigated and a recovery programme agreed with the delivery team."
        )

    # Para 7: Recommendation
    if n_neg > 0 or slip_count > 20:
        paras.append(
            "**Recommended Action:** A programme recovery plan should be prepared and agreed with the "
            "Project Director. The planner should review all negative float activities, confirm whether "
            "the logic reflects the actual delivery sequence, and identify opportunities for acceleration "
            "or resequencing. The programme should not be formally accepted in its current position "
            "without a clear recovery strategy."
        )
    elif n_nc > 0 or n_open > 10:
        paras.append(
            "**Recommended Action:** The programme team should address open logic and monitor "
            "near-critical activities on a weekly basis. Any slippage on near-critical work should "
            "be escalated immediately to prevent it from impacting the project completion date."
        )
    else:
        paras.append(
            "The programme is currently in a healthy position. "
            "Continue to monitor the critical path and maintain the current schedule performance."
        )

    # Display
    narrative_html = "".join(
        f'<p style="font-size:14px;color:#1C2B3A;line-height:1.8;margin-bottom:12px;">{p}</p>'
        for p in paras
    )
    copy_button("\n\n".join(p.replace("**","") for p in paras), "Copy Narrative", key="fp_dn_copy")
    st.markdown(
        f'<div style="background:#FFFFFF;border:1px solid #D4DCE4;border-radius:6px;'
        f'padding:24px 28px;margin-bottom:16px;">'
        f'<div style="font-size:10px;font-weight:700;color:#8A9DB0;text-transform:uppercase;'
        f'letter-spacing:1.2px;margin-bottom:16px;font-family:\'IBM Plex Mono\',monospace;">'
        f'Programme Narrative — {dd_str}</div>'
        f'{narrative_html}</div>',
        unsafe_allow_html=True,
    )

    # Word export
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor

        def _make_narrative_doc():
            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)
            title = doc.add_heading("Programme Delay Narrative", 0)
            title.runs[0].font.color.rgb = RGBColor(11,25,41)
            doc.add_paragraph(f"Programme: {pname or 'Programme'}")
            doc.add_paragraph(f"Data Date: {dd_str}")
            doc.add_paragraph("")
            for p in paras:
                # Strip markdown bold markers
                clean = p.replace("**","")
                doc.add_paragraph(clean)
                doc.add_paragraph("")
            import io
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        word_bytes = _make_narrative_doc()
        st.download_button(
            "Export Narrative to Word",
            word_bytes,
            f"delay_narrative_{datetime.now().strftime('%Y%m%d')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="fp_dn_word",
        )
    except Exception:
        pass

    # Excel export
    narrative_df = pd.DataFrame({"Paragraph": list(range(1,len(paras)+1)),
                                 "Text": [p.replace("**","") for p in paras]})
    xls = export_df_to_excel({"Delay Narrative": narrative_df})
    st.download_button("Export Narrative to Excel", xls,
                       f"delay_narrative_{datetime.now().strftime('%Y%m%d')}.xlsx",
                       "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                       key="fp_dn_excel")


# -----------------------------------------------------------------------------
# MAIN PAGE
# -----------------------------------------------------------------------------

def page_forensic_planner(data: dict, near_crit_days: float):
    """Forensic Planning & Schedule Interrogation page."""
    ctrl_bar(
        "Forensic Planner",
        "Deep-dive schedule interrogation — driving paths, delay analysis, "
        "logic integrity, float paths and delay narrative.",
    )
    mode_toggle_bar()

    tasks = data.get("tasks_df", pd.DataFrame())
    rels  = data.get("relationships_df", pd.DataFrame())

    if tasks.empty:
        empty_state("", "No Programme Loaded",
                    "Upload a programme to use the Forensic Planner.", "")
        return

    tasks_c = get_critical_threshold(tasks, near_crit_days)

    pm_note(
        "The Forensic Planner interrogates your programme in depth. "
        "Use it to understand what is really driving your key dates, "
        "where the programme risk is concentrated, and what the data says "
        "about delay and recovery."
    )

    tab_dp, tab_da, tab_ci, tab_fp, tab_li, tab_dn = st.tabs([
        "Driving Path",
        "Delay Analysis",
        "Constraint Impact",
        "Float Paths",
        "Logic Integrity",
        "Delay Narrative",
    ])

    with tab_dp:
        _fp_driving_path(tasks_c, rels, near_crit_days)

    with tab_da:
        _fp_delay_analysis(tasks_c, near_crit_days)

    with tab_ci:
        _fp_constraint_impact(tasks_c)

    with tab_fp:
        _fp_float_paths(tasks_c, rels, near_crit_days)

    with tab_li:
        _fp_logic_integrity(tasks_c, rels)

    with tab_dn:
        _fp_delay_narrative(data, near_crit_days)




# =============================================================================
# FEATURE 1 — S-CURVE / PROGRESS TRACKING
# =============================================================================

def page_scurve(data: dict, near_crit_days: float):
    ctrl_bar("S-Curve", "Planned vs forecast activity count and duration over time.")
    mode_toggle_bar()

    tasks = data.get("tasks_df", pd.DataFrame())
    if tasks.empty:
        empty_state("", "No Programme Data", "Upload a programme to view the S-Curve.", "")
        return

    pm_note(
        "The S-Curve shows how activities are distributed across the programme timeline. "
        "It lets you see whether work is front-loaded or back-loaded, and — if baseline "
        "dates are available — how much the programme has moved."
    )

    has_target = ("target_finish" in tasks.columns and
                  tasks["target_finish"].dropna().shape[0] > tasks.shape[0] * 0.1)

    tab_act, tab_dur, tab_gantt = st.tabs(
        ["Activity Count Curve", "Duration Burn-Down", "Timeline Overview"]
    )

    # -- Helper: build weekly cumulative series -------------------------------
    def _weekly_cumulative(date_col: pd.Series, label: str) -> pd.DataFrame:
        valid = date_col.dropna()
        if valid.empty:
            return pd.DataFrame()
        dates = pd.to_datetime(valid, errors="coerce").dropna()
        if dates.empty:
            return pd.DataFrame()
        min_d = dates.min().to_period("W").start_time
        max_d = dates.max().to_period("W").start_time + pd.Timedelta(weeks=1)
        weeks = pd.date_range(min_d, max_d, freq="W-MON")
        counts = [(w, int((dates <= w).sum())) for w in weeks]
        df = pd.DataFrame(counts, columns=["Week", label])
        return df

    # -- TAB 1: Activity Count Curve ------------------------------------------
    with tab_act:
        st.caption(
            "Cumulative number of activities finishing each week. "
            + ("Blue = Forecast. Grey = Baseline (target dates)." if has_target
               else "Baseline not available in this XER — showing forecast only.")
        )

        forecast_df = _weekly_cumulative(tasks["eff_finish"], "Forecast")
        fig = go.Figure()

        if has_target:
            baseline_df = _weekly_cumulative(tasks["target_finish"], "Baseline")
            if not baseline_df.empty:
                fig.add_trace(go.Scatter(
                    x=baseline_df["Week"], y=baseline_df["Baseline"],
                    name="Baseline", line=dict(color="#8AAABF", width=2, dash="dash"),
                    fill=None,
                ))

        if not forecast_df.empty:
            fig.add_trace(go.Scatter(
                x=forecast_df["Week"], y=forecast_df["Forecast"],
                name="Forecast", line=dict(color="#0B1929", width=2.5),
                fill="tozeroy", fillcolor="rgba(11,25,41,0.06)",
            ))

        fig.add_vline(
            x=datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
            line_dash="dot", line_color="#E8951D", line_width=1.5,
            annotation_text="Today", annotation_font_color="#E8951D",
        )
        fig.update_layout(
            title="Cumulative Activity Completions",
            xaxis_title="", yaxis_title="Activities Completed",
            plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
            height=380, margin=dict(l=10, r=10, t=40, b=10),
            legend=dict(orientation="h", y=1.08),
            hovermode="x unified",
        )
        st.plotly_chart(fig)

        if has_target and not forecast_df.empty and not baseline_df.empty:
            today = pd.Timestamp.now()
            f_now = int(forecast_df[forecast_df["Week"] <= today]["Forecast"].max()) if not forecast_df[forecast_df["Week"] <= today].empty else 0
            b_now = int(baseline_df[baseline_df["Week"] <= today]["Baseline"].max()) if not baseline_df[baseline_df["Week"] <= today].empty else 0
            diff  = f_now - b_now
            col1, col2, col3 = st.columns(3)
            col1.markdown(kpi_card("Forecast Completions", f_now, "to date", "navy"), unsafe_allow_html=True)
            col2.markdown(kpi_card("Baseline Completions", b_now, "to date", "grey"), unsafe_allow_html=True)
            col3.markdown(kpi_card("Variance", f"{'+' if diff>=0 else ''}{diff}", "activities ahead/behind baseline",
                                   "green" if diff >= 0 else "red"), unsafe_allow_html=True)

    # -- TAB 2: Duration Burn-Down --------------------------------------------
    with tab_dur:
        st.caption(
            "Remaining duration burns down to zero at project completion. "
            "A flat or rising curve means work is being added or delayed."
        )

        if "rem_dur_days" not in tasks.columns or "eff_finish" not in tasks.columns:
            st.info("Duration data not available.")
        else:
            tasks_t = tasks.copy()
            tasks_t["eff_finish_dt"] = pd.to_datetime(tasks_t["eff_finish"], errors="coerce")
            tasks_t = tasks_t.dropna(subset=["eff_finish_dt", "rem_dur_days"])
            tasks_t["rem_dur_days"] = tasks_t["rem_dur_days"].apply(lambda x: safe_float(x, 0))
            tasks_t = tasks_t.sort_values("eff_finish_dt")

            min_d = tasks_t["eff_finish_dt"].min().to_period("W").start_time
            max_d = tasks_t["eff_finish_dt"].max().to_period("W").start_time + pd.Timedelta(weeks=1)
            weeks = pd.date_range(min_d, max_d, freq="W-MON")

            total_rem = tasks_t["rem_dur_days"].sum()
            burn_rows = []
            for w in weeks:
                done_dur = tasks_t[tasks_t["eff_finish_dt"] <= w]["rem_dur_days"].sum()
                burn_rows.append({"Week": w, "Remaining Duration (d)": max(0, total_rem - done_dur)})

            burn_df = pd.DataFrame(burn_rows)
            fig_b = go.Figure()
            fig_b.add_trace(go.Scatter(
                x=burn_df["Week"], y=burn_df["Remaining Duration (d)"],
                name="Remaining Duration",
                line=dict(color="#C0392B", width=2.5),
                fill="tozeroy", fillcolor="rgba(192,57,43,0.07)",
            ))
            fig_b.add_vline(
                x=datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
                line_dash="dot", line_color="#E8951D", line_width=1.5,
                annotation_text="Today",
            )
            fig_b.update_layout(
                title="Remaining Duration Burn-Down",
                plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
                height=360, margin=dict(l=10, r=10, t=40, b=10),
                hovermode="x unified",
            )
            st.plotly_chart(fig_b)

    # -- TAB 3: Timeline Overview ---------------------------------------------
    with tab_gantt:
        st.caption("Top-level WBS timeline showing early start to late finish for each WBS area.")
        if "wbs_path" not in tasks.columns:
            st.info("WBS data not available.")
        else:
            tasks_t = tasks.copy()
            tasks_t["wbs_top"] = tasks_t["wbs_path"].apply(
                lambda x: str(x).split(" > ")[0].strip()
                if x and str(x).strip() not in ("","nan") else "Unknown"
            )
            tasks_t["eff_start_dt"]  = pd.to_datetime(tasks_t["eff_start"],  errors="coerce")
            tasks_t["eff_finish_dt"] = pd.to_datetime(tasks_t["eff_finish"], errors="coerce")
            wbs_agg = tasks_t.groupby("wbs_top").agg(
                Start=("eff_start_dt",  "min"),
                Finish=("eff_finish_dt","max"),
                Count=("task_code","count"),
            ).reset_index().dropna(subset=["Start","Finish"])
            wbs_agg = wbs_agg.sort_values("Start")
            wbs_agg["Label"] = wbs_agg["wbs_top"].str[:40] + " (" + wbs_agg["Count"].astype(str) + ")"

            fig_tl = px.timeline(wbs_agg, x_start="Start", x_end="Finish", y="Label",
                                 title="WBS Area Timeline",
                                 color_discrete_sequence=["#0B1929"])
            fig_tl.update_yaxes(autorange="reversed")
            fig_tl.add_vline(
                x=datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
                line_dash="dot", line_color="#E8951D",
                annotation_text="Today", annotation_font_color="#E8951D",
            )
            fig_tl.update_layout(
                height=max(300, 50 + len(wbs_agg) * 32),
                plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
                margin=dict(l=10, r=10, t=40, b=10),
            )
            st.plotly_chart(fig_tl)

    # Export
    export_df = tasks[["task_code","task_name","wbs_path","eff_start","eff_finish",
                        "orig_dur_days","rem_dur_days","phys_pct"] +
                      (["target_start","target_finish"] if has_target else [])].copy()
    for col in ["eff_start","eff_finish","target_start","target_finish"]:
        if col in export_df.columns:
            export_df[col] = export_df[col].apply(format_date)
    xls = export_df_to_excel({"S-Curve Data": export_df})
    st.download_button("Export S-Curve Data", xls,
                       f"scurve_{datetime.now().strftime('%Y%m%d')}.xlsx",
                       "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                       key="sc_export")


# =============================================================================
# FEATURE 2 — RESOURCE LEVELLING
# =============================================================================

def page_resource_levelling(data: dict, near_crit_days: float):
    ctrl_bar("Resource Levelling", "Labour demand, peak detection and double-booking flags.")
    mode_toggle_bar()

    tasks    = data.get("tasks_df",          pd.DataFrame())
    task_res = data.get("task_resources_df", pd.DataFrame())
    resources = data.get("resources_df",     pd.DataFrame())

    if task_res.empty:
        empty_state("", "No Resource Data",
                    "This programme does not contain resource loading. "
                    "Ask your planner to add resource assignments in P6 and re-export.", "")
        return

    pm_note(
        "Resource levelling shows where labour demand exceeds your team's capacity. "
        "Peaks mean multiple activities need the same resource at the same time. "
        "The double-booking table shows specific activities that overlap on the same resource."
    )

    # Merge resource names
    if not resources.empty and "rsrc_id" in task_res.columns:
        task_res = task_res.merge(
            resources[["rsrc_id","rsrc_name"]].drop_duplicates("rsrc_id"),
            on="rsrc_id", how="left"
        )

    # Merge task info
    if not tasks.empty and "task_id" in task_res.columns:
        _tc = ["task_id","task_code","task_name","wbs_path","is_critical","eff_start","eff_finish"]
        _tc = [c for c in _tc if c in tasks.columns]
        task_res = task_res.merge(
            tasks[_tc].drop_duplicates(subset="task_id"),
            on="task_id", how="left", suffixes=("","_t")
        )

    # Parse dates
    for col in ["target_start","target_finish"]:
        if col in task_res.columns:
            task_res[col] = pd.to_datetime(task_res[col], errors="coerce")

    task_res["target_qty"] = task_res["target_qty"].apply(lambda x: safe_float(x, 0))

    # Resource selector
    all_res = sorted(task_res["rsrc_name"].dropna().unique().tolist()) \
              if "rsrc_name" in task_res.columns else []

    if not all_res:
        st.info("No named resources found.")
        return

    tab_demand, tab_double, tab_table = st.tabs([
        "Weekly Demand", "Double-Booking Flags", "Resource Table"
    ])

    # -- TAB 1: Weekly Demand -------------------------------------------------
    with tab_demand:
        sel_res = st.multiselect("Filter resources", all_res,
                                 default=all_res[:min(8, len(all_res))],
                                 key="rl_res_sel", label_visibility="collapsed")
        threshold = st.number_input("Peak threshold (hrs/week)",
                                    min_value=0, value=200, step=10, key="rl_thresh")

        filtered = task_res[task_res["rsrc_name"].isin(sel_res)] if sel_res else task_res

        # Build weekly demand from individual task_resource rows
        demand_rows = []
        for _, row in filtered.iterrows():
            if pd.isna(row.get("target_start")) or pd.isna(row.get("target_finish")):
                continue
            start = row["target_start"]
            finish = row["target_finish"]
            qty    = row["target_qty"]
            if pd.isna(start) or pd.isna(finish) or qty == 0:
                continue
            weeks = max(1, int((finish - start).days / 7) + 1)
            weekly_qty = qty / weeks
            w = start.to_period("W").start_time
            while w <= finish:
                demand_rows.append({
                    "Week":     w,
                    "Resource": row.get("rsrc_name","?"),
                    "Hrs":      weekly_qty,
                })
                w += pd.Timedelta(weeks=1)

        if not demand_rows:
            st.info("No demand data could be computed from the resource dates available.")
        else:
            demand_df = pd.DataFrame(demand_rows)
            weekly    = demand_df.groupby("Week")["Hrs"].sum().reset_index()
            weekly.columns = ["Week", "Total Hrs"]
            weekly["Peak"] = weekly["Total Hrs"] > threshold
            weekly["Colour"] = weekly["Peak"].map({True: "#C0392B", False: "#0B1929"})

            fig_d = go.Figure()
            fig_d.add_trace(go.Bar(
                x=weekly["Week"], y=weekly["Total Hrs"],
                name="Demand",
                marker_color=weekly["Colour"].tolist(),
            ))
            if threshold > 0:
                fig_d.add_hline(y=threshold, line_dash="dash",
                                line_color="#E8951D", line_width=1.5,
                                annotation_text=f"Threshold ({threshold}h)",
                                annotation_font_color="#E8951D")
            fig_d.add_vline(
                x=datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
                line_dash="dot", line_color="#8AAABF", line_width=1,
                annotation_text="Today",
            )
            fig_d.update_layout(
                title="Weekly Labour Demand (hrs)",
                plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
                height=360, margin=dict(l=10, r=10, t=40, b=10),
                bargap=0.1,
            )
            st.plotly_chart(fig_d)

            n_peak_weeks = int(weekly["Peak"].sum())
            peak_max     = round(float(weekly["Total Hrs"].max()), 0)
            c1, c2, c3 = st.columns(3)
            c1.markdown(kpi_card("Peak Weeks",   n_peak_weeks, f"above {threshold}h threshold",
                                 "red" if n_peak_weeks else "green"), unsafe_allow_html=True)
            c2.markdown(kpi_card("Max Weekly",   f"{int(peak_max)}h", "single week peak",
                                 "red" if peak_max > threshold else "green"), unsafe_allow_html=True)
            c3.markdown(kpi_card("Resources",    len(sel_res), "shown", "navy"), unsafe_allow_html=True)

    # -- TAB 2: Double-Booking Flags ------------------------------------------
    with tab_double:
        st.caption(
            "Activities assigned to the same resource with overlapping dates. "
            "These represent potential resource conflicts that need to be resolved."
        )

        overlap_rows = []
        if "rsrc_name" in task_res.columns and "target_start" in task_res.columns:
            for res_name, grp in task_res.groupby("rsrc_name"):
                grp = grp.dropna(subset=["target_start","target_finish"]).sort_values("target_start")
                acts = grp.to_dict("records")
                for i in range(len(acts)):
                    for j in range(i+1, len(acts)):
                        a, b = acts[i], acts[j]
                        if b["target_start"] > a["target_finish"]:
                            break
                        overlap_days = int((min(a["target_finish"], b["target_finish"]) -
                                           max(a["target_start"],  b["target_start"])).days)
                        if overlap_days > 0:
                            overlap_rows.append({
                                "Resource":       res_name,
                                "Activity A":     str(a.get("task_code","?")),
                                "Activity A Name":str(a.get("task_name",""))[:40],
                                "Activity B":     str(b.get("task_code","?")),
                                "Activity B Name":str(b.get("task_name",""))[:40],
                                "Overlap (days)": overlap_days,
                                "A Start":        format_date(a.get("target_start")),
                                "A Finish":       format_date(a.get("target_finish")),
                                "B Start":        format_date(b.get("target_start")),
                                "B Finish":       format_date(b.get("target_finish")),
                            })

        if not overlap_rows:
            st.success("No resource double-bookings found.")
        else:
            ov_df = pd.DataFrame(overlap_rows).sort_values("Overlap (days)", ascending=False)
            st.markdown(
                f'<div class="attn attn-amber"><strong>{len(ov_df)} resource conflicts found.</strong> '
                f'Review with your planner to resolve clashes before mobilisation.</div>',
                unsafe_allow_html=True,
            )
            st.dataframe(ov_df, hide_index=True)

            xls = export_df_to_excel({"Double-Booking Flags": ov_df})
            st.download_button("Export Double-Booking Report", xls,
                               f"double_booking_{datetime.now().strftime('%Y%m%d')}.xlsx",
                               "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                               key="rl_db_export")

    # -- TAB 3: Resource Table ------------------------------------------------
    with tab_table:
        disp = task_res[["task_code","task_name","rsrc_name","target_qty",
                         "target_start","target_finish"]].copy() \
               if all(c in task_res.columns for c in
                      ["task_code","rsrc_name","target_qty"]) else task_res.copy()
        disp.columns = [c.replace("_"," ").title() for c in disp.columns]
        for col in disp.columns:
            if "start" in col.lower() or "finish" in col.lower():
                disp[col] = disp[col].apply(format_date)
        st.dataframe(disp, hide_index=True, height=440)
        xls2 = export_df_to_excel({"Resource Loading": disp})
        st.download_button("Export Resource Table", xls2,
                           f"resource_loading_{datetime.now().strftime('%Y%m%d')}.xlsx",
                           "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           key="rl_tbl_export")


# =============================================================================
# FEATURE 3 — CONTRACT MILESTONE TRACKER
# =============================================================================

_CM_STATE_KEY = "contract_milestones"


def _cm_rag(days_diff) -> tuple:
    try:
        d = float(days_diff)
        if d <= 0:  return "On Time",  "green", "#1E7A4E"
        if d <= 14: return "At Risk",  "amber", "#E8951D"
        return "Late",     "red",   "#C0392B"
    except Exception:
        return "Unknown", "grey", "#6B7280"


def page_contract_milestones(data: dict, near_crit_days: float):
    ctrl_bar("Contract Milestones",
             "Track contractual key dates against current programme forecast.")
    mode_toggle_bar()

    tasks = data.get("tasks_df", pd.DataFrame())

    pm_note(
        "Enter your contractual milestone dates here. PlanTrace will compare them "
        "against the current programme forecast and show a RAG status for each. "
        "Red = behind contract date. Amber = within 14 days. Green = on time or early."
    )

    # -- Auto-detect milestones from XER -------------------------------------
    xer_milestones = pd.DataFrame()
    if not tasks.empty and "task_type" in tasks.columns:
        xer_milestones = tasks[
            tasks["task_type"].astype(str).str.contains("Milestone|milestone", na=False)
        ][["task_code","task_name","eff_finish"]].copy()
        xer_milestones.columns = ["Activity ID","Milestone Name","Forecast Date"]
        xer_milestones["Forecast Date"] = xer_milestones["Forecast Date"].apply(format_date)

    # -- Load/save state ------------------------------------------------------
    if _CM_STATE_KEY not in st.session_state:
        st.session_state[_CM_STATE_KEY] = []

    col_auto, col_manual = st.columns([1,1], gap="large")

    # -- Left: auto-detected milestones ---------------------------------------
    with col_auto:
        st.markdown('<div class="section-label">XER Milestones — Auto-Detected</div>',
                    unsafe_allow_html=True)
        if xer_milestones.empty:
            st.info("No milestones found in this programme. Check task types in P6.")
        else:
            st.caption(f"{len(xer_milestones)} milestones found in programme.")
            st.dataframe(xer_milestones, hide_index=True)
            if st.button("Import all as Contract Milestones", key="cm_import"):
                existing_names = {m["name"] for m in st.session_state[_CM_STATE_KEY]}
                for _, row in xer_milestones.iterrows():
                    if row["Milestone Name"] not in existing_names:
                        st.session_state[_CM_STATE_KEY].append({
                            "name":          row["Milestone Name"],
                            "activity_id":   row["Activity ID"],
                            "contract_date": "",
                            "forecast_date": row["Forecast Date"],
                        })
                st.rerun()

    # -- Right: manual entry --------------------------------------------------
    with col_manual:
        st.markdown('<div class="section-label">Add Contract Milestone</div>',
                    unsafe_allow_html=True)
        with st.form("cm_add_form", clear_on_submit=True):
            name  = st.text_input("Milestone Name",        placeholder="e.g. Mechanical Completion")
            act_id= st.text_input("Activity ID (optional)", placeholder="e.g. ACT-1050")
            cdate = st.date_input("Contract Date")
            fdate_str = ""
            # Try to auto-match forecast date
            if act_id and not tasks.empty and "task_code" in tasks.columns:
                match = tasks[tasks["task_code"] == act_id]
                if not match.empty:
                    fdate_str = format_date(match.iloc[0].get("eff_finish"))
            forecast_override = st.text_input("Forecast Date (override)",
                                              value=fdate_str,
                                              placeholder="Leave blank to use programme data")
            submitted = st.form_submit_button("Add Milestone")
            if submitted and name:
                st.session_state[_CM_STATE_KEY].append({
                    "name":          name,
                    "activity_id":   act_id,
                    "contract_date": str(cdate),
                    "forecast_date": forecast_override or fdate_str,
                })
                st.rerun()

    # -- RAG Dashboard --------------------------------------------------------
    milestones = st.session_state[_CM_STATE_KEY]
    if not milestones:
        st.info("No contract milestones added yet. Use the panel above to add them.")
        return

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="section-label">Contract Milestone RAG Status</div>',
                unsafe_allow_html=True)

    # Refresh forecast dates from live programme
    for m in milestones:
        if m.get("activity_id") and not tasks.empty and "task_code" in tasks.columns:
            match = tasks[tasks["task_code"] == m["activity_id"]]
            if not match.empty:
                m["forecast_date"] = format_date(match.iloc[0].get("eff_finish"))

    # Build RAG table
    rag_rows = []
    for m in milestones:
        cdate_str = m.get("contract_date","")
        fdate_str = m.get("forecast_date","")
        days_late  = None
        try:
            if cdate_str and fdate_str and cdate_str not in ("","—","None"):
                days_late = (pd.Timestamp(fdate_str) - pd.Timestamp(cdate_str)).days
        except Exception:
            pass

        rag_label, rag_chip, rag_col = _cm_rag(days_late)
        rag_rows.append({
            "Milestone":      m["name"],
            "Activity ID":    m.get("activity_id","—"),
            "Contract Date":  cdate_str or "—",
            "Forecast Date":  fdate_str or "—",
            "Days Early/Late": (f"+{days_late}d late" if days_late and days_late > 0
                               else f"{abs(days_late)}d early" if days_late and days_late < 0
                               else "—"),
            "RAG": rag_label,
            "_col": rag_col,
        })

    # Display as visual cards
    for row in rag_rows:
        bg  = {"On Time":"#E4F5EC","At Risk":"#FDF3E0","Late":"#FAEAEA"}.get(row["RAG"],"#F5F7F9")
        brd = row["_col"]
        st.markdown(
            f'<div style="background:{bg};border:1px solid {brd};border-left:4px solid {brd};'
            f'border-radius:0 6px 6px 0;padding:12px 16px;margin-bottom:8px;'
            f'display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:10px;">'
            f'<div>'
            f'<div style="font-weight:700;color:#0B1929;font-size:14px;">{row["Milestone"]}</div>'
            f'<div style="font-size:11px;color:#6B7C8E;margin-top:2px;">'
            f'Activity: {row["Activity ID"]}</div>'
            f'</div>'
            f'<div style="display:flex;gap:20px;align-items:center;flex-wrap:wrap;">'
            f'<div style="text-align:center;">'
            f'<div style="font-size:9px;color:#8A9DB0;text-transform:uppercase;'
            f'letter-spacing:0.8px;font-family:\'IBM Plex Mono\',monospace;">Contract</div>'
            f'<div style="font-size:13px;font-weight:600;color:#0B1929;">{row["Contract Date"]}</div>'
            f'</div>'
            f'<div style="text-align:center;">'
            f'<div style="font-size:9px;color:#8A9DB0;text-transform:uppercase;'
            f'letter-spacing:0.8px;font-family:\'IBM Plex Mono\',monospace;">Forecast</div>'
            f'<div style="font-size:13px;font-weight:600;color:#0B1929;">{row["Forecast Date"]}</div>'
            f'</div>'
            f'<div style="text-align:center;">'
            f'<div style="font-size:9px;color:#8A9DB0;text-transform:uppercase;'
            f'letter-spacing:0.8px;font-family:\'IBM Plex Mono\',monospace;">Movement</div>'
            f'<div style="font-size:13px;font-weight:700;color:{brd};">{row["Days Early/Late"]}</div>'
            f'</div>'
            f'{chip(row["RAG"], ("green" if row["RAG"]=="On Time" else "amber" if row["RAG"]=="At Risk" else "red"))}'
            f'</div></div>',
            unsafe_allow_html=True,
        )

    # Clear button
    if st.button("Clear All Milestones", key="cm_clear"):
        st.session_state[_CM_STATE_KEY] = []
        st.rerun()

    # Export
    export_df = pd.DataFrame([{k: v for k, v in r.items() if not k.startswith("_")}
                               for r in rag_rows])
    xls = export_df_to_excel({"Contract Milestones": export_df})
    st.download_button("Export Milestone Tracker", xls,
                       f"contract_milestones_{datetime.now().strftime('%Y%m%d')}.xlsx",
                       "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                       key="cm_export")


# =============================================================================
# FEATURE 4 — MULTI-PROJECT DASHBOARD
# =============================================================================

_MP_KEY = "mp_programmes"


def _mp_score(prog: dict, near_crit_days: float) -> dict:
    """Compute key metrics for one programme."""
    tasks = prog.get("tasks_df", pd.DataFrame())
    rels  = prog.get("relationships_df", pd.DataFrame())
    info  = prog.get("project_info", {})
    if tasks.empty:
        return {}
    tasks_c = get_critical_threshold(tasks, near_crit_days)
    n_total = len(tasks_c)
    n_crit  = int(tasks_c["is_critical"].sum()) if "is_critical" in tasks_c.columns else 0
    n_neg   = int(tasks_c["total_float_days"].apply(lambda f: safe_float(f,0)<0).sum()) if "total_float_days" in tasks_c.columns else 0
    n_nc    = int(tasks_c["is_near_critical"].sum()) if "is_near_critical" in tasks_c.columns else 0
    n_rels  = len(rels)
    valid_ends = tasks_c["eff_finish"].dropna() if "eff_finish" in tasks_c.columns else pd.Series()
    proj_end   = format_date(valid_ends.max()) if not valid_ends.empty else "—"
    # Health score
    health = 100
    if n_total > 0:
        health -= min(40, int(n_neg / n_total * 400))
        health -= min(20, int(n_nc  / n_total * 200))
    health = max(0, health)
    rag = "red" if health < 50 else "amber" if health < 80 else "green"
    return dict(
        name=info.get("name","") or prog.get("_filename",""),
        filename=prog.get("_filename",""),
        n_total=n_total, n_rels=n_rels,
        n_crit=n_crit, n_neg=n_neg, n_nc=n_nc,
        proj_end=proj_end, health=health, rag=rag,
        data_date=format_date(info.get("data_date")) or "—",
    )


def page_multi_project(near_crit_days: float):
    ctrl_bar("Multi-Project Dashboard",
             "Side-by-side KPI comparison across multiple uploaded programmes.")
    mode_toggle_bar()

    if _MP_KEY not in st.session_state:
        st.session_state[_MP_KEY] = {}

    st.markdown('<div class="section-label">Upload Additional Programmes</div>',
                unsafe_allow_html=True)
    uploaded = st.file_uploader(
        "Upload XER files (multiple allowed)",
        type=["xer"], accept_multiple_files=True,
        key="mp_uploader", label_visibility="collapsed",
    )
    if uploaded:
        for f in uploaded:
            ck = f"{f.name}_{f.size}"
            if ck not in st.session_state[_MP_KEY]:
                with st.spinner(f"Parsing {f.name}..."):
                    try:
                        parsed = parse_xer(f.read())
                        parsed["_filename"] = f.name
                        st.session_state[_MP_KEY][ck] = parsed
                    except Exception as e:
                        st.error(f"Failed to parse {f.name}: {e}")

    # Also include the currently loaded single programme
    all_progs = {}
    if "programme" in st.session_state:
        prog = st.session_state["programme"].copy()
        prog["_filename"] = st.session_state.get("_xer_filename","Current Programme")
        all_progs["__current__"] = prog
    all_progs.update(st.session_state[_MP_KEY])

    if len(all_progs) < 1:
        empty_state("", "No Programmes Loaded",
                    "Upload at least one XER file to use the multi-project dashboard.", "")
        return

    if len(all_progs) == 1:
        st.info("Only one programme loaded. Upload more XER files above to compare side-by-side.")

    # Compute metrics for all
    metrics = [_mp_score(prog, near_crit_days) for prog in all_progs.values() if prog]
    metrics = [m for m in metrics if m]

    if not metrics:
        st.warning("Could not compute metrics for any loaded programmes.")
        return

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="section-label">Programme Comparison</div>',
                unsafe_allow_html=True)

    # Header row
    cols = st.columns(len(metrics), gap="small")
    rag_colours = {"red":"#C0392B","amber":"#E8951D","green":"#1E7A4E"}

    for col, m in zip(cols, metrics):
        rag_col = rag_colours.get(m["rag"],"#6B7280")
        col.markdown(
            f'<div style="background:#0B1929;border-radius:6px;padding:16px;'
            f'border-top:3px solid {rag_col};margin-bottom:12px;">'
            f'<div style="font-size:10px;color:#4A6070;text-transform:uppercase;'
            f'letter-spacing:1px;font-family:\'IBM Plex Mono\',monospace;margin-bottom:4px;">'
            f'Programme</div>'
            f'<div style="font-size:13px;font-weight:700;color:#FFFFFF;'
            f'white-space:nowrap;overflow:hidden;text-overflow:ellipsis;"'
            f' title="{m["name"]}">{m["name"] or m["filename"]}</div>'
            f'<div style="font-size:10px;color:#3A5265;margin-top:3px;">{m["filename"]}</div>'
            f'<div style="margin-top:10px;display:flex;justify-content:space-between;">'
            f'<span style="font-size:9px;color:#3A5265;font-family:\'IBM Plex Mono\',monospace;">'
            f'Data date</span>'
            f'<span style="font-size:10px;color:#8AAABF;font-family:\'IBM Plex Mono\',monospace;">'
            f'{m["data_date"]}</span></div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    kpi_defs = [
        ("Total Activities", "n_total",  "navy"),
        ("Critical",         "n_crit",   "amber"),
        ("Negative Float",   "n_neg",    "red"),
        ("Near-Critical",    "n_nc",     "amber"),
        ("Relationships",    "n_rels",   "navy"),
        ("Forecast End",     "proj_end", "navy"),
        ("Health Score",     "health",   "green"),
    ]

    for label, key, default_style in kpi_defs:
        cols = st.columns(len(metrics), gap="small")
        for col, m in zip(cols, metrics):
            val = m.get(key, "—")
            style = default_style
            if key == "n_neg":
                style = "red" if int(val or 0) > 0 else "green"
            elif key == "health":
                style = "green" if int(val) >= 80 else "amber" if int(val) >= 50 else "red"
                val   = f"{val}%"
            col.markdown(kpi_card(label, val, "", style), unsafe_allow_html=True)
            col.markdown("<div style='height:4px'></div>", unsafe_allow_html=True)

    # Comparison bar chart
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="section-label">Visual Comparison</div>', unsafe_allow_html=True)
    chart_df = pd.DataFrame([{
        "Programme": m["name"] or m["filename"],
        "Critical":       m["n_crit"],
        "Negative Float": m["n_neg"],
        "Near-Critical":  m["n_nc"],
    } for m in metrics])

    fig_mp = px.bar(chart_df.melt(id_vars="Programme"),
                    x="Programme", y="value", color="variable", barmode="group",
                    color_discrete_map={
                        "Critical": "#E8951D",
                        "Negative Float": "#C0392B",
                        "Near-Critical": "#8AAABF",
                    },
                    labels={"value":"Count","variable":"Metric"},
                    title="Risk Metrics Comparison")
    fig_mp.update_layout(height=340, plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
                         margin=dict(l=10, r=10, t=40, b=10),
                         legend=dict(orientation="h", y=1.08))
    st.plotly_chart(fig_mp)

    # Clear multi-project programmes
    if st.button("Clear uploaded programmes", key="mp_clear"):
        st.session_state[_MP_KEY] = {}
        st.rerun()

    # Export
    summary = pd.DataFrame([{
        "Programme":       m["name"] or m["filename"],
        "Data Date":       m["data_date"],
        "Activities":      m["n_total"],
        "Relationships":   m["n_rels"],
        "Critical":        m["n_crit"],
        "Negative Float":  m["n_neg"],
        "Near-Critical":   m["n_nc"],
        "Forecast End":    m["proj_end"],
        "Health Score":    f"{m['health']}%",
        "RAG":             m["rag"].title(),
    } for m in metrics])
    xls = export_df_to_excel({"Programme Comparison": summary})
    st.download_button("Export Comparison Report", xls,
                       f"multi_project_{datetime.now().strftime('%Y%m%d')}.xlsx",
                       "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                       key="mp_export")


# =============================================================================
# FEATURE 5 — BRANDING SETTINGS (replaces old settings page)
# =============================================================================

_BRAND_KEY = "plantrace_branding"

_BRAND_DEFAULTS = {
    "app_name":       "PlanTrace",
    "app_subtitle":   "Planning intelligence for project delivery teams",
    "primary_colour": "#0B1929",
    "accent_colour":  "#E8951D",
    "company_name":   "",
    "show_logo":      True,
}

_COLOUR_PRESETS = {
    "PlanTrace (Navy + Amber)":   ("#0B1929", "#E8951D"),
    "Slate + Teal":               ("#1E2D3D", "#1AADAA"),
    "Charcoal + Blue":            ("#1C2A3A", "#3B82F6"),
    "Dark Green + Amber":         ("#14302A", "#E8951D"),
    "Midnight + Orange":          ("#0D1117", "#F97316"),
    "Corporate Grey + Red":       ("#2C2C2C", "#C0392B"),
}

def get_brand() -> dict:
    return st.session_state.get(_BRAND_KEY, _BRAND_DEFAULTS.copy())


def page_settings(data: dict, near_crit_days: float):
    ctrl_bar("Settings", "Branding, display preferences and application settings.")

    brand = get_brand()

    tab_brand, tab_display, tab_about = st.tabs(["Branding", "Display", "About"])

    # -- BRANDING TAB ---------------------------------------------------------
    with tab_brand:
        st.markdown('<div class="section-label">Application Identity</div>',
                    unsafe_allow_html=True)

        col_form, col_preview = st.columns([2, 1], gap="large")

        with col_form:
            new_brand = brand.copy()

            new_brand["app_name"] = st.text_input(
                "Application Name", value=brand["app_name"],
                help="Shown in sidebar header and page titles",
                key="sett_name"
            )
            new_brand["app_subtitle"] = st.text_input(
                "Subtitle / Tagline", value=brand["app_subtitle"],
                key="sett_subtitle"
            )
            new_brand["company_name"] = st.text_input(
                "Company Name (optional)", value=brand.get("company_name",""),
                placeholder="e.g. Exentec Hargreaves Ltd",
                key="sett_company"
            )

            st.markdown('<div class="section-label" style="margin-top:16px;">Colour Scheme</div>',
                        unsafe_allow_html=True)

            preset = st.selectbox("Quick presets", list(_COLOUR_PRESETS.keys()),
                                  key="sett_preset")
            if st.button("Apply Preset", key="sett_apply_preset"):
                new_brand["primary_colour"], new_brand["accent_colour"] = _COLOUR_PRESETS[preset]

            c1, c2 = st.columns(2)
            with c1:
                new_brand["primary_colour"] = st.color_picker(
                    "Primary (sidebar/headers)",
                    value=brand.get("primary_colour", "#0B1929"),
                    key="sett_primary"
                )
            with c2:
                new_brand["accent_colour"] = st.color_picker(
                    "Accent (highlights/borders)",
                    value=brand.get("accent_colour", "#E8951D"),
                    key="sett_accent"
                )

            if st.button("Save Branding", key="sett_save", type="primary"):
                st.session_state[_BRAND_KEY] = new_brand
                st.success("Branding saved. Reload the page to see colour changes.")
                st.rerun()

            if st.button("Reset to Defaults", key="sett_reset"):
                st.session_state[_BRAND_KEY] = _BRAND_DEFAULTS.copy()
                st.success("Reset to defaults.")
                st.rerun()

        with col_preview:
            p_col = brand.get("primary_colour","#0B1929")
            a_col = brand.get("accent_colour","#E8951D")
            name  = brand.get("app_name","PlanTrace")
            sub   = brand.get("app_subtitle","Planning intelligence")
            comp  = brand.get("company_name","")
            st.markdown('<div class="section-label">Preview</div>', unsafe_allow_html=True)
            st.markdown(
                f'<div style="background:{p_col};border-radius:6px;padding:18px;'
                f'border-bottom:3px solid {a_col};">'
                f'<div style="font-size:18px;font-weight:700;color:#FFFFFF;">{name}</div>'
                f'<div style="font-size:11px;color:#4A6070;margin-top:3px;">{sub}</div>'
                f'{f"<div style=\'font-size:10px;color:#3A5265;margin-top:6px;\'>{comp}</div>" if comp else ""}'
                f'<div style="margin-top:14px;padding:10px 0;">'
                f'<div style="background:#0d2035;border-left:2px solid {a_col};'
                f'padding:6px 10px;border-radius:0 3px 3px 0;">'
                f'<span style="font-size:12px;color:#8AAABF;">Overview</span></div>'
                f'<div style="padding:6px 10px;">'
                f'<span style="font-size:12px;color:#3A5265;">Programme</span></div>'
                f'</div></div>',
                unsafe_allow_html=True,
            )

    # -- DISPLAY TAB ----------------------------------------------------------
    with tab_display:
        st.markdown('<div class="section-label">Display Preferences</div>',
                    unsafe_allow_html=True)
        st.info("Near-Critical Float threshold is set in the sidebar slider.")
        st.caption(f"Current threshold: {near_crit_days} days")

        st.markdown('<div class="section-label" style="margin-top:16px;">View Mode</div>',
                    unsafe_allow_html=True)
        current_mode = st.session_state.get(_MODE_KEY, "PM Mode")
        st.markdown(
            f'Current view mode: {chip(current_mode, "navy" if current_mode=="Planner Mode" else "amber")}',
            unsafe_allow_html=True,
        )
        st.caption("Change the view mode using the selector in the sidebar.")

    # -- ABOUT TAB ------------------------------------------------------------
    with tab_about:
        brand_now = get_brand()
        st.markdown(
            f'<div class="card">'
            f'<div style="font-size:22px;font-weight:700;color:#0B1929;margin-bottom:4px;">'
            f'{brand_now["app_name"]}</div>'
            f'<div style="font-size:13px;color:#6B7C8E;margin-bottom:16px;">'
            f'{brand_now["app_subtitle"]}</div>'
            f'<hr>'
            f'<div style="font-size:13px;color:#374151;line-height:1.8;">'
            f'A Primavera P6 XER programme analysis tool for Project Managers, '
            f'Project Controls Managers and Planners.<br><br>'
            f'No P6 licence required. Upload any .xer file.<br><br>'
            f'<strong>Modules:</strong> Overview &nbsp;·&nbsp; Executive Summary '
            f'&nbsp;·&nbsp; Programme &nbsp;·&nbsp; Logic &nbsp;·&nbsp; Critical Path '
            f'&nbsp;·&nbsp; Labour &nbsp;·&nbsp; Health Check &nbsp;·&nbsp; Comparison '
            f'&nbsp;·&nbsp; Forensic Planner &nbsp;·&nbsp; S-Curve &nbsp;·&nbsp; '
            f'Resource Levelling &nbsp;·&nbsp; Contract Milestones &nbsp;·&nbsp; '
            f'Multi-Project &nbsp;·&nbsp; PM Actions &nbsp;·&nbsp; Risk Register'
            f'&nbsp;·&nbsp; Reports'
            f'</div></div>',
            unsafe_allow_html=True,
        )




# =============================================================================
# DCMA 14-POINT SCHEDULE ASSESSMENT
# Industry standard. Used by MOD, US DoD, NEC, JCT auditors.
# Threshold values from DCMA 2005 guidelines, revised 2024.
# =============================================================================

_DCMA_THRESHOLDS = {
    "logic":         5.0,   # % activities with missing predecessor or successor
    "leads":         0.0,   # % relationships with negative lag (leads) — should be zero
    "lags":          5.0,   # % relationships with positive lag
    "relationship":  90.0,  # % relationships that are Finish-to-Start (FS)
    "hard_constraint":5.0,  # % activities with hard constraints
    "high_float":    5.0,   # % activities with float > 44 days
    "negative_float":0.0,   # % activities with negative float — should be zero
    "critical_pct":  10.0,  # % critical activities — >10% is acceptable, >50% is suspect
    "long_duration": 5.0,   # % activities with duration > 44 days
    "invalid_dates": 0.0,   # % activities with invalid/out-of-sequence dates
    "resources":     90.0,  # % activities with resource loading (lower = fail)
    "missed_tasks":  5.0,   # % activities that were planned to finish but haven't (baseline needed)
}

_DCMA_LABELS = {
    "logic":         "Missing Logic",
    "leads":         "Leads (Negative Lag)",
    "lags":          "Lags",
    "relationship":  "Relationship Types (% FS)",
    "hard_constraint":"Hard Constraints",
    "high_float":    "High Float (>44d)",
    "negative_float":"Negative Float",
    "critical_pct":  "Critical Activities",
    "long_duration": "Long Durations (>44d)",
    "invalid_dates": "Invalid Dates",
    "resources":     "Resource Loading",
    "missed_tasks":  "Missed Tasks",
}

_DCMA_DESCRIPTIONS = {
    "logic":          "Activities missing a predecessor, successor or both. Max 5%.",
    "leads":          "Relationships with negative lag (leads). Should be 0%.",
    "lags":           "Relationships with positive lag. Max 5%.",
    "relationship":   "Proportion of Finish-to-Start links. Should be >90%.",
    "hard_constraint":"Activities with hard constraints (Must Start/Finish On). Max 5%.",
    "high_float":     "Activities with total float >44 days. Max 5%.",
    "negative_float": "Activities with negative total float. Should be 0%.",
    "critical_pct":   "Proportion of the schedule on the critical path. 10–50% is typical.",
    "long_duration":  "Activities with original duration >44 days. Max 5%.",
    "invalid_dates":  "Activities with actual start after actual finish, or other date anomalies.",
    "resources":      "Proportion of activities with resource assignments. Target >90%.",
    "missed_tasks":   "Activities planned to finish by data date that have not. Max 5%. (Requires baseline.)",
}


def compute_dcma(tasks: pd.DataFrame, rels: pd.DataFrame) -> dict:
    """
    Compute all DCMA 14-Point Assessment metrics.
    Returns dict with keys matching _DCMA_THRESHOLDS.
    Each value: dict with pct, count, total, pass, direction (higher_is_better or lower_is_better)
    """
    results = {}
    n_tasks = len(tasks)
    if n_tasks == 0:
        return results

    # Active (incomplete) tasks only — DCMA only checks incomplete activities
    if "status" in tasks.columns:
        active = tasks[~tasks["status"].apply(
            lambda s: str(s) in ("TK_Complete","Complete","Completed")
        )]
    else:
        active = tasks
    n_active = max(len(active), 1)

    # -- 1. Missing Logic -----------------------------------------------------
    if not rels.empty and "task_id" in active.columns:
        with_pred = set(rels["succ_task_id"].dropna()) if "succ_task_id" in rels.columns else set()
        with_succ = set(rels["pred_task_id"].dropna()) if "pred_task_id" in rels.columns else set()
        # Exclude milestones and LOE
        check_tasks = active
        if "task_type" in active.columns:
            check_tasks = active[~active["task_type"].astype(str).str.contains(
                "Milestone|LOE|WBS|Level", na=False, case=False)]
        n_check = max(len(check_tasks), 1)
        n_missing = int(
            (~check_tasks["task_id"].isin(with_pred) | ~check_tasks["task_id"].isin(with_succ)).sum()
        )
        pct = round(n_missing / n_check * 100, 1)
        results["logic"] = dict(pct=pct, count=n_missing, total=n_check,
                                pass_=pct <= _DCMA_THRESHOLDS["logic"],
                                direction="lower", threshold=_DCMA_THRESHOLDS["logic"])

    # -- 2. Leads (negative lag) ----------------------------------------------
    if not rels.empty and "lag_days" in rels.columns:
        n_leads = int(rels["lag_days"].apply(lambda l: safe_float(l, 0) < 0).sum())
        n_rels  = max(len(rels), 1)
        pct     = round(n_leads / n_rels * 100, 1)
        results["leads"] = dict(pct=pct, count=n_leads, total=n_rels,
                                pass_=pct <= _DCMA_THRESHOLDS["leads"],
                                direction="lower", threshold=_DCMA_THRESHOLDS["leads"])

    # -- 3. Lags --------------------------------------------------------------
    if not rels.empty and "lag_days" in rels.columns:
        n_lags = int(rels["lag_days"].apply(lambda l: safe_float(l, 0) > 0).sum())
        n_rels = max(len(rels), 1)
        pct    = round(n_lags / n_rels * 100, 1)
        results["lags"] = dict(pct=pct, count=n_lags, total=n_rels,
                               pass_=pct <= _DCMA_THRESHOLDS["lags"],
                               direction="lower", threshold=_DCMA_THRESHOLDS["lags"])

    # -- 4. Relationship types (% FS) -----------------------------------------
    if not rels.empty and "rel_type" in rels.columns:
        n_fs  = int(rels["rel_type"].apply(
            lambda r: str(r).upper() in ("FS","PR_FS","FINISH_START","Finish to Start")
        ).sum())
        n_rels = max(len(rels), 1)
        pct_fs = round(n_fs / n_rels * 100, 1)
        results["relationship"] = dict(pct=pct_fs, count=n_fs, total=n_rels,
                                       pass_=pct_fs >= _DCMA_THRESHOLDS["relationship"],
                                       direction="higher", threshold=_DCMA_THRESHOLDS["relationship"])

    # -- 5. Hard constraints --------------------------------------------------
    if "cstr_type" in active.columns:
        HARD = {"CS_MANDFIN","CS_MANDSTART","Must Finish On","Must Start On",
                "Mandatory Start","Mandatory Finish"}
        n_hard = int(active["cstr_type"].apply(
            lambda c: str(c).strip() in HARD
        ).sum())
        pct = round(n_hard / n_active * 100, 1)
        results["hard_constraint"] = dict(pct=pct, count=n_hard, total=n_active,
                                          pass_=pct <= _DCMA_THRESHOLDS["hard_constraint"],
                                          direction="lower", threshold=_DCMA_THRESHOLDS["hard_constraint"])

    # -- 6. High float (>44 days) ---------------------------------------------
    if "total_float_days" in active.columns:
        n_hi = int(active["total_float_days"].apply(lambda f: safe_float(f, 0) > 44).sum())
        pct  = round(n_hi / n_active * 100, 1)
        results["high_float"] = dict(pct=pct, count=n_hi, total=n_active,
                                     pass_=pct <= _DCMA_THRESHOLDS["high_float"],
                                     direction="lower", threshold=_DCMA_THRESHOLDS["high_float"])

    # -- 7. Negative float ---------------------------------------------------
    if "total_float_days" in active.columns:
        n_neg = int(active["total_float_days"].apply(lambda f: safe_float(f, 999) < 0).sum())
        pct   = round(n_neg / n_active * 100, 1)
        results["negative_float"] = dict(pct=pct, count=n_neg, total=n_active,
                                         pass_=n_neg == 0,
                                         direction="lower", threshold=_DCMA_THRESHOLDS["negative_float"])

    # -- 8. Critical activities -----------------------------------------------
    if "is_critical" in active.columns:
        n_crit = int(active["is_critical"].sum())
        pct    = round(n_crit / n_active * 100, 1)
        # DCMA: >10% is acceptable; flag if >50% (schedule dominated by critical work)
        passes = 5.0 <= pct <= 50.0
        results["critical_pct"] = dict(pct=pct, count=n_crit, total=n_active,
                                       pass_=passes, direction="range",
                                       threshold=_DCMA_THRESHOLDS["critical_pct"])

    # -- 9. Long durations (>44 days) -----------------------------------------
    if "orig_dur_days" in active.columns:
        n_long = int(active["orig_dur_days"].apply(lambda d: safe_float(d, 0) > 44).sum())
        pct    = round(n_long / n_active * 100, 1)
        results["long_duration"] = dict(pct=pct, count=n_long, total=n_active,
                                        pass_=pct <= _DCMA_THRESHOLDS["long_duration"],
                                        direction="lower", threshold=_DCMA_THRESHOLDS["long_duration"])

    # -- 10. Invalid dates ----------------------------------------------------
    n_invalid = 0
    if "act_start" in active.columns and "act_finish" in active.columns:
        for _, r in active.iterrows():
            try:
                s = r.get("act_start")
                f = r.get("act_finish")
                if s and f and pd.Timestamp(s) > pd.Timestamp(f):
                    n_invalid += 1
            except Exception:
                pass
    pct = round(n_invalid / n_active * 100, 1)
    results["invalid_dates"] = dict(pct=pct, count=n_invalid, total=n_active,
                                    pass_=n_invalid == 0,
                                    direction="lower", threshold=_DCMA_THRESHOLDS["invalid_dates"])

    # -- 11. Resources --------------------------------------------------------
    task_res = st.session_state.get("programme", {}).get("task_resources_df", pd.DataFrame())
    if not task_res.empty and "task_id" in task_res.columns and "task_id" in active.columns:
        loaded_ids = set(task_res["task_id"].dropna())
        n_loaded   = int(active["task_id"].isin(loaded_ids).sum())
        pct        = round(n_loaded / n_active * 100, 1)
        results["resources"] = dict(pct=pct, count=n_loaded, total=n_active,
                                    pass_=pct >= _DCMA_THRESHOLDS["resources"],
                                    direction="higher", threshold=_DCMA_THRESHOLDS["resources"])
    else:
        results["resources"] = dict(pct=0.0, count=0, total=n_active,
                                    pass_=False, direction="higher",
                                    threshold=_DCMA_THRESHOLDS["resources"])

    # -- 12. Missed tasks (needs baseline) -----------------------------------
    if "target_finish" in active.columns and "eff_finish" in active.columns:
        ddate = st.session_state.get("programme", {}).get("project_info", {}).get("data_date")
        if ddate:
            dd = pd.Timestamp(ddate)
            planned_due = active[active["target_finish"].apply(
                lambda d: d is not None and pd.Timestamp(d) <= dd
                if d is not None else False
            )]
            n_due = max(len(planned_due), 1)
            n_missed = int(planned_due["eff_finish"].apply(
                lambda d: d is None or (d is not None and pd.Timestamp(d) > dd)
            ).sum())
            pct = round(n_missed / n_due * 100, 1)
            results["missed_tasks"] = dict(pct=pct, count=n_missed, total=n_due,
                                           pass_=pct <= _DCMA_THRESHOLDS["missed_tasks"],
                                           direction="lower", threshold=_DCMA_THRESHOLDS["missed_tasks"])

    return results


def dcma_quality_score(results: dict) -> int:
    """
    Compute a single 0-100 Schedule Quality Score from DCMA results.
    Each passed check = weighted points. Mirrors Acumen Fuse Score concept.
    """
    weights = {
        "logic": 15, "leads": 5, "lags": 8, "relationship": 10,
        "hard_constraint": 8, "high_float": 5, "negative_float": 15,
        "critical_pct": 5, "long_duration": 7, "invalid_dates": 8,
        "resources": 7, "missed_tasks": 7,
    }
    total_w = sum(weights.values())
    earned  = 0
    for key, w in weights.items():
        r = results.get(key)
        if r and r.get("pass_"):
            earned += w
        elif r and key in ("logic","lags","long_duration","hard_constraint",
                           "high_float","negative_float","leads","invalid_dates","missed_tasks"):
            # Partial credit: proportional to how close to threshold
            thresh = r.get("threshold", 5.0)
            pct    = r.get("pct", 100.0)
            if thresh > 0 and r.get("direction") == "lower":
                ratio = max(0, 1 - pct / (thresh * 3))
                earned += w * ratio
    return min(100, max(0, int(earned / total_w * 100)))


def page_dcma_assessment(data: dict, near_crit_days: float):
    """Full DCMA 14-Point Schedule Assessment page."""
    ctrl_bar(
        "DCMA 14-Point Assessment",
        "Industry-standard schedule quality assessment. "
        "Used by MOD, US DoD, NEC, JCT and major contractors worldwide.",
    )
    mode_toggle_bar()

    tasks = data.get("tasks_df", pd.DataFrame())
    rels  = data.get("relationships_df", pd.DataFrame())

    if tasks.empty:
        empty_state("", "No Programme Data",
                    "Upload a programme to run the DCMA assessment.", "")
        return

    pm_note(
        "The DCMA 14-Point Assessment is the industry-standard method for checking "
        "schedule quality. A score of 80+ is generally accepted as a healthy schedule. "
        "Each check shows Pass (green) or Fail (red) against its DCMA threshold."
    )

    with st.spinner("Running DCMA assessment..."):
        results = compute_dcma(tasks, rels)

    if not results:
        st.warning("Could not compute DCMA metrics from this programme.")
        return

    score = dcma_quality_score(results)
    n_pass = sum(1 for r in results.values() if r.get("pass_"))
    n_fail = len(results) - n_pass

    # -- Score Header ---------------------------------------------------------
    score_col = "#1E7A4E" if score >= 80 else "#E8951D" if score >= 60 else "#C0392B"
    score_label = "Healthy" if score >= 80 else "Review Required" if score >= 60 else "High Risk"

    st.markdown(
        f'<div style="background:#0B1929;border-radius:8px;padding:24px 32px;'
        f'margin-bottom:24px;border-left:4px solid {score_col};">'
        f'<div style="display:flex;justify-content:space-between;align-items:center;'
        f'flex-wrap:wrap;gap:20px;">'
        f'<div>'
        f'<div style="font-size:10px;font-weight:700;color:#3A5265;'
        f'text-transform:uppercase;letter-spacing:1.5px;margin-bottom:6px;'
        f'font-family:\'IBM Plex Mono\',monospace;">Schedule Quality Score</div>'
        f'<div style="font-size:64px;font-weight:900;color:{score_col};'
        f'font-family:\'IBM Plex Mono\',monospace;line-height:1;letter-spacing:-3px;">'
        f'{score}</div>'
        f'<div style="font-size:12px;color:{score_col};margin-top:4px;">'
        f'{score_label}</div>'
        f'</div>'
        f'<div style="display:grid;grid-template-columns:repeat(3,auto);gap:16px;">'
        f'<div style="text-align:center;background:#0d2035;border-radius:6px;padding:12px 20px;">'
        f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
        f'letter-spacing:0.8px;font-family:\'IBM Plex Mono\',monospace;">Checks Run</div>'
        f'<div style="font-size:28px;font-weight:700;color:#8AAABF;'
        f'font-family:\'IBM Plex Mono\',monospace;">{len(results)}</div>'
        f'</div>'
        f'<div style="text-align:center;background:#0d2035;border-radius:6px;padding:12px 20px;">'
        f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
        f'letter-spacing:0.8px;font-family:\'IBM Plex Mono\',monospace;">Passed</div>'
        f'<div style="font-size:28px;font-weight:700;color:#1E7A4E;'
        f'font-family:\'IBM Plex Mono\',monospace;">{n_pass}</div>'
        f'</div>'
        f'<div style="text-align:center;background:#0d2035;border-radius:6px;padding:12px 20px;">'
        f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
        f'letter-spacing:0.8px;font-family:\'IBM Plex Mono\',monospace;">Failed</div>'
        f'<div style="font-size:28px;font-weight:700;color:#C0392B;'
        f'font-family:\'IBM Plex Mono\',monospace;">{n_fail}</div>'
        f'</div></div></div></div>',
        unsafe_allow_html=True,
    )

    # -- Score Gauge Bar -------------------------------------------------------
    fig_g = go.Figure(go.Bar(
        x=[score], y=["Score"],
        orientation="h",
        marker_color=score_col,
        text=[f"{score}/100"],
        textposition="inside",
        textfont=dict(color="white", size=13, family="IBM Plex Mono"),
    ))
    # Background zones
    fig_g.add_shape(type="rect", x0=0, x1=60,  y0=-0.5, y1=0.5,
                    fillcolor="#FAEAEA", line_width=0, layer="below")
    fig_g.add_shape(type="rect", x0=60, x1=80, y0=-0.5, y1=0.5,
                    fillcolor="#FDF3E0", line_width=0, layer="below")
    fig_g.add_shape(type="rect", x0=80, x1=100,y0=-0.5, y1=0.5,
                    fillcolor="#E4F5EC", line_width=0, layer="below")
    for xv, label in [(60,"Risk"), (80,"Good")]:
        fig_g.add_vline(x=xv, line_color="#D4DCE4", line_width=1)
        fig_g.add_annotation(x=xv, y=0.55, text=label, showarrow=False,
                             font=dict(size=9, color="#8A9DB0"))
    fig_g.update_layout(
        height=70, margin=dict(l=10,r=10,t=10,b=10),
        xaxis=dict(range=[0,100], showgrid=False, showticklabels=False),
        yaxis=dict(showticklabels=False),
        plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
        showlegend=False,
    )
    st.plotly_chart(fig_g)

    # -- Results Grid ---------------------------------------------------------
    st.markdown('<div class="section-label" style="margin-top:16px;">Assessment Results</div>',
                unsafe_allow_html=True)

    # Render each check as a row card
    for key, r in results.items():
        passed   = r.get("pass_", False)
        pct      = r.get("pct", 0)
        count    = r.get("count", 0)
        total    = r.get("total", 0)
        thresh   = r.get("threshold", 0)
        direction= r.get("direction","lower")
        label    = _DCMA_LABELS.get(key, key)
        desc     = _DCMA_DESCRIPTIONS.get(key, "")

        status_bg    = "#E4F5EC" if passed else "#FAEAEA"
        status_brd   = "#1E7A4E" if passed else "#C0392B"
        status_label = "PASS" if passed else "FAIL"
        status_col   = "#1E7A4E" if passed else "#C0392B"

        # Bar showing value vs threshold
        bar_pct   = min(100, pct) if direction != "higher" else min(100, pct)
        bar_col   = status_brd
        thresh_pct= min(100, thresh)

        st.markdown(
            f'<div style="background:{status_bg};border:1px solid {status_brd};'
            f'border-left:4px solid {status_brd};border-radius:0 6px 6px 0;'
            f'padding:12px 16px;margin-bottom:6px;">'
            f'<div style="display:flex;justify-content:space-between;'
            f'align-items:flex-start;gap:12px;flex-wrap:wrap;">'
            f'<div style="flex:1;min-width:200px;">'
            f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:4px;">'
            f'<span style="font-size:11px;font-weight:800;color:{status_col};'
            f'background:white;padding:1px 7px;border-radius:3px;'
            f'border:1px solid {status_brd};">{status_label}</span>'
            f'<span style="font-size:13px;font-weight:700;color:#0B1929;">{label}</span>'
            f'</div>'
            f'<div style="font-size:11px;color:#6B7C8E;line-height:1.4;">{desc}</div>'
            f'</div>'
            f'<div style="text-align:right;flex-shrink:0;">'
            f'<div style="font-size:24px;font-weight:700;color:{status_col};'
            f'font-family:\'IBM Plex Mono\',monospace;line-height:1;">{pct}%</div>'
            f'<div style="font-size:10px;color:#8A9DB0;margin-top:2px;">'
            f'{count:,} of {total:,}</div>'
            f'<div style="font-size:10px;color:#8A9DB0;">'
            f'Threshold: {"≥" if direction=="higher" else "≤"}{thresh}%</div>'
            f'</div></div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    # -- What this means -------------------------------------------------------
    st.markdown("<br>", unsafe_allow_html=True)
    if n_fail == 0:
        st.success("This programme passes all DCMA checks. Schedule quality is high.")
    else:
        failures = [_DCMA_LABELS[k] for k, r in results.items() if not r.get("pass_")]
        st.markdown(
            f'<div style="background:#FDF3E0;border-left:3px solid #E8951D;'
            f'border-radius:0 6px 6px 0;padding:14px 18px;">'
            f'<strong>What this means:</strong> This programme fails {n_fail} DCMA check'
            f'{"s" if n_fail>1 else ""}. '
            f'Failed checks: {", ".join(failures)}. '
            f'Review these issues with your planner before submitting the programme '
            f'to a client, auditor or contract board.</div>',
            unsafe_allow_html=True,
        )

    # -- Export ----------------------------------------------------------------
    rows = []
    for key, r in results.items():
        rows.append({
            "Check":         _DCMA_LABELS.get(key, key),
            "Result":        "PASS" if r.get("pass_") else "FAIL",
            "Value (%)":     r.get("pct", "-"),
            "Count":         r.get("count", 0),
            "Total":         r.get("total", 0),
            "Threshold (%)": r.get("threshold", "-"),
            "Description":   _DCMA_DESCRIPTIONS.get(key, ""),
        })
    rows.append({
        "Check":         "SCHEDULE QUALITY SCORE",
        "Result":        f"{score}/100  ({score_label})",
        "Value (%)":     score, "Count":"","Total":"","Threshold (%)":"",
        "Description":   "Weighted composite score across all DCMA checks.",
    })
    export_df = pd.DataFrame(rows)
    xls = export_df_to_excel({"DCMA Assessment": export_df})

    col1, col2 = st.columns([1,3])
    col1.download_button(
        "Export DCMA Report",
        xls,
        f"dcma_assessment_{datetime.now().strftime('%Y%m%d')}.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key="dcma_export",
    )

    # Word export
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document()
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(11)
        h = doc.add_heading("DCMA 14-Point Schedule Assessment", 0)
        h.runs[0].font.color.rgb = RGBColor(11,25,41)
        doc.add_paragraph(
            f"Schedule Quality Score: {score}/100  ({score_label})\n"
            f"Passed: {n_pass}  |  Failed: {n_fail}  |  Checks run: {len(results)}"
        )
        doc.add_paragraph("")
        t = doc.add_table(rows=1, cols=5)
        t.style = "Table Grid"
        for cell, hdr in zip(t.rows[0].cells,
                             ["Check","Result","Value (%)","Threshold (%)","Description"]):
            cell.text = hdr
        for key, r in results.items():
            row = t.add_row().cells
            row[0].text = _DCMA_LABELS.get(key, key)
            row[1].text = "PASS" if r.get("pass_") else "FAIL"
            row[2].text = str(r.get("pct",""))
            row[3].text = str(r.get("threshold",""))
            row[4].text = _DCMA_DESCRIPTIONS.get(key,"")

        import io
        buf = io.BytesIO()
        doc.save(buf)
        col1.download_button(
            "Export DCMA to Word",
            buf.getvalue(),
            f"dcma_assessment_{datetime.now().strftime('%Y%m%d')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dcma_word",
        )
    except Exception:
        pass




# =============================================================================
# CRITICAL PATH GANTT  — colour-coded Gantt for Critical Path page
# =============================================================================

def render_critical_path_gantt(tasks: pd.DataFrame, near_crit_days: float):
    """
    Render a professional colour-coded Gantt chart for the Critical Path page.
    Colours: Critical = #C0392B, Near-Critical = #E8951D, Normal = #0B1929
    """
    if tasks.empty:
        return

    tasks_c = get_critical_threshold(tasks, near_crit_days)

    # Only show critical and near-critical activities to keep chart readable
    gantt_tasks = tasks_c[
        tasks_c["is_critical"] | tasks_c.get("is_near_critical", pd.Series(False, index=tasks_c.index))
    ].copy() if "is_critical" in tasks_c.columns else tasks_c.copy()

    if gantt_tasks.empty:
        return

    gantt_tasks = gantt_tasks.dropna(subset=["eff_start","eff_finish"]).copy()
    if gantt_tasks.empty:
        return

    gantt_tasks["eff_start_dt"]  = pd.to_datetime(gantt_tasks["eff_start"],  errors="coerce")
    gantt_tasks["eff_finish_dt"] = pd.to_datetime(gantt_tasks["eff_finish"], errors="coerce")
    gantt_tasks = gantt_tasks.dropna(subset=["eff_start_dt","eff_finish_dt"])

    # Limit to 80 activities for performance
    gantt_tasks = gantt_tasks.sort_values("total_float_days").head(80)

    gantt_tasks["Colour"] = gantt_tasks.apply(lambda r: (
        "#C0392B" if r.get("is_critical") else
        "#E8951D" if r.get("is_near_critical") else
        "#0B1929"
    ), axis=1)
    gantt_tasks["Status_Label"] = gantt_tasks.apply(lambda r: (
        "Critical" if r.get("is_critical") else
        "Near-Critical" if r.get("is_near_critical") else
        "Float"
    ), axis=1)
    gantt_tasks["Label"] = (
        gantt_tasks["task_code"].astype(str) + "  " +
        gantt_tasks["task_name"].astype(str).str[:45]
    )
    gantt_tasks["Float_str"] = gantt_tasks["total_float_days"].apply(
        lambda f: f"{round(float(f),1)}d" if f is not None and str(f) not in ("","nan") else "—"
    )

    fig = go.Figure()

    for status, colour in [("Critical","#C0392B"),("Near-Critical","#E8951D")]:
        sub = gantt_tasks[gantt_tasks["Status_Label"] == status]
        if sub.empty:
            continue
        for _, row in sub.iterrows():
            fig.add_trace(go.Bar(
                y=[row["Label"]],
                x=[(row["eff_finish_dt"] - row["eff_start_dt"]).days],
                base=[row["eff_start_dt"].timestamp() * 1000],
                orientation="h",
                name=status,
                marker_color=colour,
                marker_line_width=0,
                showlegend=False,
                hovertemplate=(
                    f"<b>{row.get('task_code','')}</b><br>"
                    f"{row.get('task_name','')}<br>"
                    f"Start: {format_date(row.get('eff_start'))}<br>"
                    f"Finish: {format_date(row.get('eff_finish'))}<br>"
                    f"Float: {row.get('Float_str','—')}<br>"
                    f"<extra></extra>"
                ),
            ))

    # Today line
    today_ms = pd.Timestamp.now().timestamp() * 1000
    fig.add_vline(
        x=datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
        line_dash="dot", line_color="#8AAABF", line_width=1.5,
        annotation_text="Today", annotation_font_color="#8AAABF", annotation_font_size=9,
    )

    # Legend manually via annotations
    fig.add_annotation(x=0, y=1.04, xref="paper", yref="paper", showarrow=False,
                       text='<span style="color:#C0392B;">■</span> Critical  '
                            '<span style="color:#E8951D;">■</span> Near-Critical',
                       font=dict(size=11), align="left")

    fig.update_layout(
        title=dict(text="Critical & Near-Critical Activities — Gantt View",
                   font=dict(size=13, color="#0B1929")),
        barmode="overlay",
        xaxis=dict(
            type="date",
            tickformat="%b %Y",
            gridcolor="#EEF1F5",
            showgrid=True,
        ),
        yaxis=dict(
            autorange="reversed",
            tickfont=dict(size=10, family="IBM Plex Mono"),
        ),
        plot_bgcolor="#FFFFFF",
        paper_bgcolor="#FFFFFF",
        height=max(300, min(900, 60 + len(gantt_tasks) * 22)),
        margin=dict(l=10, r=10, t=60, b=10),
        bargap=0.3,
    )

    st.plotly_chart(fig)


# =============================================================================
# OVERVIEW ENHANCEMENTS — Activity status doughnut + sparkline metrics
# =============================================================================

def render_activity_status_doughnut(tasks: pd.DataFrame):
    """Compact doughnut chart showing activity status breakdown."""
    if tasks.empty or "status" not in tasks.columns:
        return

    status_map = {
        "TK_NotStart":  ("Not Started", "#0B1929"),
        "Not Started":  ("Not Started", "#0B1929"),
        "TK_Active":    ("In Progress", "#E8951D"),
        "In Progress":  ("In Progress", "#E8951D"),
        "TK_Complete":  ("Complete",    "#1E7A4E"),
        "Complete":     ("Complete",    "#1E7A4E"),
        "Completed":    ("Complete",    "#1E7A4E"),
    }

    counts = {}
    for s in tasks["status"].dropna():
        label, colour = status_map.get(str(s), ("Other", "#8AAABF"))
        counts[label] = counts.get(label, {"count": 0, "colour": colour})
        counts[label]["count"] += 1

    if not counts:
        return

    labels  = list(counts.keys())
    values  = [v["count"] for v in counts.values()]
    colours = [v["colour"] for v in counts.values()]
    total   = sum(values)
    pct_complete = round(counts.get("Complete", {}).get("count", 0) / total * 100, 0) if total else 0

    fig = go.Figure(go.Pie(
        labels=labels, values=values,
        hole=0.62,
        marker=dict(colors=colours, line=dict(color="#FFFFFF", width=2)),
        textinfo="percent",
        textfont=dict(size=10, family="IBM Plex Sans"),
        hovertemplate="%{label}: %{value:,} (%{percent})<extra></extra>",
    ))
    fig.add_annotation(
        text=f"<b>{int(pct_complete)}%</b><br><span style='font-size:9px'>Complete</span>",
        x=0.5, y=0.5, showarrow=False,
        font=dict(size=13, color="#0B1929", family="IBM Plex Mono"),
    )
    fig.update_layout(
        height=220,
        margin=dict(l=5,r=5,t=5,b=5),
        paper_bgcolor="#FFFFFF",
        showlegend=True,
        legend=dict(
            orientation="v", x=1.0, y=0.5,
            font=dict(size=10, family="IBM Plex Sans"),
        ),
    )
    st.plotly_chart(fig)


def render_float_distribution_bar(tasks: pd.DataFrame, near_crit_days: float):
    """Horizontal stacked bar showing float distribution at a glance."""
    if tasks.empty or "total_float_days" not in tasks.columns:
        return

    bands = {
        "Negative": 0, "Critical (0d)": 0,
        f"Near-Crit (0–{near_crit_days}d)": 0,
        "Low Float (0–30d)": 0, "Healthy (>30d)": 0,
    }
    colours = {
        "Negative": "#C0392B", "Critical (0d)": "#E8951D",
        f"Near-Crit (0–{near_crit_days}d)": "#F5C842",
        "Low Float (0–30d)": "#5BA08A", "Healthy (>30d)": "#1E7A4E",
    }

    for f in tasks["total_float_days"].dropna():
        val = safe_float(f, 999)
        if val < 0:
            bands["Negative"] += 1
        elif val == 0:
            bands["Critical (0d)"] += 1
        elif val <= near_crit_days:
            bands[f"Near-Crit (0–{near_crit_days}d)"] += 1
        elif val <= 30:
            bands["Low Float (0–30d)"] += 1
        else:
            bands["Healthy (>30d)"] += 1

    total = max(sum(bands.values()), 1)
    fig = go.Figure()
    x_offset = 0
    for band, count in bands.items():
        if count == 0:
            continue
        pct = count / total * 100
        fig.add_trace(go.Bar(
            x=[count], y=["Float"],
            orientation="h",
            name=band,
            marker_color=colours[band],
            marker_line_width=0,
            hovertemplate=f"{band}: {count:,} ({pct:.0f}%)<extra></extra>",
        ))

    fig.update_layout(
        barmode="stack",
        height=60,
        margin=dict(l=0,r=0,t=0,b=0),
        xaxis=dict(showticklabels=False, showgrid=False),
        yaxis=dict(showticklabels=False),
        plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
        showlegend=True,
        legend=dict(orientation="h", y=2.5, x=0, font=dict(size=9)),
        bargap=0,
    )
    st.plotly_chart(fig)


# =============================================================================
# COPY TO CLIPBOARD  — JS-based, works in Streamlit via components
# =============================================================================

def copy_button(text: str, button_label: str = "Copy to Clipboard",
                key: str = "copy_btn"):
    """
    Render a copy-to-clipboard button using Streamlit's JS component.
    Works in Streamlit 1.57+.
    """
    import streamlit.components.v1 as components
    safe_text = text.replace("\\", "\\\\").replace("`", "\\`").replace("$", "\\$")
    components.html(
        f"""
        <button onclick="
            navigator.clipboard.writeText(`{safe_text}`)
                .then(() => {{
                    this.textContent = '\\u2713 Copied';
                    this.style.background = '#1E7A4E';
                    this.style.color = 'white';
                    setTimeout(() => {{
                        this.textContent = '{button_label}';
                        this.style.background = '';
                        this.style.color = '';
                    }}, 2000);
                }});
        "
        style="
            background: #FFFFFF;
            color: #0B1929;
            border: 1.5px solid #0B1929;
            border-radius: 4px;
            padding: 6px 14px;
            font-size: 12px;
            font-weight: 600;
            cursor: pointer;
            font-family: 'IBM Plex Sans', sans-serif;
            transition: all 0.12s ease;
        "
        onmouseover="this.style.background='#0B1929';this.style.color='white';"
        onmouseout="this.style.background='white';this.style.color='#0B1929';"
        >{button_label}</button>
        """,
        height=44,
    )


# =============================================================================
# SCHEDULE HEALTH SPARKLINE  — mini trend if comparison data available
# =============================================================================

def render_health_trend(near_crit_days: float):
    """
    If comparison data is available, show a mini trend of key metrics.
    Shows previous vs current float/critical count.
    """
    if "_mi_prev" not in st.session_state or "_mi_curr" not in st.session_state:
        return

    try:
        prev_t = st.session_state["_mi_prev"]["tasks_df"]
        curr_t = st.session_state["_mi_curr"]["tasks_df"]

        if prev_t.empty or curr_t.empty:
            return

        def _get_metrics(t):
            t_c = get_critical_threshold(t, near_crit_days)
            return {
                "Critical":       int(t_c["is_critical"].sum()) if "is_critical" in t_c.columns else 0,
                "Neg Float":      int(t_c["total_float_days"].apply(lambda f: safe_float(f,999)<0).sum()) if "total_float_days" in t_c.columns else 0,
                "Near-Critical":  int(t_c.get("is_near_critical", pd.Series(False, index=t_c.index)).sum()),
            }

        prev_m = _get_metrics(prev_t)
        curr_m = _get_metrics(curr_t)

        st.markdown('<div class="section-label">Trend vs Previous Revision</div>',
                    unsafe_allow_html=True)

        cols = st.columns(3)
        for col, key in zip(cols, ["Critical","Neg Float","Near-Critical"]):
            prev_v = prev_m[key]
            curr_v = curr_m[key]
            diff   = curr_v - prev_v
            arrow  = "↑" if diff > 0 else "↓" if diff < 0 else "→"
            colour = "#C0392B" if diff > 0 else "#1E7A4E" if diff < 0 else "#6B7C8E"
            col.markdown(
                f'<div class="card-flat" style="text-align:center;">'
                f'<div style="font-size:9px;font-weight:700;color:#8A9DB0;'
                f'text-transform:uppercase;letter-spacing:1px;margin-bottom:4px;'
                f'font-family:\'IBM Plex Mono\',monospace;">{key}</div>'
                f'<div style="font-size:22px;font-weight:700;color:#0B1929;'
                f'font-family:\'IBM Plex Mono\',monospace;">{curr_v}</div>'
                f'<div style="font-size:12px;font-weight:600;color:{colour};">'
                f'{arrow} {abs(diff)} vs prev</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
    except Exception:
        pass




# =============================================================================
# PAGE: AI CLIENT COMMENT REVIEW
# Upload a client programme review letter (Word or PDF).
# AI extracts every comment/finding, cross-references the live XER data,
# determines if each comment is correct, and drafts a response.
# =============================================================================

_CR_STATE_KEY  = "cr_results"          # session state: list of reviewed comments
_CR_DOC_KEY    = "cr_doc_text"         # session state: raw extracted text
_CR_RUNNING    = "cr_running"          # session state: flag


def _extract_text_from_upload(uploaded_file) -> str:
    """
    Extract plain text from a Word (.docx) or PDF upload.
    Returns the full text as a single string.
    """
    name = uploaded_file.name.lower()
    raw  = uploaded_file.read()

    if name.endswith(".docx"):
        try:
            from docx import Document
            import io
            doc  = Document(io.BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return text
        except Exception as e:
            raise ValueError(f"Could not read Word document: {e}")

    elif name.endswith(".pdf"):
        try:
            import pdfplumber, io
            text_parts = []
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Could not read PDF: {e}")

    else:
        raise ValueError("Unsupported file type. Please upload a .docx or .pdf file.")


def _build_schedule_context(data: dict, near_crit_days: float) -> str:
    """
    Build a compact schedule summary string to pass to the AI as context.
    Keeps tokens low while giving the model enough data to verify comments.
    """
    tasks = data.get("tasks_df", pd.DataFrame())
    rels  = data.get("relationships_df", pd.DataFrame())
    proj  = data.get("project_info", {})

    if tasks.empty:
        return "No programme data available."

    tasks_c = get_critical_threshold(tasks, near_crit_days)
    lines   = []

    # Project summary
    pname  = proj.get("name", "")
    ddate  = format_date(proj.get("data_date"))
    n_tot  = len(tasks_c)
    n_rels = len(rels)
    n_crit = int(tasks_c["is_critical"].sum()) if "is_critical" in tasks_c.columns else 0
    n_neg  = int(tasks_c["total_float_days"].apply(lambda f: safe_float(f, 0) < 0).sum()) \
             if "total_float_days" in tasks_c.columns else 0
    n_nc   = int(tasks_c["is_near_critical"].sum()) if "is_near_critical" in tasks_c.columns else 0

    lines.append(f"PROJECT: {pname}")
    lines.append(f"DATA DATE: {ddate}")
    lines.append(f"ACTIVITIES: {n_tot} total | {n_crit} critical | {n_neg} negative float | {n_nc} near-critical")
    lines.append(f"RELATIONSHIPS: {n_rels}")
    lines.append("")

    # Open logic
    if not rels.empty and "task_id" in tasks_c.columns:
        wp = set(rels["succ_task_id"].dropna()) if "succ_task_id" in rels.columns else set()
        ws = set(rels["pred_task_id"].dropna()) if "pred_task_id" in rels.columns else set()
        no_pred = tasks_c[~tasks_c["task_id"].isin(wp)]
        no_succ = tasks_c[~tasks_c["task_id"].isin(ws)]
        if "task_type" in tasks_c.columns:
            no_pred = no_pred[~no_pred["task_type"].astype(str).str.contains(
                "Milestone|LOE|WBS", na=False)]
            no_succ = no_succ[~no_succ["task_type"].astype(str).str.contains(
                "Milestone|LOE|WBS", na=False)]
        lines.append(f"OPEN LOGIC: {len(no_pred)} activities with no predecessor, {len(no_succ)} with no successor")

    # Negative float activities (first 20)
    if n_neg > 0 and "total_float_days" in tasks_c.columns:
        neg_df = tasks_c[tasks_c["total_float_days"].apply(
            lambda f: safe_float(f, 0) < 0
        )].sort_values("total_float_days").head(20)
        lines.append("")
        lines.append("NEGATIVE FLOAT ACTIVITIES (worst 20):")
        for _, r in neg_df.iterrows():
            lines.append(
                f"  {r.get('task_code','?')} | {str(r.get('task_name',''))[:50]} "
                f"| Float: {round(safe_float(r.get('total_float_days'), 0), 1)}d "
                f"| Finish: {format_date(r.get('eff_finish'))} "
                f"| WBS: {str(r.get('wbs_path',''))[:40]}"
            )

    # Critical activities not started (first 15)
    if "is_critical" in tasks_c.columns and "status" in tasks_c.columns:
        crit_ns = tasks_c[
            tasks_c["is_critical"] &
            tasks_c["status"].apply(lambda s: str(s) in ("TK_NotStart", "Not Started"))
        ].head(15)
        if not crit_ns.empty:
            lines.append("")
            lines.append("CRITICAL NOT STARTED (first 15):")
            for _, r in crit_ns.iterrows():
                lines.append(
                    f"  {r.get('task_code','?')} | {str(r.get('task_name',''))[:50]} "
                    f"| Start: {format_date(r.get('eff_start'))} "
                    f"| Finish: {format_date(r.get('eff_finish'))}"
                )

    # High lag relationships (first 10)
    if not rels.empty and "lag_days" in rels.columns:
        big_lag = rels[rels["lag_days"].apply(lambda l: safe_float(l, 0) > 10)].head(10)
        if not big_lag.empty:
            lines.append("")
            lines.append("EXCESSIVE LAG RELATIONSHIPS (>10 days, first 10):")
            for _, r in big_lag.iterrows():
                lines.append(
                    f"  {r.get('pred_task_code','?')} -> {r.get('succ_task_code','?')} "
                    f"| Lag: {r.get('lag_days','?')}d "
                    f"| Type: {r.get('rel_type','?')}"
                )

    # Hard constraints (first 10)
    if "cstr_type" in tasks_c.columns:
        HARD = {"CS_MANDFIN", "CS_MANDSTART", "Must Finish On",
                "Must Start On", "Mandatory Start", "Mandatory Finish"}
        hard_df = tasks_c[tasks_c["cstr_type"].apply(
            lambda c: str(c).strip() in HARD
        )].head(10)
        if not hard_df.empty:
            lines.append("")
            lines.append("HARD CONSTRAINTS (first 10):")
            for _, r in hard_df.iterrows():
                lines.append(
                    f"  {r.get('task_code','?')} | {str(r.get('task_name',''))[:50]} "
                    f"| {r.get('cstr_type','?')} | {format_date(r.get('cstr_date'))}"
                )

    # All activity codes (for reference lookups)
    all_codes = tasks_c["task_code"].dropna().tolist() if "task_code" in tasks_c.columns else []
    if all_codes:
        lines.append("")
        lines.append(f"ALL ACTIVITY CODES ({len(all_codes)} total):")
        lines.append(", ".join(str(c) for c in all_codes[:200]))
        if len(all_codes) > 200:
            lines.append(f"  ... and {len(all_codes)-200} more")

    # WBS areas
    if "wbs_path" in tasks_c.columns:
        wbs_tops = tasks_c["wbs_path"].apply(
            lambda x: str(x).split(" > ")[0].strip()
            if x and str(x).strip() not in ("","nan") else "Unknown"
        ).unique().tolist()
        lines.append("")
        lines.append(f"WBS AREAS: {', '.join(wbs_tops[:20])}")

    return "\n".join(lines)


def _get_api_key() -> str:
    """
    Resolve the Anthropic API key.
    Priority: 1. Streamlit secrets  2. Session state (user entered)
    Returns empty string if not found.
    """
    # 1. Streamlit Cloud secrets (set in app Settings → Secrets)
    try:
        key = st.secrets.get("ANTHROPIC_API_KEY", "")
        if key and str(key).startswith("sk-"):
            return str(key)
    except Exception:
        pass
    # 2. User-entered key stored in session state
    key = st.session_state.get("cr_api_key", "")
    if key and str(key).startswith("sk-"):
        return str(key)
    return ""


def _call_ai_review(doc_text: str, schedule_context: str,
                    company_name: str = "",
                    api_key: str = "") -> list:
    """
    Call the Anthropic API to extract and review client comments.
    Returns a list of dicts, one per comment finding.
    """
    import json, requests

    if not api_key:
        raise RuntimeError("NO_API_KEY")

    company_str = f" from {company_name}" if company_name else ""
    data_name   = "EHL" if not company_name else company_name

    system_prompt = f"""You are an expert project controls engineer and contract specialist
working for a construction contractor. You have full access to the current programme data.

Your job is to:
1. Read a client programme review letter{company_str}
2. Extract every distinct comment, finding, or requirement
3. Check each one against the programme data provided
4. Determine if each comment is correct, partially correct, or incorrect
5. Draft a professional but firm contractual response for each
6. Recommend what action (if any) to take in the programme

PROGRAMME DATA:
{schedule_context}

OUTPUT FORMAT:
Return ONLY a JSON array. No preamble, no markdown, no explanation outside the JSON.
Each element must have these exact keys:
- "ref": short reference code like "PC-01", "PC-02" etc
- "comment": the client's exact finding or comment (verbatim or close paraphrase)
- "category": one of: "Logic", "Float", "Constraints", "Milestones", "Resources", "Dates", "Quality", "General"
- "correct": one of: "Correct", "Partially Correct", "Incorrect", "Cannot Verify"
- "severity": one of: "High", "Medium", "Low"
- "data_evidence": what the programme data shows that supports or contradicts this comment
- "response": a professional draft response (2-4 sentences, firm but not aggressive)
- "recommended_action": specific action for the planner or PM (or "No action required")
- "priority": integer 1 (highest) to 3 (lowest)

Be precise and technical. Reference specific activity codes and float values where the data supports it.
If a comment references an activity code that exists in the programme, look it up and verify.
Do not make up data — only reference what is in the programme data provided.
Return between 3 and 20 comment objects. If the document has fewer comments, return all of them."""

    user_prompt = f"""Please review this client programme assessment letter and cross-reference it
against our schedule data. Extract all findings and provide our response.

CLIENT DOCUMENT:
{doc_text[:8000]}"""

    try:
        response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"Content-Type": "application/json", "x-api-key": api_key, "anthropic-version": "2023-06-01"},
            json={
                "model":      "claude-sonnet-4-20250514",
                "max_tokens": 4000,
                "system":     system_prompt,
                "messages":   [{"role": "user", "content": user_prompt}],
            },
            timeout=60,
        )
        response.raise_for_status()
        data = response.json()

        raw = ""
        for block in data.get("content", []):
            if block.get("type") == "text":
                raw += block.get("text", "")

        # Strip any markdown fences
        raw = raw.strip()
        if raw.startswith("```"):
            raw = "\n".join(raw.split("\n")[1:])
        if raw.endswith("```"):
            raw = "\n".join(raw.split("\n")[:-1])
        raw = raw.strip()

        results = json.loads(raw)
        if isinstance(results, list):
            return results
        return []

    except Exception as e:
        raise RuntimeError(f"AI review failed: {e}")


def page_client_review(data: dict, near_crit_days: float):
    """
    AI-powered client comment review page.
    Upload a Word or PDF programme review letter. AI extracts every comment,
    checks it against the live schedule, and drafts your response.
    """
    ctrl_bar(
        "Client Comment Review",
        "Upload a client programme review letter. AI cross-references every "
        "comment against your schedule and drafts your response.",
    )
    mode_toggle_bar()

    pm_note(
        "Upload the client's programme review letter (Word or PDF). "
        "PlanTrace will read every comment, check it against your live schedule data, "
        "and tell you which comments are correct, which are not, and what to say back."
    )

    tasks = data.get("tasks_df", pd.DataFrame())
    if tasks.empty:
        empty_state(
            "",
            "No Programme Loaded",
            "Upload your XER programme first, then upload the client's review letter.",
            "",
        )
        return

    # -- Upload panel ---------------------------------------------------------
    col_upload, col_options = st.columns([2, 1], gap="large")

    with col_upload:
        st.markdown(
            '<div class="section-label">Upload Client Review Document</div>',
            unsafe_allow_html=True,
        )
        uploaded = st.file_uploader(
            "Client document",
            type=["docx", "pdf"],
            key="cr_upload",
            label_visibility="collapsed",
            help="Upload the client's programme review letter, technical query, or assessment report.",
        )

    with col_options:
        st.markdown(
            '<div class="section-label">Options</div>',
            unsafe_allow_html=True,
        )
        company_name = st.text_input(
            "Client / Company Name",
            placeholder="e.g. Acciona, Network Rail",
            key="cr_company",
            label_visibility="collapsed",
        )
        st.caption("Used to personalise the response drafts.")

        # API key resolution
        _resolved_key = _get_api_key()
        if not _resolved_key:
            st.markdown(
                '<div style="background:#FDF3E0;border-left:3px solid #E8951D;' +
                'border-radius:0 4px 4px 0;padding:10px 12px;margin-top:8px;">' +
                '<div style="font-size:10px;font-weight:700;color:#8A5A00;' +
                'text-transform:uppercase;letter-spacing:0.8px;margin-bottom:4px;">' +
                'API Key Required</div>' +
                '<div style="font-size:11px;color:#374151;line-height:1.5;">' +
                'Enter your Anthropic API key below, or add it to Streamlit secrets ' +
                'as <code>ANTHROPIC_API_KEY</code>. ' +
                'Get a key at <a href="https://console.anthropic.com" target="_blank">' +
                'console.anthropic.com</a>. Each review costs ~£0.01–£0.03.' +
                '</div></div>',
                unsafe_allow_html=True,
            )
            entered_key = st.text_input(
                "Anthropic API Key",
                placeholder="sk-ant-...",
                type="password",
                key="cr_api_key_input",
                label_visibility="collapsed",
            )
            if entered_key and entered_key.startswith("sk-"):
                st.session_state["cr_api_key"] = entered_key
                _resolved_key = entered_key
                st.success("API key accepted.")
        else:
            key_src = "Streamlit secrets" if (
                hasattr(st, "secrets") and
                st.secrets.get("ANTHROPIC_API_KEY","").startswith("sk-")
            ) else "session"
            st.markdown(
                f'<div style="background:#E4F5EC;border-left:3px solid #1E7A4E;' +
                f'border-radius:0 4px 4px 0;padding:8px 12px;margin-top:8px;">' +
                f'<span style="font-size:11px;color:#1E7A4E;font-weight:600;">' +
                f'API key loaded ({key_src})</span></div>',
                unsafe_allow_html=True,
            )

        if st.session_state.get(_CR_STATE_KEY):
            if st.button("Clear Results", key="cr_clear"):
                st.session_state.pop(_CR_STATE_KEY, None)
                st.session_state.pop(_CR_DOC_KEY,  None)
                st.rerun()

    # -- Run Review -----------------------------------------------------------
    if uploaded and not st.session_state.get(_CR_STATE_KEY):
        # Show document preview
        with st.expander("Document Preview", expanded=False):
            try:
                doc_text = _extract_text_from_upload(uploaded)
                st.session_state[_CR_DOC_KEY] = doc_text
                st.text(doc_text[:2000] + ("\n..." if len(doc_text) > 2000 else ""))
            except Exception as e:
                st.error(str(e))
                return

        col_run, _ = st.columns([1, 3])
        if col_run.button(
            "Run AI Review",
            key="cr_run",
            type="primary",
            help="Sends the document to AI for analysis. Takes 15-30 seconds.",
        ):
            doc_text = st.session_state.get(_CR_DOC_KEY, "")
            if not doc_text.strip():
                st.error("Document appears to be empty.")
                return

            _api_key = _get_api_key()
            if not _api_key:
                st.error(
                    "No API key found. Enter your Anthropic API key in the "
                    "options panel above before running the review."
                )
                st.stop()
            with st.spinner(
                "Reading client document, checking against your schedule, "
                "drafting responses... (15-30 seconds)"
            ):
                try:
                    context  = _build_schedule_context(data, near_crit_days)
                    results  = _call_ai_review(doc_text, context, company_name, _api_key)
                    # Sort by priority then severity
                    sev_order = {"High": 0, "Medium": 1, "Low": 2}
                    results.sort(key=lambda r: (
                        int(r.get("priority", 2)),
                        sev_order.get(r.get("severity", "Low"), 2),
                    ))
                    st.session_state[_CR_STATE_KEY] = results
                    st.rerun()
                except RuntimeError as e:
                    if "NO_API_KEY" in str(e):
                        st.error(
                            "API key missing. Add ANTHROPIC_API_KEY to Streamlit secrets "
                            "or enter it in the options panel above."
                        )
                    else:
                        st.error(f"Review failed: {e}")
                    return
                except Exception as e:
                    st.error(f"Review failed: {e}")
                    return

    elif uploaded and not st.session_state.get(_CR_DOC_KEY):
        try:
            doc_text = _extract_text_from_upload(uploaded)
            st.session_state[_CR_DOC_KEY] = doc_text
        except Exception as e:
            st.error(str(e))
            return

    # -- Display Results -------------------------------------------------------
    results = st.session_state.get(_CR_STATE_KEY, [])
    if not results:
        if not uploaded:
            # Empty state
            st.markdown(
                f'<div class="empty-state">'
                f'<div style="font-size:40px;margin-bottom:12px;">📋</div>'
                f'<div style="font-size:18px;font-weight:700;color:#0B1929;'
                f'margin-bottom:8px;">Upload a Client Review Letter</div>'
                f'<div style="font-size:13px;color:#6B7C8E;max-width:440px;'
                f'margin:0 auto;line-height:1.6;">'
                f'Upload your client\'s programme review letter in Word or PDF format. '
                f'PlanTrace will extract every comment and cross-reference it against '
                f'your live schedule data.</div>'
                f'<div style="margin-top:20px;font-size:12px;color:#8A9DB0;">'
                f'Accepts: .docx &nbsp;·&nbsp; .pdf</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
        return

    # -- Summary strip ---------------------------------------------------------
    n_total    = len(results)
    n_correct  = sum(1 for r in results if r.get("correct") == "Correct")
    n_partial  = sum(1 for r in results if r.get("correct") == "Partially Correct")
    n_wrong    = sum(1 for r in results if r.get("correct") == "Incorrect")
    n_cant     = sum(1 for r in results if r.get("correct") == "Cannot Verify")
    n_high     = sum(1 for r in results if r.get("severity") == "High")

    st.markdown(
        f'<div style="background:#0B1929;border-radius:8px;padding:16px 24px;'
        f'margin-bottom:20px;display:flex;gap:16px;flex-wrap:wrap;align-items:center;">'
        f'<div style="font-size:13px;font-weight:700;color:#FFFFFF;margin-right:8px;">'
        f'{n_total} Comments Found</div>'
        f'<div style="background:#FAEAEA;border-radius:4px;padding:4px 12px;">'
        f'<span style="font-size:11px;font-weight:700;color:#C0392B;">'
        f'{n_correct} Correct</span></div>'
        f'<div style="background:#FDF3E0;border-radius:4px;padding:4px 12px;">'
        f'<span style="font-size:11px;font-weight:700;color:#8A5A00;">'
        f'{n_partial} Partial</span></div>'
        f'<div style="background:#E4F5EC;border-radius:4px;padding:4px 12px;">'
        f'<span style="font-size:11px;font-weight:700;color:#1E7A4E;">'
        f'{n_wrong} Incorrect</span></div>'
        f'<div style="background:#EEF1F5;border-radius:4px;padding:4px 12px;">'
        f'<span style="font-size:11px;font-weight:700;color:#4A5A68;">'
        f'{n_cant} Cannot Verify</span></div>'
        f'<div style="margin-left:auto;">'
        f'<span style="font-size:11px;color:#C0392B;font-weight:700;">'
        f'{n_high} High Priority</span></div>'
        f'</div>',
        unsafe_allow_html=True,
    )

    # -- Filter ----------------------------------------------------------------
    col_f1, col_f2, _ = st.columns([1, 1, 2])
    with col_f1:
        filter_correct = st.selectbox(
            "Filter by verdict",
            ["All", "Correct", "Partially Correct", "Incorrect", "Cannot Verify"],
            key="cr_filter_correct",
            label_visibility="collapsed",
        )
    with col_f2:
        filter_sev = st.selectbox(
            "Filter by severity",
            ["All Severity", "High", "Medium", "Low"],
            key="cr_filter_sev",
            label_visibility="collapsed",
        )

    filtered = results
    if filter_correct != "All":
        filtered = [r for r in filtered if r.get("correct") == filter_correct]
    if filter_sev != "All Severity":
        filtered = [r for r in filtered if r.get("severity") == filter_sev]

    st.markdown(
        f'<div class="section-label" style="margin-top:8px;">'
        f'{len(filtered)} comment{"s" if len(filtered)!=1 else ""} shown</div>',
        unsafe_allow_html=True,
    )

    # -- Comment Cards ---------------------------------------------------------
    _CORRECT_STYLE = {
        "Correct":          ("#FAEAEA", "#C0392B", "red"),
        "Partially Correct":("#FDF3E0", "#E8951D", "amber"),
        "Incorrect":        ("#E4F5EC", "#1E7A4E", "green"),
        "Cannot Verify":    ("#EEF1F5", "#6B7C8E", "grey"),
    }
    _SEV_STYLE = {
        "High":   ("red",   "#C0392B"),
        "Medium": ("amber", "#E8951D"),
        "Low":    ("grey",  "#6B7C8E"),
    }

    for i, r in enumerate(filtered):
        verdict  = r.get("correct", "Cannot Verify")
        severity = r.get("severity", "Medium")
        ref      = r.get("ref", f"PC-{i+1:02d}")
        category = r.get("category", "General")

        bg, brd, chip_style  = _CORRECT_STYLE.get(verdict, _CORRECT_STYLE["Cannot Verify"])
        sev_chip, sev_col    = _SEV_STYLE.get(severity, _SEV_STYLE["Medium"])

        with st.container():
            st.markdown(
                f'<div style="background:{bg};border:1px solid {brd};'
                f'border-left:4px solid {brd};border-radius:0 8px 8px 0;'
                f'padding:14px 18px;margin-bottom:10px;">'
                # Header row
                f'<div style="display:flex;justify-content:space-between;'
                f'align-items:flex-start;flex-wrap:wrap;gap:8px;margin-bottom:10px;">'
                f'<div style="display:flex;align-items:center;gap:10px;flex-wrap:wrap;">'
                f'<span style="font-size:11px;font-weight:800;color:{brd};'
                f'background:white;padding:1px 8px;border-radius:3px;'
                f'border:1px solid {brd};">{ref}</span>'
                f'{chip(verdict, chip_style)}'
                f'{chip(severity, sev_chip)}'
                f'<span style="font-size:10px;color:#8A9DB0;font-weight:600;">'
                f'{category}</span>'
                f'</div></div>'
                # Client comment
                f'<div style="margin-bottom:10px;">'
                f'<div style="font-size:9px;font-weight:700;color:#8A9DB0;'
                f'text-transform:uppercase;letter-spacing:1px;margin-bottom:4px;">'
                f'Client Comment</div>'
                f'<div style="font-size:13px;color:#1C2B3A;line-height:1.6;'
                f'font-style:italic;">"{r.get("comment","")}"</div>'
                f'</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

            # Data evidence + response in two columns
            ev_col, resp_col = st.columns([1, 1], gap="medium")

            with ev_col:
                st.markdown(
                    f'<div style="background:white;border:1px solid #D4DCE4;'
                    f'border-radius:6px;padding:12px 14px;height:100%;">'
                    f'<div style="font-size:9px;font-weight:700;color:#8A9DB0;'
                    f'text-transform:uppercase;letter-spacing:1px;margin-bottom:6px;">'
                    f'Schedule Evidence</div>'
                    f'<div style="font-size:12px;color:#374151;line-height:1.6;">'
                    f'{r.get("data_evidence","No data evidence available.")}'
                    f'</div></div>',
                    unsafe_allow_html=True,
                )

            with resp_col:
                st.markdown(
                    f'<div style="background:#0B1929;border:1px solid #0d2035;'
                    f'border-radius:6px;padding:12px 14px;height:100%;">'
                    f'<div style="font-size:9px;font-weight:700;color:#3A5265;'
                    f'text-transform:uppercase;letter-spacing:1px;margin-bottom:6px;">'
                    f'Draft Response</div>'
                    f'<div style="font-size:12px;color:#8AAABF;line-height:1.6;">'
                    f'{r.get("response","")}'
                    f'</div></div>',
                    unsafe_allow_html=True,
                )

            # Recommended action
            action = r.get("recommended_action", "")
            if action and action != "No action required":
                st.markdown(
                    f'<div style="background:#FFFBEB;border-left:3px solid #E8951D;'
                    f'border-radius:0 4px 4px 0;padding:8px 12px;margin-top:6px;">'
                    f'<span style="font-size:10px;font-weight:700;color:#8A5A00;'
                    f'text-transform:uppercase;letter-spacing:0.8px;">Action Required: </span>'
                    f'<span style="font-size:12px;color:#374151;">{action}</span>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
            st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)

    # -- Export ----------------------------------------------------------------
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="section-label">Export</div>', unsafe_allow_html=True)

    # Build Excel
    export_rows = [{
        "Ref":              r.get("ref", ""),
        "Category":         r.get("category", ""),
        "Severity":         r.get("severity", ""),
        "Client Comment":   r.get("comment", ""),
        "Verdict":          r.get("correct", ""),
        "Schedule Evidence":r.get("data_evidence", ""),
        "Draft Response":   r.get("response", ""),
        "Recommended Action":r.get("recommended_action", ""),
    } for r in results]

    summary_rows = [{
        "Metric": k, "Count": v
    } for k, v in [
        ("Total Comments",      n_total),
        ("Client Correct",      n_correct),
        ("Partially Correct",   n_partial),
        ("Client Incorrect",    n_wrong),
        ("Cannot Verify",       n_cant),
        ("High Priority",       n_high),
    ]]

    xls = export_df_to_excel({
        "Summary":          pd.DataFrame(summary_rows),
        "Comment Review":   pd.DataFrame(export_rows),
        "High Priority":    pd.DataFrame([r for r in export_rows
                                          if r["Severity"] == "High"]) \
                            if any(r["Severity"]=="High" for r in export_rows) \
                            else pd.DataFrame(columns=["No data"]),
    })

    dl_col, word_col, _ = st.columns([1, 1, 2])
    dl_col.download_button(
        "Export to Excel",
        xls,
        f"client_review_{datetime.now().strftime('%Y%m%d')}.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key="cr_excel_export",
    )

    # Word export — formal response letter
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io as _io

        doc = DocxDocument()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Header
        h = doc.add_heading("Programme Review — Comment Response Register", 0)
        h.runs[0].font.color.rgb = RGBColor(11, 25, 41)

        proj = data.get("project_info", {})
        doc.add_paragraph(
            f"Project: {proj.get('name','')}\n"
            f"Data Date: {format_date(proj.get('data_date'))}\n"
            f"Prepared: {datetime.now().strftime('%d %B %Y')}\n"
            f"Client: {company_name or 'Client'}\n"
            f"Total Comments: {n_total}  |  Correct: {n_correct}  |  "
            f"Partially Correct: {n_partial}  |  Incorrect: {n_wrong}"
        )
        doc.add_paragraph("")

        for r in results:
            verdict  = r.get("correct","")
            severity = r.get("severity","")
            ref      = r.get("ref","")

            h2 = doc.add_heading(f"{ref}  —  {r.get('category','')}  |  {verdict}  |  {severity}", 2)
            h2.runs[0].font.color.rgb = RGBColor(11, 25, 41)

            doc.add_paragraph("Client Comment:", style="Normal").runs[0].bold = True
            p = doc.add_paragraph(r.get("comment",""))
            p.runs[0].italic = True

            doc.add_paragraph("Schedule Evidence:", style="Normal").runs[0].bold = True
            doc.add_paragraph(r.get("data_evidence",""))

            doc.add_paragraph("Response:", style="Normal").runs[0].bold = True
            doc.add_paragraph(r.get("response",""))

            action = r.get("recommended_action","")
            if action and action != "No action required":
                doc.add_paragraph("Action Required:", style="Normal").runs[0].bold = True
                doc.add_paragraph(action)

            doc.add_paragraph("")

        buf = _io.BytesIO()
        doc.save(buf)

        word_col.download_button(
            "Export Response Letter (Word)",
            buf.getvalue(),
            f"client_response_{datetime.now().strftime('%Y%m%d')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="cr_word_export",
        )
    except Exception:
        pass




# =============================================================================
# PAGE: ACCELERATION ANALYSIS
# "How do I recover X weeks?" — the most common PM question.
# Identifies which activities on the critical/near-critical path can be
# compressed, what the compression cost/effort is, and generates a
# ranked recovery programme.
# =============================================================================

def page_acceleration_analysis(data: dict, near_crit_days: float):
    ctrl_bar(
        "Acceleration Analysis",
        "Identify which activities to compress to recover a target date. "
        "Ranked by float gain, feasibility and impact.",
    )
    mode_toggle_bar()

    tasks = data.get("tasks_df", pd.DataFrame())
    rels  = data.get("relationships_df", pd.DataFrame())
    proj  = data.get("project_info", {})

    if tasks.empty:
        empty_state("", "No Programme Data",
                    "Upload a programme to run acceleration analysis.", "")
        return

    pm_note(
        "Enter a target recovery amount (days or weeks) and PlanTrace will "
        "identify the best activities to compress. Focus is on the critical path "
        "— compressing non-critical activities wastes time and money."
    )

    tasks_c = get_critical_threshold(tasks, near_crit_days)

    # -- Current position -----------------------------------------------------
    valid_ends = tasks_c["eff_finish"].dropna() if "eff_finish" in tasks_c.columns else pd.Series()
    current_end = pd.Timestamp(valid_ends.max()) if not valid_ends.empty else None
    current_end_str = format_date(current_end) if current_end else "Unknown"

    n_crit = int(tasks_c["is_critical"].sum()) if "is_critical" in tasks_c.columns else 0
    n_neg  = int(tasks_c["total_float_days"].apply(
        lambda f: safe_float(f, 0) < 0).sum()) if "total_float_days" in tasks_c.columns else 0

    # -- Inputs ----------------------------------------------------------------
    st.markdown('<div class="section-label">Recovery Target</div>',
                unsafe_allow_html=True)

    inp_col1, inp_col2, inp_col3 = st.columns([1, 1, 2])
    with inp_col1:
        recovery_days = st.number_input(
            "Days to recover",
            min_value=1, max_value=365, value=14, step=1,
            key="acc_days",
            help="How many calendar days do you need to pull back the programme end date?",
            label_visibility="collapsed",
        )
        st.caption(f"Recovery target: **{recovery_days} days** "
                   f"({round(recovery_days/7, 1)} weeks)")

    with inp_col2:
        max_compress = st.slider(
            "Max compression per activity (%)",
            min_value=10, max_value=50, value=25, step=5,
            key="acc_compress",
            help="Maximum realistic compression for a single activity. "
                 "25% is aggressive but achievable with additional resource.",
        )
        st.caption(f"Max {max_compress}% compression per activity")

    with inp_col3:
        focus = st.radio(
            "Focus on",
            ["Critical path only", "Critical + near-critical", "All activities"],
            key="acc_focus",
            horizontal=True,
            label_visibility="collapsed",
        )

    # -- Identify compressible activities -------------------------------------
    if focus == "Critical path only":
        pool = tasks_c[tasks_c.get("is_critical", pd.Series(False, index=tasks_c.index))] \
               if "is_critical" in tasks_c.columns else tasks_c
    elif focus == "Critical + near-critical":
        mask = (tasks_c.get("is_critical", pd.Series(False, index=tasks_c.index)) |
                tasks_c.get("is_near_critical", pd.Series(False, index=tasks_c.index)))
        pool = tasks_c[mask]
    else:
        pool = tasks_c

    # Filter to not-started or in-progress only
    if "status" in pool.columns:
        pool = pool[~pool["status"].apply(
            lambda s: str(s) in ("TK_Complete", "Complete", "Completed")
        )]

    # Filter to activities with meaningful duration
    if "rem_dur_days" in pool.columns:
        pool = pool[pool["rem_dur_days"].apply(lambda d: safe_float(d, 0) >= 5)]
    elif "orig_dur_days" in pool.columns:
        pool = pool[pool["orig_dur_days"].apply(lambda d: safe_float(d, 0) >= 5)]

    if pool.empty:
        st.info("No compressible activities found with the current filter settings.")
        return

    # -- Score each activity for compression potential -------------------------
    # Score = duration × criticality weight × feasibility
    # Float penalty: negative float activities score higher (most urgent)
    candidates = []
    for _, row in pool.iterrows():
        dur = safe_float(row.get("rem_dur_days") or row.get("orig_dur_days"), 0)
        if dur < 5:
            continue

        tf        = safe_float(row.get("total_float_days"), 0)
        is_crit   = bool(row.get("is_critical", False))
        is_nc     = bool(row.get("is_near_critical", False))

        # Days that can be recovered by compressing this activity
        max_save  = round(dur * (max_compress / 100), 1)

        # Feasibility score (higher = easier to compress)
        # Long activities are easier to compress than short ones
        # Activities not yet started are easier than in-progress
        status    = str(row.get("status", ""))
        not_started = status in ("TK_NotStart", "Not Started")
        feasibility = "High" if dur > 30 and not_started else \
                      "Medium" if dur > 14 else "Low"
        feas_score = {"High": 3, "Medium": 2, "Low": 1}[feasibility]

        # Priority score
        priority_score = (
            (10 if tf < 0 else 0) +
            (8  if is_crit else 0) +
            (4  if is_nc  else 0) +
            (feas_score * 2) +
            (min(10, dur / 5))
        )

        # Compression method suggestion
        if dur > 60:
            method = "Parallel working or additional resource gang"
        elif dur > 21:
            method = "Overtime or additional resource"
        elif dur > 7:
            method = "Extended working hours"
        else:
            method = "Resource increase only — limited potential"

        candidates.append({
            "task_code":      row.get("task_code", ""),
            "task_name":      str(row.get("task_name", ""))[:55],
            "wbs_path":       str(row.get("wbs_path", ""))[:40],
            "eff_start":      format_date(row.get("eff_start")),
            "eff_finish":     format_date(row.get("eff_finish")),
            "rem_dur_days":   round(dur, 0),
            "total_float":    round(tf, 1),
            "is_critical":    is_crit,
            "is_nc":          is_nc,
            "max_save_days":  max_save,
            "feasibility":    feasibility,
            "method":         method,
            "priority_score": round(priority_score, 1),
            "status":         status,
        })

    if not candidates:
        st.info("No suitable candidates found for compression.")
        return

    # Sort by priority score descending
    cands = sorted(candidates, key=lambda x: x["priority_score"], reverse=True)

    # -- Build recovery programme ----------------------------------------------
    # Greedily select activities until recovery target is met
    selected     = []
    total_saved  = 0.0
    for c in cands:
        if total_saved >= recovery_days:
            break
        still_need = recovery_days - total_saved
        save_this  = min(c["max_save_days"], still_need)
        compress_pct = round(save_this / c["rem_dur_days"] * 100, 0) if c["rem_dur_days"] else 0
        c["proposed_save"]    = round(save_this, 1)
        c["compress_pct"]     = compress_pct
        selected.append(c)
        total_saved += save_this

    remaining_gap = max(0, recovery_days - total_saved)

    # -- Summary ---------------------------------------------------------------
    gap_col = "#1E7A4E" if remaining_gap == 0 else "#C0392B"
    st.markdown(
        f'<div style="background:#0B1929;border-radius:8px;padding:20px 28px;'
        f'margin-bottom:20px;display:flex;gap:24px;flex-wrap:wrap;align-items:center;">'
        f'<div>'
        f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
        f'letter-spacing:1.2px;font-family:\'IBM Plex Mono\',monospace;">Target Recovery</div>'
        f'<div style="font-size:36px;font-weight:800;color:#E8951D;'
        f'font-family:\'IBM Plex Mono\',monospace;line-height:1;">{recovery_days}d</div>'
        f'</div>'
        f'<div>'
        f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
        f'letter-spacing:1.2px;font-family:\'IBM Plex Mono\',monospace;">Recoverable</div>'
        f'<div style="font-size:36px;font-weight:800;color:#1E7A4E;'
        f'font-family:\'IBM Plex Mono\',monospace;line-height:1;">{round(total_saved,1)}d</div>'
        f'</div>'
        f'<div>'
        f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
        f'letter-spacing:1.2px;font-family:\'IBM Plex Mono\',monospace;">Remaining Gap</div>'
        f'<div style="font-size:36px;font-weight:800;color:{gap_col};'
        f'font-family:\'IBM Plex Mono\',monospace;line-height:1;">{round(remaining_gap,1)}d</div>'
        f'</div>'
        f'<div>'
        f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
        f'letter-spacing:1.2px;font-family:\'IBM Plex Mono\',monospace;">Activities to Compress</div>'
        f'<div style="font-size:36px;font-weight:800;color:#8AAABF;'
        f'font-family:\'IBM Plex Mono\',monospace;line-height:1;">{len(selected)}</div>'
        f'</div>'
        f'<div style="margin-left:auto;text-align:right;">'
        f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
        f'letter-spacing:1.2px;font-family:\'IBM Plex Mono\',monospace;">Current Forecast End</div>'
        f'<div style="font-size:16px;font-weight:700;color:#8AAABF;">{current_end_str}</div>'
        f'</div></div>',
        unsafe_allow_html=True,
    )

    if remaining_gap > 0:
        st.markdown(
            f'<div class="attn attn-amber">'
            f'<strong>Recovery shortfall: {round(remaining_gap,1)} days.</strong> '
            f'The programme cannot recover the full {recovery_days} days by compressing '
            f'individual activities at {max_compress}% compression. '
            f'Consider increasing the compression threshold, resequencing work, '
            f'or accepting a partial recovery of {round(total_saved,1)} days.'
            f'</div>',
            unsafe_allow_html=True,
        )

    # -- Recovery Programme Table ----------------------------------------------
    st.markdown('<div class="section-label" style="margin-top:16px;">'
                'Proposed Recovery Programme</div>', unsafe_allow_html=True)

    for i, c in enumerate(selected, 1):
        feas_chip = ("green" if c["feasibility"]=="High" else
                     "amber" if c["feasibility"]=="Medium" else "grey")
        crit_chip = ("red" if c["is_critical"] else
                     "amber" if c["is_nc"] else "blue")
        crit_label = "Critical" if c["is_critical"] else \
                     "Near-Critical" if c["is_nc"] else "Float"

        # Progress bar for compression
        bar_pct = int(c["compress_pct"])
        bar_col = ("#C0392B" if bar_pct > 30 else
                   "#E8951D" if bar_pct > 15 else "#1E7A4E")

        st.markdown(
            f'<div style="background:#FFFFFF;border:1px solid #D4DCE4;'
            f'border-left:4px solid {bar_col};border-radius:0 6px 6px 0;'
            f'padding:14px 18px;margin-bottom:8px;">'
            f'<div style="display:flex;justify-content:space-between;'
            f'align-items:flex-start;flex-wrap:wrap;gap:8px;margin-bottom:8px;">'
            f'<div style="display:flex;align-items:center;gap:8px;flex-wrap:wrap;">'
            f'<span style="font-size:12px;font-weight:800;color:white;'
            f'background:{bar_col};padding:1px 8px;border-radius:3px;">{i}</span>'
            f'<span style="font-size:13px;font-weight:700;color:#0B1929;">'
            f'{c["task_code"]}</span>'
            f'<span style="font-size:13px;color:#374151;">{c["task_name"]}</span>'
            f'</div>'
            f'<div style="display:flex;gap:6px;flex-wrap:wrap;align-items:center;">'
            f'{chip(crit_label, crit_chip)}'
            f'{chip(c["feasibility"], feas_chip)}'
            f'</div></div>'
            f'<div style="display:grid;grid-template-columns:repeat(5,1fr);'
            f'gap:8px;margin-bottom:10px;">'
            f'<div><div style="font-size:9px;color:#8A9DB0;text-transform:uppercase;'
            f'letter-spacing:0.8px;font-family:\'IBM Plex Mono\',monospace;">Rem Duration</div>'
            f'<div style="font-size:13px;font-weight:700;color:#0B1929;">'
            f'{int(c["rem_dur_days"])}d</div></div>'
            f'<div><div style="font-size:9px;color:#8A9DB0;text-transform:uppercase;'
            f'letter-spacing:0.8px;font-family:\'IBM Plex Mono\',monospace;">Compress By</div>'
            f'<div style="font-size:13px;font-weight:700;color:{bar_col};">'
            f'{int(c["compress_pct"])}%</div></div>'
            f'<div><div style="font-size:9px;color:#8A9DB0;text-transform:uppercase;'
            f'letter-spacing:0.8px;font-family:\'IBM Plex Mono\',monospace;">Days Saved</div>'
            f'<div style="font-size:13px;font-weight:700;color:#1E7A4E;">'
            f'{c["proposed_save"]}d</div></div>'
            f'<div><div style="font-size:9px;color:#8A9DB0;text-transform:uppercase;'
            f'letter-spacing:0.8px;font-family:\'IBM Plex Mono\',monospace;">Float</div>'
            f'<div style="font-size:13px;font-weight:700;'
            f'color:{"#C0392B" if c["total_float"]<0 else "#374151"};">'
            f'{c["total_float"]}d</div></div>'
            f'<div><div style="font-size:9px;color:#8A9DB0;text-transform:uppercase;'
            f'letter-spacing:0.8px;font-family:\'IBM Plex Mono\',monospace;">Finish</div>'
            f'<div style="font-size:11px;color:#374151;">{c["eff_finish"]}</div></div>'
            f'</div>'
            # Compression bar
            f'<div style="background:#EEF1F5;border-radius:3px;height:6px;'
            f'margin-bottom:8px;overflow:hidden;">'
            f'<div style="width:{min(100,bar_pct)}%;height:100%;'
            f'background:{bar_col};border-radius:3px;"></div></div>'
            # Method
            f'<div style="font-size:11px;color:#6B7C8E;">'
            f'<strong style="color:#374151;">Method:</strong> {c["method"]} &nbsp;|&nbsp; '
            f'WBS: {c["wbs_path"]}'
            f'</div></div>',
            unsafe_allow_html=True,
        )

    # -- All candidates chart --------------------------------------------------
    with st.expander(f"All {len(cands)} candidate activities (ranked)", expanded=False):
        chart_df = pd.DataFrame(cands[:40])
        if not chart_df.empty:
            chart_df["Label"] = (chart_df["task_code"].astype(str) + "  " +
                                 chart_df["task_name"].astype(str).str[:30])
            fig_acc = px.bar(
                chart_df.head(30), x="max_save_days", y="Label",
                orientation="h",
                color="feasibility",
                color_discrete_map={"High":"#1E7A4E","Medium":"#E8951D","Low":"#8AAABF"},
                title=f"Maximum Recoverable Days per Activity (top 30, {max_compress}% compression)",
                labels={"max_save_days":"Max Days Recoverable","Label":""},
            )
            fig_acc.update_yaxes(autorange="reversed")
            fig_acc.update_layout(
                height=max(300, min(800, 60 + len(chart_df.head(30)) * 22)),
                plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
                margin=dict(l=10,r=10,t=40,b=10),
                legend=dict(orientation="h", y=1.08),
            )
            st.plotly_chart(fig_acc)

    # -- Export ----------------------------------------------------------------
    selected_exp = pd.DataFrame([{
        "Priority":       i+1,
        "Activity ID":    c["task_code"],
        "Activity Name":  c["task_name"],
        "WBS":            c["wbs_path"],
        "Rem Duration (d)": int(c["rem_dur_days"]),
        "Float (d)":      c["total_float"],
        "Critical":       "Yes" if c["is_critical"] else "",
        "Compress By (%)":int(c["compress_pct"]),
        "Days Saved":     c["proposed_save"],
        "Feasibility":    c["feasibility"],
        "Method":         c["method"],
        "Start":          c["eff_start"],
        "Finish":         c["eff_finish"],
    } for i, c in enumerate(selected)])

    all_cands_exp = pd.DataFrame([{
        "Activity ID":    c["task_code"],
        "Activity Name":  c["task_name"],
        "WBS":            c["wbs_path"],
        "Rem Duration (d)": int(c["rem_dur_days"]),
        "Max Save (d)":   c["max_save_days"],
        "Float (d)":      c["total_float"],
        "Feasibility":    c["feasibility"],
        "Priority Score": c["priority_score"],
        "Method":         c["method"],
    } for c in cands])

    xls = export_df_to_excel({
        "Recovery Programme":  selected_exp,
        "All Candidates":      all_cands_exp,
    })
    st.download_button(
        "Export Recovery Programme",
        xls,
        f"acceleration_analysis_{datetime.now().strftime('%Y%m%d')}.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key="acc_export",
    )




# =============================================================================
# PAGE: COMPENSATION EVENT TRACKER
# NEC3 / NEC4 CE management linked to programme activities.
# Tracks: CE ref, clause, activity, time impact claimed vs float available,
# agreed/disputed status, programme impact.
# =============================================================================

_CE_KEY = "ce_register"   # session state

_CE_CLAUSES_NEC4 = [
    "60.1(1)  — Employer instruction changing scope",
    "60.1(2)  — Employer fails to provide access",
    "60.1(3)  — Employer fails to provide information",
    "60.1(4)  — Employer instruction to stop/not start work",
    "60.1(5)  — Employer instruction changing key date",
    "60.1(6)  — Employer does not reply to communication",
    "60.1(7)  — Employer changes decision",
    "60.1(8)  — Employer withholds acceptance without reason",
    "60.1(9)  — Employer fails to take over by due date",
    "60.1(10) — Employer causes defect",
    "60.1(11) — Testing/inspection causes unnecessary delay",
    "60.1(12) — Employer fails to make payment",
    "60.1(13) — Conditions more adverse than contractually assumed",
    "60.1(14) — Employer delays issuing certificate",
    "60.1(15) — Breach of contract not listed above",
    "60.1(16) — Clause 60.1(16) — Prevention",
    "60.1(17) — Employer-risk weather event",
    "60.1(18) — Employer-risk event (force majeure)",
    "60.1(19) — Employer does not provide facilities",
    "Other — See description",
]

_CE_STATUS_OPTS = [
    "Notified — Awaiting PM Response",
    "Under Assessment",
    "Agreed — Time + Cost",
    "Agreed — Time Only",
    "Agreed — Cost Only",
    "Disputed",
    "Withdrawn",
    "Not a CE",
]


def _ce_status_chip(status: str) -> tuple:
    """Return (chip_style, border_colour) for a CE status."""
    s = status.lower()
    if "agreed" in s:     return "green", "#1E7A4E"
    if "disputed" in s:   return "red",   "#C0392B"
    if "notified" in s:   return "amber", "#E8951D"
    if "assessment" in s: return "amber", "#E8951D"
    if "withdrawn" in s or "not a ce" in s:
        return "grey",  "#6B7C8E"
    return "blue", "#0B1929"


def page_ce_tracker(data: dict, near_crit_days: float):
    """NEC3/NEC4 Compensation Event Tracker."""
    ctrl_bar(
        "Compensation Events",
        "NEC3/NEC4 CE register — track time impact, float position and programme "
        "effect for every compensation event.",
    )
    mode_toggle_bar()

    tasks = data.get("tasks_df", pd.DataFrame())
    proj  = data.get("project_info", {})

    pm_note(
        "Log every Compensation Event here. PlanTrace checks the float on the "
        "linked activity to determine whether the CE has a programme impact. "
        "If float is available, the time impact may be absorbed. If not, "
        "it drives a claim for an extension of time."
    )

    if _CE_KEY not in st.session_state:
        st.session_state[_CE_KEY] = []

    ces = st.session_state[_CE_KEY]

    # -- Summary strip ---------------------------------------------------------
    if ces:
        n_total    = len(ces)
        n_agreed   = sum(1 for c in ces if "Agreed" in c.get("status",""))
        n_disputed = sum(1 for c in ces if "Disputed" in c.get("status",""))
        n_pending  = n_total - n_agreed - n_disputed
        total_time_claimed = sum(safe_float(c.get("time_claimed_days",0),0) for c in ces)
        total_time_agreed  = sum(safe_float(c.get("time_agreed_days",0),0)
                                  for c in ces if "Agreed" in c.get("status",""))
        total_prog_impact  = sum(safe_float(c.get("programme_impact",0),0) for c in ces)

        st.markdown(
            f'<div style="background:#0B1929;border-radius:8px;padding:16px 24px;'
            f'margin-bottom:20px;display:flex;gap:16px;flex-wrap:wrap;">'
            f'<div style="text-align:center;flex:1;min-width:80px;">'
            f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
            f'letter-spacing:1px;font-family:\'IBM Plex Mono\',monospace;">Total CEs</div>'
            f'<div style="font-size:26px;font-weight:800;color:#FFFFFF;'
            f'font-family:\'IBM Plex Mono\',monospace;">{n_total}</div>'
            f'</div>'
            f'<div style="text-align:center;flex:1;min-width:80px;">'
            f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
            f'letter-spacing:1px;font-family:\'IBM Plex Mono\',monospace;">Agreed</div>'
            f'<div style="font-size:26px;font-weight:800;color:#1E7A4E;'
            f'font-family:\'IBM Plex Mono\',monospace;">{n_agreed}</div>'
            f'</div>'
            f'<div style="text-align:center;flex:1;min-width:80px;">'
            f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
            f'letter-spacing:1px;font-family:\'IBM Plex Mono\',monospace;">Disputed</div>'
            f'<div style="font-size:26px;font-weight:800;color:#C0392B;'
            f'font-family:\'IBM Plex Mono\',monospace;">{n_disputed}</div>'
            f'</div>'
            f'<div style="text-align:center;flex:1;min-width:80px;">'
            f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
            f'letter-spacing:1px;font-family:\'IBM Plex Mono\',monospace;">Pending</div>'
            f'<div style="font-size:26px;font-weight:800;color:#E8951D;'
            f'font-family:\'IBM Plex Mono\',monospace;">{n_pending}</div>'
            f'</div>'
            f'<div style="text-align:center;flex:1;min-width:80px;">'
            f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
            f'letter-spacing:1px;font-family:\'IBM Plex Mono\',monospace;">Time Claimed</div>'
            f'<div style="font-size:26px;font-weight:800;color:#E8951D;'
            f'font-family:\'IBM Plex Mono\',monospace;">{int(total_time_claimed)}d</div>'
            f'</div>'
            f'<div style="text-align:center;flex:1;min-width:80px;">'
            f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
            f'letter-spacing:1px;font-family:\'IBM Plex Mono\',monospace;">EOT Agreed</div>'
            f'<div style="font-size:26px;font-weight:800;color:#1E7A4E;'
            f'font-family:\'IBM Plex Mono\',monospace;">{int(total_time_agreed)}d</div>'
            f'</div>'
            f'<div style="text-align:center;flex:1;min-width:80px;">'
            f'<div style="font-size:9px;color:#3A5265;text-transform:uppercase;'
            f'letter-spacing:1px;font-family:\'IBM Plex Mono\',monospace;">Prog Impact</div>'
            f'<div style="font-size:26px;font-weight:800;'
            f'color:{"#C0392B" if total_prog_impact>0 else "#1E7A4E"};'
            f'font-family:\'IBM Plex Mono\',monospace;">{int(total_prog_impact)}d</div>'
            f'</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    # -- Tabs ------------------------------------------------------------------
    tab_register, tab_add, tab_analysis = st.tabs([
        "CE Register", "Add / Edit CE", "Impact Analysis"
    ])

    # -- REGISTER TAB ---------------------------------------------------------
    with tab_register:
        if not ces:
            st.info("No compensation events logged yet. Use the 'Add / Edit CE' tab to add one.")
        else:
            # Filter
            status_filter = st.selectbox(
                "Filter by status",
                ["All"] + _CE_STATUS_OPTS,
                key="ce_filter_status",
                label_visibility="collapsed",
            )
            show_ces = ces if status_filter == "All" else \
                       [c for c in ces if c.get("status","") == status_filter]

            for i, ce in enumerate(show_ces):
                chip_style, brd_col = _ce_status_chip(ce.get("status",""))
                bg = {"green":"#E4F5EC","red":"#FAEAEA",
                      "amber":"#FDF3E0","grey":"#F5F7F9","blue":"#EEF1F5"}.get(chip_style,"#F5F7F9")

                # Float check against linked activity
                float_val  = None
                float_warn = ""
                if ce.get("activity_id") and not tasks.empty and "task_code" in tasks.columns:
                    match = tasks[tasks["task_code"] == ce["activity_id"]]
                    if not match.empty and "total_float_days" in match.columns:
                        float_val = safe_float(match.iloc[0]["total_float_days"], None)
                        tc = safe_float(ce.get("time_claimed_days",0), 0)
                        if float_val is not None:
                            if float_val < 0:
                                float_warn = f"Schedule already behind ({float_val}d float) — CE drives delay"
                            elif float_val < tc:
                                float_warn = f"Float ({float_val}d) less than claim ({tc}d) — partial EOT likely"
                            else:
                                float_warn = f"Float ({float_val}d) absorbs claim ({tc}d) — no programme impact"

                prog_impact = safe_float(ce.get("programme_impact",0),0)

                st.markdown(
                    f'<div style="background:{bg};border:1px solid {brd_col};'
                    f'border-left:4px solid {brd_col};border-radius:0 6px 6px 0;'
                    f'padding:12px 16px;margin-bottom:8px;">'
                    f'<div style="display:flex;justify-content:space-between;'
                    f'align-items:flex-start;flex-wrap:wrap;gap:6px;margin-bottom:6px;">'
                    f'<div style="display:flex;align-items:center;gap:8px;flex-wrap:wrap;">'
                    f'<span style="font-size:12px;font-weight:800;color:{brd_col};">'
                    f'{ce.get("ref","?")}</span>'
                    f'<span style="font-size:13px;font-weight:700;color:#0B1929;">'
                    f'{ce.get("title","")[:55]}</span>'
                    f'</div>'
                    f'{chip(ce.get("status","?"), chip_style)}'
                    f'</div>'
                    f'<div style="font-size:11px;color:#6B7C8E;margin-bottom:6px;">'
                    f'{ce.get("clause","")}</div>'
                    f'<div style="display:grid;grid-template-columns:repeat(4,1fr);gap:8px;">'
                    f'<div><div style="font-size:9px;color:#8A9DB0;text-transform:uppercase;'
                    f'letter-spacing:0.8px;">Activity</div>'
                    f'<div style="font-size:12px;font-weight:600;color:#374151;">'
                    f'{ce.get("activity_id","—")}</div></div>'
                    f'<div><div style="font-size:9px;color:#8A9DB0;text-transform:uppercase;'
                    f'letter-spacing:0.8px;">Time Claimed</div>'
                    f'<div style="font-size:12px;font-weight:700;color:#E8951D;">'
                    f'{int(safe_float(ce.get("time_claimed_days",0),0))}d</div></div>'
                    f'<div><div style="font-size:9px;color:#8A9DB0;text-transform:uppercase;'
                    f'letter-spacing:0.8px;">EOT Agreed</div>'
                    f'<div style="font-size:12px;font-weight:700;color:#1E7A4E;">'
                    f'{int(safe_float(ce.get("time_agreed_days",0),0))}d</div></div>'
                    f'<div><div style="font-size:9px;color:#8A9DB0;text-transform:uppercase;'
                    f'letter-spacing:0.8px;">Prog Impact</div>'
                    f'<div style="font-size:12px;font-weight:700;'
                    f'color:{"#C0392B" if prog_impact>0 else "#1E7A4E"};">'
                    f'{int(prog_impact)}d</div></div>'
                    f'</div>'
                    + (f'<div style="margin-top:6px;font-size:11px;'
                       f'color:{"#C0392B" if "behind" in float_warn or "partial" in float_warn else "#1E7A4E"};">'
                       f'Float check: {float_warn}</div>' if float_warn else "")
                    + (f'<div style="margin-top:4px;font-size:11px;color:#6B7C8E;">'
                       f'<em>{ce.get("description","")[:120]}</em></div>'
                       if ce.get("description") else "")
                    + f'</div>',
                    unsafe_allow_html=True,
                )

            if st.button("Clear All CEs", key="ce_clear_all"):
                st.session_state[_CE_KEY] = []
                st.rerun()

    # -- ADD / EDIT TAB --------------------------------------------------------
    with tab_add:
        st.markdown('<div class="section-label">Add Compensation Event</div>',
                    unsafe_allow_html=True)

        # Auto-generate next ref
        existing_refs = [c.get("ref","") for c in ces]
        next_num      = len(ces) + 1
        default_ref   = f"CE-{next_num:03d}"

        with st.form("ce_add_form", clear_on_submit=True):
            col1, col2 = st.columns(2)
            with col1:
                ce_ref    = st.text_input("CE Reference", value=default_ref)
                ce_title  = st.text_input("Title / Brief Description",
                                          placeholder="e.g. Employer instruction to change ventilation design")
                ce_clause = st.selectbox("NEC4 Clause", _CE_CLAUSES_NEC4)
                ce_status = st.selectbox("Status", _CE_STATUS_OPTS)
            with col2:
                ce_act    = st.text_input("Linked Activity ID",
                                          placeholder="e.g. ACT-1023",
                                          help="Enter the activity code most affected by this CE")
                ce_notified = st.date_input("Date Notified")
                ce_time_claimed = st.number_input(
                    "Time Impact Claimed (days)", min_value=0, value=0, step=1)
                ce_time_agreed  = st.number_input(
                    "EOT Agreed (days)", min_value=0, value=0, step=1)

            ce_prog_impact = st.number_input(
                "Programme Impact (days, 0 = absorbed by float)", min_value=0, value=0, step=1)
            ce_desc = st.text_area("Notes / Detail", height=80,
                                   placeholder="Describe the event and any agreed actions...")

            submitted = st.form_submit_button("Add CE to Register", type="primary")
            if submitted and ce_title:
                # Auto-calculate programme impact from float if activity linked
                prog_imp = ce_prog_impact
                if ce_act and not tasks.empty and "task_code" in tasks.columns:
                    match = tasks[tasks["task_code"] == ce_act]
                    if not match.empty and "total_float_days" in match.columns:
                        fv = safe_float(match.iloc[0]["total_float_days"], 0)
                        if fv < ce_time_claimed:
                            prog_imp = max(0, ce_time_claimed - max(0, fv))

                st.session_state[_CE_KEY].append({
                    "ref":                ce_ref,
                    "title":              ce_title,
                    "clause":             ce_clause,
                    "status":             ce_status,
                    "activity_id":        ce_act,
                    "date_notified":      str(ce_notified),
                    "time_claimed_days":  ce_time_claimed,
                    "time_agreed_days":   ce_time_agreed,
                    "programme_impact":   prog_imp,
                    "description":        ce_desc,
                })
                st.rerun()

    # -- IMPACT ANALYSIS TAB ---------------------------------------------------
    with tab_analysis:
        if not ces:
            st.info("Add CEs to the register to see impact analysis.")
        else:
            st.markdown('<div class="section-label">Cumulative EOT Position</div>',
                        unsafe_allow_html=True)

            # Bar chart of claimed vs agreed per CE
            chart_data = pd.DataFrame([{
                "CE":             c.get("ref",""),
                "Time Claimed":   safe_float(c.get("time_claimed_days",0),0),
                "EOT Agreed":     safe_float(c.get("time_agreed_days",0),0),
                "Prog Impact":    safe_float(c.get("programme_impact",0),0),
            } for c in ces if c.get("status","") not in ("Withdrawn","Not a CE")])

            if not chart_data.empty:
                fig_ce = go.Figure()
                fig_ce.add_trace(go.Bar(
                    x=chart_data["CE"], y=chart_data["Time Claimed"],
                    name="Claimed", marker_color="#E8951D", opacity=0.7,
                ))
                fig_ce.add_trace(go.Bar(
                    x=chart_data["CE"], y=chart_data["EOT Agreed"],
                    name="Agreed EOT", marker_color="#1E7A4E",
                ))
                fig_ce.add_trace(go.Bar(
                    x=chart_data["CE"], y=chart_data["Prog Impact"],
                    name="Programme Impact", marker_color="#C0392B", opacity=0.8,
                ))
                fig_ce.update_layout(
                    barmode="group",
                    title="CE Time Impact — Claimed vs Agreed vs Programme Impact",
                    plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
                    height=340, margin=dict(l=10,r=10,t=40,b=10),
                    legend=dict(orientation="h", y=1.08),
                    yaxis_title="Days",
                )
                st.plotly_chart(fig_ce)

            # Running totals
            total_claim  = sum(safe_float(c.get("time_claimed_days",0),0) for c in ces)
            total_agreed = sum(safe_float(c.get("time_agreed_days",0),0)
                               for c in ces if "Agreed" in c.get("status",""))
            total_impact = sum(safe_float(c.get("programme_impact",0),0) for c in ces)
            total_pending= sum(safe_float(c.get("time_claimed_days",0),0)
                               for c in ces if "Notified" in c.get("status","")
                               or "Assessment" in c.get("status",""))

            c1,c2,c3,c4 = st.columns(4)
            c1.markdown(kpi_card("Total Claimed", f"{int(total_claim)}d", "calendar days", "amber"), unsafe_allow_html=True)
            c2.markdown(kpi_card("EOT Agreed",    f"{int(total_agreed)}d","calendar days", "green"), unsafe_allow_html=True)
            c3.markdown(kpi_card("Pending",        f"{int(total_pending)}d","under assessment","amber"), unsafe_allow_html=True)
            c4.markdown(kpi_card("Prog Impact",    f"{int(total_impact)}d","net delay to completion",
                                 "red" if total_impact>0 else "green"), unsafe_allow_html=True)

            planner_note(
                "Programme impact is automatically calculated as: max(0, time_claimed − float_available). "
                "If float absorbs the full claim the programme impact is 0. "
                "Where float is negative, the full claim adds to the delay position."
            )

    # -- Export ----------------------------------------------------------------
    if ces:
        exp_rows = [{
            "CE Ref":             c.get("ref",""),
            "Title":              c.get("title",""),
            "NEC4 Clause":        c.get("clause",""),
            "Status":             c.get("status",""),
            "Activity ID":        c.get("activity_id",""),
            "Date Notified":      c.get("date_notified",""),
            "Time Claimed (d)":   int(safe_float(c.get("time_claimed_days",0),0)),
            "EOT Agreed (d)":     int(safe_float(c.get("time_agreed_days",0),0)),
            "Programme Impact (d)":int(safe_float(c.get("programme_impact",0),0)),
            "Notes":              c.get("description",""),
        } for c in ces]
        xls = export_df_to_excel({"CE Register": pd.DataFrame(exp_rows)})
        st.download_button(
            "Export CE Register",
            xls,
            f"ce_register_{datetime.now().strftime('%Y%m%d')}.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="ce_export",
        )



# -----------------------------------------------------------------------------
# SIDEBAR
# -----------------------------------------------------------------------------


# -- Palette ------------------------------------------------------------------
_NAVY    = "#07111F"
_INK     = "#0F1E2E"
_AMBER   = "#E8951D"
_TEAL_OLD = "#1AADAA"
_RED     = "#C84040"
_GREEN   = "#1E7E52"
_TEXT    = "#1a2332"
_MUTED   = "#6B7280"

# -- Nav items -----------------------------------------------------------------
_NAV = [
    ("overview",   "Overview"),
    ("executive",  "Executive Summary"),
    ("programme",  "Programme"),
    ("logic",      "Logic"),
    ("critical",   "Critical Path"),
    ("labour",     "Labour"),
    ("health",     "Health Check"),
    ("comparison", "Comparison"),
    ("pm_actions", "PM Actions"),
    ("risk",       "Risk Register"),
    ("reports",    "Reports"),
    ("settings",   "Settings"),
]

_NEEDS_PROG = {
    "overview","executive","programme","logic","critical",
    "labour","health","pm_actions","risk","reports",
}



# -- Palette (single source of truth) -----------------------------------------
_NAVY    = "#0B1929"
_NAVY2   = "#0d2035"
_AMBER   = "#E8951D"
_RED     = "#C0392B"
_GREEN   = "#1E7A4E"
_TEXT    = "#1C2B3A"
_MUTED   = "#6B7C8E"


def _sidebar() -> tuple:
    """Dark navy sidebar. Logo on dark background. Nav text visible."""
    with st.sidebar:

        # -- Brand -------------------------------------------------------------
        logo_html = _logo_b64(width=52)
        st.markdown(
            f'<div style="padding:16px 14px 14px;border-bottom:1px solid #0d2035;">'
            f'<div style="display:flex;align-items:center;gap:10px;">'
            f'<div style="background:#0d2035;border-radius:4px;padding:4px;'
            f'display:flex;align-items:center;justify-content:center;flex-shrink:0;">'
            f'{logo_html}'
            f'</div>'
            f'<div>'
            f'<div style="font-size:15px;font-weight:700;color:#FFFFFF;'
            f'letter-spacing:-0.2px;line-height:1.1;">'
            f'PlanTrace</div>'
            f'<div style="font-size:9px;color:#2D4050;text-transform:uppercase;'
            f'letter-spacing:1.8px;margin-top:3px;'
            f'font-family:\'IBM Plex Mono\',monospace;">Planning Intelligence</div>'
            f'</div></div></div>',
            unsafe_allow_html=True,
        )

        # -- Upload ------------------------------------------------------------
        st.markdown(
            '<div style="padding:12px 14px 6px;">'
            '<div style="font-size:10px;color:#2D4050;font-weight:700;'
            'letter-spacing:1px;text-transform:uppercase;margin-bottom:6px;'
            'font-family:\'IBM Plex Mono\',monospace;">Programme File</div>',
            unsafe_allow_html=True,
        )
        xer_file = st.file_uploader(
            "XER", type=["xer"],
            label_visibility="collapsed",
            key="sidebar_xer_upload",
            help="Export from P6: File > Export > Primavera P6 XER",
        )
        st.markdown("</div>", unsafe_allow_html=True)

        if xer_file is not None:
            ck = f"xer_{xer_file.name}_{xer_file.size}"
            if st.session_state.get("_xer_cache_key") != ck:
                with st.spinner("Parsing..."):
                    try:
                        parsed = parse_xer(xer_file.read())
                        st.session_state["programme"]      = parsed
                        st.session_state["_xer_cache_key"] = ck
                        st.session_state["_xer_filename"]  = xer_file.name
                    except Exception as e:
                        st.error(f"Parse error: {e}")
                        st.session_state.pop("programme", None)

        # -- Status ------------------------------------------------------------
        if "programme" in st.session_state:
            prog   = st.session_state["programme"]
            fname  = st.session_state.get("_xer_filename", "")
            ntasks = len(prog.get("tasks_df", []))
            nrels  = len(prog.get("relationships_df", []))
            ddate  = prog.get("project_info", {}).get("data_date")
            pname  = prog.get("project_info", {}).get("name", "")
            dd_s   = format_date(ddate) if ddate else "—"
            pname_html = (
                f'<div style="font-size:10px;color:#3A5265;margin-bottom:3px;'
                f'white-space:nowrap;overflow:hidden;text-overflow:ellipsis;"'
                f' title="{pname}">{pname}</div>' if pname else ""
            )
            st.markdown(
                f'<div style="margin:4px 14px 10px;background:#0d2035;'
                f'border:1px solid #142840;border-radius:4px;padding:10px 12px;">'
                f'<div style="display:flex;align-items:center;gap:5px;margin-bottom:5px;">'
                f'<div style="width:5px;height:5px;border-radius:50%;'
                f'background:{_GREEN};flex-shrink:0;"></div>'
                f'<span style="font-size:9px;font-weight:700;color:{_GREEN};'
                f'text-transform:uppercase;letter-spacing:1px;'
                f'font-family:\'IBM Plex Mono\',monospace;">Loaded</span>'
                f'</div>'
                f'<div style="font-size:11px;font-weight:600;color:#7A9AB0;'
                f'white-space:nowrap;overflow:hidden;text-overflow:ellipsis;"'
                f' title="{fname}">{fname}</div>'
                f'{pname_html}'
                f'<div style="display:grid;grid-template-columns:1fr 1fr;'
                f'gap:4px;margin-top:7px;">'
                f'<div style="background:#080f1a;border-radius:3px;padding:5px 8px;">'
                f'<div style="font-size:9px;color:#253A50;text-transform:uppercase;'
                f'letter-spacing:0.8px;font-family:\'IBM Plex Mono\',monospace;">Activities</div>'
                f'<div style="font-size:16px;font-weight:700;color:{_AMBER};'
                f'font-family:\'IBM Plex Mono\',monospace;letter-spacing:-0.5px;">'
                f'{ntasks:,}</div>'
                f'</div>'
                f'<div style="background:#080f1a;border-radius:3px;padding:5px 8px;">'
                f'<div style="font-size:9px;color:#253A50;text-transform:uppercase;'
                f'letter-spacing:0.8px;font-family:\'IBM Plex Mono\',monospace;">Rels</div>'
                f'<div style="font-size:16px;font-weight:700;color:{_AMBER};'
                f'font-family:\'IBM Plex Mono\',monospace;letter-spacing:-0.5px;">'
                f'{nrels:,}</div>'
                f'</div></div>'
                f'<div style="font-size:10px;color:#2D4050;margin-top:6px;'
                f'font-family:\'IBM Plex Mono\',monospace;">'
                f'Data date <span style="color:#3A5870;">{dd_s}</span></div>'
                f'</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                f'<div style="margin:4px 14px 10px;background:#0d2035;'
                f'border:1px solid #0f2840;border-radius:4px;padding:10px 12px;">'
                f'<div style="display:flex;align-items:center;gap:5px;">'
                f'<div style="width:5px;height:5px;border-radius:50%;'
                f'background:#253A50;flex-shrink:0;"></div>'
                f'<span style="font-size:9px;font-weight:700;color:#2D4050;'
                f'text-transform:uppercase;letter-spacing:1px;'
                f'font-family:\'IBM Plex Mono\',monospace;">No Programme</span>'
                f'</div>'
                f'<div style="font-size:11px;color:#1a2d40;margin-top:5px;'
                f'line-height:1.5;">Upload a .xer file above.</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

        # -- Navigation --------------------------------------------------------
        st.markdown(
            '<div style="padding:8px 14px 4px;">'
            '<div style="font-size:10px;color:#1a2d40;font-weight:700;'
            'letter-spacing:1px;text-transform:uppercase;margin-bottom:4px;'
            'font-family:\'IBM Plex Mono\',monospace;">Navigation</div>'
            '</div>',
            unsafe_allow_html=True,
        )

        prog_loaded = "programme" in st.session_state
        if "nav_page" not in st.session_state:
            st.session_state["nav_page"] = "overview"
        current = st.session_state["nav_page"]

        prog_loaded = "programme" in st.session_state
        if "nav_page" not in st.session_state:
            st.session_state["nav_page"] = "overview"
        current = st.session_state["nav_page"]

        for group_name, group_items in _NAV_GROUPS:
            group_html = (
                '<div style="padding:10px 14px 2px;">'
                '<div style="font-size:9px;color:#1a3550;font-weight:700;'
                "letter-spacing:1.2px;text-transform:uppercase;"
                "font-family:IBM Plex Mono,monospace;\">"
                + group_name +
                "</div></div>"
            )
            st.markdown(group_html, unsafe_allow_html=True)
            for key, label in group_items:
                is_active   = (current == key)
                is_disabled = (not prog_loaded) and (key in _NEEDS_PROG)
                if is_active:
                    div_style = "background:#0d2035;border-left:2px solid " + _AMBER + ";margin:0;"
                elif is_disabled:
                    div_style = "margin:0;opacity:0.22;"
                else:
                    div_style = "margin:0;"
                st.markdown("<div style=\"" + div_style + "\">", unsafe_allow_html=True)
                if st.button(label, key="nav_" + key,
                             use_container_width=True, disabled=is_disabled):
                    st.session_state["nav_page"] = key
                    st.rerun()
                st.markdown("</div>", unsafe_allow_html=True)


        # -- Settings ----------------------------------------------------------
        st.markdown(
            '<div style="padding:12px 14px 4px;">'
            '<div style="font-size:10px;color:#1a2d40;font-weight:700;'
            'letter-spacing:1px;text-transform:uppercase;margin-bottom:6px;'
            'font-family:\'IBM Plex Mono\',monospace;">Settings</div>'
            '</div>',
            unsafe_allow_html=True,
        )
        near_crit_days = st.slider(
            "Near-Critical Float (days)",
            min_value=1, max_value=30, value=10, step=1,
        )

        # -- View Mode ---------------------------------------------------------
        st.markdown(
            '<div style="padding:4px 14px 4px;">'
            '<div style="font-size:10px;color:#1a2d40;font-weight:700;'
            'letter-spacing:1px;text-transform:uppercase;margin-bottom:4px;'
            'font-family:\'IBM Plex Mono\',monospace;">View Mode</div>'
            '</div>',
            unsafe_allow_html=True,
        )
        _cur_mode = st.session_state.get(_MODE_KEY, "PM Mode")
        _new_mode = st.selectbox(
            "view_mode", ["PM Mode", "Planner Mode"],
            index=0 if _cur_mode == "PM Mode" else 1,
            key="sidebar_mode_sel",
            label_visibility="collapsed",
        )
        if _new_mode != _cur_mode:
            st.session_state[_MODE_KEY] = _new_mode
            st.rerun()

        # -- Footer ------------------------------------------------------------
        st.markdown(
            '<div style="padding:12px 14px;border-top:1px solid #0d2035;'
            'margin-top:12px;">'
            '<div style="font-size:10px;color:#1a2d40;line-height:1.8;'
            'font-family:\'IBM Plex Mono\',monospace;">'
            'P6 &rarr; File &rarr; Export &rarr; XER</div>'
            '</div>',
            unsafe_allow_html=True,
        )

    return st.session_state["nav_page"], near_crit_days


def _landing_page():
    """Clean, professional landing page. Navy + white + amber only."""
    # Full-width dark header bar
    logo_html = _logo_b64(width=40)
    st.markdown(
        f'<div style="background:{_NAVY};padding:16px 32px;'
        f'margin:-28px -32px 28px -32px;border-bottom:2px solid {_AMBER};">'
        f'<div style="display:flex;align-items:center;gap:12px;">'
        f'<div style="background:#0d2035;border-radius:4px;padding:4px;">'
        f'{logo_html}</div>'
        f'<div>'
        f'<div style="font-size:18px;font-weight:700;color:#FFFFFF;letter-spacing:-0.2px;">'
        f'PlanTrace</div>'
        f'<div style="font-size:11px;color:#4A6070;margin-top:1px;">'
        f'Planning intelligence for project delivery teams</div>'
        f'</div></div>'
        f'</div>',
        unsafe_allow_html=True,
    )

    # Status strip
    def _status_pill(label, value, ready):
        dot_col = _GREEN if ready else "#3A5265"
        val_col = "#FFFFFF" if ready else "#3A5265"
        return (
            f'<div style="background:#0d2035;border:1px solid #142840;'
            f'border-radius:4px;padding:8px 14px;white-space:nowrap;">'
            f'<div style="display:flex;align-items:center;gap:5px;margin-bottom:2px;">'
            f'<div style="width:5px;height:5px;border-radius:50%;'
            f'background:{dot_col};"></div>'
            f'<span style="font-size:9px;font-weight:700;color:#2D4050;'
            f'text-transform:uppercase;letter-spacing:1px;'
            f'font-family:\'IBM Plex Mono\',monospace;">{label}</span>'
            f'</div>'
            f'<div style="font-size:12px;font-weight:600;color:{val_col};'
            f'font-family:\'IBM Plex Mono\',monospace;">{value}</div>'
            f'</div>'
        )

    pills = (
        _status_pill("Programme",    "Not Loaded", False) +
        _status_pill("Relationships","—",          False) +
        _status_pill("Resources",    "—",          False) +
        _status_pill("Notes",        "Not Loaded", False) +
        _status_pill("Comparison",   "Not Loaded", False)
    )
    st.markdown(
        f'<div style="background:{_NAVY};border-radius:6px;padding:14px 18px;'
        f'margin-bottom:28px;display:flex;gap:8px;flex-wrap:wrap;">'
        f'{pills}</div>',
        unsafe_allow_html=True,
    )

    # Hero — left aligned, no centering
    hero_col, upload_col = st.columns([3, 2], gap="large")

    with hero_col:
        st.markdown(
            f'<div style="padding-top:4px;">'
            f'<div style="font-size:11px;font-weight:700;color:{_AMBER};'
            f'text-transform:uppercase;letter-spacing:1.5px;margin-bottom:12px;'
            f'font-family:\'IBM Plex Mono\',monospace;">PlanTrace Control Centre</div>'
            f'<div style="font-size:34px;font-weight:700;color:{_NAVY};'
            f'line-height:1.1;letter-spacing:-0.5px;margin-bottom:14px;">'
            f'Programme intelligence<br>without opening P6.</div>'
            f'<div style="font-size:14px;color:{_MUTED};line-height:1.7;'
            f'max-width:480px;margin-bottom:24px;">'
            f'Upload a Primavera P6 XER programme to interrogate logic, critical paths, '
            f'labour demand, schedule quality and programme movement. No licence required.'
            f'</div>'
            f'<div style="display:flex;gap:8px;flex-wrap:wrap;">'
            f'<span style="background:{_NAVY};color:#8AAABF;font-size:11px;'
            f'font-weight:600;padding:4px 10px;border-radius:3px;">Logic Trace</span>'
            f'<span style="background:{_NAVY};color:#8AAABF;font-size:11px;'
            f'font-weight:600;padding:4px 10px;border-radius:3px;">Critical Path</span>'
            f'<span style="background:{_NAVY};color:#8AAABF;font-size:11px;'
            f'font-weight:600;padding:4px 10px;border-radius:3px;">Labour Demand</span>'
            f'<span style="background:{_NAVY};color:#8AAABF;font-size:11px;'
            f'font-weight:600;padding:4px 10px;border-radius:3px;">Health Check</span>'
            f'<span style="background:{_NAVY};color:#8AAABF;font-size:11px;'
            f'font-weight:600;padding:4px 10px;border-radius:3px;">PM Actions</span>'
            f'<span style="background:{_NAVY};color:#8AAABF;font-size:11px;'
            f'font-weight:600;padding:4px 10px;border-radius:3px;">Risk Register</span>'
            f'</div></div>',
            unsafe_allow_html=True,
        )

    with upload_col:
        st.markdown(
            f'<div style="background:{_NAVY};border-radius:6px;padding:22px 24px;'
            f'border-left:3px solid {_AMBER};">'
            f'<div style="font-size:10px;font-weight:700;color:{_AMBER};'
            f'text-transform:uppercase;letter-spacing:1.2px;margin-bottom:10px;'
            f'font-family:\'IBM Plex Mono\',monospace;">Upload Programme</div>'
            f'<div style="font-size:13px;color:#4A6070;margin-bottom:14px;line-height:1.6;">'
            f'Export from Primavera P6 via<br>'
            f'<strong style="color:#8AAABF;">'
            f'File &rarr; Export &rarr; Primavera P6 XER</strong>'
            f'</div>'
            f'<div style="font-size:12px;color:#3A5265;padding:10px 12px;'
            f'background:#0d2035;border-radius:4px;border:1px dashed #142840;'
            f'line-height:1.6;">'
            f'Use the file uploader in the sidebar on the left to upload your .xer file.'
            f'</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    # Module cards — ALL use navy accent, no colour spray
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="section-label">Product Modules</div>', unsafe_allow_html=True)

    modules = [
        ("Logic",         "Logic Trace",
         "Trace predecessor and successor chains through the schedule network.",
         ["Predecessor chains", "Successor chains", "Full logic chain", "Network depth", "Export"]),
        ("Critical Path", "Critical Path Analysis",
         "Identify critical path, near-critical work and negative float.",
         ["Full critical path", "Near-critical work", "Negative float", "Driving path", "Gantt view"]),
        ("Labour",        "Labour Demand",
         "View labour histograms by week, month, resource and WBS.",
         ["Weekly histogram", "Monthly histogram", "Resource breakdown", "WBS view", "Peak demand"]),
        ("Health",        "Programme Health Check",
         "Automated quality checks covering logic, constraints and float.",
         ["Open logic", "Constraints", "Long durations", "Float issues", "11 checks"]),
    ]

    cols = st.columns(4, gap="medium")
    for col, (tag, title, desc, outputs) in zip(cols, modules):
        outputs_html = "".join(
            f'<div style="display:flex;align-items:center;gap:7px;padding:4px 0;'
            f'font-size:12px;color:#4A5A68;border-bottom:1px solid #F0F4F8;">'
            f'<div style="width:4px;height:4px;border-radius:50%;'
            f'background:#8AAABF;flex-shrink:0;"></div>{o}</div>'
            for o in outputs
        )
        col.markdown(
            f'<div style="background:#FFFFFF;border:1px solid #D4DCE4;'
            f'border-radius:6px;border-top:3px solid {_NAVY};'
            f'padding:0;overflow:hidden;'
            f'box-shadow:0 1px 3px rgba(11,25,41,0.06);height:100%;">'
            f'<div style="padding:16px 18px 14px;">'
            f'<div style="font-size:10px;font-weight:700;color:#8A9DB0;'
            f'text-transform:uppercase;letter-spacing:1px;margin-bottom:6px;'
            f'font-family:\'IBM Plex Mono\',monospace;">{tag}</div>'
            f'<div style="font-size:15px;font-weight:700;color:{_NAVY};'
            f'margin-bottom:6px;">{title}</div>'
            f'<div style="font-size:12px;color:#6B7C8E;line-height:1.5;'
            f'margin-bottom:12px;">{desc}</div>'
            f'<div style="font-size:10px;font-weight:700;color:#9AA8B5;'
            f'text-transform:uppercase;letter-spacing:0.8px;margin-bottom:6px;'
            f'font-family:\'IBM Plex Mono\',monospace;">Outputs</div>'
            f'<div>{outputs_html}</div>'
            f'</div></div>',
            unsafe_allow_html=True,
        )

    # Module directory
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="section-label">All Modules</div>', unsafe_allow_html=True)
    page_list = [
        ("Overview",       "KPI summary, PM attention panel and schedule health."),
        ("Executive Summary", "Plain-English narrative, Top 5 points, recommended actions."),
        ("Programme",      "Activity search, lookahead planner, milestones, planning notes."),
        ("Logic",          "Logic trace and path to selected activity."),
        ("Critical Path",  "Critical path, near-critical and negative float."),
        ("Labour",         "Labour histograms by week, month, resource and WBS."),
        ("Health Check",   "Eleven automated schedule quality checks."),
        ("Comparison",     "Programme movement and risk & opportunity register."),
        ("PM Actions",     "Auto-generated prioritised action list."),
        ("Risk Register",  "Auto-generated risk and opportunity register."),
        ("Reports",        "Export all data to formatted Excel workbooks."),
    ]
    lc, rc = st.columns(2, gap="large")
    for i, (mod_title, desc) in enumerate(page_list):
        col = lc if i % 2 == 0 else rc
        col.markdown(
            f'<div style="display:flex;gap:10px;padding:8px 0;'
            f'border-bottom:1px solid #EEF1F5;align-items:flex-start;">'
            f'<div style="width:3px;min-width:3px;height:16px;background:{_NAVY};'
            f'border-radius:2px;margin-top:2px;"></div>'
            f'<div>'
            f'<div style="font-weight:600;color:{_NAVY};font-size:13px;">{mod_title}</div>'
            f'<div style="font-size:12px;color:{_MUTED};line-height:1.5;">{desc}</div>'
            f'</div></div>',
            unsafe_allow_html=True,
        )

    st.markdown(
        f'<div style="margin-top:32px;padding-top:16px;border-top:1px solid #D4DCE4;">'
        f'<span style="font-size:11px;color:#8A9DB0;">'
        f'PlanTrace &nbsp;·&nbsp; Planning intelligence for project delivery '
        f'&nbsp;·&nbsp; No Primavera P6 licence required'
        f'</span></div>',
        unsafe_allow_html=True,
    )



# -----------------------------------------------------------------------------
# Guard
# -----------------------------------------------------------------------------

def _guard() -> bool:
    if "programme" not in st.session_state:
        empty_state(
            "",
            "No Programme Loaded",
            "Upload a .xer file using the Programme File uploader in the sidebar "
            "to unlock this page.",
            "Upload in Sidebar",
        )
        return True
    return False


# -----------------------------------------------------------------------------
# MAIN
# -----------------------------------------------------------------------------

def main():
    active, near_crit_days = _sidebar()
    prog_loaded = "programme" in st.session_state
    data = st.session_state.get("programme", {})

    # -- OVERVIEW ------------------------------------------------------------
    if active == "executive":
        if not prog_loaded:
            _landing_page()
        else:
            page_executive_summary(data, near_crit_days)
        return

    if active == "overview":
        if not prog_loaded:
            _landing_page()
        else:
            ctrl_bar("Overview",
                     "Programme summary, attention items and schedule health at a glance.")
            mode_toggle_bar()
            tasks = data.get("tasks_df", pd.DataFrame())
            if not tasks.empty:
                tc = get_critical_threshold(tasks, near_crit_days)
                n_crit = int(tc["is_critical"].sum()) if "is_critical" in tc.columns else 0
                n_nc   = int(tc["is_near_critical"].sum()) if "is_near_critical" in tc.columns else 0
                n_neg  = int(tc["total_float_days"].apply(lambda f: safe_float(f,0)<0).sum()) if "total_float_days" in tc.columns else 0
                n_rels = len(data.get("relationships_df", pd.DataFrame()))
                kpi_row([
                    ("Total Activities",  len(tasks), "",                   "navy"),
                    ("Critical",          n_crit,     "float <= 0",         "red"   if n_crit  else "green"),
                    ("Near-Critical",     n_nc,       f"float <= {near_crit_days}d","amber" if n_nc    else "green"),
                    ("Negative Float",    n_neg,      "beyond target date", "red"   if n_neg   else "green"),
                    ("Relationships",     n_rels,     "",                   "blue"),
                ])
                st.markdown("<br>", unsafe_allow_html=True)

            programme_readiness_check(data)
            _ov_col1, _ov_col2 = st.columns([1,2])
            with _ov_col1:
                render_activity_status_doughnut(data.get("tasks_df", pd.DataFrame()))
            with _ov_col2:
                render_float_distribution_bar(data.get("tasks_df", pd.DataFrame()), near_crit_days)
                render_health_trend(near_crit_days)


            with st.expander("WBS Risk Heatmap", expanded=False):
                render_wbs_heatmap(data, near_crit_days)

            left_col, right_col = st.columns([3,1], gap="large")
            with left_col:
                pm_attention_panel(data, near_crit_days)
            with right_col:
                data_quality_card(data)

            page_project_summary(data, near_crit_days)

    # -- PROGRAMME ------------------------------------------------------------
    elif active == "programme":
        ctrl_bar("Programme",
                 "Search activities, plan lookahead, track milestones and review planning notes.")
        mode_toggle_bar()
        pm_note("Use Activity Search to find any activity. Lookahead shows what is coming up in the next 2 to 12 weeks. Milestones tracks your key programme dates.")
        if _guard(): return
        tabs = st.tabs(["Activity Search","Lookahead","Milestones","Planning Notes"])
        with tabs[0]: page_activity_search(data, near_crit_days)
        with tabs[1]: page_lookahead(data, near_crit_days)
        with tabs[2]: page_milestone_tracker(data, near_crit_days)
        with tabs[3]: page_planning_notes(data)

    # -- LOGIC ----------------------------------------------------------------
    elif active == "logic":
        ctrl_bar("Logic",
                 "Trace predecessor and successor chains through the programme network.")
        mode_toggle_bar()
        pm_note("Logic Trace shows what is driving an activity (its predecessors) and what it will impact if delayed (its successors). Use this to understand cause and effect in the schedule.")
        if _guard(): return
        tabs = st.tabs(["Logic Trace","Path to Selected Activity"])
        with tabs[0]: page_logic_trace(data, near_crit_days)
        with tabs[1]: page_critical_path_to_activity(data, near_crit_days)

    # -- CRITICAL PATH --------------------------------------------------------
    elif active == "critical":
        ctrl_bar("Critical Path",
                 "Full critical path, near-critical activities and negative float analysis.")
        mode_toggle_bar()
        if _guard(): return
        page_critical_path(data, near_crit_days)
        with st.expander("Critical Path Gantt", expanded=False):
            render_critical_path_gantt(data.get("tasks_df", pd.DataFrame()), near_crit_days)

    # -- LABOUR ---------------------------------------------------------------
    elif active == "labour":
        ctrl_bar("Labour",
                 "Labour demand histograms by week, month, resource and WBS.")
        if _guard(): return
        page_labour_histogram(data)

    # -- HEALTH CHECK ---------------------------------------------------------
    elif active == "health":
        ctrl_bar("Health Check",
                 "Automated schedule quality checks covering logic, float, constraints and durations.")
        mode_toggle_bar()
        pm_note("These checks automatically find common schedule problems. Each issue shows a count and explains why it matters for the project.")
        if _guard(): return
        page_health_check(data, near_crit_days)

    # -- COMPARISON -----------------------------------------------------------
    elif active == "dcma":
        if _guard(): return
        page_dcma_assessment(data, near_crit_days)

    elif active == "client_review":
        if _guard(): return
        page_client_review(data, near_crit_days)

    elif active == "acceleration":
        if _guard(): return
        page_acceleration_analysis(data, near_crit_days)

    elif active == "ce_tracker":
        if _guard(): return
        page_ce_tracker(data, near_crit_days)

    elif active == "forensic":
        ctrl_bar("Forensic Planner", "")
        if _guard(): return
        data = st.session_state.get("programme", {})
        page_forensic_planner(data, near_crit_days)

    elif active == "comparison":
        ctrl_bar("Comparison",
                 "Programme movement intelligence between two XER revisions.")
        tabs = st.tabs(["Programme Movement","Risk & Opportunity"])
        with tabs[0]: page_programme_comparison()
        with tabs[1]:
            if not _guard(): page_risk_register(data, near_crit_days)

    # -- PM ACTIONS -----------------------------------------------------------
    elif active == "pm_actions":
        ctrl_bar("PM Actions",
                 "Auto-generated prioritised action list based on programme analysis.")
        if _guard(): return
        page_pm_actions(data, near_crit_days)

    # -- RISK REGISTER --------------------------------------------------------
    elif active == "risk":
        ctrl_bar("Risk Register",
                 "Auto-generated risk and opportunity register from programme data.")
        if _guard(): return
        page_risk_register(data, near_crit_days)

    # -- REPORTS --------------------------------------------------------------
    elif active == "reports":
        ctrl_bar("Reports",
                 "Export all programme data and analysis to formatted Excel workbooks.")
        mode_toggle_bar()
        if _guard(): return
        page_export_reports(data, near_crit_days)

    # -- NEW PAGES --------------------------------------------------------
    elif active == "scurve":
        if _guard(): return
        page_scurve(data, near_crit_days)

    elif active == "resources":
        if _guard(): return
        page_resource_levelling(data, near_crit_days)

    elif active == "milestones":
        if _guard(): return
        page_contract_milestones(data, near_crit_days)

    elif active == "multiproject":
        page_multi_project(near_crit_days)

    # -- SETTINGS --------------------------------------------------------
    elif active == "settings":
        page_settings(data, near_crit_days)
